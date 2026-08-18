from __future__ import annotations

import re
import shutil
import unicodedata
from collections import defaultdict
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.field_utils import compact_dropdown_options
from app.placeholder_scanner import PLACEHOLDER_PATTERN, scan_docx_fields
from app.smart_template import suggest_field_type
from app.word_control_utils import classify_native_control, get_control_identifier
from app.document_understanding import (
    annotate_document_records,
    postprocess_candidates,
    semantic_label,
    semantic_section,
)


# The automatic detector is deliberately conservative. Explicit tags and Word
# form controls remain authoritative; this module only proposes additions.
X_PLACEHOLDER_PATTERN = re.compile(
    r"(?<![A-Za-z0-9])(?:\(\d{2}\)\s*)?[Xx]{4,}(?:\s*[@./()\-]\s*[Xx]{2,})*(?![A-Za-z0-9])"
)
# Underscore masks are often short (``UF: __`` or ``Banco: ___``) and can
# include punctuation (``__/__/____``, ``___.___.___-__``).  Match the whole
# visual mask instead of only the longest underscore fragment so the inserted
# tag replaces the complete fill area.
UNDERSCORE_PLACEHOLDER_PATTERN = re.compile(
    r"(?<![A-Za-z0-9])_{2,}(?:\s*[/.:\-]\s*_{2,})*(?![A-Za-z0-9])"
)
ZERO_PHONE_PLACEHOLDER_PATTERN = re.compile(
    r"(?<!\d)\(\s*0{2}\s*\)\s*0{4,5}\s*-\s*0{4}(?!\d)"
)
ZERO_CPF_PLACEHOLDER_PATTERN = re.compile(
    r"(?<!\d)0{3}\s*\.\s*0{3}\s*\.\s*0{3}\s*-\s*0{2}(?!\d)"
)
# Monetary fill masks occur frequently in institutional documents as
# ``R$ XXX.XXX,XX``, ``R$ XXX.XXX.XX``, ``R$ 000.000,00`` or underline
# variants. Match the currency prefix together with the visual mask so the
# generated currency value does not leave a fixed ``R$`` behind and become
# ``R$ R$ 1.000,00``. Decimal/group punctuation is deliberately tolerant
# because hand-authored templates are not always consistent.
CURRENCY_PLACEHOLDER_PATTERN = re.compile(
    r"(?i)(?<![A-Za-z0-9])R\$\s*"
    r"(?:"
    r"[Xx]{2,}(?:\s*[.,]\s*[Xx]{2,3}){0,3}"
    r"|0{2,}(?:\s*[.,]\s*0{2,3}){0,3}"
    r"|_{2,}(?:\s*[.,]\s*_{2,3}){0,3}"
    r")"
    r"(?![A-Za-z0-9])"
)
SAMPLE_EMAIL_PLACEHOLDER_PATTERN = re.compile(
    r"(?i)(?<![A-Za-z0-9._%+\-])"
    r"(?:contato|email|e-mail|exemplo|teste|usuario|usu[aá]rio|user|x{4,})"
    r"@(?:empresa|exemplo|example|dominio|dom[ií]nio|x{4,})"
    r"(?:\.[A-Za-zx]{2,}){1,3}"
    r"(?![A-Za-z0-9._%+\-])"
)
# Legacy/custom form documents often use a single-braced token such as
# ``{descricao.demanda}`` instead of Padroniza's authoritative ``{{...}}``
# syntax.  Treat these as assisted placeholders only when context strongly
# supports a field interpretation; ordinary prose using braces must remain
# static. Unicode word characters are accepted here because Brazilian forms
# frequently contain accents inside these legacy markers.
LEGACY_BRACED_PLACEHOLDER_PATTERN = re.compile(
    r"(?<!\{)\{([^\W\d_][\w.-]{0,95})\}(?!\})",
    re.UNICODE,
)
CHOICE_SEPARATOR_PATTERN = re.compile(r"^\s*OU\s*$", re.IGNORECASE)
INSTRUCTION_PATTERN = re.compile(
    r"^\s*(?:informar|informe|descrever|descreva|detalhar|detalhe|"
    r"indicar|indique|justificar|justifique|preencher|preencha)\b",
    re.IGNORECASE,
)
GENERIC_DROPDOWN_PATTERN = re.compile(
    r"^\s*(?:escolher|selecione|selecionar)\s+(?:um|uma)\s+(?:item|op[cç][aã]o)\.?\s*$",
    re.IGNORECASE,
)
CHECKBOX_LINE_PATTERN = re.compile(r"^\s*(?:☐|□|☑|☒|\(\s*\))\s*(.+?)\s*$")
CHECKBOX_TOKEN_PATTERN = re.compile(r"(?:☐|□|☑|☒|\(\s*\))")
# Checked Word forms are sometimes rendered as a bare check mark in an otherwise
# empty narrow cell instead of a Unicode checked-box character. Keep these
# glyphs restricted to the isolated-cell heuristic so they are not mistaken for
# ordinary mathematical or prose characters elsewhere in the document.
ISOLATED_CHECK_MARK_PATTERN = re.compile(r"(?:✓|✔|√)")
FOLLOWUP_AREA_PATTERN = re.compile(
    r"^\s*(?:observa[cç][aã]o(?:\s*/\s*justificativa)?|justificativa|"
    r"complemento|detalhamento|informa[cç][oõ]es? complementares?)\b",
    re.IGNORECASE,
)
SECTION_NUMBER_PATTERN = re.compile(r"^\s*\d+(?:\.\d+)*\.?\s*")
LABEL_TAIL_PATTERN = re.compile(r"([^:;|]{2,120})\s*[:：]\s*$")


_SOURCE_LABELS = {
    "long_choice": "Alternativas separadas por OU",
    "repeatable_table": "Tabela com linhas repetíveis",
    "inline_placeholder": "Texto de preenchimento (XXXX ou sublinhado)",
    "legacy_placeholder": "Marcador legado entre chaves simples",
    "instruction": "Texto instrucional substituível",
    "empty_cell": "Célula vazia ao lado de um rótulo",
    "dropdown_prompt": "Indicação 'Escolher um item'",
    "sample_value": "Valor de exemplo após o rótulo",
    "checkbox_choice": "Opções com caixas de seleção",
    "checkbox_single": "Caixa de seleção independente",
    "consistency_repair": "Reparo por consistência do formulário",
    "prefilled_text": "Texto existente possivelmente editável",
}


class AutomaticDetectionError(ValueError):
    """Raised when accepted automatic detections cannot be applied safely."""


class _ParagraphRecord:
    __slots__ = (
        "ordinal",
        "paragraph",
        "text",
        "story",
        "table_index",
        "row_index",
        "cell_index",
        "cell",
        "table",
        "understanding",
    )

    def __init__(
        self,
        *,
        ordinal: int,
        paragraph: Paragraph,
        story: str,
        table_index: int | None = None,
        row_index: int | None = None,
        cell_index: int | None = None,
        cell: _Cell | None = None,
        table: Table | None = None,
    ) -> None:
        self.ordinal = ordinal
        self.paragraph = paragraph
        self.text = paragraph.text or ""
        self.story = story
        self.table_index = table_index
        self.row_index = row_index
        self.cell_index = cell_index
        self.cell = cell
        self.table = table
        self.understanding = None


def detect_docx_field_candidates(
    docx_path: Path,
    *,
    existing_field_ids: Iterable[str] | None = None,
    existing_fields: Iterable[dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    """Return conservative fill-field suggestions for an untagged DOCX.

    The result is safe to display in a review dialog. It does not modify the
    document. Explicit ``{{tags}}`` and native Word controls are excluded from
    automatic replacement candidates.
    """

    path = Path(docx_path)
    if not path.exists() or not path.is_file() or path.suffix.lower() != ".docx":
        raise AutomaticDetectionError("Selecione um arquivo DOCX válido.")

    document = Document(str(path))
    records = _collect_paragraph_records(document)
    annotate_document_records(records)
    by_ordinal = {record.ordinal: record for record in records}

    existing_field_list = [
        dict(field)
        for field in (existing_fields or [])
        if isinstance(field, dict)
    ]
    known_ids = {
        str(field_id).strip()
        for field_id in (existing_field_ids or [])
        if str(field_id).strip()
    }
    known_ids.update(
        str(field.get("id", "")).strip()
        for field in existing_field_list
        if str(field.get("id", "")).strip()
    )
    try:
        scanned_existing_fields = [
            dict(field)
            for field in scan_docx_fields(path)
            if isinstance(field, dict)
            and str(field.get("id", "")).strip()
        ]
        known_ids.update(
            str(field.get("id", "")).strip()
            for field in scanned_existing_fields
        )

        # Keep the scanner authoritative even when callers only provide IDs or
        # no existing metadata at all.  Richer editor metadata wins for the
        # same ID; newly scanned explicit/native fields fill the gaps.
        existing_ids = {
            str(field.get("id", "")).strip()
            for field in existing_field_list
            if str(field.get("id", "")).strip()
        }
        for field in scanned_existing_fields:
            field_id = str(field.get("id", "")).strip()
            if field_id not in existing_ids:
                existing_field_list.append(field)
                existing_ids.add(field_id)
    except Exception:
        # Automatic detection should still be usable when an unrelated
        # malformed native control exists. The normal scanner will report
        # that issue before the model can be saved.
        pass

    candidates: list[dict[str, Any]] = []
    reserved_ordinals: set[int] = set()

    long_choices = _detect_long_choice_blocks(
        document,
        records,
        known_ids,
    )
    for candidate in long_choices:
        candidates.append(candidate)
        reserved_ordinals.update(
            int(value)
            for value in candidate.get("location", {}).get("paragraphs", [])
        )
        known_ids.add(str(candidate.get("field_id", "")))

    repeatable_tables = _detect_repeatable_tables(
        document,
        records,
        known_ids,
        reserved_ordinals,
    )
    for candidate in repeatable_tables:
        candidates.append(candidate)
        reserved_ordinals.update(
            int(value)
            for value in candidate.get("location", {}).get("paragraphs", [])
        )
        known_ids.add(str(candidate.get("field_id", "")))

    editable_sheets = _detect_editable_sheet_tables(
        document,
        records,
        known_ids,
        reserved_ordinals,
    )
    for candidate in editable_sheets:
        candidates.append(candidate)
        reserved_ordinals.update(
            int(value)
            for value in candidate.get("location", {}).get("paragraphs", [])
        )
        known_ids.add(str(candidate.get("field_id", "")))

    checkbox_choices = _detect_checkbox_choice_groups(
        document,
        records,
        known_ids,
        reserved_ordinals,
    )
    for candidate in checkbox_choices:
        candidates.append(candidate)
        reserved_ordinals.update(
            int(value)
            for value in candidate.get("location", {}).get("paragraphs", [])
        )
        for field in candidate.get("fields", []) or []:
            known_ids.add(str(field.get("id", "")))

    single_checkboxes = _detect_standalone_checkboxes(
        records,
        known_ids,
        reserved_ordinals,
    )
    for candidate in single_checkboxes:
        candidates.append(candidate)
        reserved_ordinals.add(int(candidate.get("location", {}).get("paragraph", -1)))
        known_ids.add(str(candidate.get("field_id", "")))

    followup_areas = _detect_blank_followup_areas(
        records,
        known_ids,
        reserved_ordinals,
    )
    for candidate in followup_areas:
        candidates.append(candidate)
        known_ids.add(str(candidate.get("field_id", "")))

    for record in records:
        if record.ordinal in reserved_ordinals:
            continue
        if _contains_authoritative_marker(record.paragraph):
            continue

        dropdown_prompt = _detect_dropdown_prompt(
            record,
            records,
            known_ids,
        )
        if dropdown_prompt is not None:
            candidates.append(dropdown_prompt)
            known_ids.add(str(dropdown_prompt.get("field_id", "")))
            continue

        sample_value = _detect_labeled_sample_value(
            record,
            known_ids,
        )
        if sample_value is not None:
            candidates.append(sample_value)
            known_ids.add(str(sample_value.get("field_id", "")))
            continue

        labeled_instruction = _detect_labeled_instruction(
            record,
            known_ids,
        )
        if labeled_instruction is not None:
            candidates.append(labeled_instruction)
            known_ids.add(str(labeled_instruction.get("field_id", "")))
            continue

        inline = _detect_inline_placeholders(
            record,
            records,
            known_ids,
        )
        if inline:
            candidates.extend(inline)
            known_ids.update(str(item.get("field_id", "")) for item in inline)
            continue

        text = _normalize_space(record.text)
        if not text:
            continue

        if _is_instruction_candidate(record):
            label = _context_label_for_record(record, records)
            field_type = "multiline" if len(text) >= 70 else suggest_field_type(label or text)
            if field_type == "text" and len(text) >= 70:
                field_type = "multiline"
            field_id = _unique_field_id(
                _make_field_id(label or text[:60]),
                known_ids,
            )
            known_ids.add(field_id)
            candidates.append(
                _candidate(
                    field_id=field_id,
                    label=label or _instruction_label(text),
                    field_type=field_type,
                    confidence=0.84 if _paragraph_is_red(record.paragraph) else 0.76,
                    source="instruction",
                    preview=text,
                    location={
                        "kind": "paragraph",
                        "paragraph": record.ordinal,
                    },
                )
            )
            continue

        prefilled_text = _detect_prefilled_written_text(
            record,
            records,
            known_ids,
        )
        if prefilled_text is not None:
            candidates.append(prefilled_text)
            known_ids.add(str(prefilled_text.get("field_id", "")))
            continue

        adjacent_sample = _detect_adjacent_sample_value(
            record,
            records,
            known_ids,
        )
        if adjacent_sample is not None:
            candidates.append(adjacent_sample)
            known_ids.add(str(adjacent_sample.get("field_id", "")))
            continue

        label_only = _detect_label_only_field(
            record,
            records,
            known_ids,
        )
        if label_only is not None:
            candidates.append(label_only)
            known_ids.add(str(label_only.get("field_id", "")))

    candidates.extend(
        _detect_empty_cells(
            document,
            records,
            known_ids,
            reserved_ordinals,
        )
    )

    candidates.extend(
        _detect_consistency_repair_fields(
            records,
            candidates,
            known_ids,
        )
    )

    source_kind = (
        "pdf_reconstruction"
        if "convertido de pdf" in str(document.core_properties.subject or "").casefold()
        else "docx"
    )
    candidates = postprocess_candidates(candidates, records, source_kind=source_kind)
    candidates = _suppress_authoritative_semantic_duplicates(
        candidates,
        existing_field_list,
    )

    # Stable order keeps the review screen aligned with the source document.
    candidates.sort(
        key=lambda item: (
            _candidate_first_ordinal(item),
            -float(item.get("confidence", 0.0)),
            str(item.get("field_id", "")),
        )
    )
    for index, candidate in enumerate(candidates, start=1):
        candidate["candidate_id"] = f"candidate_{index:04d}"
    return candidates



def _suppress_authoritative_semantic_duplicates(
    candidates: list[dict[str, Any]],
    existing_fields: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Drop assisted suggestions already owned by authoritative fields.

    Explicit Padroniza tags, native Word/PDF controls and manually configured
    fields are stronger than assisted detection.  This matters especially for
    a Word form laid out as ``Rótulo | {{campo}}``: the empty-cell detector can
    otherwise create ``auto.rotulo`` from the label cell even though the
    adjacent tagged cell already owns the same semantic field.

    Suppression remains conservative.  We primarily match label + section.
    When an authoritative field has no section metadata yet (common while a
    DOCX is first being scanned), label-only ownership is allowed only for a
    unique, non-generic label.
    """

    authoritative: set[tuple[str, str]] = set()
    authoritative_label_counts: dict[str, int] = {}
    sectionless_labels: set[str] = set()
    generic_label_keys = {
        "campo", "data", "nome", "tipo", "valor", "item", "quantidade",
        "descricao", "observacao", "responsavel", "matricula", "situacao",
    }

    for field in existing_fields:
        field_id = str(field.get("id", "")).strip()
        source = str(field.get("detection_source") or "").strip().casefold()

        # Existing assisted fields must not suppress a newer interpretation of
        # the document.  Everything else is authoritative here: explicit tags
        # usually have no detection_source, while native controls identify
        # themselves explicitly.
        is_assisted = source in {"automatic", "assisted", "auto_detection"}
        if not is_assisted and not source and field_id.casefold().startswith("auto."):
            # Defensive compatibility with very old auto-detected models that
            # did not persist detection_source.
            is_assisted = True
        if is_assisted:
            continue

        label_key = _slug(str(field.get("label", "")))
        if not label_key:
            continue
        section_key = _slug(str(field.get("section", "")))
        authoritative.add((label_key, section_key))
        authoritative_label_counts[label_key] = authoritative_label_counts.get(label_key, 0) + 1
        if not section_key:
            sectionless_labels.add(label_key)

    if not authoritative:
        return candidates

    unique_sectionless_labels = {
        label_key
        for label_key in sectionless_labels
        if authoritative_label_counts.get(label_key, 0) == 1
        and label_key not in generic_label_keys
        and len(label_key) >= 6
    }

    result: list[dict[str, Any]] = []
    for candidate in candidates:
        label_keys = {
            key
            for key in (
                _slug(str(candidate.get("label", ""))),
                _slug(str(candidate.get("semantic_label_suggestion", ""))),
            )
            if key
        }
        section_key = _slug(str(candidate.get("section", "")))

        duplicate = any((label_key, section_key) in authoritative for label_key in label_keys)

        # Native/explicit fields can be discovered before section inference has
        # run.  A unique descriptive label is still strong enough to own the
        # semantic field across that short metadata gap.
        if not duplicate and label_keys:
            duplicate = any(label_key in unique_sectionless_labels for label_key in label_keys)

        if duplicate:
            continue
        result.append(candidate)
    return result

def apply_docx_field_candidates(
    source_docx: Path,
    destination_docx: Path,
    candidates: Iterable[dict[str, Any]],
) -> Path:
    """Apply approved suggestions to a copy by converting them into tags.

    This is the key safety property of assisted detection: after approval the
    existing tag scanner and DOCX engine handle the model exactly like a
    manually tagged template.
    """

    source = Path(source_docx)
    destination = Path(destination_docx)
    accepted = [deepcopy(item) for item in candidates if isinstance(item, dict)]
    if not accepted:
        raise AutomaticDetectionError("Nenhuma sugestão foi selecionada.")

    destination.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, destination)
    document = Document(str(destination))
    records = _collect_paragraph_records(document)
    by_ordinal = {record.ordinal: record for record in records}

    _validate_accepted_candidates(accepted)

    # Whole-block operations are applied first. They do not invalidate the
    # stored Paragraph XML objects used by the span replacements below.
    for candidate in accepted:
        kind = str(candidate.get("location", {}).get("kind", ""))
        if kind == "repeatable_table":
            _apply_repeatable_table(candidate, document)
        elif kind == "paragraph_block":
            _apply_paragraph_block(candidate, by_ordinal)
        elif kind == "checkbox_group":
            _apply_checkbox_group(candidate, by_ordinal)
        elif kind == "checkbox_group_inline":
            _apply_inline_checkbox_group(candidate, by_ordinal)
        elif kind == "checkbox_group_multi_cell":
            _apply_multi_cell_checkbox_group(candidate, by_ordinal)

    # Apply text spans from right to left inside each paragraph so offsets stay
    # valid even when one line contains several placeholders.
    spans_by_paragraph: dict[int, list[dict[str, Any]]] = defaultdict(list)
    for candidate in accepted:
        location = candidate.get("location", {}) or {}
        if str(location.get("kind", "")) == "text_span":
            spans_by_paragraph[int(location.get("paragraph", -1))].append(candidate)

    for ordinal, paragraph_candidates in spans_by_paragraph.items():
        record = by_ordinal.get(ordinal)
        if record is None:
            raise AutomaticDetectionError(
                "A estrutura do DOCX mudou durante a detecção automática. "
                "Execute a análise novamente."
            )
        replacements: list[tuple[int, int, str]] = []
        for candidate in paragraph_candidates:
            location = candidate.get("location", {}) or {}
            replacements.append(
                (
                    int(location.get("start", 0)),
                    int(location.get("end", 0)),
                    _tag_for_candidate(candidate),
                )
            )
        _replace_paragraph_spans(record.paragraph, replacements)

    for candidate in accepted:
        location = candidate.get("location", {}) or {}
        kind = str(location.get("kind", ""))
        if kind == "append_tag":
            ordinal = int(location.get("paragraph", -1))
            record = by_ordinal.get(ordinal)
            if record is None:
                raise AutomaticDetectionError(
                    "Não foi possível localizar uma área aprovada no DOCX. "
                    "Execute a análise novamente."
                )
            _append_tag_to_paragraph(record.paragraph, _tag_for_candidate(candidate))
            continue
        if kind not in {"paragraph", "empty_cell"}:
            continue
        ordinal = int(location.get("paragraph", -1))
        record = by_ordinal.get(ordinal)
        if record is None:
            raise AutomaticDetectionError(
                "Não foi possível localizar uma área aprovada no DOCX. "
                "Execute a análise novamente."
            )
        _replace_entire_paragraph(record.paragraph, _tag_for_candidate(candidate))

    document.save(str(destination))
    return destination


def candidate_field_definitions(
    candidates: Iterable[dict[str, Any]],
) -> list[dict[str, Any]]:
    """Convert approved candidates to normal editable field definitions."""

    fields: list[dict[str, Any]] = []
    for candidate in candidates:
        if not isinstance(candidate, dict):
            continue
        if str(candidate.get("source", "")) == "repeatable_table":
            field_id = str(candidate.get("field_id", "")).strip()
            if not field_id:
                continue
            fields.append(
                {
                    "id": field_id,
                    "label": str(candidate.get("label", "")).strip() or field_id,
                    "type": "repeatable_table",
                    "columns": [
                        dict(column)
                        for column in candidate.get("columns", []) or []
                        if isinstance(column, dict)
                    ],
                    "minimum_rows": 1,
                    "numbering_padding": 2,
                    "required": True,
                    "label_source": "automatic_detection",
                    "type_source": "automatic_detection",
                    "detection_source": "automatic",
                    "detection_confidence": float(candidate.get("confidence", 0.0)),
                    "detection_confidence_band": str(candidate.get("confidence_band", "")),
                    "detection_evidence": deepcopy(candidate.get("evidence", []) or []),
                    "detector_version": int(candidate.get("detector_version", 1) or 1),
                    "full_width": True,
                }
            )
            if str(candidate.get("section", "")).strip():
                fields[-1]["section"] = str(candidate.get("section", "")).strip()
            continue
        if str(candidate.get("source", "")) == "checkbox_choice":
            for raw_field in candidate.get("fields", []) or []:
                field = dict(raw_field)
                field.setdefault("required", False)
                field.setdefault("label_source", "automatic_detection")
                field.setdefault("type_source", "automatic_detection")
                field["detection_source"] = "automatic"
                field["detection_confidence"] = float(candidate.get("confidence", 0.0))
                field["detection_confidence_band"] = str(candidate.get("confidence_band", ""))
                field["detection_evidence"] = deepcopy(candidate.get("evidence", []) or [])
                field["detector_version"] = int(candidate.get("detector_version", 1) or 1)
                fields.append(field)
            continue

        field_id = str(candidate.get("field_id", "")).strip()
        if not field_id:
            continue
        field: dict[str, Any] = {
            "id": field_id,
            "label": str(candidate.get("label", "")).strip(),
            "type": str(candidate.get("type", "text")).strip() or "text",
            "required": str(candidate.get("type", "text")) != "checkbox",
            "label_source": "automatic_detection",
            "type_source": "automatic_detection",
            "detection_source": "automatic",
            "detection_confidence": float(candidate.get("confidence", 0.0)),
            "detection_confidence_band": str(candidate.get("confidence_band", "")),
            "detection_evidence": deepcopy(candidate.get("evidence", []) or []),
            "detector_version": int(candidate.get("detector_version", 1) or 1),
        }
        options = compact_dropdown_options(candidate.get("options", []))
        if options:
            field["options"] = options
        placeholder = str(candidate.get("placeholder", "")).strip()
        if placeholder:
            field["placeholder"] = placeholder
        if "default_value" in candidate:
            field["default_value"] = deepcopy(candidate.get("default_value"))
        if str(candidate.get("layout", "")) == "choice":
            group = str(candidate.get("layout_group", f"auto_choice_{field_id}"))
            field.update(
                {
                    "layout": "choice",
                    "layout_group": group,
                    "layout_group_label": field["label"],
                    "group": group,
                    "selection": "single",
                    "choice_required": True,
                    "tag_type": "single_choice",
                }
            )
        # Automatically detected dates represent visible fill areas in the
        # source document (for example ``Data: __/__/____``). They must stay
        # editable instead of being silently replaced with today's date.
        if field["type"] == "date":
            field["automatic"] = False
        if field["type"] == "multiline":
            field["full_width"] = True
        fields.append(field)
    return fields


def candidate_source_label(candidate: dict[str, Any]) -> str:
    return _SOURCE_LABELS.get(
        str(candidate.get("source", "")),
        "Sugestão automática",
    )


def _candidate(
    *,
    field_id: str,
    label: str,
    field_type: str,
    confidence: float,
    source: str,
    preview: str,
    location: dict[str, Any],
    options: Iterable[Any] | None = None,
    default_selected: bool | None = None,
    requires_configuration: bool = False,
    layout: str = "",
    layout_group: str = "",
    placeholder: str = "",
    default_value: Any | None = None,
) -> dict[str, Any]:
    confidence = max(0.0, min(float(confidence), 1.0))
    if default_selected is None:
        default_selected = confidence >= 0.80 and not requires_configuration
    result: dict[str, Any] = {
        "field_id": field_id,
        "label": _normalize_space(label) or field_id,
        "type": field_type,
        "confidence": confidence,
        "source": source,
        "preview": _normalize_space(preview)[:420],
        "location": dict(location),
        "selected": bool(default_selected),
        "requires_configuration": bool(requires_configuration),
    }
    cleaned_options = compact_dropdown_options(options or [])
    if cleaned_options:
        result["options"] = cleaned_options
    if layout:
        result["layout"] = layout
    if layout_group:
        result["layout_group"] = layout_group
    if str(placeholder).strip():
        result["placeholder"] = str(placeholder).strip()
    if default_value is not None:
        result["default_value"] = deepcopy(default_value)
    return result


def _collect_paragraph_records(document: _Document) -> list[_ParagraphRecord]:
    records: list[_ParagraphRecord] = []
    seen_cells: set[int] = set()

    def add_paragraph(
        paragraph: Paragraph,
        *,
        story: str,
        table_index: int | None = None,
        row_index: int | None = None,
        cell_index: int | None = None,
        cell: _Cell | None = None,
        table: Table | None = None,
    ) -> None:
        records.append(
            _ParagraphRecord(
                ordinal=len(records),
                paragraph=paragraph,
                story=story,
                table_index=table_index,
                row_index=row_index,
                cell_index=cell_index,
                cell=cell,
                table=table,
            )
        )

    def walk_table(table: Table, *, story: str, index: int) -> None:
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                for paragraph in cell.paragraphs:
                    add_paragraph(
                        paragraph,
                        story=story,
                        table_index=index,
                        row_index=row_index,
                        cell_index=cell_index,
                        cell=cell,
                        table=table,
                    )
                for nested in cell.tables:
                    nonlocal_table_index[0] += 1
                    walk_table(
                        nested,
                        story=story,
                        index=nonlocal_table_index[0],
                    )

    nonlocal_table_index = [-1]
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            add_paragraph(Paragraph(child, document), story="body")
        elif child.tag == qn("w:tbl"):
            nonlocal_table_index[0] += 1
            walk_table(Table(child, document), story="body", index=nonlocal_table_index[0])

    # Headers and footers are included for ordinary placeholder suggestions,
    # but long choice-block detection remains focused on body tables.
    for section_index, section in enumerate(document.sections):
        for story_name, story in (
            (f"header_{section_index}", section.header),
            (f"footer_{section_index}", section.footer),
        ):
            for paragraph in story.paragraphs:
                add_paragraph(paragraph, story=story_name)
            for table in story.tables:
                nonlocal_table_index[0] += 1
                walk_table(table, story=story_name, index=nonlocal_table_index[0])

    return records


def _detect_long_choice_blocks(
    document: _Document,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> list[dict[str, Any]]:
    del document
    by_cell: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if record.cell is not None and record.story == "body":
            by_cell[id(record.cell._tc)].append(record)

    result: list[dict[str, Any]] = []
    for cell_records in by_cell.values():
        cell_records.sort(key=lambda item: item.ordinal)
        texts = [_normalize_space(record.text) for record in cell_records]
        separators = [
            index
            for index, text in enumerate(texts)
            if CHOICE_SEPARATOR_PATTERN.match(text)
        ]
        if len(separators) < 2:
            continue
        if any(_contains_authoritative_marker(record.paragraph) for record in cell_records):
            continue

        segment_ranges: list[tuple[int, int]] = []
        start = 0
        for separator in separators:
            segment_ranges.append((start, separator))
            start = separator + 1
        segment_ranges.append((start, len(cell_records)))

        options: list[dict[str, str]] = []
        first_content_index: int | None = None
        last_content_index: int | None = None
        for range_start, range_end in segment_ranges:
            segment_records = [
                record
                for record in cell_records[range_start:range_end]
                if _normalize_space(record.text)
            ]
            if not segment_records:
                continue
            value = _normalize_space("\n".join(record.text for record in segment_records))
            if len(value) < 5:
                continue
            # Avoid interpreting page headers repeated inside a malformed cell
            # as an option.
            if _looks_like_page_header(value):
                continue
            label = _short_choice_label(value)
            options.append({"label": label, "value": value})
            segment_first = cell_records.index(segment_records[0])
            segment_last = cell_records.index(segment_records[-1])
            first_content_index = (
                segment_first
                if first_content_index is None
                else min(first_content_index, segment_first)
            )
            last_content_index = (
                segment_last
                if last_content_index is None
                else max(last_content_index, segment_last)
            )

        if len(options) < 3 or len(options) > 10:
            continue
        if first_content_index is None or last_content_index is None:
            continue

        included = cell_records[first_content_index : last_content_index + 1]
        included_ordinals = [record.ordinal for record in included]
        # Include separators between the first and last option in the replaced
        # block, but never remove text outside that block.
        label = _context_label_for_record(included[0], records)
        field_id = _unique_field_id(
            _make_field_id(label or "justificativa"),
            known_ids,
        )
        group = f"auto_choice_{field_id}"
        result.append(
            _candidate(
                field_id=field_id,
                label=label or "Escolha uma alternativa",
                field_type="dropdown",
                confidence=0.94,
                source="long_choice",
                preview=" OU ".join(option["label"] for option in options),
                location={
                    "kind": "paragraph_block",
                    "paragraphs": included_ordinals,
                },
                options=options,
                layout="choice",
                layout_group=group,
            )
        )
    return result


def _detect_repeatable_tables(
    document: _Document,
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Detect conservative numbered tables that represent repeated records.

    The automatic detector only proposes a repeatable table when the evidence
    is strong: a header row, at least two numbered data rows, and at least two
    editable columns.  This intentionally avoids turning questionnaire tables
    (which also have repeated visual rows) into repeatable item editors.
    """

    top_level_index = {
        id(table._tbl): index
        for index, table in enumerate(document.tables)
    }
    by_table: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    table_refs: dict[int, Table] = {}
    for record in records:
        if record.story != "body" or record.table is None or record.table_index is None:
            continue
        table_key = id(record.table._tbl)
        if table_key not in top_level_index:
            continue
        by_table[table_key].append(record)
        table_refs[table_key] = record.table

    result: list[dict[str, Any]] = []
    for table_key, table_records in by_table.items():
        table = table_refs[table_key]
        if len(table.rows) < 3:
            continue
        if any(
            record.ordinal in reserved_ordinals
            or _contains_authoritative_marker(record.paragraph)
            for record in table_records
        ):
            continue

        header_cells = _unique_row_cells(table.rows[0])
        if len(header_cells) < 3:
            continue
        headers = [_clean_label(cell.text) for cell in header_cells]
        if sum(_is_reasonable_label(value, maximum=80) for value in headers) < 3:
            continue

        data_rows: list[int] = []
        number_values: list[int] = []
        for row_index in range(1, len(table.rows)):
            cells = _unique_row_cells(table.rows[row_index])
            if len(cells) != len(header_cells):
                break
            number_text = _normalize_space(cells[0].text)
            if not re.fullmatch(r"0*\d{1,4}", number_text):
                break
            data_rows.append(row_index)
            number_values.append(int(number_text))

        if len(data_rows) < 2:
            continue
        expected = list(range(number_values[0], number_values[0] + len(number_values)))
        if number_values != expected:
            continue

        number_header = _slug(headers[0])
        if number_header not in {"n", "no", "numero", "item", "n_item"}:
            continue

        editable_columns: list[int] = []
        for column_index in range(1, len(header_cells)):
            values = [
                _normalize_space(_unique_row_cells(table.rows[row_index])[column_index].text)
                for row_index in data_rows
            ]
            if any(_looks_like_fill_area_text(value) or not value for value in values):
                editable_columns.append(column_index)
        if len(editable_columns) < 2:
            continue

        first_record = min(table_records, key=lambda item: item.ordinal)
        section_title = _nearest_section_title(first_record, records, preserve_number=True)
        label = _clean_label(section_title) if section_title else "Itens da tabela"
        field_id = _unique_field_id(_make_field_id(label), known_ids)

        columns: list[dict[str, Any]] = [
            {
                "id": "item",
                "label": headers[0] or "Item",
                "type": "auto_number",
                "required": False,
            }
        ]
        used_column_ids = {"item"}
        for column_index in editable_columns:
            header = headers[column_index] or f"Coluna {column_index + 1}"
            column_values = [
                _normalize_space(_unique_row_cells(table.rows[row_index])[column_index].text)
                for row_index in data_rows
            ]
            column_id = _slug(header) or f"coluna_{column_index + 1}"
            base_column_id = column_id
            suffix = 2
            while column_id in used_column_ids:
                column_id = f"{base_column_id}_{suffix}"
                suffix += 1
            used_column_ids.add(column_id)
            columns.append(
                {
                    "id": column_id,
                    "label": header,
                    "type": _repeatable_column_type(header, column_values),
                    "required": True,
                    "column_index": column_index,
                }
            )

        # Region ownership: once this physical table segment is classified as a
        # repeatable table, its header and model rows belong to that high-level
        # interpretation.  Lower-level detectors must not reinterpret header
        # cells such as ``Unidade | Quantidade`` as ordinary label/value fields.
        #
        # Keep the ownership information explicit in the candidate as a second
        # safety layer: post-processing can suppress overlapping interpretations
        # even when a future detector runs before the reservation pass.
        header_row = 0
        owned_rows = [header_row, *data_rows]
        data_ordinals = sorted(
            record.ordinal
            for record in table_records
            if record.row_index in data_rows
        )
        owned_ordinals = sorted(
            record.ordinal
            for record in table_records
            if record.row_index in owned_rows
        )
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type="repeatable_table",
                confidence=0.95,
                source="repeatable_table",
                preview=" | ".join(headers)
                + f" — {len(data_rows)} linha(s) modelo detectada(s)",
                location={
                    "kind": "repeatable_table",
                    "document_table_index": top_level_index[table_key],
                    "table_index": first_record.table_index,
                    "header_row": header_row,
                    "template_row": data_rows[0],
                    "data_rows": data_rows,
                    "owned_rows": owned_rows,
                    "data_paragraphs": data_ordinals,
                    "owned_paragraphs": owned_ordinals,
                    # ``paragraphs`` is the generic reservation contract used
                    # by the rest of the detector. Include the complete owned
                    # region, not only the rows that will receive tags.
                    "paragraphs": owned_ordinals,
                },
            )
        )
        result[-1]["region_owner"] = "repeatable_table"
        result[-1]["columns"] = columns
        if section_title:
            result[-1]["section"] = section_title.rstrip(":").strip()
        result[-1]["selected"] = True

    return result




def _detect_editable_sheet_tables(
    document: _Document,
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Detect a spreadsheet header that has no editable data row yet.

    A common institutional Word pattern is a multi-column header followed by a
    merged narrative row, for example::

        Item | Quantidade | Unidade | Especificação | Valor
        [ merged explanatory / prefilled paragraph                  ]

    The header describes a real worksheet even though the source document does
    not provide numbered model rows.  Earlier detector versions preserved the
    visual header but offered no editable cells.  This detector creates a
    repeatable-table interpretation with a *synthetic* model row inserted
    between the header and the merged note when the approved suggestions are
    applied.  All header columns are editable; the merged note remains available
    to the normal prefilled-text detector as a separate full-width field.

    The rule is intentionally narrow: at least three short header cells are
    required and the immediately following row must be one merged/full-width
    cell containing either substantial prose or an empty/fill area.  This keeps
    ordinary label/value form grids out of the spreadsheet path.
    """

    top_level_index = {
        id(table._tbl): index
        for index, table in enumerate(document.tables)
    }
    by_table: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    table_refs: dict[int, Table] = {}
    for record in records:
        if record.story != "body" or record.table is None or record.table_index is None:
            continue
        table_key = id(record.table._tbl)
        if table_key not in top_level_index:
            continue
        by_table[table_key].append(record)
        table_refs[table_key] = record.table

    result: list[dict[str, Any]] = []
    for table_key, table_records in by_table.items():
        table = table_refs[table_key]
        if len(table.rows) < 2:
            continue

        for header_row_index in range(0, len(table.rows) - 1):
            header_cells = _unique_row_cells(table.rows[header_row_index])
            if len(header_cells) < 3:
                continue
            headers = [_clean_label(cell.text) for cell in header_cells]
            if sum(_is_reasonable_label(value, maximum=100) for value in headers) < 3:
                continue
            if any(len(value) > 120 for value in headers if value):
                continue

            next_row_index = header_row_index + 1
            next_cells = _unique_row_cells(table.rows[next_row_index])
            if len(next_cells) != 1:
                continue
            merged_text = _normalize_space(next_cells[0].text)
            if merged_text:
                # Short merged rows are often totals, signatures, or section
                # separators.  Long prose (or a visual fill area) is the sheet
                # + merged-note pattern we want.
                if len(merged_text) < 55 and not _looks_like_fill_area_text(merged_text):
                    continue

            header_records = [
                record
                for record in table_records
                if record.row_index == header_row_index
            ]
            if not header_records:
                continue
            if any(
                record.ordinal in reserved_ordinals
                or _contains_authoritative_marker(record.paragraph)
                for record in header_records
            ):
                continue

            first_record = min(header_records, key=lambda item: item.ordinal)
            section_title = _nearest_section_title(
                first_record,
                records,
                preserve_number=True,
            )
            label = _clean_label(section_title) if section_title else "Itens da planilha"
            field_id = _unique_field_id(_make_field_id(label), known_ids)

            used_column_ids: set[str] = set()
            columns: list[dict[str, Any]] = []
            for column_index, header in enumerate(headers):
                display = header or f"Coluna {column_index + 1}"
                column_id = _slug(display) or f"coluna_{column_index + 1}"
                base_column_id = column_id
                suffix = 2
                while column_id in used_column_ids:
                    column_id = f"{base_column_id}_{suffix}"
                    suffix += 1
                used_column_ids.add(column_id)

                header_key = _slug(display)
                if header_key in {"item", "codigo", "código"}:
                    # ``Item`` in user-authored spreadsheets is not assumed to
                    # be an automatic row number; the user can edit it.
                    column_type = "text"
                elif any(token in header_key for token in ("descricao", "especificacao", "detalhamento")):
                    column_type = "multiline"
                elif header_key in {"valor", "preco", "preço", "custo", "montante"} or header_key.endswith("_valor"):
                    column_type = "currency"
                else:
                    column_type = _repeatable_column_type(display, [])

                columns.append(
                    {
                        "id": column_id,
                        "label": display,
                        "type": column_type,
                        "required": False,
                        "column_index": column_index,
                    }
                )

            if len(columns) < 3:
                continue

            owned_ordinals = sorted(record.ordinal for record in header_records)
            candidate = _candidate(
                field_id=field_id,
                label=label,
                field_type="repeatable_table",
                confidence=0.93,
                source="repeatable_table",
                preview=" | ".join(headers) + " — planilha editável detectada",
                location={
                    "kind": "repeatable_table",
                    "document_table_index": top_level_index[table_key],
                    "table_index": first_record.table_index,
                    "header_row": header_row_index,
                    # No source model row exists. _apply_repeatable_table will
                    # create one immediately before the merged narrative row.
                    "template_row": -1,
                    "synthetic_template_row": True,
                    "insert_before_row": next_row_index,
                    "data_rows": [],
                    "owned_rows": [header_row_index],
                    "owned_paragraphs": owned_ordinals,
                    "paragraphs": owned_ordinals,
                },
            )
            candidate["region_owner"] = "repeatable_table"
            candidate["sheet_generated_model_row"] = True
            candidate["columns"] = columns
            candidate["minimum_rows"] = 1
            candidate["numbering_padding"] = 2
            candidate["selected"] = True
            if section_title:
                candidate["section"] = section_title.rstrip(":").strip()
            result.append(candidate)
            # A single Word table should have one primary sheet header for this
            # pattern.  Stop after the first convincing match.
            break

    return result


def _detect_checkbox_choice_groups(
    document: _Document,
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    del document
    result: list[dict[str, Any]] = []
    used_ordinals: set[int] = set(reserved_ordinals)

    # 1) Several checkbox options on the same visual line/cell, e.g.
    # ``Natureza: ☐ Material  ☐ Serviço  ☐ Material e serviço``.
    for record in records:
        if record.ordinal in used_ordinals or _contains_authoritative_marker(record.paragraph):
            continue
        parsed = _inline_checkbox_options(record.text or "")
        if parsed is None:
            continue
        prefix, options, token_spans = parsed
        label = _local_label(prefix) or _context_label_for_record(record, records)
        candidate = _checkbox_candidate(
            label=label,
            options=options,
            known_ids=known_ids,
            confidence=0.92,
            location={
                "kind": "checkbox_group_inline",
                "paragraph": record.ordinal,
                "checkbox_spans": [list(span) for span in token_spans],
            },
        )
        result.append(candidate)
        used_ordinals.add(record.ordinal)

    # 1b) One checkbox option in each cell of the same visual row.  Word forms
    # often use this for alternatives such as ``Entrega imediata`` versus
    # ``Entrega parcelada``.  The cell may also contain explanatory text after
    # a manual line break, so the ordinary whole-line checkbox regex cannot
    # safely recognize it.  Keep that explanatory text in the document and
    # use it as additional context in the client-facing option label.
    by_row: dict[tuple[str, int, int], list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if (
            record.table_index is None
            or record.row_index is None
            or record.ordinal in used_ordinals
            or _contains_authoritative_marker(record.paragraph)
        ):
            continue
        by_row[(record.story, int(record.table_index), int(record.row_index))].append(record)

    for row_records in by_row.values():
        parsed_cells: list[tuple[_ParagraphRecord, str, tuple[int, int]]] = []
        seen_cells: set[int] = set()
        for record in sorted(row_records, key=lambda item: (item.cell_index or 0, item.ordinal)):
            if record.cell is None:
                continue
            cell_key = id(record.cell._tc)
            if cell_key in seen_cells:
                continue
            parsed = _single_checkbox_cell_option(record.text or "")
            if parsed is None:
                continue
            option_label, token_span = parsed
            seen_cells.add(cell_key)
            parsed_cells.append((record, option_label, token_span))

        if len(parsed_cells) < 2 or len(parsed_cells) > 6:
            continue

        label = _nearest_section_title(parsed_cells[0][0], records)
        if not label:
            label = _context_label_for_record(parsed_cells[0][0], records)
        candidate = _checkbox_candidate(
            label=label,
            options=[option_label for _record, option_label, _span in parsed_cells],
            known_ids=known_ids,
            confidence=0.94,
            location={
                "kind": "checkbox_group_multi_cell",
                "paragraphs": [record.ordinal for record, _label, _span in parsed_cells],
                "checkbox_spans": [list(span) for _record, _label, span in parsed_cells],
            },
        )
        result.append(candidate)
        used_ordinals.update(record.ordinal for record, _label, _span in parsed_cells)

    # 1c) Checkbox marker isolated in a narrow cell with the option text in
    # the adjacent cell. This pattern is common in institutional forms where
    # the left column contains only a square and the right column contains a
    # numbered occurrence/condition plus explanatory text.
    by_table_rows: dict[int, dict[int, list[_ParagraphRecord]]] = defaultdict(
        lambda: defaultdict(list)
    )
    for record in records:
        if (
            record.table_index is None
            or record.row_index is None
            or record.ordinal in used_ordinals
        ):
            continue
        by_table_rows[int(record.table_index)][int(record.row_index)].append(record)

    for rows in by_table_rows.values():
        row_options: dict[
            int,
            tuple[_ParagraphRecord, str, tuple[int, int] | None, str],
        ] = {}
        row_cells: dict[int, dict[int, list[_ParagraphRecord]]] = {}
        explicit_marker_columns: set[int] = set()

        # First pass: collect rows whose marker is structurally present in the
        # narrow cell.  This is the normal case (Unicode marker, Word control,
        # symbol or drawing).
        for row_index, row_records in rows.items():
            by_cell_index: dict[int, list[_ParagraphRecord]] = defaultdict(list)
            for record in row_records:
                if record.cell_index is None:
                    continue
                by_cell_index[int(record.cell_index)].append(record)
            row_cells[row_index] = by_cell_index

            ordered_cells = sorted(by_cell_index)
            for cell_index in ordered_cells:
                marker_records = sorted(
                    by_cell_index[cell_index],
                    key=lambda item: item.ordinal,
                )
                marker = _isolated_checkbox_marker(marker_records)
                if marker is None:
                    continue
                marker_record, token_span, marker_mode = marker

                # Require the immediately adjacent visual cell. This avoids
                # pairing decorative checkboxes with unrelated text elsewhere
                # in a wide row.
                adjacent_records = sorted(
                    by_cell_index.get(cell_index + 1, []),
                    key=lambda item: item.ordinal,
                )
                option_text = _adjacent_checkbox_option_text(adjacent_records)
                if not option_text:
                    continue

                row_options[row_index] = (
                    marker_record,
                    option_text,
                    token_span,
                    marker_mode,
                )
                explicit_marker_columns.add(cell_index)
                break

        # Second pass: some institutional DOCX files draw a checked square as
        # an absolutely-positioned floating text box.  Visually it sits in the
        # narrow marker cell, but the actual table cell is completely empty.
        # Once this table has established a checkbox marker column from another
        # row, an empty cell in that same column followed by option-like text is
        # a strong signal that the row belongs to the same choice group.
        #
        # This intentionally does *not* infer a checkbox in arbitrary empty
        # cells: a structural marker must already exist in the same table and
        # the marker column must be narrow relative to its adjacent text cell.
        for row_index, by_cell_index in row_cells.items():
            if row_index in row_options:
                continue
            for cell_index in sorted(explicit_marker_columns):
                marker_records = sorted(
                    by_cell_index.get(cell_index, []),
                    key=lambda item: item.ordinal,
                )
                adjacent_records = sorted(
                    by_cell_index.get(cell_index + 1, []),
                    key=lambda item: item.ordinal,
                )
                if not marker_records or not adjacent_records:
                    continue
                if not _is_blank_checkbox_marker_cell(marker_records, adjacent_records):
                    continue
                option_text = _adjacent_checkbox_option_text(adjacent_records)
                if not option_text or not _looks_like_adjacent_choice_option(adjacent_records):
                    continue

                # Reuse the real blank paragraph in the marker cell. During
                # application it will be replaced with the normal checkbox tag.
                row_options[row_index] = (
                    marker_records[0],
                    option_text,
                    None,
                    "inferred_blank",
                )
                break

        ordered_rows = sorted(row_options)
        runs: list[list[int]] = []
        current: list[int] = []
        for row_index in ordered_rows:
            if current and row_index != current[-1] + 1:
                if len(current) >= 2:
                    runs.append(current)
                current = []
            current.append(row_index)
        if len(current) >= 2:
            runs.append(current)

        for row_run in runs:
            matched = [row_options[row_index] for row_index in row_run]
            if len(matched) > 8:
                matched = matched[:8]
            if any(record.ordinal in used_ordinals for record, _label, _span, _mode in matched):
                continue

            first_record = matched[0][0]
            label = _adjacent_checkbox_group_label(first_record, records)
            inferred_count = sum(
                1 for _record, _label, _span, mode in matched if mode == "inferred_blank"
            )
            candidate = _checkbox_candidate(
                label=label,
                options=[option for _record, option, _span, _mode in matched],
                known_ids=known_ids,
                confidence=0.89 if inferred_count else 0.93,
                location={
                    "kind": "checkbox_group_multi_cell",
                    "paragraphs": [record.ordinal for record, _label, _span, _mode in matched],
                    "checkbox_spans": [
                        list(span) if span is not None else [-1, -1]
                        for _record, _label, span, _mode in matched
                    ],
                    "checkbox_marker_modes": [
                        mode for _record, _label, _span, mode in matched
                    ],
                    "inferred_blank_markers": inferred_count,
                },
            )
            result.append(candidate)
            used_ordinals.update(record.ordinal for record, _label, _span, _mode in matched)

    # 2) Several checkbox paragraphs inside the same Word cell.
    by_cell: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if record.cell is not None and record.ordinal not in used_ordinals:
            by_cell[id(record.cell._tc)].append(record)

    for cell_records in by_cell.values():
        matched: list[tuple[_ParagraphRecord, re.Match[str]]] = []
        for record in cell_records:
            if record.ordinal in used_ordinals or _contains_authoritative_marker(record.paragraph):
                continue
            match = CHECKBOX_LINE_PATTERN.match(record.text or "")
            if match:
                matched.append((record, match))
        if len(matched) < 2 or len(matched) > 8:
            continue

        label = _context_label_for_record(matched[0][0], records)
        candidate = _checkbox_candidate(
            label=label,
            options=[_normalize_space(match.group(1)) for _record, match in matched],
            known_ids=known_ids,
            confidence=0.90,
            location={
                "kind": "checkbox_group",
                "paragraphs": [record.ordinal for record, _match in matched],
            },
        )
        result.append(candidate)
        used_ordinals.update(record.ordinal for record, _match in matched)

    # 3) A checkbox on each consecutive row of a simple one-column table.
    # This is common for declarations/acknowledgements and previously made an
    # entire section disappear from automatic detection.
    by_table: dict[int, dict[int, list[_ParagraphRecord]]] = defaultdict(lambda: defaultdict(list))
    for record in records:
        if (
            record.table_index is None
            or record.row_index is None
            or record.ordinal in used_ordinals
        ):
            continue
        by_table[int(record.table_index)][int(record.row_index)].append(record)

    for rows in by_table.values():
        matching_rows: dict[int, tuple[_ParagraphRecord, re.Match[str]]] = {}
        for row_index, row_records in rows.items():
            nonempty = [record for record in row_records if _normalize_space(record.text)]
            matches: list[tuple[_ParagraphRecord, re.Match[str]]] = []
            for record in nonempty:
                if _contains_authoritative_marker(record.paragraph):
                    continue
                match = CHECKBOX_LINE_PATTERN.match(record.text or "")
                if match:
                    matches.append((record, match))
            if len(nonempty) == 1 and len(matches) == 1:
                matching_rows[row_index] = matches[0]

        ordered_rows = sorted(matching_rows)
        run: list[int] = []
        runs: list[list[int]] = []
        for row_index in ordered_rows:
            if run and row_index != run[-1] + 1:
                if len(run) >= 2:
                    runs.append(run)
                run = []
            run.append(row_index)
        if len(run) >= 2:
            runs.append(run)

        for row_run in runs:
            matched = [matching_rows[row_index] for row_index in row_run]
            if len(matched) > 8:
                matched = matched[:8]
            if any(record.ordinal in used_ordinals for record, _match in matched):
                continue
            label = _context_label_for_record(matched[0][0], records)
            candidate = _checkbox_candidate(
                label=label,
                options=[_normalize_space(match.group(1)) for _record, match in matched],
                known_ids=known_ids,
                confidence=0.87,
                location={
                    "kind": "checkbox_group",
                    "paragraphs": [record.ordinal for record, _match in matched],
                },
            )
            result.append(candidate)
            used_ordinals.update(record.ordinal for record, _match in matched)

    return result


def _detect_standalone_checkboxes(
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Detect one independent checkbox embedded in an otherwise meaningful line.

    Institutional forms commonly use a declaration such as
    ``Declaro que ... ☐ Li e concordo``.  The multi-option detector correctly
    ignores it because there is only one checkbox, but the checkbox is still a
    real user input.  Keep this rule intentionally narrow: exactly one visible
    checkbox token, meaningful surrounding text, and no authoritative tag or
    Word control.
    """

    result: list[dict[str, Any]] = []
    for record in records:
        if record.ordinal in reserved_ordinals or _contains_authoritative_marker(record.paragraph):
            continue
        text = str(record.text or "")
        matches = list(CHECKBOX_TOKEN_PATTERN.finditer(text))
        if len(matches) != 1:
            continue
        match = matches[0]
        before = _normalize_space(text[: match.start()]).strip(" :：;|–—-")
        after = _normalize_space(text[match.end() :]).strip(" :：;|–—-")
        if not before and not after:
            continue
        if len(before) > 240 or len(after) > 160:
            continue

        # A single marker that is merely decorative beside an empty area is too
        # ambiguous.  Require actual declaration/option wording on at least one
        # side of the marker.
        semantic = after or before
        if len(semantic) < 2:
            continue

        if before and after:
            label = f"{before.rstrip('.;:')} — {after}"
        else:
            label = after or before
        label = _clean_label(label)
        if not label:
            continue

        field_id = _unique_field_id(_make_field_id(label), known_ids)
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type="checkbox",
                confidence=0.92,
                source="checkbox_single",
                preview=_normalize_space(text),
                location={
                    "kind": "text_span",
                    "paragraph": record.ordinal,
                    "start": match.start(),
                    "end": match.end(),
                    "original": match.group(0),
                },
            )
        )
    return result


def _isolated_checkbox_marker(
    records: list[_ParagraphRecord],
) -> tuple[_ParagraphRecord, tuple[int, int] | None, str] | None:
    """Return an isolated checkbox marker from a narrow table cell.

    Real-world Word forms use several representations for the same visual
    square: Unicode box characters, bare check marks, unnamed content-control
    checkboxes, legacy form fields, Wingdings symbols and sometimes a small
    drawing/VML shape. This heuristic is intentionally limited to the narrow
    marker cell beside an option-description cell.
    """

    # Controls are inspected before visible text so a named Word control can
    # never be duplicated by automatic detection merely because its displayed
    # result happens to look like a checkbox character.
    for record in records:
        element = record.paragraph._p
        unnamed_controls = 0
        named_controls = 0
        for sdt in element.xpath(".//w:sdt"):
            properties = sdt.find(qn("w:sdtPr"))
            if properties is None:
                continue
            control_type, _control = classify_native_control(properties)
            if control_type != "checkbox":
                continue
            if get_control_identifier(sdt):
                named_controls += 1
            else:
                unnamed_controls += 1
        if named_controls:
            return None
        if unnamed_controls == 1:
            return record, None, "paragraph"

        unnamed_legacy = 0
        named_legacy = 0
        for fld_char in element.xpath(".//w:fldChar"):
            ff_data = fld_char.find(qn("w:ffData"))
            if ff_data is None or ff_data.find(qn("w:checkBox")) is None:
                continue
            name = ff_data.find(qn("w:name"))
            name_value = "" if name is None else str(name.get(qn("w:val"), "")).strip()
            if name_value:
                named_legacy += 1
            else:
                unnamed_legacy += 1
        if named_legacy:
            return None
        if unnamed_legacy == 1:
            return record, None, "paragraph"

    nonempty = [record for record in records if _normalize_space(record.text)]

    # Normal Unicode marker or a standalone check mark.
    if len(nonempty) == 1:
        record = nonempty[0]
        if _contains_authoritative_marker(record.paragraph):
            return None
        value = record.text or ""
        matches = list(CHECKBOX_TOKEN_PATTERN.finditer(value))
        if len(matches) == 1:
            match = matches[0]
            if not value[: match.start()].strip() and not value[match.end() :].strip():
                return record, (match.start(), match.end()), "text_span"

        check_matches = list(ISOLATED_CHECK_MARK_PATTERN.finditer(value))
        if len(check_matches) == 1:
            match = check_matches[0]
            if not value[: match.start()].strip() and not value[match.end() :].strip():
                return record, (match.start(), match.end()), "text_span"

    # Symbols and drawings often do not surface through paragraph.text.
    for record in records:
        element = record.paragraph._p
        if _normalize_space(record.text):
            continue

        symbols = element.xpath(".//w:sym")
        if len(symbols) == 1:
            font_name = str(symbols[0].get(qn("w:font"), "")).casefold()
            if any(token in font_name for token in ("wingdings", "webdings")):
                return record, None, "paragraph"

        # Old templates can use a tiny VML/Word drawing as the square/check.
        # Treat a single drawing in this marker-only cell as a checkbox signal;
        # the surrounding row-group requirement prevents isolated artwork from
        # becoming a field by itself.
        # Word commonly stores one drawing twice inside ``mc:AlternateContent``:
        # a DrawingML ``w:drawing`` choice plus a VML ``w:pict`` fallback.
        # Count that pair as one semantic marker rather than two independent
        # drawings. This is how many real institutional templates represent
        # an empty checkbox rectangle.
        alternate_contents = element.xpath(".//*[local-name()='AlternateContent']")
        if len(alternate_contents) == 1:
            alt = alternate_contents[0]
            alt_text = "".join((item.text or "") for item in alt.iter(qn("w:t"))).strip()
            alt_drawings = alt.xpath(".//*[local-name()='drawing' or local-name()='pict']")
            if alt_drawings and not alt_text:
                return record, None, "paragraph"

        drawings = element.xpath(".//w:drawing | .//w:pict")
        if len(drawings) == 1:
            return record, None, "paragraph"

    return None


def _cell_width_dxa(record: _ParagraphRecord) -> int | None:
    if record.cell is None:
        return None
    tc_pr = record.cell._tc.tcPr
    if tc_pr is None or tc_pr.tcW is None:
        return None
    raw = tc_pr.tcW.get(qn("w:w"))
    try:
        return int(raw) if raw is not None else None
    except (TypeError, ValueError):
        return None


def _is_blank_checkbox_marker_cell(
    marker_records: list[_ParagraphRecord],
    adjacent_records: list[_ParagraphRecord],
) -> bool:
    """Return whether an empty narrow cell can safely stand for a floating box.

    Some Word templates place a drawn checkbox in a floating shape whose XML
    is anchored outside the table. The table cell underneath is genuinely
    empty. We infer that row only after another row has established the same
    marker column as a checkbox column.
    """

    if not marker_records or not adjacent_records:
        return False
    if any(_normalize_space(record.text) for record in marker_records):
        return False
    for record in marker_records:
        element = record.paragraph._p
        if element.xpath(".//w:sdt | .//w:fldChar | .//w:sym | .//w:drawing | .//w:pict"):
            return False

    marker_width = _cell_width_dxa(marker_records[0])
    adjacent_width = _cell_width_dxa(adjacent_records[0])
    if marker_width is not None and adjacent_width is not None:
        if marker_width > 1800 or marker_width * 3 > adjacent_width:
            return False
    return True


def _looks_like_adjacent_choice_option(records: list[_ParagraphRecord]) -> bool:
    """Conservative evidence that the adjacent cell is an option, not prose."""

    visible = [record for record in records if _normalize_space(record.text)]
    if not visible:
        return False
    first = visible[0]
    value = _normalize_space(first.text)
    if re.match(r"^\d{1,3}[.)]\s*\S", value):
        return True
    p_pr = first.paragraph._p.pPr
    if p_pr is not None and p_pr.numPr is not None:
        return True
    # A short bold lead paragraph is also common for institutional options.
    if len(value) <= 180 and first.paragraph.runs:
        significant_runs = [run for run in first.paragraph.runs if (run.text or "").strip()]
        if significant_runs and all(bool(run.bold) for run in significant_runs):
            return True
    return False


def _remove_floating_checkmark_shapes(root: Any, *, limit: int) -> int:
    """Remove standalone floating checked-box artwork from an assisted copy.

    Word can render a checked square as a floating text box completely outside
    the table that visually contains it. When we materialize an inferred blank
    marker cell as a real tagged checkbox, that old floating artwork must be
    removed or it would remain permanently checked in the generated document.
    """

    if limit <= 0:
        return 0
    mc_tag = "{http://schemas.openxmlformats.org/markup-compatibility/2006}AlternateContent"
    tc_tag = qn("w:tc")
    check_chars = set("✓✔√☑☒")
    removed = 0
    for node in list(root.iter(mc_tag)):
        # Never remove artwork structurally inside a table cell here; regular
        # marker-cell replacement already handles those nodes.
        ancestor = node.getparent()
        inside_cell = False
        while ancestor is not None:
            if ancestor.tag == tc_tag:
                inside_cell = True
                break
            ancestor = ancestor.getparent()
        if inside_cell:
            continue

        text = "".join((item.text or "") for item in node.iter(qn("w:t")))
        compact = "".join(ch for ch in text if not ch.isspace())
        if not compact or any(ch not in check_chars for ch in compact):
            continue
        parent = node.getparent()
        if parent is None:
            continue
        parent.remove(node)
        removed += 1
        if removed >= limit:
            break
    return removed


def _adjacent_checkbox_option_text(records: list[_ParagraphRecord]) -> str:
    """Build the visible option label from the cell beside an isolated box."""

    parts: list[str] = []
    for record in records:
        value = _normalize_space(record.text)
        if not value:
            continue
        # A follow-up justification/observation prompt is a separate fill area,
        # not part of the choice label itself.
        if FOLLOWUP_AREA_PATTERN.match(value):
            break
        if _looks_like_fill_area_text(value):
            continue
        parts.append(value)
    if not parts:
        return ""
    option = " — ".join(parts)
    if len(option) > 360:
        option = option[:357].rstrip(" ,;:-") + "…"
    return _normalize_space(option)


def _adjacent_checkbox_group_label(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
) -> str:
    """Find a concise group prompt for a vertical checkbox+description table."""

    prompt_map = {
        "ocorrência": "Ocorrência verificada",
        "ocorrências": "Ocorrências verificadas",
        "situacao": "Situação verificada",
        "situação": "Situação verificada",
        "situações": "Situações verificadas",
        "condição": "Condição verificada",
        "condições": "Condições verificadas",
        "alternativa": "Alternativa",
        "alternativas": "Alternativas",
        "opção": "Opção",
        "opções": "Opções",
    }
    tail_pattern = re.compile(
        r"(?:seguinte|seguintes)\s+"
        r"(ocorr[eê]ncias?|situa[cç][aã]o(?:ões)?|condi[cç][aã]o(?:ões)?|"
        r"alternativas?|op[cç][aã]o(?:ões)?)\s*:?\s*$",
        re.IGNORECASE,
    )
    for previous in reversed(records[: record.ordinal]):
        value = _normalize_space(previous.text)
        if not value:
            continue
        if previous.table_index == record.table_index:
            continue
        match = tail_pattern.search(value)
        if match:
            token = match.group(1).casefold()
            token = token.replace("ocorrencia", "ocorrência")
            return prompt_map.get(token, _clean_label(match.group(1)))
        if SECTION_NUMBER_PATTERN.match(value) and len(value) <= 190:
            return _clean_label(value)
        if value.endswith((":", "：")) and len(value) <= 110:
            return _clean_label(value)
        # Stop once we reach a normal prose paragraph outside this table. A
        # long declaration is context, but should not become the field label.
        if len(value) > 110:
            break
    return _nearest_section_title(record, records) or "Ocorrência verificada"


def _detect_blank_followup_areas(
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Detect blank paragraphs directly after observation/justification prompts."""

    by_cell: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if record.cell is not None and record.story == "body":
            by_cell[id(record.cell._tc)].append(record)

    result: list[dict[str, Any]] = []
    for cell_records in by_cell.values():
        cell_records.sort(key=lambda item: item.ordinal)
        for index, record in enumerate(cell_records[:-1]):
            if record.ordinal in reserved_ordinals:
                continue
            prompt = _normalize_space(record.text)
            if not prompt or not FOLLOWUP_AREA_PATTERN.match(prompt):
                continue
            # Use the first truly blank paragraph after the prompt, but do not
            # jump over another visible paragraph.
            target: _ParagraphRecord | None = None
            for following in cell_records[index + 1 :]:
                if _normalize_space(following.text):
                    break
                target = following
                break
            if target is None or target.ordinal in reserved_ordinals:
                continue
            label = _clean_label(prompt)
            field_id = _unique_field_id(_make_field_id(label), known_ids)
            result.append(
                _candidate(
                    field_id=field_id,
                    label=label,
                    field_type="multiline",
                    confidence=0.82,
                    source="empty_cell",
                    preview=prompt,
                    location={
                        "kind": "paragraph",
                        "paragraph": target.ordinal,
                    },
                )
            )
    return result


def _inline_checkbox_options(
    text: str,
) -> tuple[str, list[str], list[tuple[int, int]]] | None:
    matches = list(CHECKBOX_TOKEN_PATTERN.finditer(str(text or "")))
    if len(matches) < 2 or len(matches) > 8:
        return None

    options: list[str] = []
    spans: list[tuple[int, int]] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        option = _normalize_space(text[match.end() : end]).strip(" |;–—-")
        if not option or len(option) > 220:
            return None
        options.append(option)
        spans.append((match.start(), match.end()))
    return text[: matches[0].start()], options, spans


def _single_checkbox_cell_option(
    text: str,
) -> tuple[str, tuple[int, int]] | None:
    """Parse one checkbox-led option from a table cell.

    The first visual line is the option itself.  Any following lines are kept
    as explanatory context in the UI label while remaining untouched in the
    DOCX when the checkbox token is replaced by a tag.
    """

    value = str(text or "")
    tokens = list(CHECKBOX_TOKEN_PATTERN.finditer(value))
    if len(tokens) != 1:
        return None
    token = tokens[0]
    if value[: token.start()].strip():
        return None

    remainder = value[token.end() :].strip()
    if not remainder:
        return None
    lines = [_normalize_space(line) for line in remainder.splitlines() if _normalize_space(line)]
    if not lines:
        return None
    option = lines[0].strip(" |;–—-")
    if not option or len(option) > 220:
        return None

    context = " ".join(line.strip() for line in lines[1:] if line.strip())
    if context:
        context = context.lstrip("*•-–— ").strip()
        if context and context.casefold() not in option.casefold():
            option = f"{option} — {context}"
    return _normalize_space(option), (token.start(), token.end())


def _checkbox_candidate(
    *,
    label: str,
    options: list[str],
    known_ids: set[str],
    confidence: float,
    location: dict[str, Any],
) -> dict[str, Any]:
    clean_options = [_normalize_space(option) for option in options if _normalize_space(option)]
    selection = _infer_checkbox_selection(label, clean_options)
    group_id = _unique_field_id(_make_field_id(label or "escolha"), known_ids)
    known_ids.add(group_id)
    ui_group = f"auto_checkbox_{group_id}"
    fields: list[dict[str, Any]] = []

    for index, option_text in enumerate(clean_options, start=1):
        option_id = _unique_field_id(
            f"{group_id}.{_slug(option_text)[:34] or f'opcao_{index}'}",
            known_ids,
        )
        known_ids.add(option_id)
        field: dict[str, Any] = {
            "id": option_id,
            "label": option_text,
            "type": "checkbox",
            "required": False,
            "selection": selection,
        }
        if selection == "single":
            field.update(
                {
                    "layout": "choice",
                    "layout_group": ui_group,
                    "layout_group_label": label or "Escolha uma opção",
                    "group": ui_group,
                    "choice_required": True,
                }
            )
        fields.append(field)

    return {
        "field_id": group_id,
        "label": label or ("Selecione as opções aplicáveis" if selection == "multiple" else "Escolha uma opção"),
        "type": "checkbox_group",
        "selection": selection,
        "confidence": confidence,
        "source": "checkbox_choice",
        "preview": " | ".join(clean_options),
        "location": dict(location),
        "fields": fields,
        "selected": True,
        "requires_configuration": False,
    }


def _infer_checkbox_selection(label: str, options: list[str]) -> str:
    combined = " ".join([str(label or ""), *options]).casefold()
    multiple_tokens = (
        "declaro",
        "declaramos",
        "autorizo",
        "autorizamos",
        "confirmo",
        "confirmamos",
        "aceito os termos",
        "ciência das declarações",
    )
    if any(token in combined for token in multiple_tokens):
        return "multiple"
    return "single"


def _detect_inline_placeholders(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> list[dict[str, Any]]:
    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text):
        return []

    matches = list(CURRENCY_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(X_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(UNDERSCORE_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(ZERO_PHONE_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(ZERO_CPF_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(SAMPLE_EMAIL_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(LEGACY_BRACED_PLACEHOLDER_PATTERN.finditer(text))
    matches.sort(key=lambda match: (match.start(), -(match.end() - match.start())))
    # Some patterns intentionally overlap (for example an xxxxx@example style
    # e-mail also matches the generic X placeholder). Keep only the widest
    # non-overlapping span so one visual fill area becomes one field.
    non_overlapping: list[re.Match[str]] = []
    for match in matches:
        if any(
            match.start() < existing.end() and existing.start() < match.end()
            for existing in non_overlapping
        ):
            continue
        non_overlapping.append(match)
    matches = sorted(non_overlapping, key=lambda match: match.start())
    if not matches:
        return []

    result: list[dict[str, Any]] = []
    previous_end = 0
    for match in matches:
        is_legacy_braced = LEGACY_BRACED_PLACEHOLDER_PATTERN.fullmatch(match.group(0)) is not None
        local_context = text[previous_end : match.start()]
        # PDF-to-DOCX reconstruction frequently groups several visual PDF
        # lines into one Word paragraph separated by manual line breaks.
        # Prefer the text on the current visual line so a paragraph such as
        # ``Placa: ABC1D23\nData: __/__/____\nHorário: __:__`` yields
        # ``Data`` and ``Horário`` instead of labels polluted by the previous
        # line. Fall back to the complete local context for ordinary DOCX.
        visual_line = local_context.rsplit("\n", 1)[-1]
        label = _local_label(visual_line)
        if not label and visual_line != local_context:
            label = _local_label(local_context)
        if not label:
            label = _context_label_for_record(record, records)

        if is_legacy_braced:
            inner_token = str(match.group(1) or "").strip()
            # A single-braced token is only safe to claim automatically when
            # it has a strong field signal: a local/adjacent label, an inline
            # label ending in ':', or an identifier-like dotted/underscored
            # token.  This deliberately ignores prose such as
            # ``Use chaves {assim} no exemplo``.
            has_inline_colon = bool(re.search(r"[:：]\s*$", local_context))
            token_is_structured = any(separator in inner_token for separator in (".", "_", "-"))
            is_isolated_table_value = (
                record.table is not None
                and _normalize_space(text) == _normalize_space(match.group(0))
                and bool(label)
            )
            if not (label and (has_inline_colon or is_isolated_table_value)) and not token_is_structured:
                previous_end = match.end()
                continue

        if not label and CURRENCY_PLACEHOLDER_PATTERN.fullmatch(match.group(0)):
            # A bare monetary mask is semantically stronger than an anonymous
            # ``Campo XX`` suggestion. Parenthetical text such as
            # ``(valor por extenso)`` remains untouched as contextual text.
            label = "Valor"
        if is_legacy_braced:
            inner_token = str(match.group(1) or "").strip()
            field_id_seed = _legacy_placeholder_field_id(inner_token)
            field_id = _unique_field_id(field_id_seed, known_ids)
        else:
            field_id = _unique_field_id(
                _make_field_id(label or f"campo_{record.ordinal + 1}"),
                known_ids,
            )
        known_ids.add(field_id)
        field_type = _detected_placeholder_type(label or field_id, match.group(0))
        candidate = _candidate(
                field_id=field_id,
                label=label or _humanize_id(field_id),
                field_type=field_type,
                confidence=(0.98 if label else 0.82) if is_legacy_braced else (0.91 if label else 0.74),
                source="legacy_placeholder" if is_legacy_braced else "inline_placeholder",
                preview=match.group(0),
                location={
                    "kind": "text_span",
                    "paragraph": record.ordinal,
                    "start": match.start(),
                    "end": match.end(),
                    "original": match.group(0),
                },
            )
        if is_legacy_braced:
            candidate["legacy_marker"] = match.group(0)
            candidate["legacy_marker_id"] = str(match.group(1) or "").strip()
        result.append(candidate)
        previous_end = match.end()
    return result


def _detect_dropdown_prompt(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text):
        return None

    prompt_pattern = re.compile(
        r"(?:escolher|selecione|selecionar)\s+(?:um|uma)\s+"
        r"(?:item|op[cç][aã]o)\.?\s*$",
        re.IGNORECASE,
    )
    match = prompt_pattern.search(text)
    if match is None:
        return None

    prefix = text[: match.start()]
    label = _local_label(prefix) or _context_label_for_record(record, records)
    field_id = _unique_field_id(_make_field_id(label or "opcao"), known_ids)
    return _candidate(
        field_id=field_id,
        label=label or "Selecione uma opção",
        field_type="dropdown",
        confidence=0.82 if label else 0.70,
        source="dropdown_prompt",
        preview=match.group(0),
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": match.start(),
            "end": match.end(),
            "original": match.group(0),
        },
        options=[],
        default_selected=False,
        requires_configuration=True,
    )


def _detect_labeled_sample_value(
    record: _ParagraphRecord,
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect short example/default values that should still be editable.

    This is deliberately whitelist-based.  Institutional fixed text such as
    ``Órgão: Secretaria ...`` must stay static, while values such as
    ``País: Brasil`` are commonly examples/defaults that help the user
    understand what belongs in the field.
    """

    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text) or ":" not in text:
        return None

    colon = text.find(":")
    label = _clean_label(text[:colon])
    value_start = colon + 1
    while value_start < len(text) and text[value_start].isspace():
        value_start += 1
    value = _normalize_space(text[value_start:])
    if not value or len(value) > 80:
        return None

    normalized_label = _slug(label)
    editable_example_labels = {
        "pais",
        "nacionalidade",
        # Common short defaults/examples in reconstructed PDF forms. These are
        # deliberately label-whitelisted so institutional prose such as
        # ``Órgão: Secretaria ...`` remains static.
        "placa",
        "setor_responsavel",
    }
    if normalized_label not in editable_example_labels:
        return None
    if (
        INSTRUCTION_PATTERN.match(value)
        or GENERIC_DROPDOWN_PATTERN.match(value)
        or CHECKBOX_TOKEN_PATTERN.search(value)
        or X_PLACEHOLDER_PATTERN.search(value)
        or UNDERSCORE_PLACEHOLDER_PATTERN.search(value)
        or ZERO_PHONE_PLACEHOLDER_PATTERN.search(value)
        or ZERO_CPF_PLACEHOLDER_PATTERN.search(value)
        or CURRENCY_PLACEHOLDER_PATTERN.search(value)
        or SAMPLE_EMAIL_PLACEHOLDER_PATTERN.search(value)
    ):
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=suggest_field_type(label),
        confidence=0.86,
        source="sample_value",
        preview=value,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": value_start,
            "end": len(text),
            "original": text[value_start:],
        },
        placeholder=value,
    )


def _detect_labeled_instruction(
    record: _ParagraphRecord,
    known_ids: set[str],
) -> dict[str, Any] | None:
    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text) or ":" not in text:
        return None
    colon = text.find(":")
    label = _clean_label(text[:colon])
    tail_start = colon + 1
    while tail_start < len(text) and text[tail_start].isspace():
        tail_start += 1
    tail = text[tail_start:]
    if not _is_reasonable_label(label, maximum=120) or not INSTRUCTION_PATTERN.match(tail):
        return None
    if len(_normalize_space(tail)) < 12 or len(_normalize_space(tail)) > 800:
        return None

    field_type = suggest_field_type(label)
    if field_type == "text" and len(_normalize_space(tail)) >= 70:
        field_type = "multiline"
    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=field_type,
        confidence=0.90 if _paragraph_is_red(record.paragraph) else 0.82,
        source="instruction",
        preview=tail,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": tail_start,
            "end": len(text),
            "original": tail,
        },
    )


_PREFILLED_TEXT_SECTION_PATTERN = re.compile(
    r"\b(?:justificativa|fundamenta[cç][aã]o|descri[cç][aã]o|detalhamento|"
    r"necessidade|motiva[cç][aã]o|objeto|especifica[cç][aã]o|observa[cç][aã]o|"
    r"provid[eê]ncia|parecer|an[aá]lise|considera[cç][oõ]es|informa[cç][oõ]es)\b",
    re.IGNORECASE,
)
_PREFILLED_TEXT_STATIC_PREFIX_PATTERN = re.compile(
    r"^\s*(?:texto\s+fixo|nota|aten[cç][aã]o|aviso|instru[cç][aã]o|"
    r"orienta[cç][aã]o|observa[cç][aã]o\s+fixa|rodap[eé])\s*[:\-]",
    re.IGNORECASE,
)


def _detect_prefilled_written_text(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect existing prose that likely represents an editable answer.

    Some institutional templates are distributed *already filled* with a
    previous/requester-authored justification instead of a visual blank.  A
    placeholder-only detector necessarily treats that prose as fixed text.
    Detector V2.9 recognizes the strongest structural versions of this pattern
    and converts the prose into a multiline field whose initial value is the
    original text.

    The rule is intentionally conservative.  Long prose is not enough on its
    own: it must live under a numbered response-like section or inside a
    full-width/merged table response row.  Explicitly fixed notes and ordinary
    headers remain static.
    """

    raw_text = str(record.text or "")
    text = _normalize_space(raw_text)
    if (
        len(text) < 45
        or len(text) > 2200
        or record.story != "body"
        or _contains_authoritative_marker(record.paragraph)
        or _looks_like_section_label(text)
        or _PREFILLED_TEXT_STATIC_PREFIX_PATTERN.match(text)
        or CHECKBOX_TOKEN_PATTERN.search(text)
        or GENERIC_DROPDOWN_PATTERN.fullmatch(text)
        or _is_pure_fill_area_text(text)
    ):
        return None

    # A paragraph made mostly of a short heading followed by punctuation is
    # not a written response, even when Word wrapped it visually.
    if len(text.split()) < 7 or sum(ch.isalpha() for ch in text) < 24:
        return None

    section = _clean_label(semantic_section(record))
    section_is_response = bool(section and _PREFILLED_TEXT_SECTION_PATTERN.search(section))

    # A full-width/merged table row is another common way users store a written
    # answer below a heading or below a table header.  Determine this from
    # unique XML cells rather than the apparent ``row.cells`` count because
    # python-docx repeats references for merged cells.
    full_width_response_row = False
    table_context_label = ""
    if record.table is not None and record.row_index is not None and record.cell is not None:
        try:
            row = record.table.rows[int(record.row_index)]
            unique_cells = _unique_row_cells(row)
            current_key = id(record.cell._tc)
            if len(unique_cells) == 1 and id(unique_cells[0]._tc) == current_key:
                # Require meaningful structure above the prose: either a
                # numbered section already in scope or a preceding header row
                # with at least two short cells. This prevents random one-cell
                # narrative tables from becoming editable by accident.
                has_header_row = False
                for previous_index in range(int(record.row_index) - 1, -1, -1):
                    previous_cells = _unique_row_cells(record.table.rows[previous_index])
                    values = [_normalize_space(cell.text) for cell in previous_cells]
                    values = [value for value in values if value]
                    if not values:
                        continue
                    if len(values) >= 2 and all(len(value) <= 90 for value in values):
                        has_header_row = True
                        break
                    if len(values) == 1 and _looks_like_section_label(values[0]):
                        table_context_label = _clean_label(values[0])
                        break
                full_width_response_row = bool(section or has_header_row or table_context_label)
        except (IndexError, AttributeError, TypeError, ValueError):
            full_width_response_row = False

    if not section_is_response and not full_width_response_row:
        return None

    label = section or table_context_label or _context_label_for_record(record, records)
    label = _clean_label(label)
    if not _is_reasonable_label(label, maximum=180):
        label = "Texto editável"

    # Full-width narrative rows are a little more ambiguous than explicit
    # Justificativa/Descrição sections, so keep their confidence lower.  Both
    # remain reviewable in the assisted-detection screen.
    confidence = 0.90 if section_is_response else 0.82
    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type="multiline",
        confidence=confidence,
        source="prefilled_text",
        preview=text,
        location={
            "kind": "paragraph",
            "paragraph": record.ordinal,
            "table_index": record.table_index,
            "row_index": record.row_index,
            "cell_index": record.cell_index,
            "prefilled_text": True,
        },
        default_value=raw_text.strip(),
    )


def _detect_adjacent_sample_value(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect editable example/default values stored in the next table cell.

    Many institutional forms use four physical cells per row::

        E-mail: | servidor@orgao.gov.br | Telefone: | (83) 99999-9999

    The values are examples/defaults, not fixed institutional prose.  Keep this
    heuristic label-whitelisted so rows such as ``Órgão: Secretaria ...`` stay
    read-only.
    """

    text = _normalize_space(record.text)
    if (
        not text
        or record.cell is None
        or record.table is None
        or record.row_index is None
        or len(record.cell.paragraphs) != 1
        or _contains_authoritative_marker(record.paragraph)
    ):
        return None

    # Do not require a trailing colon here. Real Word forms frequently use
    # alternating cells such as ``E-mail | exemplo@orgao.gov.br | Telefone |
    # (83) 99999-9999``. The physical adjacency is the delimiter. Keeping the
    # semantic whitelist below prevents ordinary institutional prose from
    # being converted merely because it sits next to another cell.
    explicit_label = text.endswith((":", "："))
    label = _clean_label(text)
    if not _is_reasonable_label(label, maximum=100):
        return None

    normalized_label = _slug(label)
    allowed_plain_labels = {
        "unidade",
        "lotacao",
        "setor",
        "municipio",
        "cidade",
        "pais",
        "nacionalidade",
    }
    email_labels = {"email", "e_mail", "correio_eletronico"}
    phone_labels = {"telefone", "celular", "fone"}

    unique_cells = _unique_row_cells(record.table.rows[record.row_index])
    current_index = next(
        (index for index, cell in enumerate(unique_cells) if id(cell._tc) == id(record.cell._tc)),
        -1,
    )
    if current_index < 0 or current_index + 1 >= len(unique_cells):
        return None
    value_cell = unique_cells[current_index + 1]
    value_records = [item for item in records if item.cell is not None and id(item.cell._tc) == id(value_cell._tc)]
    non_empty = [item for item in value_records if _normalize_space(item.text)]
    if len(non_empty) != 1 or _contains_authoritative_marker(non_empty[0].paragraph):
        return None

    value_record = non_empty[0]
    value = _normalize_space(value_record.text)
    if not value or len(value) > 100 or _is_pure_fill_area_text(value):
        return None

    looks_editable = False
    if normalized_label in email_labels:
        looks_editable = bool(
            re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", value, re.IGNORECASE)
        )
    elif normalized_label in phone_labels:
        looks_editable = bool(
            re.fullmatch(r"\(?\d{2}\)?\s*\d{4,5}\s*-\s*\d{4}", value)
        )
    elif normalized_label in allowed_plain_labels:
        # Short human-readable defaults such as ``Diretoria Administrativa``
        # are useful placeholders. Avoid values that look like sentences.
        looks_editable = (
            len(value) <= 60
            and not value.endswith((".", ";"))
            and len(value.split()) <= 6
            and not CHECKBOX_TOKEN_PATTERN.search(value)
        )

    if not looks_editable:
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=suggest_field_type(label),
        confidence=0.88 if explicit_label else 0.85,
        source="sample_value",
        preview=value,
        location={
            "kind": "text_span",
            "paragraph": value_record.ordinal,
            "start": 0,
            "end": len(value_record.text),
            "original": value_record.text,
        },
        placeholder=value,
    )


def _detect_label_only_field(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    text = _normalize_space(record.text)
    if (
        not text.endswith(":")
        or record.cell is None
        or len(record.cell.paragraphs) != 1
        or _contains_authoritative_marker(record.paragraph)
    ):
        return None
    label = _clean_label(text)
    if not _is_reasonable_label(label, maximum=100):
        return None
    if _looks_like_section_label(text) and SECTION_NUMBER_PATTERN.match(text):
        return None

    # A bare label can represent a fill area inside its own cell (for example
    # ``Responsável legal:``), but a very common Word grid stores the label in
    # one cell and the actual mask/value in the immediately following cell. In
    # that case the following cell owns the field and creating another tag in
    # the label cell causes duplicate inputs and staircase layouts.
    has_form_neighbor = False
    if record.table is not None and record.row_index is not None:
        try:
            unique_cells = _unique_row_cells(record.table.rows[record.row_index])
            current_index = next(
                (
                    index
                    for index, cell in enumerate(unique_cells)
                    if id(cell._tc) == id(record.cell._tc)
                ),
                -1,
            )
            if current_index >= 0 and current_index + 1 < len(unique_cells):
                immediate_text = _normalize_space(unique_cells[current_index + 1].text)
                if not immediate_text:
                    return None
                if _is_pure_fill_area_text(immediate_text):
                    return None

            for cell in unique_cells:
                if id(cell._tc) == id(record.cell._tc):
                    continue
                neighbor = _normalize_space(cell.text)
                if not neighbor:
                    return None
                if _looks_like_fill_area_text(neighbor):
                    has_form_neighbor = True
                    break
        except (IndexError, AttributeError):
            pass
    if not has_form_neighbor:
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=suggest_field_type(label),
        confidence=0.82,
        source="empty_cell",
        preview="Área vazia após o rótulo",
        location={
            "kind": "append_tag",
            "paragraph": record.ordinal,
        },
    )


def _detect_empty_cells(
    document: _Document,
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    del document
    result: list[dict[str, Any]] = []
    by_cell: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if record.cell is not None:
            by_cell[id(record.cell._tc)].append(record)

    visited_cells: set[int] = set()
    for record in records:
        if record.cell is None or record.table is None or record.row_index is None or record.cell_index is None:
            continue
        cell_key = id(record.cell._tc)
        if cell_key in visited_cells:
            continue
        visited_cells.add(cell_key)
        cell_records = by_cell.get(cell_key, [])
        if not cell_records or any(item.ordinal in reserved_ordinals for item in cell_records):
            continue
        if any(_normalize_space(item.text) for item in cell_records):
            continue
        if record.cell.tables:
            continue
        if record.cell_index <= 0:
            continue
        try:
            previous_cell = record.table.rows[record.row_index].cells[record.cell_index - 1]
        except (IndexError, AttributeError):
            continue
        previous_text = _normalize_space(previous_cell.text)
        label = previous_text.strip(" :：–—-")
        if not _is_reasonable_label(label):
            continue
        explicit_label = previous_text.rstrip().endswith((":", "："))
        ordinal = cell_records[0].ordinal
        field_id = _unique_field_id(_make_field_id(label), known_ids)
        known_ids.add(field_id)
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type=suggest_field_type(label),
                confidence=0.84 if explicit_label else 0.68,
                source="empty_cell",
                preview="Célula vazia",
                location={
                    "kind": "empty_cell",
                    "paragraph": ordinal,
                },
                default_selected=explicit_label,
            )
        )
    return result


def _detect_consistency_repair_fields(
    records: list[_ParagraphRecord],
    candidates: list[dict[str, Any]],
    known_ids: set[str],
) -> list[dict[str, Any]]:
    """Suggest a missing sibling field when the surrounding table is consistent.

    This is intentionally a *repair* pass rather than another primary heuristic.
    It only fires when at least two peer rows already establish that the same
    visual column is editable. This helps arbitrary user-made matrices where one
    row uses a slightly different blank representation without requiring a fixed
    document template.
    """

    by_ordinal = {record.ordinal: record for record in records}
    occupied_cells: set[tuple[int, int, int]] = set()
    column_rows: dict[tuple[int, int], set[int]] = defaultdict(set)
    repeatable_tables: set[int] = set()

    for candidate in candidates:
        location = candidate.get("location", {}) or {}
        if str(candidate.get("source", "")) == "repeatable_table":
            try:
                repeatable_tables.add(int(location.get("table_index", -1)))
            except (TypeError, ValueError):
                pass
        ordinals: list[int] = []
        if "paragraph" in location:
            try:
                ordinals.append(int(location.get("paragraph", -1)))
            except (TypeError, ValueError):
                pass
        for value in location.get("paragraphs", []) or []:
            try:
                ordinals.append(int(value))
            except (TypeError, ValueError):
                continue
        for ordinal in ordinals:
            record = by_ordinal.get(ordinal)
            if (
                record is None
                or record.table_index is None
                or record.row_index is None
                or record.cell_index is None
            ):
                continue
            key = (int(record.table_index), int(record.row_index), int(record.cell_index))
            occupied_cells.add(key)
            column_rows[(key[0], key[2])].add(key[1])

    established_columns = {
        key for key, rows in column_rows.items()
        if len(rows) >= 2 and key[0] not in repeatable_tables
    }
    if not established_columns:
        return []

    by_cell: dict[tuple[int, int, int], list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if (
            record.table_index is None
            or record.row_index is None
            or record.cell_index is None
        ):
            continue
        by_cell[(int(record.table_index), int(record.row_index), int(record.cell_index))].append(record)

    result: list[dict[str, Any]] = []
    for (table_index, row_index, cell_index), cell_records in by_cell.items():
        if (table_index, cell_index) not in established_columns:
            continue
        if (table_index, row_index, cell_index) in occupied_cells:
            continue
        if any(_contains_authoritative_marker(record.paragraph) for record in cell_records):
            continue

        texts = [record.text or "" for record in cell_records]
        combined = _normalize_space(" ".join(texts))
        if combined and not all(_is_pure_fill_area_text(text) for text in texts):
            continue

        target = next((record for record in cell_records if record.paragraph is not None), None)
        if target is None:
            continue
        label, label_source, label_confidence = semantic_label(target)
        if not label or label_confidence < 0.74 or label_source == "section_fallback":
            continue

        field_id = _unique_field_id(_make_field_id(label), known_ids)
        known_ids.add(field_id)
        preview = combined or "área vazia"
        field_type = _detected_placeholder_type(label, combined)
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type=field_type,
                confidence=0.68,
                source="consistency_repair",
                preview=preview,
                location={
                    "kind": "empty_cell" if not combined else "paragraph",
                    "paragraph": target.ordinal,
                    "repair_basis": "peer_column",
                    "table_index": table_index,
                    "row_index": row_index,
                    "cell_index": cell_index,
                },
                default_selected=False,
            )
        )

    return result


def _unique_row_cells(row) -> list[_Cell]:
    result: list[_Cell] = []
    seen: set[int] = set()
    for cell in row.cells:
        key = id(cell._tc)
        if key in seen:
            continue
        seen.add(key)
        result.append(cell)
    return result


def _is_pure_fill_area_text(value: str) -> bool:
    """Return True when a cell consists only of the visual fill control.

    This is stricter than :func:`_looks_like_fill_area_text`: ``CPF: ___`` is
    *not* pure because the same cell contains another field's label, while
    ``___``, ``R$ ____`` and ``☐ Sim ☐ Não`` are pure fill areas.
    """

    text = _normalize_space(value)
    if not text:
        return True
    if X_PLACEHOLDER_PATTERN.fullmatch(text):
        return True
    if UNDERSCORE_PLACEHOLDER_PATTERN.fullmatch(text):
        return True
    if ZERO_PHONE_PLACEHOLDER_PATTERN.fullmatch(text):
        return True
    if ZERO_CPF_PLACEHOLDER_PATTERN.fullmatch(text):
        return True
    if CURRENCY_PLACEHOLDER_PATTERN.fullmatch(text):
        return True
    if SAMPLE_EMAIL_PLACEHOLDER_PATTERN.fullmatch(text):
        return True
    if GENERIC_DROPDOWN_PATTERN.fullmatch(text):
        return True
    if CHECKBOX_TOKEN_PATTERN.match(text) is not None:
        return True
    return False


def _looks_like_fill_area_text(value: str) -> bool:
    text = str(value or "")
    if not _normalize_space(text):
        return True
    return bool(
        PLACEHOLDER_PATTERN.search(text)
        or X_PLACEHOLDER_PATTERN.search(text)
        or UNDERSCORE_PLACEHOLDER_PATTERN.search(text)
        or ZERO_PHONE_PLACEHOLDER_PATTERN.search(text)
        or ZERO_CPF_PLACEHOLDER_PATTERN.search(text)
        or CURRENCY_PLACEHOLDER_PATTERN.search(text)
        or SAMPLE_EMAIL_PLACEHOLDER_PATTERN.search(text)
        or GENERIC_DROPDOWN_PATTERN.search(text)
        or CHECKBOX_TOKEN_PATTERN.search(text)
    )


def _nearest_section_title(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    *,
    preserve_number: bool = False,
) -> str:
    understood = semantic_section(record)
    if understood:
        return understood.rstrip(":").strip() if preserve_number else _clean_label(understood)
    for previous in reversed(records[: record.ordinal]):
        value = _normalize_space(previous.text)
        if not value:
            continue
        if SECTION_NUMBER_PATTERN.match(value) and len(value) <= 190:
            return value.rstrip(":").strip() if preserve_number else _clean_label(value)
        if previous.table_index == record.table_index:
            continue
        if _looks_like_section_label(value) and len(value) <= 190:
            return value.rstrip(":").strip() if preserve_number else _clean_label(value)
    return ""


def _repeatable_column_type(header: str, values: list[str]) -> str:
    for value in values:
        compact = _normalize_space(value)
        if re.fullmatch(r"_{2,}\s*/\s*_{2,}\s*/\s*_{2,}", compact):
            return "date"
        if ZERO_PHONE_PLACEHOLDER_PATTERN.fullmatch(compact):
            return "phone"
        if ZERO_CPF_PLACEHOLDER_PATTERN.fullmatch(compact):
            return "cpf"
        if CURRENCY_PLACEHOLDER_PATTERN.fullmatch(compact):
            return "currency"
        if SAMPLE_EMAIL_PLACEHOLDER_PATTERN.fullmatch(compact):
            return "email"
    return suggest_field_type(header)


def _detected_placeholder_type(label: str, preview: str) -> str:
    normalized_label = _slug(label)
    compact = _normalize_space(preview)
    if "observacao_curta" in normalized_label:
        return "text"
    # Questionnaire/matrix PDFs often reconstruct the short observation cell
    # as a modest underline next to a fixed row label. Keep those compact
    # instead of turning every label ending in ``Observação`` into a large
    # multiline editor. Long blank areas still become multiline elsewhere.
    if (
        normalized_label.endswith("_observacao")
        and re.fullmatch(r"_{2,}", compact)
        and len(compact) <= 20
    ):
        return "text"
    if re.fullmatch(r"_{2,}\s*/\s*_{2,}\s*/\s*_{2,}", compact):
        return "date"
    if ZERO_PHONE_PLACEHOLDER_PATTERN.fullmatch(compact):
        return "phone"
    if ZERO_CPF_PLACEHOLDER_PATTERN.fullmatch(compact):
        return "cpf"
    if CURRENCY_PLACEHOLDER_PATTERN.fullmatch(compact):
        return "currency"
    if SAMPLE_EMAIL_PLACEHOLDER_PATTERN.fullmatch(compact):
        return "email"
    return suggest_field_type(label)


def _same_cell_previous_label(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
) -> str:
    if record.cell is None:
        return ""
    cell_key = id(record.cell._tc)
    for previous in reversed(records[: record.ordinal]):
        if previous.cell is None or id(previous.cell._tc) != cell_key:
            continue
        value = _normalize_space(previous.text)
        if not value or _looks_like_fill_area_text(value):
            continue
        if value.endswith((":", "：")) and _is_reasonable_label(value, maximum=140):
            return _clean_label(value)
        if _is_reasonable_label(value, maximum=100) and len(value.split()) <= 12:
            return _clean_label(value)
    return ""


def _table_axis_context(record: _ParagraphRecord) -> tuple[str, str]:
    """Return (row label, column label) for a field inside a table grid."""

    if record.table is None or record.row_index is None or record.cell_index is None:
        return "", ""
    try:
        row = record.table.rows[record.row_index]
        current_key = id(record.cell._tc) if record.cell is not None else None
        row_label = ""
        seen: set[int] = set()
        for cell_index, cell in enumerate(row.cells):
            key = id(cell._tc)
            if key in seen:
                continue
            seen.add(key)
            if current_key is not None and key == current_key:
                break
            value = _normalize_space(cell.text)
            if (
                _is_reasonable_label(value, maximum=120)
                and not _looks_like_fill_area_text(value)
                and not re.fullmatch(r"\d+", value)
            ):
                row_label = _clean_label(value)

        column_label = ""
        for row_index in range(record.row_index - 1, -1, -1):
            earlier = record.table.rows[row_index]
            if record.cell_index >= len(earlier.cells):
                continue
            value = _normalize_space(earlier.cells[record.cell_index].text)
            if not value or _looks_like_fill_area_text(value):
                continue
            cleaned = _clean_label(value)
            if (
                _is_reasonable_label(cleaned, maximum=90)
                and not SECTION_NUMBER_PATTERN.match(cleaned)
            ):
                column_label = cleaned
                break
        return row_label, column_label
    except (IndexError, AttributeError):
        return "", ""


def _context_label_for_record(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
) -> str:
    understood_label, understood_source, understood_confidence = semantic_label(record)
    if understood_label and understood_confidence >= 0.72 and understood_source != "section_fallback":
        return _clean_label(understood_label)

    text = record.text or ""
    if ":" in text:
        before = text.split(":", 1)[0]
        if _is_reasonable_label(before):
            return _clean_label(before)

    same_cell = _same_cell_previous_label(record, records)
    if same_cell:
        return same_cell

    # Same table: prefer the previous cell in the same row.
    if record.table is not None and record.row_index is not None and record.cell_index is not None:
        try:
            row = record.table.rows[record.row_index]
            if record.cell_index > 0:
                previous = _normalize_space(row.cells[record.cell_index - 1].text)
                if _is_reasonable_label(previous) and not _looks_like_fill_area_text(previous):
                    return _clean_label(previous)
        except (IndexError, AttributeError):
            pass

        row_label, column_label = _table_axis_context(record)
        if row_label and column_label and row_label.casefold() != column_label.casefold():
            return f"{row_label} — {column_label}"
        if column_label:
            return column_label
        if row_label:
            return row_label

        # Search earlier rows for a section-like merged title or a short label.
        try:
            for row_index in range(record.row_index - 1, -1, -1):
                row = record.table.rows[row_index]
                unique_texts: list[str] = []
                seen_tc: set[int] = set()
                for cell in row.cells:
                    key = id(cell._tc)
                    if key in seen_tc:
                        continue
                    seen_tc.add(key)
                    value = _normalize_space(cell.text)
                    if value and value not in unique_texts:
                        unique_texts.append(value)
                if not unique_texts:
                    continue
                for value in unique_texts:
                    cleaned = _clean_label(value)
                    if _looks_like_section_label(cleaned):
                        return cleaned
                if len(unique_texts) == 1 and _is_reasonable_label(unique_texts[0], maximum=180):
                    return _clean_label(unique_texts[0])
        except (IndexError, AttributeError):
            pass

    # Document order fallback.
    for previous in reversed(records[: record.ordinal]):
        value = _normalize_space(previous.text)
        if not value:
            continue
        cleaned = _clean_label(value)
        if _looks_like_section_label(cleaned):
            return cleaned
        if _is_reasonable_label(cleaned, maximum=120):
            return cleaned
        break
    return ""


def _is_instruction_candidate(record: _ParagraphRecord) -> bool:
    text = _normalize_space(record.text)
    if not INSTRUCTION_PATTERN.match(text):
        return False
    if len(text) < 12 or len(text) > 800:
        return False
    # A paragraph containing several sentences of ordinary policy text can
    # begin with an imperative. Red text or table placement increases safety.
    return _paragraph_is_red(record.paragraph) or record.cell is not None


def _paragraph_is_red(paragraph: Paragraph) -> bool:
    for run in paragraph.runs:
        color = run.font.color.rgb
        if color is None:
            continue
        try:
            red, green, blue = int(color[0]), int(color[1]), int(color[2])
        except Exception:
            text = str(color)
            if len(text) == 6:
                try:
                    red, green, blue = int(text[:2], 16), int(text[2:4], 16), int(text[4:], 16)
                except ValueError:
                    continue
            else:
                continue
        if red >= 150 and red > green * 1.35 and red > blue * 1.35:
            return True
    return False


def _contains_authoritative_marker(paragraph: Paragraph) -> bool:
    text = paragraph.text or ""
    if PLACEHOLDER_PATTERN.search(text):
        return True
    element = paragraph._p
    return bool(element.xpath(".//w:sdt | .//w:fldChar"))


def _tag_for_candidate(candidate: dict[str, Any]) -> str:
    field_id = str(candidate.get("field_id", "")).strip()
    field_type = str(candidate.get("type", "text")).strip().casefold()
    if not field_id:
        raise AutomaticDetectionError("Uma sugestão selecionada não possui ID de campo.")

    if field_type == "date":
        return f"{{{{date:{field_id}}}}}"
    if field_type == "checkbox":
        return f"{{{{checkbox:{field_id}}}}}"
    if field_type == "dropdown":
        options = compact_dropdown_options(candidate.get("options", []))
        if len(options) < 2:
            raise AutomaticDetectionError(
                f"A lista '{candidate.get('label', field_id)}' precisa de pelo menos duas opções."
            )
        encoded = []
        for option in options:
            if isinstance(option, dict):
                label = _safe_tag_option(option.get("label", ""))
                value = _safe_tag_option(option.get("value", ""))
                encoded.append(value if label == value else f"{label} => {value}")
            else:
                encoded.append(_safe_tag_option(option))
        prefix = "single_choice" if str(candidate.get("layout", "")) == "choice" else "dropdown"
        return "{{" + prefix + ":" + field_id + "|" + "|".join(encoded) + "}}"
    return f"{{{{{field_id}}}}}"


def _apply_paragraph_block(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    ordinals = [
        int(value)
        for value in candidate.get("location", {}).get("paragraphs", [])
    ]
    records = [by_ordinal.get(value) for value in ordinals]
    records = [record for record in records if record is not None]
    if not records:
        raise AutomaticDetectionError("Bloco de alternativas não encontrado no DOCX.")
    first = records[0].paragraph
    _replace_entire_paragraph(first, _tag_for_candidate(candidate))
    for record in records[1:]:
        _remove_paragraph(record.paragraph)


def _apply_repeatable_table(
    candidate: dict[str, Any],
    document: _Document,
) -> None:
    location = candidate.get("location", {}) or {}
    try:
        table_index = int(location.get("document_table_index", -1))
        template_row_index = int(location.get("template_row", -1))
        data_rows = sorted(
            {int(value) for value in location.get("data_rows", []) or []}
        )
        insert_before_row = int(location.get("insert_before_row", -1))
    except (TypeError, ValueError):
        raise AutomaticDetectionError("Tabela repetível detectada com posição inválida.")

    if table_index < 0 or table_index >= len(document.tables):
        raise AutomaticDetectionError("A tabela repetível detectada não foi encontrada no DOCX.")
    table = document.tables[table_index]
    synthetic_template = bool(location.get("synthetic_template_row", False))
    if not synthetic_template and (
        template_row_index < 0 or template_row_index >= len(table.rows)
    ):
        raise AutomaticDetectionError("A linha modelo da tabela repetível não foi encontrada.")
    if synthetic_template and (
        insert_before_row < 0 or insert_before_row >= len(table.rows)
    ):
        raise AutomaticDetectionError(
            "A posição da nova linha editável da planilha não foi encontrada."
        )

    field_id = str(candidate.get("field_id", "")).strip()
    columns = [
        dict(column)
        for column in candidate.get("columns", []) or []
        if isinstance(column, dict)
    ]
    if not field_id or len(columns) < 2:
        raise AutomaticDetectionError("A tabela repetível detectada está incompleta.")

    if synthetic_template:
        # ``python-docx`` can append a correctly sized row using the table grid.
        # Move that XML row immediately before the merged narrative/note row so
        # the original document keeps its visual order: header -> editable rows
        # -> note.  Using a fresh row avoids copying header shading/bold styles.
        row = table.add_row()
        target_row = table.rows[insert_before_row]
        target_row._tr.addprevious(row._tr)
    else:
        row = table.rows[template_row_index]

    cells = _unique_row_cells(row)
    repeat_marker_written = False
    for column in columns:
        column_type = str(column.get("type", "text")).strip().casefold()
        column_id = str(column.get("id", "")).strip()
        if not column_id:
            continue
        if column_type == "auto_number":
            column_index = 0
            text = f"{{{{repeat:{field_id}}}}} {{{{row.number}}}}"
            repeat_marker_written = True
        else:
            try:
                column_index = int(column.get("column_index", -1))
            except (TypeError, ValueError):
                column_index = -1
            if column_index < 0:
                continue
            child_id = f"{field_id}.{column_id}"
            if column_type == "date":
                text = f"{{{{date:{child_id}}}}}"
            elif column_type == "checkbox":
                text = f"{{{{checkbox:{child_id}}}}}"
            else:
                text = f"{{{{{child_id}}}}}"

        if column_index >= len(cells):
            raise AutomaticDetectionError(
                "A estrutura de colunas da tabela repetível mudou após a análise."
            )
        _replace_cell_with_text(cells[column_index], text)

    if not repeat_marker_written:
        # The detector currently requires an auto-number column, but keep this
        # fallback so reviewed candidates remain robust if that rule evolves.
        first_editable = next(
            (
                column
                for column in columns
                if str(column.get("type", "")).casefold() != "auto_number"
            ),
            None,
        )
        if first_editable is None:
            raise AutomaticDetectionError("A tabela repetível não possui coluna editável.")
        column_index = int(first_editable.get("column_index", 0))
        paragraph = cells[column_index].paragraphs[0]
        paragraph.insert_paragraph_before(f"{{{{repeat:{field_id}}}}}")

    # Keep only the first detected data row as the Word model row. The DOCX
    # engine duplicates it according to the rows entered by the client. A
    # synthetic sheet has no source data rows to remove.
    if not synthetic_template:
        for row_index in sorted(data_rows, reverse=True):
            if row_index == template_row_index or row_index >= len(table.rows):
                continue
            table._tbl.remove(table.rows[row_index]._tr)


def _replace_cell_with_text(cell: _Cell, text: str) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph(text)
        return
    _replace_entire_paragraph(paragraphs[0], text)
    for paragraph in paragraphs[1:]:
        _remove_paragraph(paragraph)


def _apply_checkbox_group(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    ordinals = [
        int(value)
        for value in candidate.get("location", {}).get("paragraphs", [])
    ]
    fields = [dict(field) for field in candidate.get("fields", []) or []]
    if len(ordinals) != len(fields):
        raise AutomaticDetectionError("Grupo de caixas de seleção inconsistente.")
    for ordinal, field in zip(ordinals, fields):
        record = by_ordinal.get(ordinal)
        if record is None:
            raise AutomaticDetectionError("Opção de caixa de seleção não encontrada.")
        text = record.paragraph.text or ""
        match = CHECKBOX_LINE_PATTERN.match(text)
        if not match:
            raise AutomaticDetectionError("O texto de uma opção mudou após a análise.")
        replacement = f"{{{{checkbox:{field['id']}}}}} {match.group(1).strip()}"
        _replace_entire_paragraph(record.paragraph, replacement)


def _apply_inline_checkbox_group(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    location = candidate.get("location", {}) or {}
    ordinal = int(location.get("paragraph", -1))
    record = by_ordinal.get(ordinal)
    if record is None:
        raise AutomaticDetectionError("Grupo de caixas de seleção não encontrado.")
    fields = [dict(field) for field in candidate.get("fields", []) or []]
    spans = [
        tuple(int(value) for value in span)
        for span in location.get("checkbox_spans", []) or []
        if isinstance(span, (list, tuple)) and len(span) == 2
    ]
    if len(fields) != len(spans) or not fields:
        raise AutomaticDetectionError("Grupo de caixas de seleção inconsistente.")
    replacements = [
        (start, end, f"{{{{checkbox:{field['id']}}}}}")
        for (start, end), field in zip(spans, fields, strict=True)
    ]
    _replace_paragraph_spans(record.paragraph, replacements)


def _apply_multi_cell_checkbox_group(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    location = candidate.get("location", {}) or {}
    ordinals = [int(value) for value in location.get("paragraphs", []) or []]
    spans = [
        tuple(int(value) for value in span)
        for span in location.get("checkbox_spans", []) or []
        if isinstance(span, (list, tuple)) and len(span) == 2
    ]
    marker_modes = [str(value or "text_span") for value in location.get("checkbox_marker_modes", []) or []]
    fields = [dict(field) for field in candidate.get("fields", []) or []]
    if not marker_modes:
        marker_modes = ["text_span"] * len(fields)
    if not fields or len(fields) != len(ordinals) or len(fields) != len(spans) or len(fields) != len(marker_modes):
        raise AutomaticDetectionError("Grupo de caixas de seleção entre células inconsistente.")

    inferred_count = sum(1 for mode in marker_modes if mode == "inferred_blank")
    if inferred_count:
        first_record = by_ordinal.get(ordinals[0]) if ordinals else None
        if first_record is not None:
            _remove_floating_checkmark_shapes(
                first_record.paragraph._p.getroottree().getroot(),
                limit=inferred_count,
            )

    for ordinal, span, mode, field in zip(ordinals, spans, marker_modes, fields, strict=True):
        record = by_ordinal.get(ordinal)
        if record is None:
            raise AutomaticDetectionError("Opção de caixa de seleção não encontrada.")
        replacement = f"{{{{checkbox:{field['id']}}}}}"
        if mode in {"paragraph", "inferred_blank"}:
            _replace_all_paragraph_content(record.paragraph, replacement)
            continue

        start, end = span
        text = record.paragraph.text or ""
        if start < 0 or end <= start or end > len(text):
            raise AutomaticDetectionError("A posição de uma opção mudou após a análise.")
        marker_text = text[start:end]
        if not (
            CHECKBOX_TOKEN_PATTERN.fullmatch(marker_text)
            or ISOLATED_CHECK_MARK_PATTERN.fullmatch(marker_text)
        ):
            raise AutomaticDetectionError("O marcador de uma opção mudou após a análise.")
        _replace_paragraph_spans(
            record.paragraph,
            [(start, end, replacement)],
        )


def _replace_all_paragraph_content(paragraph: Paragraph, text: str) -> None:
    """Replace runs, controls, fields and symbols while preserving paragraph properties."""

    element = paragraph._p
    for child in list(element):
        if child.tag == qn("w:pPr"):
            continue
        element.remove(child)
    paragraph.add_run(text)


def _replace_entire_paragraph(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _append_tag_to_paragraph(paragraph: Paragraph, tag: str) -> None:
    current = paragraph.text or ""
    separator = "" if not current or current.endswith((" ", "\t")) else " "
    paragraph.add_run(separator + str(tag))


def _replace_paragraph_spans(
    paragraph: Paragraph,
    replacements: Iterable[tuple[int, int, str]],
) -> None:
    # Candidate offsets are measured against ``Paragraph.text``. That text
    # includes manual line breaks (w:br/w:cr) as ``\n`` and tabs (w:tab) as
    # ``\t``. The old replacement code concatenated only w:t nodes, making
    # every span after a break or tab drift to the right. This is especially
    # common in DOCX files reconstructed from PDFs.
    segments = _paragraph_position_segments(paragraph._p)
    original_text = "".join(text for _element, text in segments)
    editable_spans: list[tuple[Any, int, int]] = []
    cursor = 0
    for element, text in segments:
        segment_end = cursor + len(text)
        if element is not None and text:
            editable_spans.append((element, cursor, segment_end))
        cursor = segment_end

    if not editable_spans:
        raise AutomaticDetectionError("O trecho detectado não contém texto editável no DOCX.")

    for start, end, replacement in sorted(replacements, key=lambda item: item[0], reverse=True):
        if start < 0 or end <= start or end > len(original_text):
            raise AutomaticDetectionError("Posição de preenchimento inválida no DOCX.")
        start_index = _span_index_for_position(editable_spans, start)
        end_index = _span_index_for_position(editable_spans, end - 1)
        if start_index is None or end_index is None:
            raise AutomaticDetectionError("Não foi possível localizar o trecho detectado no XML do DOCX.")

        start_element, start_offset, _ = editable_spans[start_index]
        end_element, end_offset, _ = editable_spans[end_index]
        local_start = start - start_offset
        local_end = end - end_offset

        if start_index == end_index:
            current = start_element.text or ""
            _set_text_element_value(
                start_element,
                current[:local_start] + replacement + current[local_end:],
            )
            continue

        start_text = start_element.text or ""
        end_text = end_element.text or ""
        _set_text_element_value(start_element, start_text[:local_start] + replacement)
        for element, _node_start, _node_end in editable_spans[start_index + 1 : end_index]:
            _set_text_element_value(element, "")
        _set_text_element_value(end_element, end_text[local_end:])


def _paragraph_position_segments(paragraph_element) -> list[tuple[Any | None, str]]:
    """Return paragraph content in the same coordinate space as Paragraph.text.

    Text/instruction nodes are editable and are returned with their XML node.
    Manual line breaks and tabs occupy positions too, but are represented by a
    ``None`` node because automatic field replacement should never overwrite
    those structural elements. Nested paragraphs are processed separately.
    """

    segments: list[tuple[Any | None, str]] = []

    def walk(element) -> None:
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                continue
            if child.tag in {qn("w:t"), qn("w:instrText")}:
                segments.append((child, child.text or ""))
                continue
            if child.tag in {qn("w:br"), qn("w:cr")} :
                segments.append((None, "\n"))
                continue
            if child.tag == qn("w:tab"):
                segments.append((None, "\t"))
                continue
            walk(child)

    walk(paragraph_element)
    return segments


def _paragraph_text_elements(paragraph_element) -> list[Any]:
    elements: list[Any] = []

    def walk(element) -> None:
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                continue
            if child.tag in {qn("w:t"), qn("w:instrText")} :
                elements.append(child)
                continue
            walk(child)

    walk(paragraph_element)
    return elements


def _span_index_for_position(
    spans: list[tuple[Any, int, int]],
    position: int,
) -> int | None:
    for index, (_element, start, end) in enumerate(spans):
        if start <= position < end:
            return index
    return None


def _set_text_element_value(text_element, value: Any) -> None:
    normalized = str(value or "")
    text_element.text = normalized
    space_attribute = "{http://www.w3.org/XML/1998/namespace}space"
    if normalized.startswith(" ") or normalized.endswith(" "):
        text_element.set(space_attribute, "preserve")
    else:
        text_element.attrib.pop(space_attribute, None)


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def _validate_accepted_candidates(candidates: list[dict[str, Any]]) -> None:
    ids: list[str] = []
    for candidate in candidates:
        source = str(candidate.get("source", ""))
        if source == "checkbox_choice":
            for field in candidate.get("fields", []) or []:
                field_id = str(field.get("id", "")).strip()
                if not field_id:
                    raise AutomaticDetectionError("Uma opção detectada não possui ID.")
                ids.append(field_id)
            continue
        field_id = str(candidate.get("field_id", "")).strip()
        if not field_id:
            raise AutomaticDetectionError("Uma sugestão selecionada não possui ID.")
        ids.append(field_id)
        if str(candidate.get("source", "")) == "repeatable_table":
            columns = [
                dict(column)
                for column in candidate.get("columns", []) or []
                if isinstance(column, dict)
            ]
            column_ids = [str(column.get("id", "")).strip() for column in columns]
            if len(columns) < 2 or any(not column_id for column_id in column_ids):
                raise AutomaticDetectionError(
                    f"A tabela repetível '{candidate.get('label', field_id)}' não possui colunas válidas."
                )
            if len(set(column_ids)) != len(column_ids):
                raise AutomaticDetectionError(
                    f"A tabela repetível '{candidate.get('label', field_id)}' possui colunas repetidas."
                )
        if str(candidate.get("type", "")) == "dropdown":
            if len(compact_dropdown_options(candidate.get("options", []))) < 2:
                raise AutomaticDetectionError(
                    f"Configure pelo menos duas opções para '{candidate.get('label', field_id)}'."
                )
    duplicates = sorted({field_id for field_id in ids if ids.count(field_id) > 1})
    if duplicates:
        raise AutomaticDetectionError(
            "IDs repetidos nas sugestões: " + ", ".join(duplicates)
        )


def _safe_tag_option(value: Any) -> str:
    text = _normalize_space(value)
    if not text:
        raise AutomaticDetectionError("Uma opção detectada está vazia.")
    if "}}" in text:
        raise AutomaticDetectionError(
            "Uma opção contém '}}', sequência reservada para fechar tags."
        )
    return text.replace("|", " / ")


def _candidate_first_ordinal(candidate: dict[str, Any]) -> int:
    location = candidate.get("location", {}) or {}
    if "paragraph" in location:
        try:
            return int(location["paragraph"])
        except (TypeError, ValueError):
            return 10**9
    paragraphs = location.get("paragraphs", []) or []
    try:
        return min(int(value) for value in paragraphs)
    except (TypeError, ValueError):
        return 10**9


def _short_choice_label(value: str) -> str:
    cleaned = _clean_label(value)
    lowered = cleaned.casefold()
    rules = (
        (("não se aplica", "nao se aplica"), "Não se aplica"),
        (("não ser superior", "nao ser superior", "provável valor"), "Valor abaixo do limite"),
        (("demandas supervenientes", "art. 13"), "Demanda superveniente"),
        (("emergencial", "calamidade pública", "calamidade publica"), "Emergência ou calamidade"),
        (("sigilos", "sigilosa", "sigilosas"), "Informações sigilosas"),
        (("verificado posteriormente", "setor administrativo"), "Análise administrativa posterior"),
    )
    for tokens, label in rules:
        if any(token in lowered for token in tokens):
            return label
    first_sentence = re.split(r"(?<=[.!?;])\s+", cleaned, maxsplit=1)[0]
    if len(first_sentence) <= 88:
        return first_sentence
    return first_sentence[:85].rstrip(" ,;:-") + "…"


def _looks_like_page_header(value: str) -> bool:
    lowered = value.casefold()
    tokens = (
        "secretaria de estado",
        "governo da paraíba",
        "governo da paraiba",
        "sistema integrado de gerenciamento",
        "cep:",
    )
    return sum(token in lowered for token in tokens) >= 2


def _local_label(value: str) -> str:
    text = _normalize_space(value)
    match = LABEL_TAIL_PATTERN.search(text)
    if match:
        return _clean_label(match.group(1))
    # After another placeholder the remaining text normally starts with the
    # next local label, e.g. " Matrícula: ".
    text = text.strip(" |;–—-")
    if _is_reasonable_label(text):
        return _clean_label(text)
    return ""


def _instruction_label(text: str) -> str:
    cleaned = _normalize_space(text)
    cleaned = re.sub(
        r"^(?:informar|informe|descrever|descreva|detalhar|detalhe|"
        r"indicar|indique|justificar|justifique|preencher|preencha)\s+",
        "",
        cleaned,
        flags=re.IGNORECASE,
    )
    if len(cleaned) > 76:
        cleaned = cleaned[:73].rstrip(" ,;:-") + "…"
    return cleaned[:1].upper() + cleaned[1:] if cleaned else "Campo a preencher"


def _legacy_placeholder_field_id(token: str) -> str:
    """Return a stable valid assisted ID for a legacy ``{token}`` marker."""

    normalized = unicodedata.normalize("NFKD", str(token or ""))
    normalized = "".join(
        character
        for character in normalized
        if not unicodedata.combining(character)
    ).casefold()
    normalized = re.sub(r"[^a-z0-9_.-]+", "_", normalized)
    normalized = re.sub(r"_+", "_", normalized).strip("._-")
    if not normalized:
        normalized = "campo"
    if not normalized[0].isalpha():
        normalized = "campo_" + normalized
    return f"auto.{normalized[:72]}"


def _make_field_id(label: str) -> str:
    slug = _slug(label)
    slug = SECTION_NUMBER_PATTERN.sub("", slug)
    slug = slug.strip("._-")
    if not slug:
        slug = "campo"
    if slug[0].isdigit():
        slug = "campo_" + slug
    return f"auto.{slug[:72]}"


def _unique_field_id(base: str, used: set[str]) -> str:
    base = str(base or "auto.campo").strip()
    candidate = base
    index = 2
    while candidate in used:
        candidate = f"{base}_{index}"
        index += 1
    return candidate


def _slug(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", str(value or ""))
    normalized = "".join(
        character
        for character in normalized
        if not unicodedata.combining(character)
    ).casefold()
    normalized = SECTION_NUMBER_PATTERN.sub("", normalized)
    normalized = re.sub(r"[^a-z0-9]+", "_", normalized)
    return normalized.strip("_")


def _humanize_id(field_id: str) -> str:
    text = str(field_id).split(".")[-1].replace("_", " ").replace("-", " ")
    return text[:1].upper() + text[1:]


def _clean_label(value: str) -> str:
    text = _normalize_space(value)
    text = SECTION_NUMBER_PATTERN.sub("", text)
    text = text.strip(" :：–—-\t\r\n")
    return text


def _looks_like_section_label(value: str) -> bool:
    raw = _normalize_space(value)
    if not raw or len(raw) > 190:
        return False
    return bool(re.match(r"^\d+(?:\.\d+)*\.?\s+", raw)) or raw.endswith(":")


def _is_reasonable_label(value: str, *, maximum: int = 150) -> bool:
    text = _normalize_space(value).strip(" :：–—-")
    if len(text) < 2 or len(text) > maximum:
        return False
    if PLACEHOLDER_PATTERN.search(text):
        return False
    if CHOICE_SEPARATOR_PATTERN.match(text):
        return False
    return sum(character.isalpha() for character in text) >= 2


def _normalize_space(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()
