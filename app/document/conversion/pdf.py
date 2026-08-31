from __future__ import annotations

from dataclasses import dataclass
from html import escape
from importlib.util import find_spec
from io import BytesIO
import os
import tempfile
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any, Callable, Iterable

from lxml import etree as LET


class PdfConversionError(RuntimeError):
    """Raised when a DOCX cannot be converted to PDF."""


class DocxConversionError(RuntimeError):
    """Raised when a PDF cannot be converted to DOCX."""


class ConversionCancelledError(RuntimeError):
    """Raised when the user cancels a cooperative conversion."""


def _ensure_not_cancelled(cancel_check: Callable[[], bool] | None) -> None:
    if cancel_check is not None and cancel_check():
        raise ConversionCancelledError("Conversão cancelada pelo usuário.")


def repair_legacy_pdf_docx_layout(source_path: Path, destination_path: Path) -> bool:
    """Repair paragraph-width artifacts created by older Padroniza PDF imports.

    Older PDF->DOCX reconstruction treated the visible PDF text bbox as the
    logical Word paragraph width and wrote both left and right indents.  Short
    labels therefore became extremely narrow paragraphs.  Only documents
    explicitly marked by Padroniza as ``Convertido de PDF`` are eligible.

    The source is never mutated unless ``source_path == destination_path``; in
    that case the repaired OOXML package is published atomically after the ZIP
    handle is closed.  Returns ``True`` only when layout attributes changed.
    """

    source = Path(source_path).expanduser().resolve()
    destination = Path(destination_path).expanduser().resolve()
    if not source.is_file() or source.suffix.casefold() != ".docx":
        return False

    core_ns = "http://purl.org/dc/elements/1.1/"
    cp_ns = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    ns = {"w": w_ns, "dc": core_ns, "cp": cp_ns}

    try:
        with zipfile.ZipFile(source, "r") as archive:
            names = set(archive.namelist())
            if "word/document.xml" not in names or "docProps/core.xml" not in names:
                return False
            core = LET.fromstring(archive.read("docProps/core.xml"))
            subject_nodes = core.xpath("//dc:subject", namespaces=ns)
            subject = "".join(node.text or "" for node in subject_nodes).strip().casefold()
            if "convertido de pdf" not in subject:
                return False

            document_root = LET.fromstring(archive.read("word/document.xml"))
            changed = False
            left_attr = f"{{{w_ns}}}left"
            right_attr = f"{{{w_ns}}}right"
            val_attr = f"{{{w_ns}}}val"

            for paragraph in document_root.xpath(".//w:p", namespaces=ns):
                ppr = paragraph.find(f"{{{w_ns}}}pPr")
                if ppr is None:
                    continue
                indent = ppr.find(f"{{{w_ns}}}ind")
                if indent is None:
                    continue
                justification = ppr.find(f"{{{w_ns}}}jc")
                alignment = (
                    justification.get(val_attr, "").casefold()
                    if justification is not None
                    else ""
                )

                # Mirror the corrected converter: centered paragraphs use the
                # normal text width; right-aligned blocks keep only their right
                # positional edge; ordinary/left blocks keep only the left edge.
                attrs_to_remove: tuple[str, ...]
                if alignment in {"center", "both", "distribute"}:
                    attrs_to_remove = (left_attr, right_attr)
                elif alignment in {"right", "end"}:
                    attrs_to_remove = (left_attr,)
                else:
                    attrs_to_remove = (right_attr,)

                for attr in attrs_to_remove:
                    if attr in indent.attrib:
                        del indent.attrib[attr]
                        changed = True
                if not indent.attrib:
                    ppr.remove(indent)

            if not changed:
                return False

            replacement_xml = LET.tostring(
                document_root,
                encoding="utf-8",
                xml_declaration=True,
                standalone=True,
            )
            entries = [(info, archive.read(info.filename)) for info in archive.infolist()]

        destination.parent.mkdir(parents=True, exist_ok=True)
        temporary: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(
                prefix=f".{destination.stem}-layout-",
                suffix=".docx",
                dir=str(destination.parent),
                delete=False,
            ) as handle:
                temporary = Path(handle.name)
            with zipfile.ZipFile(temporary, "w", compression=zipfile.ZIP_DEFLATED) as output:
                for info, payload in entries:
                    output.writestr(
                        info,
                        replacement_xml if info.filename == "word/document.xml" else payload,
                    )
            with zipfile.ZipFile(temporary, "r") as check:
                broken = check.testzip()
                if broken:
                    raise DocxConversionError(
                        f"A correção de layout produziu uma entrada DOCX inválida: {broken}"
                    )
            os.replace(temporary, destination)
        finally:
            if temporary is not None:
                temporary.unlink(missing_ok=True)
        return True
    except (OSError, zipfile.BadZipFile, LET.XMLSyntaxError) as exc:
        raise DocxConversionError(
            f"Não foi possível corrigir o layout do DOCX convertido de PDF: {exc}"
        ) from exc


@dataclass(frozen=True)
class ConverterCapabilities:
    docx_to_pdf: bool
    pdf_to_docx: bool
    description: str


def available_converter() -> str:
    """Return the built-in converter label used by the user interface."""

    return 'Conversor integrado'


def converter_capabilities() -> ConverterCapabilities:
    """Describe the conversion engine without requiring external applications."""

    docx_to_pdf = find_spec("reportlab") is not None
    pdf_to_docx = find_spec("fitz") is not None

    return ConverterCapabilities(
        docx_to_pdf=docx_to_pdf,
        pdf_to_docx=pdf_to_docx,
        description='Mecanismo integrado de conversão com ReportLab e PyMuPDF',
    )


def convert_docx_to_pdf(
    docx_path: Path,
    pdf_path: Path | None = None,
    *,
    warnings: list[str] | None = None,
    cancel_check: Callable[[], bool] | None = None,
) -> Path:
    """
    Convert a DOCX to PDF without Microsoft Word or LibreOffice.

    The converter preserves common paragraphs, headings, tables, images,
    alignment, margins, page breaks, and basic run formatting. Very complex
    Word-only layout features may be simplified.
    """

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.table import Table as DocxTable
        from docx.text.paragraph import Paragraph as DocxParagraph
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import (
            Image,
            PageBreak,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except Exception as exc:
        raise PdfConversionError(
            "Os componentes integrados de conversão de DOCX para PDF não estão disponíveis. "
            "Instale os pacotes listados em requirements.txt."
        ) from exc

    _ensure_not_cancelled(cancel_check)
    source = Path(docx_path).expanduser().resolve()
    destination = Path(pdf_path or source.with_suffix(".pdf")).expanduser().resolve()

    if not source.exists():
        raise PdfConversionError(f"O arquivo DOCX não existe: {source}")
    if source.suffix.lower() != ".docx":
        raise PdfConversionError('O arquivo de entrada deve usar a extensão .docx.')

    warning_list = warnings if warnings is not None else []

    _ensure_not_cancelled(cancel_check)
    try:
        document = Document(str(source))
    except Exception as exc:
        raise PdfConversionError(f"Não foi possível abrir o arquivo DOCX: {exc}") from exc

    destination.parent.mkdir(parents=True, exist_ok=True)

    section = document.sections[0]
    page_width = _length_points(section.page_width, 595.28)
    page_height = _length_points(section.page_height, 841.89)
    left_margin = _length_points(section.left_margin, 54.0)
    right_margin = _length_points(section.right_margin, 54.0)
    top_margin = _length_points(section.top_margin, 54.0)
    bottom_margin = _length_points(section.bottom_margin, 54.0)
    content_width = max(72.0, page_width - left_margin - right_margin)
    content_height = max(72.0, page_height - top_margin - bottom_margin)

    header_lines = _container_text_lines(section.header.paragraphs)
    footer_lines = _container_text_lines(section.footer.paragraphs)
    page_graphics = _page_anchor_graphics(
        section,
        page_width=page_width,
        page_height=page_height,
        left_margin=left_margin,
        top_margin=top_margin,
        warning_list=warning_list,
        qn=qn,
    )

    def page_decoration(canvas, _doc) -> None:
        canvas.saveState()
        _draw_page_anchor_graphics(canvas, page_graphics, colors, ImageReader)
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#555555"))

        header_y = page_height - max(18.0, top_margin * 0.55)
        for index, line in enumerate(header_lines[:3]):
            canvas.drawCentredString(
                page_width / 2,
                header_y - (index * 9),
                _pdf_safe_text(line),
            )

        footer_y = max(12.0, bottom_margin * 0.35)
        for index, line in enumerate(reversed(footer_lines[-3:])):
            canvas.drawCentredString(
                page_width / 2,
                footer_y + (index * 9),
                _pdf_safe_text(line),
            )

        canvas.restoreState()

    pdf_document = SimpleDocTemplate(
        str(destination),
        pagesize=(page_width, page_height),
        leftMargin=left_margin,
        rightMargin=right_margin,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        title=document.core_properties.title or source.stem,
        author=document.core_properties.author or "",
        subject=document.core_properties.subject or "",
    )

    body_story: list[Any] = []
    image_streams: list[BytesIO] = []

    paragraph_counter: dict[tuple[str, int], int] = {}

    for block in _iter_docx_blocks(document, DocxParagraph, DocxTable):
        _ensure_not_cancelled(cancel_check)
        if isinstance(block, DocxParagraph):
            page_break_before = bool(block.paragraph_format.page_break_before)
            if page_break_before and body_story:
                body_story.append(PageBreak())

            paragraph_markup = _paragraph_markup(block)
            images = _paragraph_images(
                block,
                content_width=content_width,
                content_height=content_height,
                image_class=Image,
                image_reader=ImageReader,
                qn=qn,
                streams=image_streams,
                warning_list=warning_list,
            )

            has_page_break = bool(
                block._element.xpath(
                    ".//w:br[@w:type='page'] | .//w:lastRenderedPageBreak"
                )
            )

            if paragraph_markup.strip():
                style = _reportlab_paragraph_style(
                    block,
                    ParagraphStyle,
                    WD_ALIGN_PARAGRAPH,
                    TA_LEFT,
                    TA_CENTER,
                    TA_RIGHT,
                    TA_JUSTIFY,
                )
                bullet_text = _paragraph_bullet_text(block, paragraph_counter)
                try:
                    body_story.append(
                        Paragraph(
                            paragraph_markup,
                            style,
                            bulletText=bullet_text,
                        )
                    )
                except Exception:
                    body_story.append(
                        Paragraph(
                            escape(_pdf_safe_text(block.text)),
                            style,
                            bulletText=bullet_text,
                        )
                    )
            elif not images and not has_page_break:
                spacing = max(
                    2.0,
                    _length_points(block.paragraph_format.space_after, 6.0),
                )
                body_story.append(Spacer(1, spacing))

            for image in images:
                body_story.append(image)
                body_story.append(Spacer(1, 4))

            if has_page_break:
                body_story.append(PageBreak())

        else:
            converted_table = _convert_docx_table(
                block,
                content_width=content_width,
                paragraph_class=Paragraph,
                paragraph_style_class=ParagraphStyle,
                table_class=Table,
                table_style_class=TableStyle,
                colors_module=colors,
            )
            if converted_table is not None:
                body_story.append(converted_table)
                body_story.append(Spacer(1, 7))

    if not body_story:
        empty_style = ParagraphStyle(
            "EmptyDocument",
            fontName="Helvetica",
            fontSize=10,
            leading=13,
        )
        body_story.append(Paragraph("", empty_style))

    _ensure_not_cancelled(cancel_check)
    try:
        pdf_document.build(
            body_story,
            onFirstPage=page_decoration,
            onLaterPages=page_decoration,
        )
    except Exception as exc:
        raise PdfConversionError(f"Não foi possível criar o PDF: {exc}") from exc

    _ensure_not_cancelled(cancel_check)
    if not destination.exists() or destination.stat().st_size == 0:
        raise PdfConversionError('O conversor terminou sem criar um arquivo PDF.')

    return destination


def convert_pdf_to_docx(
    pdf_path: Path,
    docx_path: Path | None = None,
    *,
    warnings: list[str] | None = None,
    cancel_check: Callable[[], bool] | None = None,
) -> Path:
    """
    Convert a PDF to an editable DOCX without an external office suite.

    Text-based PDFs become editable paragraphs and tables. Pages without
    extractable text are inserted as page images so their visible content is
    preserved, though those scanned pages are not editable without OCR.
    """

    try:
        import fitz
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except Exception as exc:
        raise DocxConversionError(
            "Os componentes integrados de conversão de PDF para DOCX não estão disponíveis. "
            "Instale os pacotes listados em requirements.txt."
        ) from exc

    _ensure_not_cancelled(cancel_check)
    source = Path(pdf_path).expanduser().resolve()
    destination = Path(docx_path or source.with_suffix(".docx")).expanduser().resolve()

    if not source.exists():
        raise DocxConversionError(f"O arquivo PDF não existe: {source}")
    if source.suffix.lower() != ".pdf":
        raise DocxConversionError('O arquivo de entrada deve usar a extensão .pdf.')

    warning_list = warnings if warnings is not None else []

    try:
        pdf = fitz.open(str(source))
    except Exception as exc:
        raise DocxConversionError(f"Não foi possível abrir o arquivo PDF: {exc}") from exc

    if pdf.needs_pass:
        pdf.close()
        raise DocxConversionError('PDFs protegidos por senha não são compatíveis.')

    if pdf.page_count == 0:
        pdf.close()
        raise DocxConversionError('O PDF não contém páginas.')

    document = Document()
    document.core_properties.title = source.stem
    document.core_properties.subject = "Convertido de PDF"

    scanned_pages = 0
    table_count = 0

    with TemporaryDirectory(prefix="padroniza-pdf-to-docx-") as temp_folder:
        temp_dir = Path(temp_folder)

        try:
            for page_index in range(pdf.page_count):
                _ensure_not_cancelled(cancel_check)
                page = pdf.load_page(page_index)

                if page_index == 0:
                    section = document.sections[0]
                else:
                    section = document.add_section(WD_SECTION.NEW_PAGE)

                _configure_docx_section(
                    section,
                    page.rect.width,
                    page.rect.height,
                    Inches,
                )

                page_dict = page.get_text("dict", sort=True)
                blocks = list(page_dict.get("blocks", []))
                text_blocks = [
                    block
                    for block in blocks
                    if block.get("type") == 0
                    and _block_text(block).strip()
                ]

                try:
                    page_drawings = list(page.get_drawings())
                except Exception:
                    page_drawings = []

                table_items: list[dict[str, Any]] = []
                table_boxes: list[tuple[float, float, float, float]] = []
                try:
                    finder = page.find_tables()
                    for found_table in finder.tables:
                        extracted = found_table.extract()
                        if not extracted:
                            continue
                        bbox = tuple(float(value) for value in found_table.bbox)
                        table_items.append(
                            {
                                "kind": "table",
                                "bbox": bbox,
                                "rows": extracted,
                                "table": found_table,
                            }
                        )
                        table_boxes.append(bbox)
                except Exception:
                    table_items = []
                    table_boxes = []

                if not text_blocks:
                    scanned_pages += 1
                    image_path = temp_dir / f"page-{page_index + 1}.png"
                    matrix = fitz.Matrix(1.6, 1.6)
                    pixmap = page.get_pixmap(matrix=matrix, alpha=False)
                    pixmap.save(str(image_path))

                    content_width_inches = max(
                        1.0,
                        (page.rect.width - 54.0) / 72.0,
                    )
                    document.add_picture(
                        str(image_path),
                        width=Inches(content_width_inches),
                    )
                    continue

                items: list[dict[str, Any]] = []

                for block in blocks:
                    block_type = block.get("type")
                    bbox = tuple(float(value) for value in block.get("bbox", (0, 0, 0, 0)))

                    if block_type == 0:
                        if not _block_text(block).strip():
                            continue
                        if any(_bbox_center_inside(bbox, table_box) for table_box in table_boxes):
                            continue
                        items.append(
                            {
                                "kind": "text",
                                "bbox": bbox,
                                "block": block,
                            }
                        )
                    elif block_type == 1 and block.get("image"):
                        items.append(
                            {
                                "kind": "image",
                                "bbox": bbox,
                                "block": block,
                            }
                        )

                items.extend(table_items)
                items.sort(key=lambda item: (item["bbox"][1], item["bbox"][0]))

                for item_index, item in enumerate(items):
                    kind = item["kind"]

                    if kind == "text":
                        _append_pdf_text_block(
                            document,
                            item["block"],
                            page_width=page.rect.width,
                            page_drawings=page_drawings,
                            Inches=Inches,
                            Pt=Pt,
                            RGBColor=RGBColor,
                            alignment_enum=WD_ALIGN_PARAGRAPH,
                        )
                    elif kind == "table":
                        _append_pdf_table(
                            document,
                            item["rows"],
                            Pt=Pt,
                            page=page,
                            table_geometry=item.get("table"),
                            page_drawings=page_drawings,
                        )
                        table_count += 1
                    elif kind == "image":
                        block = item["block"]
                        image_bytes = block.get("image")
                        extension = str(block.get("ext") or "png").lower()
                        if extension == "jpeg":
                            extension = "jpg"
                        image_path = temp_dir / (
                            f"page-{page_index + 1}-image-{item_index + 1}.{extension}"
                        )
                        image_path.write_bytes(image_bytes)

                        bbox = item["bbox"]
                        width_inches = max(0.5, (bbox[2] - bbox[0]) / 72.0)
                        max_width_inches = max(1.0, (page.rect.width - 72.0) / 72.0)
                        document.add_picture(
                            str(image_path),
                            width=Inches(min(width_inches, max_width_inches)),
                        )

        except ConversionCancelledError:
            pdf.close()
            raise
        except Exception as exc:
            pdf.close()
            raise DocxConversionError(f"Não foi possível converter o PDF: {exc}") from exc

        pdf.close()

        if scanned_pages:
            warning_list.append(
                f"{scanned_pages} página(s) não continham texto extraível e foram "
                "inseridas como imagens. Essas páginas não são editáveis sem OCR."
            )
        if table_count:
            warning_list.append(
                f"Foram detectadas e recriadas {table_count} tabela(s). Células mescladas complexas podem ser simplificadas."
            )

        destination.parent.mkdir(parents=True, exist_ok=True)
        try:
            _ensure_not_cancelled(cancel_check)
            document.save(str(destination))
        except ConversionCancelledError:
            raise
        except Exception as exc:
            raise DocxConversionError(f"Não foi possível salvar o DOCX: {exc}") from exc

    _ensure_not_cancelled(cancel_check)
    if not destination.exists() or destination.stat().st_size == 0:
        raise DocxConversionError('O conversor terminou sem criar um arquivo DOCX.')

    return destination


def _iter_docx_blocks(document, paragraph_class, table_class) -> Iterable[Any]:
    for child in document.element.body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "p":
            yield paragraph_class(child, document)
        elif tag == "tbl":
            yield table_class(child, document)


def _page_anchor_graphics(
    section: Any,
    *,
    page_width: float,
    page_height: float,
    left_margin: float,
    top_margin: float,
    warning_list: list[str],
    qn: Any,
) -> list[dict[str, Any]]:
    """Extract anchored header/footer pictures and simple filled shapes.

    The integrated PDF backend cannot ask Word to render floating DrawingML.
    This small bridge preserves the official Padroniza letterhead (and similar
    anchored page decorations) using the same Word positions on every page.
    """

    graphics: list[dict[str, Any]] = []
    containers = (
        (section.header, "header"),
        (section.footer, "footer"),
    )
    header_distance = _length_points(getattr(section, "header_distance", None), 36.0)

    for container, kind in containers:
        try:
            anchors = container._element.xpath(".//wp:anchor")
        except Exception:
            anchors = []
        for anchor in anchors:
            try:
                extent = anchor.xpath("./*[local-name()='extent']")
                if not extent:
                    continue
                width = float(extent[0].get("cx", 0) or 0) / 12700.0
                height = float(extent[0].get("cy", 0) or 0) / 12700.0
                if width <= 0 or height <= 0:
                    continue

                x = _anchor_position_points(
                    anchor,
                    axis="H",
                    page_extent=page_width,
                    margin_start=left_margin,
                    paragraph_base=left_margin,
                )
                y_top = _anchor_position_points(
                    anchor,
                    axis="V",
                    page_extent=page_height,
                    margin_start=top_margin,
                    paragraph_base=header_distance if kind == "header" else 0.0,
                )

                blips = anchor.xpath(".//*[local-name()='blip']")
                if blips:
                    relationship_id = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    if relationship_id:
                        related = container.part.related_parts[relationship_id]
                        graphics.append(
                            {
                                "kind": "image",
                                "blob": related.blob,
                                "x": x,
                                "y": page_height - y_top - height,
                                "width": width,
                                "height": height,
                            }
                        )
                        continue

                fill_nodes = anchor.xpath(
                    ".//*[local-name()='spPr']/*[local-name()='solidFill']/*[local-name()='srgbClr']"
                )
                path_nodes = anchor.xpath(
                    ".//*[local-name()='spPr']/*[local-name()='custGeom']/*[local-name()='pathLst']/*[local-name()='path']"
                )
                if fill_nodes and path_nodes:
                    color = str(fill_nodes[0].get("val", "808080") or "808080")
                    path_node = path_nodes[0]
                    path_width = float(path_node.get("w", 0) or 0)
                    path_height = float(path_node.get("h", 0) or 0)
                    if path_width <= 0 or path_height <= 0:
                        continue
                    points: list[tuple[float, float]] = []
                    for command in list(path_node):
                        point_nodes = command.xpath("./*[local-name()='pt']")
                        if not point_nodes:
                            continue
                        px = float(point_nodes[0].get("x", 0) or 0)
                        py = float(point_nodes[0].get("y", 0) or 0)
                        points.append(
                            (
                                x + (px / path_width) * width,
                                page_height - y_top - (py / path_height) * height,
                            )
                        )
                    if len(points) >= 3:
                        graphics.append(
                            {
                                "kind": "polygon",
                                "points": points,
                                "color": color,
                            }
                        )
            except Exception as exc:
                warning_list.append(f"Uma decoração de cabeçalho/rodapé foi simplificada: {exc}")

    return graphics


def _anchor_position_points(
    anchor: Any,
    *,
    axis: str,
    page_extent: float,
    margin_start: float,
    paragraph_base: float,
) -> float:
    position = anchor.xpath(f"./*[local-name()='position{axis}']")
    if not position:
        return margin_start
    node = position[0]
    relative = str(node.get("relativeFrom", "page") or "page").casefold()
    offset_nodes = node.xpath("./*[local-name()='posOffset']")
    offset = 0.0
    if offset_nodes and offset_nodes[0].text:
        try:
            offset = float(offset_nodes[0].text) / 12700.0
        except (TypeError, ValueError):
            offset = 0.0

    if relative == "page":
        base = 0.0
    elif relative in {"margin", "column"}:
        base = margin_start
    elif relative in {"paragraph", "line", "character"}:
        base = paragraph_base
    else:
        base = margin_start
    return max(-page_extent, min(page_extent * 2.0, base + offset))


def _draw_page_anchor_graphics(
    canvas: Any,
    graphics: list[dict[str, Any]],
    colors_module: Any,
    image_reader: Any,
) -> None:
    for item in graphics:
        if item.get("kind") == "image":
            try:
                canvas.drawImage(
                    image_reader(BytesIO(item["blob"])),
                    float(item["x"]),
                    float(item["y"]),
                    width=float(item["width"]),
                    height=float(item["height"]),
                    preserveAspectRatio=True,
                    mask="auto",
                )
            except Exception:
                continue
        elif item.get("kind") == "polygon":
            points = list(item.get("points", []))
            if len(points) < 3:
                continue
            try:
                canvas.setFillColor(colors_module.HexColor(f"#{item.get('color', '808080')}"))
                path = canvas.beginPath()
                path.moveTo(*points[0])
                for point in points[1:]:
                    path.lineTo(*point)
                path.close()
                canvas.drawPath(path, stroke=0, fill=1)
            except Exception:
                continue


def _length_points(value: Any, default: float) -> float:
    try:
        return float(value.pt)
    except Exception:
        return float(default)


def _container_text_lines(paragraphs: Iterable[Any]) -> list[str]:
    return [
        paragraph.text.strip()
        for paragraph in paragraphs
        if paragraph.text.strip()
    ]


def _pdf_safe_text(value: str) -> str:
    return (
        str(value)
        .replace("☑", "[X]")
        .replace("☒", "[X]")
        .replace("☐", "[ ]")
        .replace("–", "-")
        .replace("—", "-")
        .replace("\u00a0", " ")
    )


def _paragraph_markup(paragraph: Any) -> str:
    markup: list[str] = []

    for run in paragraph.runs:
        text = _pdf_safe_text(run.text)
        if not text:
            continue

        text = escape(text).replace("\n", "<br/>").replace("\t", "&nbsp;&nbsp;&nbsp;&nbsp;")

        style_font = getattr(getattr(paragraph, "style", None), "font", None)
        size = run.font.size or getattr(style_font, "size", None)
        color = run.font.color.rgb if run.font.color is not None else None

        opening: list[str] = []
        closing: list[str] = []

        bold = run.bold
        if bold is None:
            bold = bool(getattr(style_font, "bold", False))
        italic = run.italic
        if italic is None:
            italic = bool(getattr(style_font, "italic", False))

        if bold:
            opening.append("<b>")
            closing.insert(0, "</b>")
        if italic:
            opening.append("<i>")
            closing.insert(0, "</i>")
        if run.underline:
            opening.append("<u>")
            closing.insert(0, "</u>")
        if run.font.superscript:
            opening.append("<super>")
            closing.insert(0, "</super>")
        elif run.font.subscript:
            opening.append("<sub>")
            closing.insert(0, "</sub>")

        font_attributes: list[str] = []
        if size is not None:
            try:
                font_attributes.append(f'size="{max(5.0, min(float(size.pt), 72.0)):.1f}"')
            except Exception:
                pass
        if color is not None:
            font_attributes.append(f'color="#{color}"')

        if font_attributes:
            opening.append(f"<font {' '.join(font_attributes)}>")
            closing.insert(0, "</font>")

        markup.append("".join(opening) + text + "".join(closing))

    if markup:
        return "".join(markup)
    return escape(_pdf_safe_text(paragraph.text))


def _reportlab_paragraph_style(
    paragraph: Any,
    paragraph_style_class: Any,
    alignment_enum: Any,
    ta_left: int,
    ta_center: int,
    ta_right: int,
    ta_justify: int,
) -> Any:
    style_font = getattr(getattr(paragraph, "style", None), "font", None)
    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "")

    size_value = getattr(style_font, "size", None)
    font_size = _length_points(size_value, 10.5) if size_value is not None else 10.5
    bold = bool(getattr(style_font, "bold", False))

    if style_name.casefold().startswith("heading"):
        bold = True
        if "1" in style_name:
            font_size = max(font_size, 16.0)
        elif "2" in style_name:
            font_size = max(font_size, 14.0)
        else:
            font_size = max(font_size, 12.0)

    alignment_map = {
        alignment_enum.LEFT: ta_left,
        alignment_enum.CENTER: ta_center,
        alignment_enum.RIGHT: ta_right,
        alignment_enum.JUSTIFY: ta_justify,
        alignment_enum.DISTRIBUTE: ta_justify,
    }
    alignment = alignment_map.get(paragraph.alignment, ta_left)

    formatting = paragraph.paragraph_format
    line_spacing = formatting.line_spacing
    if hasattr(line_spacing, "pt"):
        leading = max(font_size * 1.05, float(line_spacing.pt))
    elif isinstance(line_spacing, (int, float)):
        leading = max(font_size * 1.05, font_size * float(line_spacing))
    else:
        leading = font_size * 1.22

    font_name = "Helvetica-Bold" if bold else "Helvetica"

    return paragraph_style_class(
        f"DocxParagraph-{id(paragraph)}",
        fontName=font_name,
        fontSize=font_size,
        leading=leading,
        alignment=alignment,
        leftIndent=_length_points(formatting.left_indent, 0.0),
        rightIndent=_length_points(formatting.right_indent, 0.0),
        firstLineIndent=_length_points(formatting.first_line_indent, 0.0),
        spaceBefore=_length_points(formatting.space_before, 0.0),
        spaceAfter=_length_points(formatting.space_after, 6.0),
        keepWithNext=bool(formatting.keep_with_next),
        splitLongWords=True,
        allowWidows=1,
        allowOrphans=1,
    )


def _paragraph_bullet_text(paragraph: Any, counters: dict[tuple[str, int], int]) -> str | None:
    properties = paragraph._p.pPr
    if properties is None or properties.numPr is None:
        return None

    num_id_element = properties.numPr.numId
    level_element = properties.numPr.ilvl
    num_id = str(num_id_element.val if num_id_element is not None else "0")
    level = int(level_element.val if level_element is not None else 0)
    key = (num_id, level)
    counters[key] = counters.get(key, 0) + 1

    style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").casefold()
    if "bullet" in style_name:
        return "•"
    return f"{counters[key]}."


def _paragraph_images(
    paragraph: Any,
    *,
    content_width: float,
    content_height: float,
    image_class: Any,
    image_reader: Any,
    qn: Any,
    streams: list[BytesIO],
    warning_list: list[str],
) -> list[Any]:
    images: list[Any] = []

    for run in paragraph.runs:
        blips = run._element.xpath(".//a:blip")
        extents = run._element.xpath(".//wp:extent")

        for index, blip in enumerate(blips):
            relationship_id = blip.get(qn("r:embed"))
            if not relationship_id:
                continue

            try:
                part = run.part.related_parts[relationship_id]
                stream = BytesIO(part.blob)
                streams.append(stream)
                reader = image_reader(stream)
                intrinsic_width, intrinsic_height = reader.getSize()

                if index < len(extents):
                    width = float(extents[index].get("cx", 0)) / 12700.0
                    height = float(extents[index].get("cy", 0)) / 12700.0
                else:
                    width = float(intrinsic_width) * 0.75
                    height = float(intrinsic_height) * 0.75

                if width <= 0 or height <= 0:
                    width = min(content_width, float(intrinsic_width) * 0.75)
                    height = width * (float(intrinsic_height) / max(1.0, float(intrinsic_width)))

                scale = min(
                    1.0,
                    content_width / max(width, 1.0),
                    content_height / max(height, 1.0),
                )
                images.append(
                    image_class(
                        stream,
                        width=width * scale,
                        height=height * scale,
                    )
                )
            except Exception as exc:
                warning_list.append(f"Uma imagem incorporada foi ignorada: {exc}")

    return images


def _convert_docx_table(
    docx_table: Any,
    *,
    content_width: float,
    paragraph_class: Any,
    paragraph_style_class: Any,
    table_class: Any,
    table_style_class: Any,
    colors_module: Any,
) -> Any | None:
    rows: list[list[Any]] = []

    for row in docx_table.rows:
        converted_row: list[Any] = []
        for cell in row.cells:
            text_parts = [
                _pdf_safe_text(paragraph.text)
                for paragraph in cell.paragraphs
            ]
            cell_text = "<br/>".join(escape(value) for value in text_parts)
            cell_style = paragraph_style_class(
                f"TableCell-{id(cell)}",
                fontName="Helvetica",
                fontSize=8.5,
                leading=10.5,
                spaceAfter=0,
            )
            converted_row.append(paragraph_class(cell_text or " ", cell_style))
        rows.append(converted_row)

    if not rows:
        return None

    column_count = max(len(row) for row in rows)
    if column_count <= 0:
        return None

    for row in rows:
        while len(row) < column_count:
            row.append("")

    column_widths = [content_width / column_count] * column_count

    converted = table_class(
        rows,
        colWidths=column_widths,
        repeatRows=1 if len(rows) > 1 else 0,
        hAlign="LEFT",
    )
    converted.setStyle(
        table_style_class(
            [
                ("GRID", (0, 0), (-1, -1), 0.5, colors_module.HexColor("#777777")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("BACKGROUND", (0, 0), (-1, 0), colors_module.HexColor("#eeeeee")),
            ]
        )
    )
    return converted


def _configure_docx_section(section: Any, width_points: float, height_points: float, Inches: Any) -> None:
    section.page_width = Inches(max(1.0, width_points / 72.0))
    section.page_height = Inches(max(1.0, height_points / 72.0))
    section.top_margin = Inches(0.38)
    section.bottom_margin = Inches(0.38)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)


def _xml_safe_text(value: Any) -> str:
    """Remove characters that XML 1.0 / python-docx cannot serialize.

    Some SIAGOV PDFs contain embedded control bytes in otherwise extractable
    text. PyMuPDF faithfully returns them, while ``python-docx`` rejects the
    whole conversion with ``All strings must be XML compatible``. Preserve all
    printable Unicode and the XML whitespace characters instead of failing the
    user's scan because of one invisible byte.
    """

    text = str(value or "")
    return "".join(
        char
        for char in text
        if (
            char in {"\t", "\n", "\r"}
            or "\x20" <= char <= "\ud7ff"
            or "\ue000" <= char <= "\ufffd"
            or "\U00010000" <= char <= "\U0010ffff"
        )
    )


def _block_text(block: dict[str, Any]) -> str:
    values: list[str] = []
    for line in block.get("lines", []):
        values.append("".join(str(span.get("text", "")) for span in line.get("spans", [])))
    return "\n".join(values)


def _bbox_center_inside(
    bbox: tuple[float, float, float, float],
    container: tuple[float, float, float, float],
) -> bool:
    center_x = (bbox[0] + bbox[2]) / 2.0
    center_y = (bbox[1] + bbox[3]) / 2.0
    return container[0] <= center_x <= container[2] and container[1] <= center_y <= container[3]


def _append_pdf_text_block(
    document: Any,
    block: dict[str, Any],
    *,
    page_width: float,
    page_drawings: list[dict[str, Any]],
    Inches: Any,
    Pt: Any,
    RGBColor: Any,
    alignment_enum: Any,
) -> None:
    lines = list(block.get("lines", []) or [])

    # PyMuPDF often stores several visually horizontal form cells as separate
    # ``lines`` inside one text block.  Preserve that horizontal relationship
    # as a one-row Word table instead of turning it into manual line breaks.
    # This is important for PDF forms such as ``Placa | Data | Horário`` and
    # ``Próxima revisão | Responsável``.
    if _pdf_lines_form_one_visual_row(lines):
        _append_pdf_visual_row_table(
            document,
            lines,
            page_width=page_width,
            page_drawings=page_drawings,
            Inches=Inches,
            Pt=Pt,
        )
        return

    paragraph = document.add_paragraph()
    bbox = tuple(float(value) for value in block.get("bbox", (0, 0, 0, 0)))
    block_width = max(1.0, bbox[2] - bbox[0])
    center = (bbox[0] + bbox[2]) / 2.0

    if abs(center - (page_width / 2.0)) < page_width * 0.06 and block_width < page_width * 0.82:
        paragraph.alignment = alignment_enum.CENTER
    elif bbox[0] > page_width * 0.55:
        paragraph.alignment = alignment_enum.RIGHT
    else:
        paragraph.alignment = alignment_enum.LEFT

    # A PDF text block's bounding box describes the pixels actually occupied by
    # its current text; it does *not* describe the width of the logical paragraph.
    # Using both bbox edges as Word paragraph indents therefore constrains short
    # labels/headings to a tiny column (e.g. ``1. Solicitante``), and any changed
    # or generated value wraps one word/character per line.  Preserve only the
    # positional edge implied by the alignment and leave the opposite side free
    # to use the section's normal text width.
    edge_margin = 24.0
    if paragraph.alignment == alignment_enum.CENTER:
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.right_indent = None
    elif paragraph.alignment == alignment_enum.RIGHT:
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.right_indent = Inches(
            max(0.0, (page_width - bbox[2] - edge_margin) / 72.0)
        )
    else:
        paragraph.paragraph_format.left_indent = Inches(
            max(0.0, (bbox[0] - edge_margin) / 72.0)
        )
        paragraph.paragraph_format.right_indent = None
    paragraph.paragraph_format.space_after = Pt(1.5)
    paragraph.paragraph_format.line_spacing = 1.0

    for line_index, line in enumerate(lines):
        line_text = _pdf_enriched_line_text(line, page_drawings)
        spans = list(line.get("spans", []) or [])
        original_text = "".join(str(span.get("text", "")) for span in spans).strip()
        if spans and line_text == original_text:
            # Preserve the original PDF formatting whenever no synthetic form
            # marker or underline had to be inserted.
            for span in spans:
                text = str(span.get("text", ""))
                if not text:
                    continue

                run = paragraph.add_run(_xml_safe_text(text))
                font_name = str(span.get("font", ""))
                flags = int(span.get("flags", 0) or 0)
                run.bold = bool(flags & 16) or "bold" in font_name.casefold()
                run.italic = bool(flags & 2) or "italic" in font_name.casefold() or "oblique" in font_name.casefold()
                run.font.size = Pt(max(5.0, min(float(span.get("size", 10.0)), 72.0)))

                color_value = int(span.get("color", 0) or 0)
                red = (color_value >> 16) & 255
                green = (color_value >> 8) & 255
                blue = color_value & 255
                run.font.color.rgb = RGBColor(red, green, blue)
        else:
            run = paragraph.add_run(_xml_safe_text(line_text))
            if spans:
                first = spans[0]
                font_name = str(first.get("font", ""))
                flags = int(first.get("flags", 0) or 0)
                run.bold = bool(flags & 16) or "bold" in font_name.casefold()
                run.italic = bool(flags & 2) or "italic" in font_name.casefold() or "oblique" in font_name.casefold()
                run.font.size = Pt(max(5.0, min(float(first.get("size", 10.0)), 72.0)))
                color_value = int(first.get("color", 0) or 0)
                run.font.color.rgb = RGBColor(
                    (color_value >> 16) & 255,
                    (color_value >> 8) & 255,
                    color_value & 255,
                )

        if line_index < len(lines) - 1:
            paragraph.add_run().add_break()


def _pdf_lines_form_one_visual_row(lines: list[dict[str, Any]]) -> bool:
    if len(lines) < 2:
        return False
    boxes = [tuple(float(v) for v in line.get("bbox", (0, 0, 0, 0))) for line in lines]
    centers = [(box[1] + box[3]) / 2.0 for box in boxes]
    if max(centers) - min(centers) > 2.8:
        return False
    ordered = sorted(boxes, key=lambda box: box[0])
    return all(
        ordered[index + 1][0] - ordered[index][2] >= 12.0
        for index in range(len(ordered) - 1)
    )


def _append_pdf_visual_row_table(
    document: Any,
    lines: list[dict[str, Any]],
    *,
    page_width: float,
    page_drawings: list[dict[str, Any]],
    Inches: Any,
    Pt: Any,
) -> None:
    ordered = sorted(lines, key=lambda line: float(line.get("bbox", (0, 0, 0, 0))[0]))
    table = document.add_table(rows=1, cols=len(ordered))
    table.autofit = False

    starts = [float(line.get("bbox", (0, 0, 0, 0))[0]) for line in ordered]
    right_edge = max(
        float(ordered[-1].get("bbox", (0, 0, page_width, 0))[2]),
        page_width - 42.0,
    )
    boundaries = starts[1:] + [right_edge]

    for index, line in enumerate(ordered):
        cell = table.cell(0, index)
        width_points = max(42.0, boundaries[index] - starts[index])
        cell.width = Inches(width_points / 72.0)
        text = _pdf_enriched_line_text(line, page_drawings)
        cell.text = _xml_safe_text(text)
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(8.5)
                spans = list(line.get("spans", []) or [])
                if spans:
                    first = spans[0]
                    font_name = str(first.get("font", ""))
                    flags = int(first.get("flags", 0) or 0)
                    run.bold = bool(flags & 16) or "bold" in font_name.casefold()
                    run.italic = bool(flags & 2) or "italic" in font_name.casefold() or "oblique" in font_name.casefold()


def _pdf_enriched_line_text(
    line: dict[str, Any],
    page_drawings: list[dict[str, Any]],
) -> str:
    spans = list(line.get("spans", []) or [])
    if not spans:
        return ""
    bbox = tuple(float(value) for value in line.get("bbox", (0, 0, 0, 0)))
    components: list[tuple[float, str]] = []
    for span in spans:
        text = str(span.get("text", ""))
        if text:
            span_bbox = tuple(float(v) for v in span.get("bbox", bbox))
            components.append((span_bbox[0], text))

    for marker in _pdf_checkbox_rectangles(page_drawings):
        marker_center_y = (marker[1] + marker[3]) / 2.0
        line_center_y = (bbox[1] + bbox[3]) / 2.0
        if abs(marker_center_y - line_center_y) > max(5.0, (bbox[3] - bbox[1]) * 0.75):
            continue
        if marker[2] <= bbox[0] + 4.0 and bbox[0] - marker[2] <= 18.0:
            components.append((marker[0], "☐ "))

    components.sort(key=lambda item: item[0])
    text = "".join(value for _x, value in components).strip()

    fill = _pdf_following_fill_line(bbox, page_drawings)
    if fill is not None:
        width = max(12.0, fill[2] - fill[0])
        underscore_count = max(4, min(48, int(width / 6.0)))
        text = f"{text} {'_' * underscore_count}".strip()
    return text


def _pdf_checkbox_rectangles(
    page_drawings: list[dict[str, Any]],
) -> list[tuple[float, float, float, float]]:
    markers: list[tuple[float, float, float, float]] = []
    for drawing in page_drawings:
        rect = drawing.get("rect")
        if rect is None:
            continue
        box = tuple(float(value) for value in rect)
        width = box[2] - box[0]
        height = box[3] - box[1]
        if 5.0 <= width <= 16.0 and 5.0 <= height <= 16.0 and abs(width - height) <= 3.0:
            markers.append(box)
    return markers


def _pdf_horizontal_fill_lines(
    page_drawings: list[dict[str, Any]],
) -> list[tuple[float, float, float, float]]:
    lines: list[tuple[float, float, float, float]] = []
    for drawing in page_drawings:
        rect = drawing.get("rect")
        if rect is None:
            continue
        box = tuple(float(value) for value in rect)
        width = box[2] - box[0]
        height = box[3] - box[1]
        if width >= 20.0 and height <= 1.6:
            lines.append(box)
    return lines


def _pdf_following_fill_line(
    text_bbox: tuple[float, float, float, float],
    page_drawings: list[dict[str, Any]],
) -> tuple[float, float, float, float] | None:
    candidates: list[tuple[float, tuple[float, float, float, float]]] = []
    text_center_y = (text_bbox[1] + text_bbox[3]) / 2.0
    for line in _pdf_horizontal_fill_lines(page_drawings):
        line_y = (line[1] + line[3]) / 2.0
        if line[0] < text_bbox[2] - 3.0:
            continue
        if line[0] - text_bbox[2] > 28.0:
            continue
        if abs(line_y - text_center_y) > 8.5:
            continue
        candidates.append((line[0] - text_bbox[2], line))
    if not candidates:
        return None
    return min(candidates, key=lambda item: item[0])[1]


def _pdf_cell_text_with_controls(
    page: Any,
    cell_bbox: tuple[float, float, float, float],
    fallback: str,
    page_drawings: list[dict[str, Any]],
) -> str:
    x0, y0, x1, y1 = cell_bbox
    spans: list[tuple[float, float, str]] = []
    try:
        page_dict = page.get_text("dict", sort=True)
        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    span_bbox = tuple(float(v) for v in span.get("bbox", (0, 0, 0, 0)))
                    cx = (span_bbox[0] + span_bbox[2]) / 2.0
                    cy = (span_bbox[1] + span_bbox[3]) / 2.0
                    if x0 <= cx <= x1 and y0 <= cy <= y1:
                        text = str(span.get("text", "")).strip()
                        if text:
                            spans.append((span_bbox[0], span_bbox[1], text))
    except Exception:
        spans = []

    markers = [
        marker
        for marker in _pdf_checkbox_rectangles(page_drawings)
        if x0 <= (marker[0] + marker[2]) / 2.0 <= x1
        and y0 <= (marker[1] + marker[3]) / 2.0 <= y1
    ]

    if markers and spans:
        parts: list[str] = []
        for marker in sorted(markers, key=lambda item: item[0]):
            right_candidates = [span for span in spans if span[0] >= marker[2] - 1.0]
            if not right_candidates:
                continue
            nearest = min(
                right_candidates,
                key=lambda span: (span[0] - marker[2], abs(span[1] - marker[1])),
            )
            label = nearest[2]
            if label not in parts:
                parts.append(label)
        if parts:
            return " ".join(f"☐ {part}" for part in parts)

    text = _xml_safe_text(fallback).strip()
    if not text and any(
        x0 <= (line[0] + line[2]) / 2.0 <= x1
        and y0 <= (line[1] + line[3]) / 2.0 <= y1
        for line in _pdf_horizontal_fill_lines(page_drawings)
    ):
        width = max(24.0, x1 - x0 - 12.0)
        return "_" * max(4, min(40, int(width / 7.0)))
    return text


def _append_pdf_table(
    document: Any,
    rows: list[list[Any]],
    *,
    Pt: Any,
    page: Any | None = None,
    table_geometry: Any | None = None,
    page_drawings: list[dict[str, Any]] | None = None,
) -> None:
    cleaned_rows = [
        ["" if value is None else str(value) for value in row]
        for row in rows
    ]
    if not cleaned_rows:
        return

    column_count = max(len(row) for row in cleaned_rows)
    table = document.add_table(rows=len(cleaned_rows), cols=column_count)
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    geometry_rows = (
        list(getattr(table_geometry, "rows", []) or [])
        if table_geometry is not None
        else []
    )
    drawings = list(page_drawings or [])

    for row_index, row in enumerate(cleaned_rows):
        for column_index in range(column_count):
            value = row[column_index] if column_index < len(row) else ""
            if page is not None and row_index < len(geometry_rows):
                cells = list(getattr(geometry_rows[row_index], "cells", []) or [])
                if column_index < len(cells) and cells[column_index] is not None:
                    value = _pdf_cell_text_with_controls(
                        page,
                        tuple(float(v) for v in cells[column_index]),
                        value,
                        drawings,
                    )

            cell = table.cell(row_index, column_index)
            cell.text = _xml_safe_text(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)
                    if row_index == 0:
                        run.bold = True


