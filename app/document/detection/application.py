from __future__ import annotations

import os
import shutil
import tempfile
from collections import defaultdict
from copy import deepcopy
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.paragraph import Paragraph

from app.domain.field_metadata import compact_dropdown_options
from app.document.detection.identifiers import _normalize_space
from app.document.detection.models import AutomaticDetectionError, ParagraphRecord as _ParagraphRecord
from app.document.detection.patterns import CHECKBOX_LINE_PATTERN, CHECKBOX_TOKEN_PATTERN, ISOLATED_CHECK_MARK_PATTERN
from app.document.detection.records import _collect_paragraph_records
from app.document.detection.word_helpers import _remove_floating_checkmark_shapes, _unique_row_cells

def apply_docx_field_candidates(
    source_docx: Path,
    destination_docx: Path,
    candidates: Iterable[dict[str, Any]],
) -> Path:
    """Apply approved suggestions to a copy by converting them into tags.

    This is the key safety property of assisted detection: after approval the
    existing tag scanner and DOCX engine handle the model exactly like a
    manually tagged template.
    """

    source = Path(source_docx)
    destination = Path(destination_docx)
    accepted = [deepcopy(item) for item in candidates if isinstance(item, dict)]
    if not accepted:
        raise AutomaticDetectionError("Nenhuma sugestão foi selecionada.")

    destination.parent.mkdir(parents=True, exist_ok=True)
    _validate_accepted_candidates(accepted)

    fd, staged_name = tempfile.mkstemp(
        prefix=f".{destination.stem}-detect-",
        suffix=".docx",
        dir=str(destination.parent),
    )
    os.close(fd)
    staged = Path(staged_name)
    try:
        shutil.copy2(source, staged)
        document = Document(str(staged))
        records = _collect_paragraph_records(document)
        by_ordinal = {record.ordinal: record for record in records}

        # Whole-block operations are applied first. They do not invalidate the
        # stored Paragraph XML objects used by the span replacements below.
        for candidate in accepted:
            kind = str(candidate.get("location", {}).get("kind", ""))
            if kind == "repeatable_table":
                _apply_repeatable_table(candidate, document)
            elif kind == "paragraph_block":
                _apply_paragraph_block(candidate, by_ordinal)
            elif kind == "paragraph_list":
                _apply_paragraph_list(candidate, by_ordinal)
            elif kind == "checkbox_group":
                _apply_checkbox_group(candidate, by_ordinal)
            elif kind == "checkbox_group_inline":
                _apply_inline_checkbox_group(candidate, by_ordinal)
            elif kind == "checkbox_group_multi_cell":
                _apply_multi_cell_checkbox_group(candidate, by_ordinal)

        # Apply text spans from right to left inside each paragraph so offsets stay
        # valid even when one line contains several placeholders.
        spans_by_paragraph: dict[int, list[dict[str, Any]]] = defaultdict(list)
        for candidate in accepted:
            location = candidate.get("location", {}) or {}
            kind = str(location.get("kind", ""))
            if kind == "text_span":
                spans_by_paragraph[int(location.get("paragraph", -1))].append(candidate)
            elif kind == "text_spans":
                for span in location.get("spans", []) or []:
                    if not isinstance(span, dict):
                        continue
                    clone = deepcopy(candidate)
                    clone["location"] = dict(span, kind="text_span")
                    spans_by_paragraph[int(span.get("paragraph", -1))].append(clone)

        for ordinal, paragraph_candidates in spans_by_paragraph.items():
            record = by_ordinal.get(ordinal)
            if record is None:
                raise AutomaticDetectionError(
                    "A estrutura do DOCX mudou durante a detecção automática. "
                    "Execute a análise novamente."
                )
            replacements: list[tuple[int, int, str]] = []
            for candidate in paragraph_candidates:
                location = candidate.get("location", {}) or {}
                replacements.append(
                    (
                        int(location.get("start", 0)),
                        int(location.get("end", 0)),
                        _tag_for_candidate(candidate),
                    )
                )
            _replace_paragraph_spans(record.paragraph, replacements)

        for candidate in accepted:
            location = candidate.get("location", {}) or {}
            kind = str(location.get("kind", ""))
            if kind == "append_tag":
                ordinal = int(location.get("paragraph", -1))
                record = by_ordinal.get(ordinal)
                if record is None:
                    raise AutomaticDetectionError(
                        "Não foi possível localizar uma área aprovada no DOCX. "
                        "Execute a análise novamente."
                    )
                _append_tag_to_paragraph(record.paragraph, _tag_for_candidate(candidate))
                continue
            if kind not in {"paragraph", "empty_cell"}:
                continue
            ordinal = int(location.get("paragraph", -1))
            record = by_ordinal.get(ordinal)
            if record is None:
                raise AutomaticDetectionError(
                    "Não foi possível localizar uma área aprovada no DOCX. "
                    "Execute a análise novamente."
                )
            _replace_entire_paragraph(record.paragraph, _tag_for_candidate(candidate))

        document.save(str(staged))
        _validate_detection_roundtrip(staged, accepted)
        os.replace(staged, destination)
        return destination
    except Exception:
        try:
            staged.unlink(missing_ok=True)
        except Exception:
            pass
        raise

def _validate_detection_roundtrip(
    prepared_docx: Path,
    accepted: list[dict[str, Any]],
) -> None:
    """Strictly re-scan our own output before publishing it.

    Assisted detection is never allowed to produce a DOCX that the normal tag
    scanner cannot understand. This catches duplicate repeatable columns, bad
    prefixes, malformed dropdowns and future writer/scanner disagreements at
    the boundary where they are cheapest to recover.
    """

    from app.document.docx.scanner import clear_docx_scan_cache, scan_docx_fields

    clear_docx_scan_cache()
    try:
        scanned = [dict(field) for field in scan_docx_fields(Path(prepared_docx))]
    except Exception as exc:
        raise AutomaticDetectionError(
            "A validação de ida-e-volta rejeitou as tags geradas automaticamente: "
            f"{exc}"
        ) from exc

    by_id = {str(field.get("id", "")).strip(): field for field in scanned}
    expected_ids: set[str] = set()
    for candidate in accepted:
        source = str(candidate.get("source", ""))
        if source == "checkbox_choice":
            expected_ids.update(
                str(field.get("id", "")).strip()
                for field in candidate.get("fields", []) or []
                if str(field.get("id", "")).strip()
            )
        else:
            field_id = str(candidate.get("field_id", "")).strip()
            if field_id:
                expected_ids.add(field_id)

    missing = sorted(expected_ids - set(by_id))
    if missing:
        raise AutomaticDetectionError(
            "A validação de ida-e-volta não encontrou os campos gerados: "
            + ", ".join(missing)
        )

    for candidate in accepted:
        if str(candidate.get("type", "")).casefold() == "repeatable_list":
            field_id = str(candidate.get("field_id", "")).strip()
            scanned_field = by_id.get(field_id, {})
            if str(scanned_field.get("type", "")).casefold() != "repeatable_list":
                raise AutomaticDetectionError(
                    f"A lista repetível '{field_id}' mudou de tipo ao ser reescaneada."
                )
            expected_style = str(candidate.get("list_style", "bullet") or "bullet").casefold()
            expected_punctuation = str(
                candidate.get("list_punctuation", "semicolon") or "semicolon"
            ).casefold()
            if str(scanned_field.get("list_style", "bullet")).casefold() != expected_style:
                raise AutomaticDetectionError(
                    f"A lista repetível '{field_id}' mudou de estilo ao ser reescaneada."
                )
            if (
                str(scanned_field.get("list_punctuation", "semicolon")).casefold()
                != expected_punctuation
            ):
                raise AutomaticDetectionError(
                    f"A lista repetível '{field_id}' mudou de pontuação ao ser reescaneada."
                )

        if str(candidate.get("source", "")) != "repeatable_table":
            continue
        field_id = str(candidate.get("field_id", "")).strip()
        scanned_field = by_id.get(field_id, {})
        scanned_columns = [
            str(column.get("id", "")).strip()
            for column in scanned_field.get("columns", []) or []
            if isinstance(column, dict)
        ]
        expected_columns = [
            str(column.get("id", "")).strip()
            for column in candidate.get("columns", []) or []
            if isinstance(column, dict)
        ]
        # auto-number can be represented as ``item`` by both sides; compare the
        # stable set while preserving duplicate detection in the strict scanner.
        if [value for value in scanned_columns if value] != [value for value in expected_columns if value]:
            raise AutomaticDetectionError(
                f"A tabela '{field_id}' mudou de colunas ao ser reescaneada. "
                f"Esperado: {expected_columns}; encontrado: {scanned_columns}."
            )


def _tag_for_candidate(candidate: dict[str, Any]) -> str:
    field_id = str(candidate.get("field_id", "")).strip()
    field_type = str(candidate.get("type", "text")).strip().casefold()
    if not field_id:
        raise AutomaticDetectionError("Uma sugestão selecionada não possui ID de campo.")

    location = dict(candidate.get("location", {}) or {})
    if str(location.get("render", "")).casefold() == "currency_words":
        return f"{{{{currency_words:{field_id}}}}}"

    if field_type == "date":
        return f"{{{{date:{field_id}}}}}"
    if field_type == "checkbox":
        return f"{{{{checkbox:{field_id}}}}}"
    if field_type == "repeatable_list":
        style = str(candidate.get("list_style", "bullet") or "bullet").casefold()
        punctuation = str(candidate.get("list_punctuation", "semicolon") or "semicolon").casefold()
        return "{{repeat_list:" + field_id + "|" + style + "|" + punctuation + "}}"
    if field_type == "dropdown":
        options = compact_dropdown_options(candidate.get("options", []))
        if len(options) < 2:
            raise AutomaticDetectionError(
                f"A lista '{candidate.get('label', field_id)}' precisa de pelo menos duas opções."
            )
        encoded = []
        for option in options:
            if isinstance(option, dict):
                label = _safe_tag_option(option.get("label", ""))
                value = _safe_tag_option(option.get("value", ""))
                encoded.append(value if label == value else f"{label} => {value}")
            else:
                encoded.append(_safe_tag_option(option))
        prefix = "single_choice" if str(candidate.get("layout", "")) == "choice" else "dropdown"
        return "{{" + prefix + ":" + field_id + "|" + "|".join(encoded) + "}}"
    return f"{{{{{field_id}}}}}"


def _apply_paragraph_block(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    ordinals = [
        int(value)
        for value in candidate.get("location", {}).get("paragraphs", [])
    ]
    records = [by_ordinal.get(value) for value in ordinals]
    records = [record for record in records if record is not None]
    if not records:
        raise AutomaticDetectionError("Bloco de alternativas não encontrado no DOCX.")
    first = records[0].paragraph
    _replace_entire_paragraph(first, _tag_for_candidate(candidate))
    for record in records[1:]:
        _remove_paragraph(record.paragraph)


def _apply_paragraph_list(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    ordinals = [
        int(value)
        for value in candidate.get("location", {}).get("paragraphs", []) or []
    ]
    records = [by_ordinal.get(value) for value in ordinals]
    records = [record for record in records if record is not None]
    if not records:
        raise AutomaticDetectionError("A lista detectada não foi encontrada no DOCX.")
    _replace_entire_paragraph(records[0].paragraph, _tag_for_candidate(candidate))
    for record in records[1:]:
        _remove_paragraph(record.paragraph)


def _apply_repeatable_table(
    candidate: dict[str, Any],
    document: _Document,
) -> None:
    location = candidate.get("location", {}) or {}
    try:
        table_index = int(location.get("document_table_index", -1))
        template_row_index = int(location.get("template_row", -1))
        data_rows = sorted(
            {int(value) for value in location.get("data_rows", []) or []}
        )
        insert_before_row = int(location.get("insert_before_row", -1))
    except (TypeError, ValueError):
        raise AutomaticDetectionError("Tabela repetível detectada com posição inválida.")

    if table_index < 0 or table_index >= len(document.tables):
        raise AutomaticDetectionError("A tabela repetível detectada não foi encontrada no DOCX.")
    table = document.tables[table_index]
    synthetic_template = bool(location.get("synthetic_template_row", False))
    if not synthetic_template and (
        template_row_index < 0 or template_row_index >= len(table.rows)
    ):
        raise AutomaticDetectionError("A linha modelo da tabela repetível não foi encontrada.")
    if synthetic_template and (
        insert_before_row < 0 or insert_before_row >= len(table.rows)
    ):
        raise AutomaticDetectionError(
            "A posição da nova linha editável da planilha não foi encontrada."
        )

    field_id = str(candidate.get("field_id", "")).strip()
    columns = [
        dict(column)
        for column in candidate.get("columns", []) or []
        if isinstance(column, dict)
    ]
    if not field_id or len(columns) < 2:
        raise AutomaticDetectionError("A tabela repetível detectada está incompleta.")

    if synthetic_template:
        # ``python-docx`` can append a correctly sized row using the table grid.
        # Move that XML row immediately before the merged narrative/note row so
        # the original document keeps its visual order: header -> editable rows
        # -> note.  Using a fresh row avoids copying header shading/bold styles.
        row = table.add_row()
        target_row = table.rows[insert_before_row]
        target_row._tr.addprevious(row._tr)
    else:
        row = table.rows[template_row_index]

    cells = _unique_row_cells(row)
    repeat_marker_written = False
    for column in columns:
        column_type = str(column.get("type", "text")).strip().casefold()
        column_id = str(column.get("id", "")).strip()
        if not column_id:
            continue
        if column_type == "auto_number":
            column_index = 0
            text = f"{{{{repeat:{field_id}}}}} {{{{row.number}}}}"
            repeat_marker_written = True
        else:
            try:
                column_index = int(column.get("column_index", -1))
            except (TypeError, ValueError):
                column_index = -1
            if column_index < 0:
                continue
            child_id = f"{field_id}.{column_id}"
            if column_type == "date":
                text = f"{{{{date:{child_id}}}}}"
            elif column_type == "checkbox":
                text = f"{{{{checkbox:{child_id}}}}}"
            elif column_type == "dropdown":
                options = compact_dropdown_options(column.get("options", []))
                if len(options) < 2:
                    raise AutomaticDetectionError(
                        f"A coluna '{column.get('label', column_id)}' precisa de pelo menos duas opções."
                    )
                encoded: list[str] = []
                for option in options:
                    if isinstance(option, dict):
                        label = _safe_tag_option(option.get("label", ""))
                        value = _safe_tag_option(option.get("value", ""))
                        encoded.append(value if label == value else f"{label} => {value}")
                    else:
                        encoded.append(_safe_tag_option(option))
                text = "{{dropdown:" + child_id + "|" + "|".join(encoded) + "}}"
            else:
                text = f"{{{{{child_id}}}}}"

        if column_index >= len(cells):
            raise AutomaticDetectionError(
                "A estrutura de colunas da tabela repetível mudou após a análise."
            )
        _replace_cell_with_text(cells[column_index], text)

    if not repeat_marker_written:
        # The detector currently requires an auto-number column, but keep this
        # fallback so reviewed candidates remain robust if that rule evolves.
        first_editable = next(
            (
                column
                for column in columns
                if str(column.get("type", "")).casefold() != "auto_number"
            ),
            None,
        )
        if first_editable is None:
            raise AutomaticDetectionError("A tabela repetível não possui coluna editável.")
        column_index = int(first_editable.get("column_index", 0))
        paragraph = cells[column_index].paragraphs[0]
        paragraph.insert_paragraph_before(f"{{{{repeat:{field_id}}}}}")

    # Keep only the first detected data row as the Word model row. The DOCX
    # engine duplicates it according to the rows entered by the client. A
    # synthetic sheet has no source data rows to remove.
    if not synthetic_template:
        for row_index in sorted(data_rows, reverse=True):
            if row_index == template_row_index or row_index >= len(table.rows):
                continue
            table._tbl.remove(table.rows[row_index]._tr)


def _replace_cell_with_text(cell: _Cell, text: str) -> None:
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph(text)
        return
    _replace_entire_paragraph(paragraphs[0], text)
    for paragraph in paragraphs[1:]:
        _remove_paragraph(paragraph)


def _apply_checkbox_group(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    ordinals = [
        int(value)
        for value in candidate.get("location", {}).get("paragraphs", [])
    ]
    fields = [dict(field) for field in candidate.get("fields", []) or []]
    if len(ordinals) != len(fields):
        raise AutomaticDetectionError("Grupo de caixas de seleção inconsistente.")
    for ordinal, field in zip(ordinals, fields):
        record = by_ordinal.get(ordinal)
        if record is None:
            raise AutomaticDetectionError("Opção de caixa de seleção não encontrada.")
        text = record.paragraph.text or ""
        match = CHECKBOX_LINE_PATTERN.match(text)
        if not match:
            raise AutomaticDetectionError("O texto de uma opção mudou após a análise.")
        replacement = f"{{{{checkbox:{field['id']}}}}} {match.group(1).strip()}"
        _replace_entire_paragraph(record.paragraph, replacement)


def _apply_inline_checkbox_group(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    location = candidate.get("location", {}) or {}
    ordinal = int(location.get("paragraph", -1))
    record = by_ordinal.get(ordinal)
    if record is None:
        raise AutomaticDetectionError("Grupo de caixas de seleção não encontrado.")
    fields = [dict(field) for field in candidate.get("fields", []) or []]
    spans = [
        tuple(int(value) for value in span)
        for span in location.get("checkbox_spans", []) or []
        if isinstance(span, (list, tuple)) and len(span) == 2
    ]
    if len(fields) != len(spans) or not fields:
        raise AutomaticDetectionError("Grupo de caixas de seleção inconsistente.")
    replacements = [
        (start, end, f"{{{{checkbox:{field['id']}}}}}")
        for (start, end), field in zip(spans, fields, strict=True)
    ]
    _replace_paragraph_spans(record.paragraph, replacements)


def _apply_multi_cell_checkbox_group(
    candidate: dict[str, Any],
    by_ordinal: dict[int, _ParagraphRecord],
) -> None:
    location = candidate.get("location", {}) or {}
    ordinals = [int(value) for value in location.get("paragraphs", []) or []]
    spans = [
        tuple(int(value) for value in span)
        for span in location.get("checkbox_spans", []) or []
        if isinstance(span, (list, tuple)) and len(span) == 2
    ]
    marker_modes = [str(value or "text_span") for value in location.get("checkbox_marker_modes", []) or []]
    fields = [dict(field) for field in candidate.get("fields", []) or []]
    if not marker_modes:
        marker_modes = ["text_span"] * len(fields)
    if not fields or len(fields) != len(ordinals) or len(fields) != len(spans) or len(fields) != len(marker_modes):
        raise AutomaticDetectionError("Grupo de caixas de seleção entre células inconsistente.")

    inferred_count = sum(1 for mode in marker_modes if mode == "inferred_blank")
    if inferred_count:
        first_record = by_ordinal.get(ordinals[0]) if ordinals else None
        if first_record is not None:
            _remove_floating_checkmark_shapes(
                first_record.paragraph._p.getroottree().getroot(),
                limit=inferred_count,
            )

    for ordinal, span, mode, field in zip(ordinals, spans, marker_modes, fields, strict=True):
        record = by_ordinal.get(ordinal)
        if record is None:
            raise AutomaticDetectionError("Opção de caixa de seleção não encontrada.")
        replacement = f"{{{{checkbox:{field['id']}}}}}"
        if mode in {"paragraph", "inferred_blank"}:
            _replace_all_paragraph_content(record.paragraph, replacement)
            continue

        start, end = span
        text = record.paragraph.text or ""
        if start < 0 or end <= start or end > len(text):
            raise AutomaticDetectionError("A posição de uma opção mudou após a análise.")
        marker_text = text[start:end]
        if not (
            CHECKBOX_TOKEN_PATTERN.fullmatch(marker_text)
            or ISOLATED_CHECK_MARK_PATTERN.fullmatch(marker_text)
        ):
            raise AutomaticDetectionError("O marcador de uma opção mudou após a análise.")
        _replace_paragraph_spans(
            record.paragraph,
            [(start, end, replacement)],
        )


def _replace_all_paragraph_content(paragraph: Paragraph, text: str) -> None:
    """Replace runs, controls, fields and symbols while preserving paragraph properties."""

    element = paragraph._p
    for child in list(element):
        if child.tag == qn("w:pPr"):
            continue
        element.remove(child)
    paragraph.add_run(text)


def _replace_entire_paragraph(paragraph: Paragraph, text: str) -> None:
    runs = list(paragraph.runs)
    if runs:
        runs[0].text = text
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _append_tag_to_paragraph(paragraph: Paragraph, tag: str) -> None:
    current = paragraph.text or ""
    separator = "" if not current or current.endswith((" ", "\t")) else " "
    paragraph.add_run(separator + str(tag))


def _replace_paragraph_spans(
    paragraph: Paragraph,
    replacements: Iterable[tuple[int, int, str]],
) -> None:
    # Candidate offsets are measured against ``Paragraph.text``. That text
    # includes manual line breaks (w:br/w:cr) as ``\n`` and tabs (w:tab) as
    # ``\t``. The old replacement code concatenated only w:t nodes, making
    # every span after a break or tab drift to the right. This is especially
    # common in DOCX files reconstructed from PDFs.
    segments = _paragraph_position_segments(paragraph._p)
    original_text = "".join(text for _element, text in segments)
    editable_spans: list[tuple[Any, int, int]] = []
    cursor = 0
    for element, text in segments:
        segment_end = cursor + len(text)
        if element is not None and text:
            editable_spans.append((element, cursor, segment_end))
        cursor = segment_end

    if not editable_spans:
        raise AutomaticDetectionError("O trecho detectado não contém texto editável no DOCX.")

    for start, end, replacement in sorted(replacements, key=lambda item: item[0], reverse=True):
        if start < 0 or end <= start or end > len(original_text):
            raise AutomaticDetectionError("Posição de preenchimento inválida no DOCX.")
        start_index = _span_index_for_position(editable_spans, start)
        end_index = _span_index_for_position(editable_spans, end - 1)
        if start_index is None or end_index is None:
            raise AutomaticDetectionError("Não foi possível localizar o trecho detectado no XML do DOCX.")

        start_element, start_offset, _ = editable_spans[start_index]
        end_element, end_offset, _ = editable_spans[end_index]
        local_start = start - start_offset
        local_end = end - end_offset

        if start_index == end_index:
            current = start_element.text or ""
            _set_text_element_value(
                start_element,
                current[:local_start] + replacement + current[local_end:],
            )
            continue

        start_text = start_element.text or ""
        end_text = end_element.text or ""
        _set_text_element_value(start_element, start_text[:local_start] + replacement)
        for element, _node_start, _node_end in editable_spans[start_index + 1 : end_index]:
            _set_text_element_value(element, "")
        _set_text_element_value(end_element, end_text[local_end:])


def _paragraph_position_segments(paragraph_element) -> list[tuple[Any | None, str]]:
    """Return paragraph content in the same coordinate space as Paragraph.text.

    Text/instruction nodes are editable and are returned with their XML node.
    Manual line breaks and tabs occupy positions too, but are represented by a
    ``None`` node because automatic field replacement should never overwrite
    those structural elements. Nested paragraphs are processed separately.
    """

    segments: list[tuple[Any | None, str]] = []

    def walk(element) -> None:
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                continue
            if child.tag in {qn("w:t"), qn("w:instrText")}:
                segments.append((child, child.text or ""))
                continue
            if child.tag in {qn("w:br"), qn("w:cr")} :
                segments.append((None, "\n"))
                continue
            if child.tag == qn("w:tab"):
                segments.append((None, "\t"))
                continue
            walk(child)

    walk(paragraph_element)
    return segments


def _paragraph_text_elements(paragraph_element) -> list[Any]:
    elements: list[Any] = []

    def walk(element) -> None:
        for child in element.iterchildren():
            if child.tag == qn("w:p"):
                continue
            if child.tag in {qn("w:t"), qn("w:instrText")} :
                elements.append(child)
                continue
            walk(child)

    walk(paragraph_element)
    return elements


def _span_index_for_position(
    spans: list[tuple[Any, int, int]],
    position: int,
) -> int | None:
    for index, (_element, start, end) in enumerate(spans):
        if start <= position < end:
            return index
    return None


def _set_text_element_value(text_element, value: Any) -> None:
    normalized = str(value or "")
    text_element.text = normalized
    space_attribute = "{http://www.w3.org/XML/1998/namespace}space"
    if normalized.startswith(" ") or normalized.endswith(" "):
        text_element.set(space_attribute, "preserve")
    else:
        text_element.attrib.pop(space_attribute, None)


def _remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    paragraph._p = paragraph._element = None


def _validate_accepted_candidates(candidates: list[dict[str, Any]]) -> None:
    ids: list[str] = []
    for candidate in candidates:
        source = str(candidate.get("source", ""))
        if source == "checkbox_choice":
            for field in candidate.get("fields", []) or []:
                field_id = str(field.get("id", "")).strip()
                if not field_id:
                    raise AutomaticDetectionError("Uma opção detectada não possui ID.")
                ids.append(field_id)
            continue
        field_id = str(candidate.get("field_id", "")).strip()
        if not field_id:
            raise AutomaticDetectionError("Uma sugestão selecionada não possui ID.")
        ids.append(field_id)
        if str(candidate.get("source", "")) == "repeatable_table":
            columns = [
                dict(column)
                for column in candidate.get("columns", []) or []
                if isinstance(column, dict)
            ]
            column_ids = [str(column.get("id", "")).strip() for column in columns]
            if len(columns) < 2 or any(not column_id for column_id in column_ids):
                raise AutomaticDetectionError(
                    f"A tabela repetível '{candidate.get('label', field_id)}' não possui colunas válidas."
                )
            if len(set(column_ids)) != len(column_ids):
                raise AutomaticDetectionError(
                    f"A tabela repetível '{candidate.get('label', field_id)}' possui colunas repetidas."
                )
        if str(candidate.get("type", "")) == "dropdown":
            if len(compact_dropdown_options(candidate.get("options", []))) < 2:
                raise AutomaticDetectionError(
                    f"Configure pelo menos duas opções para '{candidate.get('label', field_id)}'."
                )
    duplicates = sorted({field_id for field_id in ids if ids.count(field_id) > 1})
    if duplicates:
        raise AutomaticDetectionError(
            "IDs repetidos nas sugestões: " + ", ".join(duplicates)
        )


def _safe_tag_option(value: Any) -> str:
    text = _normalize_space(value)
    if not text:
        raise AutomaticDetectionError("Uma opção detectada está vazia.")
    if "}}" in text:
        raise AutomaticDetectionError(
            "Uma opção contém '}}', sequência reservada para fechar tags."
        )
    return text.replace("|", " / ")
