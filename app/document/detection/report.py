from __future__ import annotations

from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

from app.document.detection.records import _collect_paragraph_records
from app.document.detection.roles import ContentRole, classify_record_role
from app.document.detection.structure import DocumentStructure, extract_document_structure


@dataclass
class DetectionReport:
    scanner_version: int
    sections: list[str] = field(default_factory=list)
    tables: list[dict[str, Any]] = field(default_factory=list)
    roles: dict[str, int] = field(default_factory=dict)
    candidate_count: int = 0
    ignored_ambiguous_tables: int = 0
    protected_tables: int = 0
    warnings: list[str] = field(default_factory=list)
    invariant_issues: list[dict[str, str]] = field(default_factory=list)

    def as_dict(self) -> dict[str, Any]:
        return {
            "scanner_version": self.scanner_version,
            "sections": list(self.sections),
            "tables": [dict(value) for value in self.tables],
            "roles": dict(self.roles),
            "candidate_count": self.candidate_count,
            "ignored_ambiguous_tables": self.ignored_ambiguous_tables,
            "protected_tables": self.protected_tables,
            "warnings": list(self.warnings),
            "invariant_issues": [dict(value) for value in self.invariant_issues],
        }


def build_detection_report(
    docx_path: Path,
    candidates: list[dict[str, Any]],
    *,
    invariant_issues: list[Any] | None = None,
) -> DetectionReport:
    document = Document(str(Path(docx_path)))
    records = _collect_paragraph_records(document)
    structure = extract_document_structure(document, records)
    roles = Counter(
        classify_record_role(record, records, structure).value
        for record in records
    )
    tables = [
        {
            "index": table.table_index,
            "kind": table.kind,
            "section": table.section,
            "title": table.structure.title,
            "columns": int(table.structure.total_columns),
            "header_labels": list(table.structure.header_labels),
            "header_groups": list(table.structure.header_groups),
            "data_rows": list(table.structure.data_rows),
            "confidence": round(float(table.structure.confidence), 3),
            "protected": table.protected,
            "protection_reason": table.protection_reason,
            "reasons": list(table.structure.reasons),
        }
        for table in structure.tables
    ]
    issues = []
    for issue in invariant_issues or []:
        issues.append({
            "severity": str(getattr(issue, "severity", "warning")),
            "code": str(getattr(issue, "code", "")),
            "message": str(getattr(issue, "message", issue)),
            "field_id": str(getattr(issue, "field_id", "")),
        })
    report_warnings = list(structure.warnings)
    for table in structure.tables:
        if (
            table.kind == "unknown"
            and len(table.structure.data_rows) >= 2
            and table.structure.total_columns >= 3
        ):
            report_warnings.append(
                f"Tabela {table.table_index + 1}: estrutura tabular ambígua preservada sem achatamento; revise manualmente se ela for preenchível."
            )

    return DetectionReport(
        scanner_version=structure.version,
        sections=[section.full_title for section in structure.sections],
        tables=tables,
        roles=dict(roles),
        candidate_count=len(candidates),
        ignored_ambiguous_tables=sum(
            table.kind == "unknown"
            and len(table.structure.data_rows) >= 2
            and table.structure.total_columns >= 3
            for table in structure.tables
        ),
        protected_tables=sum(table.protected for table in structure.tables),
        warnings=report_warnings,
        invariant_issues=issues,
    )
