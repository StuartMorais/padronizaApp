from __future__ import annotations

from docx.document import Document as _Document
from docx.oxml.ns import qn
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.document.detection.models import ParagraphRecord as _ParagraphRecord

def _collect_paragraph_records(document: _Document) -> list[_ParagraphRecord]:
    records: list[_ParagraphRecord] = []
    seen_cells: set[int] = set()

    def add_paragraph(
        paragraph: Paragraph,
        *,
        story: str,
        table_index: int | None = None,
        row_index: int | None = None,
        cell_index: int | None = None,
        cell: _Cell | None = None,
        table: Table | None = None,
    ) -> None:
        records.append(
            _ParagraphRecord(
                ordinal=len(records),
                paragraph=paragraph,
                story=story,
                table_index=table_index,
                row_index=row_index,
                cell_index=cell_index,
                cell=cell,
                table=table,
            )
        )

    def walk_table(table: Table, *, story: str, index: int) -> None:
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                cell_key = id(cell._tc)
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                for paragraph in cell.paragraphs:
                    add_paragraph(
                        paragraph,
                        story=story,
                        table_index=index,
                        row_index=row_index,
                        cell_index=cell_index,
                        cell=cell,
                        table=table,
                    )
                for nested in cell.tables:
                    nonlocal_table_index[0] += 1
                    walk_table(
                        nested,
                        story=story,
                        index=nonlocal_table_index[0],
                    )

    nonlocal_table_index = [-1]
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            add_paragraph(Paragraph(child, document), story="body")
        elif child.tag == qn("w:tbl"):
            nonlocal_table_index[0] += 1
            walk_table(Table(child, document), story="body", index=nonlocal_table_index[0])

    def add_textbox_paragraphs(root, *, story: str, parent) -> None:
        """Expose Word textbox paragraphs to the same scanner pipeline.

        ``python-docx`` intentionally exposes ordinary body/table paragraphs but
        not paragraphs nested under ``w:txbxContent`` (VML/DrawingML text boxes).
        Institutional forms frequently place editable labels/values inside those
        shapes.  Keep them as normal paragraph records so detection and approved
        tag application use the same stable ordinal map.
        """

        try:
            elements = root.xpath(".//w:txbxContent//w:p")
        except Exception:
            elements = []
        seen: set[int] = set()
        for element in elements:
            key = id(element)
            if key in seen:
                continue
            seen.add(key)
            add_paragraph(Paragraph(element, parent), story=story)

    # Text boxes are descendants of normal body paragraphs and therefore are
    # not returned by the direct-child body walk above. Add only their nested
    # text paragraphs here, after the ordinary body/table story, so ordinals are
    # deterministic in both detection and application.
    add_textbox_paragraphs(document.element, story="body_textbox", parent=document)

    # Headers and footers are included for ordinary placeholder suggestions,
    # but long choice-block detection remains focused on body tables.
    for section_index, section in enumerate(document.sections):
        for story_name, story in (
            (f"header_{section_index}", section.header),
            (f"footer_{section_index}", section.footer),
        ):
            for paragraph in story.paragraphs:
                add_paragraph(paragraph, story=story_name)
            for table in story.tables:
                nonlocal_table_index[0] += 1
                walk_table(table, story=story_name, index=nonlocal_table_index[0])
            add_textbox_paragraphs(
                story._element,
                story=f"{story_name}_textbox",
                parent=story,
            )

    return records
