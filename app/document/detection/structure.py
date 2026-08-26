from __future__ import annotations

import re
from dataclasses import dataclass, field
from enum import Enum
from typing import Any

from docx.document import Document as _Document

from app.document.detection.models import ParagraphRecord
from app.document.detection.table_structure import TableKind, TableStructure, analyze_word_table
from app.document.docx.tags import PLACEHOLDER_PATTERN


SCANNER_STRUCTURE_VERSION = 4
_SECTION_RE = re.compile(
    r"^\s*(?P<number>\d+(?:\.\d+)*)(?P<delimiter>[.)])?\s+(?P<title>\S.*?)(?:\s*:)?\s*$"
)
_REPEAT_MARKER_RE = re.compile(r"\{\{\s*repeat\s*:\s*[^{}]+\}\}", re.IGNORECASE)


class StoryZone(str, Enum):
    BODY = "body"
    HEADER = "header"
    FOOTER = "footer"


@dataclass(frozen=True)
class SectionNode:
    number: str
    title: str
    full_title: str
    level: int
    start_ordinal: int
    end_ordinal: int
    parent_number: str = ""


@dataclass
class TableInfo:
    table_index: int
    story: str
    structure: TableStructure
    section: str = ""
    record_ordinals: list[int] = field(default_factory=list)
    protected: bool = False
    protection_reason: str = ""

    @property
    def kind(self) -> str:
        return self.structure.kind.value


@dataclass(frozen=True)
class RecordOwner:
    ordinal: int
    zone: StoryZone
    section: str = ""
    section_number: str = ""
    table_index: int | None = None
    table_kind: str = ""
    tagged: bool = False
    protected: bool = False


@dataclass
class DocumentStructure:
    version: int
    sections: list[SectionNode]
    tables: list[TableInfo]
    owners: dict[int, RecordOwner]
    protected_ordinals: set[int]
    warnings: list[str] = field(default_factory=list)

    def owner_for(self, ordinal: int) -> RecordOwner | None:
        return self.owners.get(int(ordinal))

    def section_for(self, ordinal: int) -> str:
        owner = self.owner_for(ordinal)
        return owner.section if owner is not None else ""

    def table_for(self, table_index: int | None) -> TableInfo | None:
        if table_index is None:
            return None
        for info in self.tables:
            if info.table_index == int(table_index):
                return info
        return None

    def summary(self) -> dict[str, Any]:
        table_kinds: dict[str, int] = {}
        for info in self.tables:
            table_kinds[info.kind] = table_kinds.get(info.kind, 0) + 1
        return {
            "scanner_structure_version": self.version,
            "sections": len(self.sections),
            "tables": len(self.tables),
            "table_kinds": table_kinds,
            "protected_regions": sum(1 for table in self.tables if table.protected),
            "warnings": list(self.warnings),
        }


def extract_document_structure(
    document: _Document,
    records: list[ParagraphRecord],
) -> DocumentStructure:
    """Build a source-of-truth structural model before field heuristics run.

    The detector used to infer fields while it was still learning what owned the
    surrounding content.  This pass reverses that order: sections, story zones,
    physical Word tables and manually-tagged/protected regions are resolved
    first, and every paragraph receives an owner before any field candidate is
    allowed to exist.
    """

    sections = _resolve_sections(records)
    section_by_ordinal = _section_lookup(sections, records)
    table_records: dict[int, list[ParagraphRecord]] = {}
    table_objects: dict[int, Any] = {}
    for record in records:
        if record.table_index is None or record.table is None:
            continue
        table_records.setdefault(int(record.table_index), []).append(record)
        table_objects.setdefault(int(record.table_index), record.table)

    tables: list[TableInfo] = []
    protected_ordinals: set[int] = set()
    warnings: list[str] = []
    for table_index in sorted(table_records):
        rows = table_records[table_index]
        table = table_objects[table_index]
        try:
            analyzed = analyze_word_table(table)
        except Exception as exc:  # structure analysis must never disable scanning
            analyzed = TableStructure(
                TableKind.UNKNOWN,
                max(1, len(table.columns)),
                confidence=0.0,
                reasons=[f"Falha ao analisar a grade física: {type(exc).__name__}: {exc}"],
            )
            warnings.append(
                f"Tabela {table_index + 1}: não foi possível classificar a estrutura física."
            )

        ordinals = sorted({int(record.ordinal) for record in rows})
        has_repeat_marker = any(_has_repeat_marker(record) for record in rows)
        physical_title = analyzed.title.strip()
        section = (
            physical_title
            if parse_numbered_section_heading(physical_title) is not None
            else ""
        )
        if not section and ordinals:
            section = section_by_ordinal.get(ordinals[0], "")

        protected = bool(has_repeat_marker)
        protection_reason = "repeat_tag" if protected else ""
        if protected:
            protected_ordinals.update(ordinals)

        tables.append(
            TableInfo(
                table_index=table_index,
                story=rows[0].story if rows else "body",
                structure=analyzed,
                section=section,
                record_ordinals=ordinals,
                protected=protected,
                protection_reason=protection_reason,
            )
        )

    table_by_index = {info.table_index: info for info in tables}
    owners: dict[int, RecordOwner] = {}
    for record in records:
        section = section_by_ordinal.get(record.ordinal, "")
        section_number = _section_number(section)
        tagged = _record_has_authoritative_content(record)
        explicit_tag = _record_has_explicit_tag(record)
        table_info = table_by_index.get(record.table_index) if record.table_index is not None else None
        protected = bool(explicit_tag or (table_info is not None and table_info.protected))
        if explicit_tag:
            protected_ordinals.add(record.ordinal)
        owner = RecordOwner(
            ordinal=record.ordinal,
            zone=_story_zone(record.story),
            section=table_info.section if table_info and table_info.section else section,
            section_number=_section_number(table_info.section if table_info and table_info.section else section),
            table_index=record.table_index,
            table_kind=table_info.kind if table_info is not None else "",
            tagged=tagged,
            protected=protected,
        )
        owners[record.ordinal] = owner
        try:
            record.structure = owner
        except Exception:
            pass

    return DocumentStructure(
        version=SCANNER_STRUCTURE_VERSION,
        sections=sections,
        tables=tables,
        owners=owners,
        protected_ordinals=protected_ordinals,
        warnings=warnings,
    )


def _resolve_sections(records: list[ParagraphRecord]) -> list[SectionNode]:
    starts: list[tuple[int, str, str, int, str]] = []
    stack: list[tuple[str, int]] = []
    for record in records:
        if _story_zone(record.story) is not StoryZone.BODY:
            continue
        parsed = parse_numbered_section_heading(record.text, paragraph=record.paragraph)
        if parsed is None:
            continue
        number, title, full_title = parsed
        level = number.count(".") + 1
        while stack and stack[-1][1] >= level:
            stack.pop()
        parent = stack[-1][0] if stack else ""
        starts.append((record.ordinal, number, full_title, level, parent))
        stack.append((number, level))

    body_ordinals = [record.ordinal for record in records if _story_zone(record.story) is StoryZone.BODY]
    max_ordinal = max(body_ordinals, default=-1)
    sections: list[SectionNode] = []
    for index, (start, number, full_title, level, parent) in enumerate(starts):
        next_start = starts[index + 1][0] if index + 1 < len(starts) else max_ordinal + 1
        title = full_title[len(number):].lstrip(".) ").rstrip(": ")
        sections.append(
            SectionNode(
                number=number,
                title=title,
                full_title=full_title.rstrip(": "),
                level=level,
                start_ordinal=start,
                end_ordinal=max(start, next_start - 1),
                parent_number=parent,
            )
        )
    return sections


def _section_lookup(
    sections: list[SectionNode],
    records: list[ParagraphRecord],
) -> dict[int, str]:
    lookup: dict[int, str] = {}
    body_sections = sorted(sections, key=lambda item: item.start_ordinal)
    current = ""
    section_index = 0
    for record in records:
        if _story_zone(record.story) is not StoryZone.BODY:
            continue
        while section_index < len(body_sections) and body_sections[section_index].start_ordinal <= record.ordinal:
            current = body_sections[section_index].full_title
            section_index += 1
        lookup[record.ordinal] = current
    return lookup


def parse_numbered_section_heading(value: str, *, paragraph=None) -> tuple[str, str, str] | None:
    text = re.sub(r"\s+", " ", str(value or "")).strip()
    if not text or len(text) > 220:
        return None
    match = _SECTION_RE.match(text)
    if match is None:
        return None
    number = match.group("number")
    delimiter = str(match.group("delimiter") or "")
    # ``1) ...`` / ``2) ...`` are overwhelmingly instruction-list items in
    # institutional forms. True sections use a dot, a hierarchical number
    # such as ``1.1`` or are strongly formatted headings.
    if delimiter == ")":
        return None
    if not delimiter and "." not in number:
        return None
    title = match.group("title").rstrip(": ")
    if not title or not any(character.isalpha() for character in title):
        return None
    if paragraph is not None and not _paragraph_looks_like_heading(paragraph, text):
        return None
    return number, title, f"{number}. {title}"


def is_numbered_section_heading(value: str, *, paragraph=None) -> bool:
    return parse_numbered_section_heading(value, paragraph=paragraph) is not None


def _paragraph_looks_like_heading(paragraph, text: str) -> bool:
    runs = [run for run in getattr(paragraph, "runs", []) if str(run.text or "").strip()]
    if not runs:
        return True
    total = sum(max(1, len(str(run.text or ""))) for run in runs)
    bold = sum(max(1, len(str(run.text or ""))) for run in runs if run.bold is True)
    red = 0
    for run in runs:
        color = getattr(getattr(run.font, "color", None), "rgb", None)
        if color is None:
            continue
        try:
            value = str(color)
            r, g, b = int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16)
        except Exception:
            continue
        if r >= 170 and g <= 120 and b <= 120:
            red += max(1, len(str(run.text or "")))
    bold_ratio = bold / max(1, total)
    red_ratio = red / max(1, total)
    # Red numbered alternatives/instructions such as ``5.1.1. Não se aplica``
    # are content, not navigation sections. Normal section headings in the
    # target corpus are strongly bold and neutral colored.
    if red_ratio >= 0.55 and bold_ratio < 0.85:
        return False
    if text.rstrip().endswith(":"):
        return bold_ratio >= 0.35 or red_ratio < 0.20
    return bold_ratio >= 0.55


def _section_number(value: str) -> str:
    parsed = parse_numbered_section_heading(value)
    return parsed[0] if parsed else ""


def _story_zone(story: str) -> StoryZone:
    value = str(story or "").casefold()
    if value.startswith("header"):
        return StoryZone.HEADER
    if value.startswith("footer"):
        return StoryZone.FOOTER
    return StoryZone.BODY


def _record_has_authoritative_content(record: ParagraphRecord) -> bool:
    text = record.paragraph.text or ""
    if PLACEHOLDER_PATTERN.search(text):
        return True
    element = record.paragraph._p
    return bool(element.xpath(".//w:sdt | .//w:fldChar"))


def _record_has_explicit_tag(record: ParagraphRecord) -> bool:
    return bool(PLACEHOLDER_PATTERN.search(record.paragraph.text or ""))


def _has_repeat_marker(record: ParagraphRecord) -> bool:
    return bool(_REPEAT_MARKER_RE.search(record.paragraph.text or ""))
