from __future__ import annotations

import re
from collections import defaultdict
from typing import Any

from docx.document import Document as _Document

from app.document.docx.tags import PLACEHOLDER_PATTERN
from app.document.detection.candidates import _candidate
from app.document.detection.context_helpers import (
    _contains_authoritative_marker, _context_label_for_record, _detected_placeholder_type,
    _is_pure_fill_area_text, _looks_like_fill_area_text, _paragraph_is_red, _run_color_candidates, _run_is_red,
)
from app.document.detection.identifiers import (
    _clean_label, _humanize_id, _is_reasonable_label, _legacy_placeholder_field_id,
    _local_label, _looks_like_section_label, _make_field_id, _normalize_space, _slug,
    _unique_field_id,
)
from app.document.detection.models import ParagraphRecord as _ParagraphRecord
from app.document.detection.patterns import (
    CHECKBOX_TOKEN_PATTERN, CURRENCY_PLACEHOLDER_PATTERN, FOLLOWUP_AREA_PATTERN,
    GENERIC_DROPDOWN_PATTERN, INSTRUCTION_PATTERN, LEGACY_BRACED_PLACEHOLDER_PATTERN,
    SAMPLE_EMAIL_PLACEHOLDER_PATTERN, SECTION_NUMBER_PATTERN, UNDERSCORE_PLACEHOLDER_PATTERN,
    X_PLACEHOLDER_PATTERN, ZERO_CPF_PLACEHOLDER_PATTERN, ZERO_PHONE_PLACEHOLDER_PATTERN,
)
from app.document.detection.word_helpers import _unique_row_cells
from app.document.understanding.semantic import semantic_label, semantic_section
from app.document.understanding.smart_template import suggest_field_type

_PREFILLED_TEXT_SECTION_PATTERN = re.compile(
    r"\b(?:justificativa|fundamenta[cç][aã]o|descri[cç][aã]o|detalhamento|"
    r"necessidade|motiva[cç][aã]o|objeto|especifica[cç][aã]o|observa[cç][aã]o|"
    r"provid[eê]ncia|parecer|an[aá]lise|considera[cç][oõ]es|informa[cç][oõ]es)\b",
    re.IGNORECASE,
)
_PREFILLED_TEXT_STATIC_PREFIX_PATTERN = re.compile(
    r"^\s*(?:texto\s+fixo|nota|aten[cç][aã]o|aviso|instru[cç][aã]o|"
    r"orienta[cç][aã]o|observa[cç][aã]o\s+fixa|rodap[eé])\s*[:\-]",
    re.IGNORECASE,
)
_COLORED_STATIC_GUIDANCE_PATTERN = re.compile(
    r"^\s*(?:notas?|importante|aten[cç][aã]o|aviso|orienta[cç][aã]o|exemplo|"
    r"veja\s+que|a\s+reda[cç][aã]o\s+acima|observa[cç][aã]o\s+fixa)\b",
    re.IGNORECASE,
)
_COLORED_COLUMN_GUIDANCE_PATTERN = re.compile(
    r"^\s*[\[(]?(?:preencher|informar)\s+(?:a\s+)?coluna\b",
    re.IGNORECASE,
)
_ANGLE_PLACEHOLDER_PATTERN = re.compile(r"^\s*<\s*(.+?)\s*>\s*[.,;:]?\s*$", re.DOTALL)
_GENERIC_INSTRUCTION_PLACEHOLDER_PATTERN = re.compile(
    r"(?:\[\s*(?P<bracket>(?:informar|informe|descrever|descreva|detalhar|detalhe|"
    r"indicar|indique|justificar|justifique|preencher|preencha|especificar|especifique|"
    r"inserir|insira|registrar|registre|definir|defina|identificar|identifique|"
    r"relacionar|relacione|mencionar|mencione)\b[^\]\n]{2,320})\s*\]"
    r"|<\s*(?P<angle>(?:informar|informe|descrever|descreva|detalhar|detalhe|"
    r"indicar|indique|justificar|justifique|preencher|preencha|especificar|especifique|"
    r"inserir|insira|registrar|registre|definir|defina|identificar|identifique|"
    r"relacionar|relacione|mencionar|mencione)\b[^>\n]{2,320})\s*>)",
    re.IGNORECASE,
)
_GENERIC_STATIC_LABEL_PATTERN = re.compile(
    r"^\s*(?:lei|decreto|portaria|resolu[cç][aã]o|instru[cç][aã]o normativa|"
    r"art(?:igo)?\.?|inciso|par[aá]grafo|cap[ií]tulo|se[cç][aã]o|anexo)\b",
    re.IGNORECASE,
)
_GENERIC_VALUE_SENTENCE_PATTERN = re.compile(
    r"(?:[.!?]\s|\b(?:considerando|declaramos|solicitamos|informamos|conforme|"
    r"mediante|observado|observada|devendo|dever[aá]|ser[aá]|fica|ficam)\b)",
    re.IGNORECASE,
)
_FIELD_STYLE_HINT_PATTERN = re.compile(
    r"(?:campo|field|edit[aá]vel|editar|preench|input|placeholder|resposta|vari[aá]vel|formul[aá]rio)",
    re.IGNORECASE,
)



def _detect_blank_followup_areas(
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Detect blank paragraphs directly after observation/justification prompts."""

    by_cell: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if record.cell is not None and record.story == "body":
            by_cell[id(record.cell._tc)].append(record)

    result: list[dict[str, Any]] = []
    for cell_records in by_cell.values():
        cell_records.sort(key=lambda item: item.ordinal)
        for index, record in enumerate(cell_records[:-1]):
            if record.ordinal in reserved_ordinals:
                continue
            prompt = _normalize_space(record.text)
            if not prompt or not FOLLOWUP_AREA_PATTERN.match(prompt):
                continue
            # Use the first truly blank paragraph after the prompt, but do not
            # jump over another visible paragraph.
            target: _ParagraphRecord | None = None
            for following in cell_records[index + 1 :]:
                if _normalize_space(following.text):
                    break
                target = following
                break
            if target is None or target.ordinal in reserved_ordinals:
                continue
            label = _clean_label(prompt)
            field_id = _unique_field_id(_make_field_id(label), known_ids)
            result.append(
                _candidate(
                    field_id=field_id,
                    label=label,
                    field_type="multiline",
                    confidence=0.82,
                    source="empty_cell",
                    preview=prompt,
                    location={
                        "kind": "paragraph",
                        "paragraph": target.ordinal,
                    },
                )
            )
    return result


def _detect_inline_placeholders(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> list[dict[str, Any]]:
    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text):
        return []

    matches = list(CURRENCY_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(X_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(UNDERSCORE_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(ZERO_PHONE_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(ZERO_CPF_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(SAMPLE_EMAIL_PLACEHOLDER_PATTERN.finditer(text))
    matches.extend(LEGACY_BRACED_PLACEHOLDER_PATTERN.finditer(text))
    matches.sort(key=lambda match: (match.start(), -(match.end() - match.start())))
    # Some patterns intentionally overlap (for example an xxxxx@example style
    # e-mail also matches the generic X placeholder). Keep only the widest
    # non-overlapping span so one visual fill area becomes one field.
    non_overlapping: list[re.Match[str]] = []
    for match in matches:
        if any(
            match.start() < existing.end() and existing.start() < match.end()
            for existing in non_overlapping
        ):
            continue
        non_overlapping.append(match)
    matches = sorted(non_overlapping, key=lambda match: match.start())
    if not matches:
        return []

    result: list[dict[str, Any]] = []
    previous_end = 0
    for match in matches:
        is_legacy_braced = LEGACY_BRACED_PLACEHOLDER_PATTERN.fullmatch(match.group(0)) is not None
        local_context = text[previous_end : match.start()]
        # PDF-to-DOCX reconstruction frequently groups several visual PDF
        # lines into one Word paragraph separated by manual line breaks.
        # Prefer the text on the current visual line so a paragraph such as
        # ``Placa: ABC1D23\nData: __/__/____\nHorário: __:__`` yields
        # ``Data`` and ``Horário`` instead of labels polluted by the previous
        # line. Fall back to the complete local context for ordinary DOCX.
        visual_line = local_context.rsplit("\n", 1)[-1]
        label = _local_label(visual_line)
        if not label and visual_line != local_context:
            label = _local_label(local_context)
        if not label:
            label = _context_label_for_record(record, records)

        if is_legacy_braced:
            inner_token = str(match.group(1) or "").strip()
            # A single-braced token is only safe to claim automatically when
            # it has a strong field signal: a local/adjacent label, an inline
            # label ending in ':', or an identifier-like dotted/underscored
            # token.  This deliberately ignores prose such as
            # ``Use chaves {assim} no exemplo``.
            has_inline_colon = bool(re.search(r"[:：]\s*$", local_context))
            token_is_structured = any(separator in inner_token for separator in (".", "_", "-"))
            is_isolated_table_value = (
                record.table is not None
                and _normalize_space(text) == _normalize_space(match.group(0))
                and bool(label)
            )
            if not (label and (has_inline_colon or is_isolated_table_value)) and not token_is_structured:
                previous_end = match.end()
                continue

        if not label and CURRENCY_PLACEHOLDER_PATTERN.fullmatch(match.group(0)):
            # A bare monetary mask is semantically stronger than an anonymous
            # ``Campo XX`` suggestion. Parenthetical text such as
            # ``(valor por extenso)`` remains untouched as contextual text.
            label = "Valor"
        if is_legacy_braced:
            inner_token = str(match.group(1) or "").strip()
            field_id_seed = _legacy_placeholder_field_id(inner_token)
            field_id = _unique_field_id(field_id_seed, known_ids)
        else:
            field_id = _unique_field_id(
                _make_field_id(label or f"campo_{record.ordinal + 1}"),
                known_ids,
            )
        known_ids.add(field_id)
        field_type = _detected_placeholder_type(label or field_id, match.group(0))
        candidate = _candidate(
                field_id=field_id,
                label=label or _humanize_id(field_id),
                field_type=field_type,
                confidence=(0.98 if label else 0.82) if is_legacy_braced else (0.91 if label else 0.74),
                source="legacy_placeholder" if is_legacy_braced else "inline_placeholder",
                preview=match.group(0),
                location={
                    "kind": "text_span",
                    "paragraph": record.ordinal,
                    "start": match.start(),
                    "end": match.end(),
                    "original": match.group(0),
                },
            )
        if is_legacy_braced:
            candidate["legacy_marker"] = match.group(0)
            candidate["legacy_marker_id"] = str(match.group(1) or "").strip()
        result.append(candidate)
        previous_end = match.end()
    return result


def _detect_dropdown_prompt(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text):
        return None

    prompt_pattern = re.compile(
        r"(?:escolher|selecione|selecionar)\s+(?:um|uma)\s+"
        r"(?:item|op[cç][aã]o)\.?\s*$",
        re.IGNORECASE,
    )
    match = prompt_pattern.search(text)
    if match is None:
        return None

    prefix = text[: match.start()]
    label = _local_label(prefix) or _context_label_for_record(record, records)
    field_id = _unique_field_id(_make_field_id(label or "opcao"), known_ids)
    return _candidate(
        field_id=field_id,
        label=label or "Selecione uma opção",
        field_type="dropdown",
        confidence=0.82 if label else 0.70,
        source="dropdown_prompt",
        preview=match.group(0),
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": match.start(),
            "end": match.end(),
            "original": match.group(0),
        },
        options=[],
        default_selected=False,
        requires_configuration=True,
    )


def _detect_labeled_sample_value(
    record: _ParagraphRecord,
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect short example/default values that should still be editable.

    This is deliberately whitelist-based.  Institutional fixed text such as
    ``Órgão: Secretaria ...`` must stay static, while values such as
    ``País: Brasil`` are commonly examples/defaults that help the user
    understand what belongs in the field.
    """

    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text) or ":" not in text:
        return None

    colon = text.find(":")
    label = _clean_label(text[:colon])
    value_start = colon + 1
    while value_start < len(text) and text[value_start].isspace():
        value_start += 1
    value = _normalize_space(text[value_start:])
    if not value or len(value) > 80:
        return None

    normalized_label = _slug(label)
    editable_example_labels = {
        "pais",
        "nacionalidade",
        # Common short defaults/examples in reconstructed PDF forms. These are
        # deliberately label-whitelisted so institutional prose such as
        # ``Órgão: Secretaria ...`` remains static.
        "placa",
        "setor_responsavel",
    }
    if normalized_label not in editable_example_labels:
        return None
    if (
        INSTRUCTION_PATTERN.match(value)
        or GENERIC_DROPDOWN_PATTERN.match(value)
        or CHECKBOX_TOKEN_PATTERN.search(value)
        or X_PLACEHOLDER_PATTERN.search(value)
        or UNDERSCORE_PLACEHOLDER_PATTERN.search(value)
        or ZERO_PHONE_PLACEHOLDER_PATTERN.search(value)
        or ZERO_CPF_PLACEHOLDER_PATTERN.search(value)
        or CURRENCY_PLACEHOLDER_PATTERN.search(value)
        or SAMPLE_EMAIL_PLACEHOLDER_PATTERN.search(value)
    ):
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=suggest_field_type(label),
        confidence=0.86,
        source="sample_value",
        preview=value,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": value_start,
            "end": len(text),
            "original": text[value_start:],
        },
        placeholder=value,
    )


def _detect_labeled_instruction(
    record: _ParagraphRecord,
    known_ids: set[str],
) -> dict[str, Any] | None:
    text = record.text or ""
    if not text or PLACEHOLDER_PATTERN.search(text) or ":" not in text:
        return None
    colon = text.find(":")
    label = _clean_label(text[:colon])
    tail_start = colon + 1
    while tail_start < len(text) and text[tail_start].isspace():
        tail_start += 1
    tail = text[tail_start:]
    if not _is_reasonable_label(label, maximum=120) or not INSTRUCTION_PATTERN.match(tail):
        return None
    if len(_normalize_space(tail)) < 12 or len(_normalize_space(tail)) > 800:
        return None

    field_type = suggest_field_type(label)
    if field_type == "text" and len(_normalize_space(tail)) >= 70:
        field_type = "multiline"
    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=field_type,
        confidence=0.90 if _paragraph_is_red(record.paragraph) else 0.82,
        source="instruction",
        preview=tail,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": tail_start,
            "end": len(text),
            "original": tail,
        },
    )


def _detect_prefilled_written_text(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect existing prose that likely represents an editable answer.

    Some institutional templates are distributed *already filled* with a
    previous/requester-authored justification instead of a visual blank.  A
    placeholder-only detector necessarily treats that prose as fixed text.
    Detector V2.9 recognizes the strongest structural versions of this pattern
    and converts the prose into a multiline field whose initial value is the
    original text.

    The rule is intentionally conservative.  Long prose is not enough on its
    own: it must live under a numbered response-like section or inside a
    full-width/merged table response row.  Explicitly fixed notes and ordinary
    headers remain static.
    """

    raw_text = str(record.text or "")
    text = _normalize_space(raw_text)
    role = str(getattr(getattr(record, "understanding", None), "role", "") or "")
    if role in {
        "instruction", "note", "heading", "table_title", "table_header",
        "table_reference", "signature", "example", "tagged",
    }:
        return None
    owner = getattr(record, "structure", None)
    if getattr(owner, "table_kind", "") in {"fixed_form", "reference"}:
        return None
    if (
        len(text) < 45
        or len(text) > 2200
        or record.story != "body"
        or _contains_authoritative_marker(record.paragraph)
        or _looks_like_section_label(text)
        or _PREFILLED_TEXT_STATIC_PREFIX_PATTERN.match(text)
        or CHECKBOX_TOKEN_PATTERN.search(text)
        or GENERIC_DROPDOWN_PATTERN.fullmatch(text)
        or _is_pure_fill_area_text(text)
    ):
        return None

    # A paragraph made mostly of a short heading followed by punctuation is
    # not a written response, even when Word wrapped it visually.
    if len(text.split()) < 7 or sum(ch.isalpha() for ch in text) < 24:
        return None

    section = _clean_label(semantic_section(record))
    section_is_response = bool(section and _PREFILLED_TEXT_SECTION_PATTERN.search(section))

    # A full-width/merged table row is another common way users store a written
    # answer below a heading or below a table header.  Determine this from
    # unique XML cells rather than the apparent ``row.cells`` count because
    # python-docx repeats references for merged cells.
    full_width_response_row = False
    table_context_label = ""
    if record.table is not None and record.row_index is not None and record.cell is not None:
        try:
            row = record.table.rows[int(record.row_index)]
            unique_cells = _unique_row_cells(row)
            current_key = id(record.cell._tc)
            if len(unique_cells) == 1 and id(unique_cells[0]._tc) == current_key:
                # Require meaningful structure above the prose: either a
                # numbered section already in scope or a preceding header row
                # with at least two short cells. This prevents random one-cell
                # narrative tables from becoming editable by accident.
                has_header_row = False
                for previous_index in range(int(record.row_index) - 1, -1, -1):
                    previous_cells = _unique_row_cells(record.table.rows[previous_index])
                    values = [_normalize_space(cell.text) for cell in previous_cells]
                    values = [value for value in values if value]
                    if not values:
                        continue
                    if len(values) >= 2 and all(len(value) <= 90 for value in values):
                        has_header_row = True
                        break
                    if len(values) == 1 and _looks_like_section_label(values[0]):
                        table_context_label = _clean_label(values[0])
                        break
                full_width_response_row = bool(section or has_header_row or table_context_label)
        except (IndexError, AttributeError, TypeError, ValueError):
            full_width_response_row = False

    if not section_is_response and not full_width_response_row:
        return None

    label = section or table_context_label or _context_label_for_record(record, records)
    label = _clean_label(label)
    if not _is_reasonable_label(label, maximum=180):
        label = "Texto editável"

    # Full-width narrative rows are a little more ambiguous than explicit
    # Justificativa/Descrição sections, so keep their confidence lower.  Both
    # remain reviewable in the assisted-detection screen.
    confidence = 0.90 if section_is_response else 0.82
    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type="multiline",
        confidence=confidence,
        source="prefilled_text",
        preview=text,
        location={
            "kind": "paragraph",
            "paragraph": record.ordinal,
            "table_index": record.table_index,
            "row_index": record.row_index,
            "cell_index": record.cell_index,
            "prefilled_text": True,
        },
        default_value=raw_text.strip(),
    )


def _detect_adjacent_sample_value(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect editable example/default values stored in the next table cell.

    Many institutional forms use four physical cells per row::

        E-mail: | servidor@orgao.gov.br | Telefone: | (83) 99999-9999

    The values are examples/defaults, not fixed institutional prose.  Keep this
    heuristic label-whitelisted so rows such as ``Órgão: Secretaria ...`` stay
    read-only.
    """

    text = _normalize_space(record.text)
    role = str(getattr(getattr(record, "understanding", None), "role", "") or "")
    owner = getattr(record, "structure", None)
    if role in {"table_header", "table_title", "table_reference", "heading", "tagged"}:
        return None
    if (
        not text
        or record.cell is None
        or record.table is None
        or record.row_index is None
        or len(record.cell.paragraphs) != 1
        or _contains_authoritative_marker(record.paragraph)
    ):
        return None

    # Do not require a trailing colon here. Real Word forms frequently use
    # alternating cells such as ``E-mail | exemplo@orgao.gov.br | Telefone |
    # (83) 99999-9999``. The physical adjacency is the delimiter. Keeping the
    # semantic whitelist below prevents ordinary institutional prose from
    # being converted merely because it sits next to another cell.
    explicit_label = text.endswith((":", "："))
    label = _clean_label(text)
    if not _is_reasonable_label(label, maximum=100):
        return None

    normalized_label = _slug(label)
    allowed_plain_labels = {
        "unidade",
        "lotacao",
        "setor",
        "municipio",
        "cidade",
        "pais",
        "nacionalidade",
    }
    email_labels = {"email", "e_mail", "correio_eletronico"}
    phone_labels = {"telefone", "celular", "fone"}

    unique_cells = _unique_row_cells(record.table.rows[record.row_index])
    current_index = next(
        (index for index, cell in enumerate(unique_cells) if id(cell._tc) == id(record.cell._tc)),
        -1,
    )
    if current_index < 0 or current_index + 1 >= len(unique_cells):
        return None
    value_cell = unique_cells[current_index + 1]
    value_records = [item for item in records if item.cell is not None and id(item.cell._tc) == id(value_cell._tc)]
    non_empty = [item for item in value_records if _normalize_space(item.text)]
    if len(non_empty) != 1 or _contains_authoritative_marker(non_empty[0].paragraph):
        return None

    value_record = non_empty[0]
    value = _normalize_space(value_record.text)
    if not value or len(value) > 100 or _is_pure_fill_area_text(value):
        return None

    looks_editable = False
    if normalized_label in email_labels:
        looks_editable = bool(
            re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", value, re.IGNORECASE)
        )
    elif normalized_label in phone_labels:
        looks_editable = bool(
            re.fullmatch(r"\(?\d{2}\)?\s*\d{4,5}\s*-\s*\d{4}", value)
        )
    elif normalized_label in allowed_plain_labels:
        # Short human-readable defaults such as ``Diretoria Administrativa``
        # are useful placeholders. Avoid values that look like sentences.
        looks_editable = (
            len(value) <= 60
            and not value.endswith((".", ";"))
            and len(value.split()) <= 6
            and not CHECKBOX_TOKEN_PATTERN.search(value)
        )

    if not looks_editable:
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=suggest_field_type(label),
        confidence=0.88 if explicit_label else 0.85,
        source="sample_value",
        preview=value,
        location={
            "kind": "text_span",
            "paragraph": value_record.ordinal,
            "start": 0,
            "end": len(value_record.text),
            "original": value_record.text,
        },
        placeholder=value,
    )


def _detect_label_only_field(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    text = _normalize_space(record.text)
    if (
        not text.endswith(":")
        or record.cell is None
        or len(record.cell.paragraphs) != 1
        or _contains_authoritative_marker(record.paragraph)
    ):
        return None
    label = _clean_label(text)
    if not _is_reasonable_label(label, maximum=100):
        return None
    if _looks_like_section_label(text) and SECTION_NUMBER_PATTERN.match(text):
        return None

    # A bare label can represent a fill area inside its own cell (for example
    # ``Responsável legal:``), but a very common Word grid stores the label in
    # one cell and the actual mask/value in the immediately following cell. In
    # that case the following cell owns the field and creating another tag in
    # the label cell causes duplicate inputs and staircase layouts.
    has_form_neighbor = False
    if record.table is not None and record.row_index is not None:
        try:
            unique_cells = _unique_row_cells(record.table.rows[record.row_index])
            current_index = next(
                (
                    index
                    for index, cell in enumerate(unique_cells)
                    if id(cell._tc) == id(record.cell._tc)
                ),
                -1,
            )
            if current_index >= 0 and current_index + 1 < len(unique_cells):
                immediate_text = _normalize_space(unique_cells[current_index + 1].text)
                if not immediate_text:
                    return None
                if _is_pure_fill_area_text(immediate_text):
                    return None

            for cell in unique_cells:
                if id(cell._tc) == id(record.cell._tc):
                    continue
                neighbor = _normalize_space(cell.text)
                if not neighbor:
                    return None
                if _looks_like_fill_area_text(neighbor):
                    has_form_neighbor = True
                    break
        except (IndexError, AttributeError):
            pass
    if not has_form_neighbor:
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=suggest_field_type(label),
        confidence=0.82,
        source="empty_cell",
        preview="Área vazia após o rótulo",
        location={
            "kind": "append_tag",
            "paragraph": record.ordinal,
        },
    )


def _detect_empty_cells(
    document: _Document,
    records: list[_ParagraphRecord],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    del document
    result: list[dict[str, Any]] = []
    by_cell: dict[int, list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if record.cell is not None:
            by_cell[id(record.cell._tc)].append(record)

    visited_cells: set[int] = set()
    for record in records:
        if record.cell is None or record.table is None or record.row_index is None or record.cell_index is None:
            continue
        cell_key = id(record.cell._tc)
        if cell_key in visited_cells:
            continue
        visited_cells.add(cell_key)
        cell_records = by_cell.get(cell_key, [])
        if not cell_records or any(item.ordinal in reserved_ordinals for item in cell_records):
            continue
        if any(_normalize_space(item.text) for item in cell_records):
            continue
        if record.cell.tables:
            continue
        if record.cell_index <= 0:
            continue
        try:
            previous_cell = record.table.rows[record.row_index].cells[record.cell_index - 1]
        except (IndexError, AttributeError):
            continue
        previous_text = _normalize_space(previous_cell.text)
        label = previous_text.strip(" :：–—-")
        if not _is_reasonable_label(label):
            continue
        explicit_label = previous_text.rstrip().endswith((":", "："))
        ordinal = cell_records[0].ordinal
        field_id = _unique_field_id(_make_field_id(label), known_ids)
        known_ids.add(field_id)
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type=suggest_field_type(label),
                confidence=0.84 if explicit_label else 0.68,
                source="empty_cell",
                preview="Célula vazia",
                location={
                    "kind": "empty_cell",
                    "paragraph": ordinal,
                },
                default_selected=explicit_label,
            )
        )
    return result


def _detect_consistency_repair_fields(
    records: list[_ParagraphRecord],
    candidates: list[dict[str, Any]],
    known_ids: set[str],
) -> list[dict[str, Any]]:
    """Suggest a missing sibling field when the surrounding table is consistent.

    This is intentionally a *repair* pass rather than another primary heuristic.
    It only fires when at least two peer rows already establish that the same
    visual column is editable. This helps arbitrary user-made matrices where one
    row uses a slightly different blank representation without requiring a fixed
    document template.
    """

    by_ordinal = {record.ordinal: record for record in records}
    occupied_cells: set[tuple[int, int, int]] = set()
    column_rows: dict[tuple[int, int], set[int]] = defaultdict(set)
    repeatable_tables: set[int] = set()

    for candidate in candidates:
        location = candidate.get("location", {}) or {}
        if str(candidate.get("source", "")) == "repeatable_table":
            try:
                repeatable_tables.add(int(location.get("table_index", -1)))
            except (TypeError, ValueError):
                pass
        ordinals: list[int] = []
        if "paragraph" in location:
            try:
                ordinals.append(int(location.get("paragraph", -1)))
            except (TypeError, ValueError):
                pass
        for value in location.get("paragraphs", []) or []:
            try:
                ordinals.append(int(value))
            except (TypeError, ValueError):
                continue
        for ordinal in ordinals:
            record = by_ordinal.get(ordinal)
            if (
                record is None
                or record.table_index is None
                or record.row_index is None
                or record.cell_index is None
            ):
                continue
            key = (int(record.table_index), int(record.row_index), int(record.cell_index))
            occupied_cells.add(key)
            column_rows[(key[0], key[2])].add(key[1])

    established_columns = {
        key for key, rows in column_rows.items()
        if len(rows) >= 2 and key[0] not in repeatable_tables
    }
    if not established_columns:
        return []

    by_cell: dict[tuple[int, int, int], list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        if (
            record.table_index is None
            or record.row_index is None
            or record.cell_index is None
        ):
            continue
        by_cell[(int(record.table_index), int(record.row_index), int(record.cell_index))].append(record)

    result: list[dict[str, Any]] = []
    for (table_index, row_index, cell_index), cell_records in by_cell.items():
        if (table_index, cell_index) not in established_columns:
            continue
        if (table_index, row_index, cell_index) in occupied_cells:
            continue
        if any(_contains_authoritative_marker(record.paragraph) for record in cell_records):
            continue

        texts = [record.text or "" for record in cell_records]
        combined = _normalize_space(" ".join(texts))
        if combined and not all(_is_pure_fill_area_text(text) for text in texts):
            continue

        target = next((record for record in cell_records if record.paragraph is not None), None)
        if target is None:
            continue
        label, label_source, label_confidence = semantic_label(target)
        if not label or label_confidence < 0.74 or label_source == "section_fallback":
            continue

        field_id = _unique_field_id(_make_field_id(label), known_ids)
        known_ids.add(field_id)
        preview = combined or "área vazia"
        field_type = _detected_placeholder_type(label, combined)
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type=field_type,
                confidence=0.68,
                source="consistency_repair",
                preview=preview,
                location={
                    "kind": "empty_cell" if not combined else "paragraph",
                    "paragraph": target.ordinal,
                    "repair_basis": "peer_column",
                    "table_index": table_index,
                    "row_index": row_index,
                    "cell_index": cell_index,
                },
                default_selected=False,
            )
        )

    return result


def _detect_terminal_prompt(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
    structure,
) -> dict[str, Any] | None:
    """Detect the final prompt after an instructional block.

    Institutional forms often put several note paragraphs and the actual input
    prompt in the same merged Word cell.  There may be no blank cell or visual
    placeholder after the prompt, so a placeholder-only detector misses the
    entire numbered section.  The prompt is preserved and the accepted tag is
    appended after it rather than replacing the printed label.
    """

    from app.document.detection.field_inference import infer_field_type
    from app.document.detection.roles import terminal_prompt_score

    if _contains_authoritative_marker(record.paragraph):
        return None
    score, reasons = terminal_prompt_score(record, records, structure)
    if score < 0.65:
        return None

    label = _clean_label(record.text)
    if not _is_reasonable_label(label, maximum=190):
        return None
    owner = structure.owner_for(record.ordinal) if structure is not None else None
    section = owner.section if owner is not None else ""
    inference = infer_field_type(label, section=section, preview=record.text)
    field_id = _unique_field_id(_make_field_id(label), known_ids)
    candidate = _candidate(
        field_id=field_id,
        label=label,
        field_type=inference.field_type,
        confidence=min(0.96, max(score, inference.confidence * 0.88)),
        source="terminal_prompt",
        preview=record.text,
        location={
            "kind": "append_tag",
            "paragraph": record.ordinal,
        },
    )
    if section:
        candidate["section"] = section
    candidate["terminal_prompt_reasons"] = reasons
    candidate["type_inference"] = {
        "confidence": inference.confidence,
        "reasons": list(inference.reasons),
    }
    return candidate



def _instruction_placeholder_label(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    match: re.Match[str],
) -> str:
    prefix = (record.text or "")[: match.start()].rstrip()
    if ":" in prefix:
        local = prefix.rsplit("\n", 1)[-1].rsplit(":", 1)[0]
        cleaned = _clean_label(local[-150:])
        if _is_reasonable_label(cleaned, maximum=120):
            return cleaned

    inner = _normalize_space(match.group("bracket") or match.group("angle") or "")
    cleaned = re.sub(
        r"^(?:informar|informe|descrever|descreva|detalhar|detalhe|indicar|indique|"
        r"justificar|justifique|preencher|preencha|especificar|especifique|inserir|insira|"
        r"registrar|registre|definir|defina|identificar|identifique|relacionar|relacione|"
        r"mencionar|mencione)\s+",
        "",
        inner,
        flags=re.IGNORECASE,
    )
    cleaned = _clean_label(cleaned.strip(" .,:;"))
    if _is_reasonable_label(cleaned, maximum=150):
        return cleaned

    context = _context_label_for_record(record, records)
    if _is_reasonable_label(context, maximum=150):
        return _clean_label(context)
    return "Campo a preencher"


def _detect_instruction_placeholders(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> list[dict[str, Any]]:
    """Detect explicit ``[INFORMAR ...]`` / ``<INFORMAR ...>`` instructions.

    Unlike generic square-bracket prose, the contents must begin with a fill
    instruction verb.  This keeps legal omissions/citations such as ``[...]``
    static while allowing unfamiliar author instructions in any font color.
    """

    text = record.text or ""
    if not text or _contains_authoritative_marker(record.paragraph):
        return []
    result: list[dict[str, Any]] = []
    for match in _GENERIC_INSTRUCTION_PLACEHOLDER_PATTERN.finditer(text):
        raw = match.group(0)
        inner = _normalize_space(match.group("bracket") or match.group("angle") or "")
        if len(inner) < 5 or len(inner) > 320:
            continue
        label = _instruction_placeholder_label(record, records, match)
        field_id = _unique_field_id(_make_field_id(label), known_ids)
        known_ids.add(field_id)
        field_type = suggest_field_type(label)
        if field_type == "text" and len(inner) >= 100:
            field_type = "multiline"
        result.append(
            _candidate(
                field_id=field_id,
                label=label,
                field_type=field_type,
                confidence=0.86 if label != "Campo a preencher" else 0.76,
                source="instruction_placeholder",
                preview=raw,
                location={
                    "kind": "text_span",
                    "paragraph": record.ordinal,
                    "start": match.start(),
                    "end": match.end(),
                    "original": raw,
                },
                default_selected=False,
            )
        )
    return result


def _record_role(record: _ParagraphRecord) -> str:
    return str(getattr(getattr(record, "understanding", None), "role", "") or "")


def _looks_like_static_generic_context(record: _ParagraphRecord, label: str = "") -> bool:
    role = _record_role(record)
    if role in {
        "heading", "instruction", "note", "signature", "header_footer",
        "table_title", "table_header", "table_reference", "example", "tagged",
    }:
        return True
    cleaned = _clean_label(label)
    return bool(cleaned and _GENERIC_STATIC_LABEL_PATTERN.match(cleaned))


def _strong_value_shape(value: str) -> bool:
    text = _normalize_space(value)
    if not text:
        return False
    return bool(
        re.fullmatch(r"R\$\s*[\d.]+(?:,\d{2})?", text, re.IGNORECASE)
        or re.fullmatch(r"\d{1,3}(?:\.\d{3})*/\d{4}-\d{2}", text)
        or re.fullmatch(r"\d{2}\.\d{3}\.\d{3}/\d{4}-\d{1,2}", text)
        or re.fullmatch(r"\d{3}\.\d{3}\.\d{3}-\d{2}", text)
        or re.fullmatch(r"\(?\d{2}\)?\s*\d{4,5}-\d{4}", text)
        or re.fullmatch(r"[^\s@]+@[^\s@]+\.[^\s@]+", text)
        or re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}", text)
        or re.fullmatch(r"\d+(?:[.,]\d+)?%", text)
        or re.fullmatch(r"[A-ZÁÉÍÓÚÂÊÔÃÕÇ]{2,12}(?:-[A-Z0-9ÁÉÍÓÚÂÊÔÃÕÇ]{2,12})*-\d{4}/\d+", text)
    )


def _looks_like_short_existing_value(value: str, *, maximum: int = 180) -> bool:
    text = _normalize_space(value)
    if not text or len(text) > maximum:
        return False
    if (
        _is_pure_fill_area_text(text)
        or _GENERIC_INSTRUCTION_PLACEHOLDER_PATTERN.search(text)
        or INSTRUCTION_PATTERN.match(text)
        or GENERIC_DROPDOWN_PATTERN.match(text)
        or CHECKBOX_TOKEN_PATTERN.search(text)
        or text.casefold().startswith(("http://", "https://", "www."))
    ):
        return False
    if _strong_value_shape(text):
        return True
    if len(text.split()) > 18:
        return False
    if _GENERIC_VALUE_SENTENCE_PATTERN.search(text):
        return False
    # Short values commonly include people, organizational units, identifiers,
    # names, codes and current sample/default text.  Commas/semicolons and a
    # terminal full stop are much more typical of prose than of one value.
    if text.endswith((".", ";")) or ";" in text:
        return False
    if text.count(",") >= 2:
        return False
    return sum(ch.isalpha() for ch in text) >= 2 or any(ch.isdigit() for ch in text)


def _detect_generic_inline_choice(
    record: _ParagraphRecord,
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect compact black/plain alternatives after an explicit label."""

    text = record.text or ""
    if not text or ":" not in text or _contains_authoritative_marker(record.paragraph):
        return None
    colon = text.find(":")
    label = _clean_label(text[:colon])
    if not _is_reasonable_label(label, maximum=120) or _looks_like_static_generic_context(record, label):
        return None
    start = colon + 1
    while start < len(text) and text[start].isspace():
        start += 1
    tail = _normalize_space(text[start:])
    if not tail or len(tail) > 220 or tail.endswith("."):
        return None
    parts = _visual_choice_parts(tail)
    if not parts:
        return None
    if any(len(part.split()) > 9 for part in parts):
        return None
    # A bare use of the conjunction "ou" inside ordinary prose is common.
    # Requiring compact alternatives keeps this a field-like expression.
    if re.search(r"\s+\bou\b\s+", tail, re.IGNORECASE) and any(len(part) > 90 for part in parts):
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type="dropdown",
        confidence=0.82,
        source="generic_choice",
        preview=tail,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": start,
            "end": len(text),
            "original": text[start:],
        },
        options=parts,
        default_selected=False,
    )


def _detect_generic_labeled_value(
    record: _ParagraphRecord,
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Review-only fallback for unfamiliar ``Label: current value`` fields."""

    text = record.text or ""
    if not text or ":" not in text or _contains_authoritative_marker(record.paragraph):
        return None
    colon = text.find(":")
    label = _clean_label(text[:colon])
    if not _is_reasonable_label(label, maximum=120) or _looks_like_static_generic_context(record, label):
        return None
    start = colon + 1
    while start < len(text) and text[start].isspace():
        start += 1
    raw_value = text[start:]
    value = _normalize_space(raw_value)
    if not _looks_like_short_existing_value(value):
        return None
    # Compact alternatives belong to the dropdown fallback, not a text field.
    if _visual_choice_parts(value):
        return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    field_type = suggest_field_type(label)
    confidence = 0.82 if _strong_value_shape(value) else 0.74
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=field_type,
        confidence=confidence,
        source="generic_labeled_value",
        preview=value,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": start,
            "end": len(text),
            "original": raw_value,
        },
        placeholder=value,
        default_selected=False,
    )


def _detect_generic_adjacent_value(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Broaden table ``label | current value`` discovery without a label whitelist.

    The older detector only trusts a few labels.  This fallback accepts an
    unfamiliar label when the physical row relationship is strong, but leaves
    the result unchecked for human review.
    """

    text = _normalize_space(record.text)
    if (
        not text
        or record.cell is None
        or record.table is None
        or record.row_index is None
        or _contains_authoritative_marker(record.paragraph)
    ):
        return None
    label = _clean_label(text)
    if not _is_reasonable_label(label, maximum=100) or _looks_like_static_generic_context(record, label):
        return None

    try:
        unique_cells = _unique_row_cells(record.table.rows[record.row_index])
        current_index = next(
            index for index, cell in enumerate(unique_cells)
            if id(cell._tc) == id(record.cell._tc)
        )
    except (StopIteration, IndexError, AttributeError, TypeError):
        return None
    if current_index + 1 >= len(unique_cells):
        return None
    value_cell = unique_cells[current_index + 1]
    value_records = [
        item for item in records
        if item.cell is not None and id(item.cell._tc) == id(value_cell._tc)
    ]
    non_empty = [item for item in value_records if _normalize_space(item.text)]
    if len(non_empty) != 1:
        return None
    value_record = non_empty[0]
    if _contains_authoritative_marker(value_record.paragraph):
        return None
    value = _normalize_space(value_record.text)
    if not _looks_like_short_existing_value(value, maximum=140):
        return None

    explicit_label = text.endswith((":", "："))
    # Without ':' require another strong signal so ordinary two-column
    # reference tables do not become an avalanche of editable suggestions.
    if not explicit_label and not _strong_value_shape(value):
        if len(label.split()) > 7 or len(value.split()) > 8:
            return None
        value_visual = any(_run_visual_intent_signals(run) for run in value_record.paragraph.runs if run.text)
        if not value_visual:
            return None

    field_id = _unique_field_id(_make_field_id(label), known_ids)
    field_type = suggest_field_type(label)
    return _candidate(
        field_id=field_id,
        label=label,
        field_type=field_type,
        confidence=0.80 if explicit_label or _strong_value_shape(value) else 0.72,
        source="generic_labeled_value",
        preview=value,
        location={
            "kind": "text_span",
            "paragraph": value_record.ordinal,
            "start": 0,
            "end": len(value_record.text),
            "original": value_record.text,
            "table_index": value_record.table_index,
            "row_index": value_record.row_index,
            "cell_index": value_record.cell_index,
        },
        placeholder=value,
        default_selected=False,
    )


def _color_tuple(color) -> tuple[int, int, int] | None:
    try:
        return int(color[0]), int(color[1]), int(color[2])
    except Exception:
        value = str(color or "")
        if len(value) != 6:
            return None
        try:
            return int(value[:2], 16), int(value[2:4], 16), int(value[4:], 16)
        except ValueError:
            return None


def _run_visual_intent_signals(run) -> set[str]:
    """Return strong non-red formatting signals that can mark editable text."""

    text = run.text or ""
    if not _normalize_space(text):
        return set()
    style_name = ""
    try:
        style_name = str(run.style.name or "")
    except Exception:
        pass
    if style_name.casefold() == "hyperlink" or text.casefold().startswith(("http://", "https://", "www.")):
        return set()

    signals: set[str] = set()
    try:
        if run.font.highlight_color is not None:
            signals.add("highlight")
    except Exception:
        pass
    try:
        if run.underline not in {None, False}:
            signals.add("underline")
    except Exception:
        pass
    if style_name and _FIELD_STYLE_HINT_PATTERN.search(style_name):
        signals.add("field_style")
    try:
        paragraph_style = str(run._parent.style.name or "")
    except Exception:
        paragraph_style = ""
    if paragraph_style and _FIELD_STYLE_HINT_PATTERN.search(paragraph_style):
        signals.add("field_style")

    # Red already has its specialized detector. This branch handles blue,
    # green, purple, etc. and explicit accent/theme colors. Neutral black/gray
    # is ignored because ordinary body text often carries it explicitly.
    if not _run_is_red(run):
        for color in _run_color_candidates(run):
            rgb = _color_tuple(color)
            if rgb is None:
                continue
            red, green, blue = rgb
            if max(rgb) - min(rgb) >= 45 and max(rgb) >= 90:
                signals.add("nondefault_color")
                break
        try:
            theme = run.font.color.theme_color
        except Exception:
            theme = None
        if theme is not None and any(token in str(theme).casefold() for token in ("accent", "followed_hyperlink")):
            signals.add("nondefault_color")

    # Run shading is another common hand-authored "fill this" convention.
    try:
        shading = run._r.xpath("./w:rPr/w:shd")
    except Exception:
        shading = []
    for node in shading:
        fill = str(node.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", "") or "")
        if fill and fill.casefold() not in {"auto", "ffffff", "000000"}:
            signals.add("shading")
            break
    return signals


def _visual_intent_run_groups(paragraph) -> list[tuple[int, int, str, set[str]]]:
    groups: list[tuple[int, int, str, set[str]]] = []
    offset = 0
    active_start: int | None = None
    active_text: list[str] = []
    active_signals: set[str] = set()
    for run in paragraph.runs:
        run_text = run.text or ""
        start = offset
        end = start + len(run_text)
        signals = _run_visual_intent_signals(run)
        # Leave red regions to the mature colored fallback so only new visual
        # conventions are handled here.
        if signals and not _run_is_red(run):
            if active_start is None:
                active_start = start
            active_text.append(run_text)
            active_signals.update(signals)
        elif active_start is not None:
            groups.append((active_start, start, "".join(active_text), set(active_signals)))
            active_start = None
            active_text = []
            active_signals.clear()
        offset = end
    if active_start is not None:
        groups.append((active_start, offset, "".join(active_text), set(active_signals)))
    return groups


def _neighbor_record(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    *,
    direction: int,
) -> _ParagraphRecord | None:
    index = record.ordinal + direction
    while 0 <= index < len(records):
        other = records[index]
        index += direction
        if other.story != record.story:
            continue
        if record.cell is not None:
            if other.cell is None or id(other.cell._tc) != id(record.cell._tc):
                continue
        elif other.cell is not None:
            continue
        if _normalize_space(other.text):
            return other
    return None


def _standalone_visual_context_is_fieldlike(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    visible: str,
    signals: set[str],
) -> bool:
    """Require context before treating a fully formatted paragraph as a field."""

    previous = _neighbor_record(record, records, direction=-1)
    following = _neighbor_record(record, records, direction=1)
    if any(_record_role(item) == "signature" for item in (previous, following) if item is not None):
        return False
    if "field_style" in signals:
        return True
    if previous is not None:
        previous_text = _normalize_space(previous.text)
        if _record_role(previous) == "field_prompt" or (
            previous_text.endswith((":", "："))
            and _is_reasonable_label(_clean_label(previous_text), maximum=160)
        ):
            return True
    if INSTRUCTION_PATTERN.match(visible):
        return True
    # Strong highlight/shading can mark a standalone fill value, but without a
    # nearby prompt it is equally common on signatures and headings. Keep those
    # ambiguous cases silent rather than flooding review with document chrome.
    return False


def _detect_unclaimed_visual_intent_fields(
    records: list[_ParagraphRecord],
    existing_candidates: list[dict[str, Any]],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Surface highlighted/underlined/styled/non-red regions nobody claimed."""

    whole_claims, span_claims = _candidate_claims(existing_candidates)
    result: list[dict[str, Any]] = []
    for record in records:
        if record.ordinal in reserved_ordinals or record.ordinal in whole_claims:
            continue
        if _contains_authoritative_marker(record.paragraph):
            continue
        role = _record_role(record)
        if role in {"heading", "instruction", "note", "signature", "header_footer", "table_title", "table_header", "table_reference", "tagged"}:
            continue
        full = record.text or ""
        for start, end, raw, signals in _visual_intent_run_groups(record.paragraph):
            visible = _normalize_space(raw)
            if not visible or len(visible) > 420:
                continue
            if _span_overlaps_claimed(record.ordinal, start, end, whole_claims, span_claims):
                continue
            if visible.casefold().startswith(("http://", "https://", "www.")):
                continue
            prefix = full[:start].rstrip()
            suffix = full[end:].lstrip()
            inline_label = bool(
                ":" in prefix
                and _is_reasonable_label(_clean_label(prefix.rsplit(":", 1)[0][-150:]), maximum=120)
            )
            standalone = start == 0 and end >= len(full)
            if standalone and not inline_label and not _standalone_visual_context_is_fieldlike(
                record, records, visible, signals
            ):
                continue
            # A color-only prefix such as ``Indicação de marcas ou modelos``
            # followed by a black legal citation is normally a heading, not a
            # partial editable span. Require an inline/adjacent prompt or an
            # explicitly field-named style for such mixed formatting.
            if (
                start == 0
                and suffix
                and signals <= {"nondefault_color", "underline"}
                and not inline_label
                and "field_style" not in signals
            ):
                previous = _neighbor_record(record, records, direction=-1)
                if previous is None or _record_role(previous) != "field_prompt":
                    continue
            # Underline by itself is weak and common in legal prose. Require a
            # compact value next to a label.
            if signals == {"underline"}:
                if not inline_label or len(visible) > 120 or len(visible.split()) > 14:
                    continue
            if len(visible.split()) > 28 and not ({"highlight", "field_style", "shading"} & signals):
                continue

            label = _visual_label_from_span(record, records, raw, start)
            if _looks_like_static_generic_context(record, label):
                continue
            parts = _visual_choice_parts(visible)
            has_explicit_choice_separator = bool(
                "|" in visible
                or re.search(r"\s+/\s+", visible)
                or ";" in visible
                or re.search(r"\bOU\b", visible)
                or inline_label
            )
            is_choice = bool(
                parts
                and has_explicit_choice_separator
                and all(len(part.split()) <= 10 for part in parts)
            )
            source = "visual_choice" if is_choice else "visual_field"
            field_type = "dropdown" if is_choice else suggest_field_type(label)
            if field_type == "text" and len(visible) >= 140:
                field_type = "multiline"
            field_id = _unique_field_id(_make_field_id(label), known_ids)
            known_ids.add(field_id)
            candidate = _candidate(
                field_id=field_id,
                label=label,
                field_type=field_type,
                confidence=0.82 if {"highlight", "field_style", "shading"} & signals else 0.74,
                source=source,
                preview=visible,
                location={
                    "kind": "text_span",
                    "paragraph": record.ordinal,
                    "start": start,
                    "end": end,
                    "original": raw,
                },
                options=parts if is_choice else None,
                placeholder=visible if not is_choice else "",
                default_selected=False,
            )
            candidate["visual_intent_signals"] = sorted(signals)
            result.append(candidate)
            span_claims[record.ordinal].append((start, end))
    return result


def _generic_candidate_overlaps_claims(
    candidate: dict[str, Any],
    whole_claims: set[int],
    span_claims: dict[int, list[tuple[int, int]]],
) -> bool:
    location = dict(candidate.get("location", {}) or {})
    kind = str(location.get("kind", ""))
    if kind == "text_span":
        try:
            return _span_overlaps_claimed(
                int(location.get("paragraph", -1)),
                int(location.get("start", 0)),
                int(location.get("end", 0)),
                whole_claims,
                span_claims,
            )
        except (TypeError, ValueError):
            return True
    if "paragraph" in location:
        try:
            return int(location.get("paragraph", -1)) in whole_claims
        except (TypeError, ValueError):
            return True
    return False


def _register_generic_claim(
    candidate: dict[str, Any],
    whole_claims: set[int],
    span_claims: dict[int, list[tuple[int, int]]],
) -> None:
    location = dict(candidate.get("location", {}) or {})
    kind = str(location.get("kind", ""))
    try:
        ordinal = int(location.get("paragraph", -1))
    except (TypeError, ValueError):
        return
    if kind == "text_span":
        try:
            span_claims[ordinal].append(
                (int(location.get("start", 0)), int(location.get("end", 0)))
            )
        except (TypeError, ValueError):
            pass
    elif ordinal >= 0:
        whole_claims.add(ordinal)


def _detect_generic_intent_fields(
    records: list[_ParagraphRecord],
    existing_candidates: list[dict[str, Any]],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Last-resort review candidates for unfamiliar editable intent.

    This pass is intentionally color-agnostic and runs only after stronger
    structural, semantic and visual detectors.  It broadens recall without
    stealing ownership from a detector that already understood the region.
    """

    whole_claims, span_claims = _candidate_claims(existing_candidates)
    result: list[dict[str, Any]] = []
    for record in records:
        if record.ordinal in reserved_ordinals or record.ordinal in whole_claims:
            continue
        if _contains_authoritative_marker(record.paragraph):
            continue

        # Use an isolated ID set while probing. A rejected overlapping fallback
        # must not consume an identifier and force suffixes on later candidates.
        temporary_ids = set(known_ids)
        proposals: list[dict[str, Any]] = []

        placeholders = _detect_instruction_placeholders(record, records, temporary_ids)
        proposals.extend(placeholders)
        if not placeholders:
            choice = _detect_generic_inline_choice(record, temporary_ids)
            if choice is not None:
                proposals.append(choice)
            else:
                labeled = _detect_generic_labeled_value(record, temporary_ids)
                if labeled is not None:
                    proposals.append(labeled)

        # A label in one cell can point to a value in the next cell, so this
        # proposal may target another record ordinal. Check its physical claim
        # rather than assuming the current record owns the candidate.
        if not proposals:
            adjacent = _detect_generic_adjacent_value(record, records, temporary_ids)
            if adjacent is not None:
                proposals.append(adjacent)

        for candidate in proposals:
            if _generic_candidate_overlaps_claims(candidate, whole_claims, span_claims):
                continue
            result.append(candidate)
            field_id = str(candidate.get("field_id", "")).strip()
            if field_id:
                known_ids.add(field_id)
            _register_generic_claim(candidate, whole_claims, span_claims)
    return result


def _colored_run_groups(paragraph) -> list[tuple[int, int, str]]:
    """Return contiguous red run spans using paragraph-text offsets."""

    groups: list[tuple[int, int, str]] = []
    offset = 0
    active_start: int | None = None
    active_text: list[str] = []
    for run in paragraph.runs:
        run_text = run.text or ""
        run_start = offset
        run_end = run_start + len(run_text)
        is_red = bool(run_text) and _run_is_red(run)
        if is_red:
            if active_start is None:
                active_start = run_start
            active_text.append(run_text)
        elif active_start is not None:
            groups.append((active_start, run_start, "".join(active_text)))
            active_start = None
            active_text = []
        offset = run_end
    if active_start is not None:
        groups.append((active_start, offset, "".join(active_text)))
    return groups


def _visual_choice_parts(value: str) -> list[str]:
    text = _normalize_space(value)
    if not text:
        return []

    # Delimiters are deliberately evaluated separately. Semicolons are common
    # punctuation in legal clauses, so only treat them as option separators in
    # short, list-like text.
    if re.search(r"\s+\bou\b\s+", text, flags=re.IGNORECASE):
        raw_parts = re.split(r"\s+\bou\b\s+", text, flags=re.IGNORECASE)
    elif "|" in text:
        raw_parts = re.split(r"\s*\|\s*", text)
    elif re.search(r"\s+/\s+", text):
        raw_parts = re.split(r"\s+/\s+", text)
    elif ";" in text and len(text) <= 180 and "." not in text:
        raw_parts = re.split(r"\s*;\s*", text)
    else:
        return []

    parts = [_normalize_space(part) for part in raw_parts]
    parts = [part for part in parts if part]
    if len(parts) < 2 or len(parts) > 8:
        return []
    if any(len(part) < 2 or len(part) > 180 for part in parts):
        return []
    if len({part.casefold() for part in parts}) != len(parts):
        return []
    return parts


def _looks_like_static_colored_guidance(value: str) -> bool:
    text = _normalize_space(value)
    if not text:
        return True
    if text.casefold() in {"ou", "e/ou", "ou:"}:
        return True
    if _COLORED_STATIC_GUIDANCE_PATTERN.match(text):
        return True
    if _COLORED_COLUMN_GUIDANCE_PATTERN.match(text):
        return True
    if text.startswith(("“", '"')) and text.endswith(("”", '"')):
        return True
    return False


def _visual_label_from_span(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    raw: str,
    start: int,
) -> str:
    """Build a useful review label for a colored fill span."""

    angle = _ANGLE_PLACEHOLDER_PATTERN.match(raw)
    if angle:
        text = _normalize_space(angle.group(1))
        text = re.sub(
            r"^(?:informar|informe|descrever|descreva|detalhar|detalhe|"
            r"indicar|indique|justificar|justifique|preencher|preencha|"
            r"especificar|especifique|inserir|insira)\s+",
            "",
            text,
            flags=re.IGNORECASE,
        )
        text = re.split(
            r"\s*,?\s+(?:constante|conforme|em conson[aâ]ncia)\s+",
            text,
            maxsplit=1,
            flags=re.IGNORECASE,
        )[0]
        cleaned = _clean_label(text)
        if cleaned:
            return cleaned

    full_text = record.text or ""
    prefix = full_text[:start].rstrip()
    if ":" in prefix:
        tail = prefix.rsplit(":", 1)[0]
        tail = re.split(r"[.;!?]\s+", tail)[-1]
        cleaned = _clean_label(tail[-140:])
        if _is_reasonable_label(cleaned, maximum=120):
            return cleaned

    context = _context_label_for_record(record, records)
    if context and _is_reasonable_label(context, maximum=150):
        return _clean_label(context)

    cleaned_raw = _clean_label(_normalize_space(raw).strip("<>[]() "))
    return cleaned_raw or "Campo destacado"


def _candidate_claims(
    candidates: list[dict[str, Any]],
) -> tuple[set[int], dict[int, list[tuple[int, int]]]]:
    whole: set[int] = set()
    spans: dict[int, list[tuple[int, int]]] = defaultdict(list)
    for candidate in candidates:
        location = dict(candidate.get("location", {}) or {})
        kind = str(location.get("kind", ""))
        if kind in {"paragraph", "empty_cell", "append_tag"}:
            try:
                whole.add(int(location.get("paragraph", -1)))
            except (TypeError, ValueError):
                pass
        elif kind in {"paragraph_block", "paragraph_list"}:
            for value in location.get("paragraphs", []) or []:
                try:
                    whole.add(int(value))
                except (TypeError, ValueError):
                    pass
        elif kind == "text_span":
            try:
                ordinal = int(location.get("paragraph", -1))
                spans[ordinal].append((int(location.get("start", 0)), int(location.get("end", 0))))
            except (TypeError, ValueError):
                pass
        elif kind == "text_spans":
            for span in location.get("spans", []) or []:
                if not isinstance(span, dict):
                    continue
                try:
                    ordinal = int(span.get("paragraph", -1))
                    spans[ordinal].append((int(span.get("start", 0)), int(span.get("end", 0))))
                except (TypeError, ValueError):
                    pass
    return whole, spans


def _span_overlaps_claimed(
    ordinal: int,
    start: int,
    end: int,
    whole_claims: set[int],
    span_claims: dict[int, list[tuple[int, int]]],
) -> bool:
    if ordinal in whole_claims:
        return True
    return any(start < other_end and other_start < end for other_start, other_end in span_claims.get(ordinal, []))


def _detect_colored_choice_blocks(
    records: list[_ParagraphRecord],
    existing_candidates: list[dict[str, Any]],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Find review-only red alternatives split across consecutive paragraphs.

    This complements the older table-cell ``long_choice`` detector.  Many
    SIAGOV Word forms use red clauses in the document body, separated by a red
    ``OU`` paragraph, or use the same pattern in compact table headers.
    """

    whole_claims, _ = _candidate_claims(existing_candidates)
    by_region: dict[tuple[Any, ...], list[_ParagraphRecord]] = defaultdict(list)
    for record in records:
        cell_key = id(record.cell._tc) if record.cell is not None else None
        by_region[(record.story, record.table_index, cell_key)].append(record)

    result: list[dict[str, Any]] = []
    seen_sequences: set[tuple[int, ...]] = set()
    for region_records in by_region.values():
        region_records.sort(key=lambda item: item.ordinal)
        sequence: list[_ParagraphRecord] = []

        def flush() -> None:
            nonlocal sequence
            if not sequence:
                return
            meaningful = [item for item in sequence if _normalize_space(item.text)]
            sequence = []
            if len(meaningful) < 3:
                return
            if not any(_normalize_space(item.text).casefold() == "ou" for item in meaningful):
                return
            ordinals = tuple(item.ordinal for item in meaningful)
            if ordinals in seen_sequences:
                return
            if any(item.ordinal in reserved_ordinals or item.ordinal in whole_claims for item in meaningful):
                return
            if any(_contains_authoritative_marker(item.paragraph) for item in meaningful):
                return

            segments: list[list[_ParagraphRecord]] = [[]]
            for item in meaningful:
                if _normalize_space(item.text).casefold() == "ou":
                    if segments[-1]:
                        segments.append([])
                    continue
                segments[-1].append(item)
            segments = [segment for segment in segments if segment]
            if len(segments) < 2 or len(segments) > 8:
                return

            options: list[dict[str, str]] = []
            for segment in segments:
                value = _normalize_space("\n".join(item.text for item in segment))
                if len(value) < 2 or len(value) > 1400:
                    return
                label = value
                numbered = re.sub(r"^\s*\d+(?:\.\d+)*\.?\s*", "", label)
                if numbered:
                    label = numbered
                if len(label) > 90:
                    label = label[:87].rstrip() + "..."
                options.append({"label": label, "value": value})

            first = meaningful[0]
            context = _context_label_for_record(first, records)
            label = _clean_label(context or "Escolha uma alternativa")
            field_id = _unique_field_id(_make_field_id(label), known_ids)
            known_ids.add(field_id)
            seen_sequences.add(ordinals)
            result.append(
                _candidate(
                    field_id=field_id,
                    label=label,
                    field_type="dropdown",
                    confidence=0.86,
                    source="colored_choice_block",
                    preview=" OU ".join(option["label"] for option in options),
                    location={
                        "kind": "paragraph_block",
                        "paragraphs": [item.ordinal for item in meaningful],
                    },
                    options=options,
                    layout="choice",
                    layout_group=f"auto_choice_{field_id}",
                )
            )

        active_section = ""
        for record in region_records:
            text = _normalize_space(record.text)
            section = str(semantic_section(record) or "").strip()
            if sequence and section and active_section and section != active_section:
                flush()
            if section:
                active_section = section

            if not text:
                if sequence:
                    sequence.append(record)
                continue
            if _paragraph_is_red(record.paragraph):
                # Red section headings organize optional clauses; they are not
                # themselves an alternative and should terminate the previous
                # choice window rather than being swallowed into it.
                if record.cell is None and _looks_like_section_label(text):
                    flush()
                    active_section = section
                    continue
                sequence.append(record)
                continue
            flush()
        flush()
    return result


def _detect_unclaimed_colored_fields(
    records: list[_ParagraphRecord],
    existing_candidates: list[dict[str, Any]],
    known_ids: set[str],
    reserved_ordinals: set[int],
) -> list[dict[str, Any]]:
    """Surface unclaimed red spans instead of silently dropping them.

    These findings are deliberately review-only.  They are the scanner's
    visual-intent safety net for new institutional forms whose labels have
    never appeared in Padroniza before.
    """

    from app.document.detection.field_inference import infer_field_type

    whole_claims, span_claims = _candidate_claims(existing_candidates)
    result: list[dict[str, Any]] = []
    for record in records:
        if record.ordinal in reserved_ordinals or record.ordinal in whole_claims:
            continue
        if _contains_authoritative_marker(record.paragraph):
            continue
        groups = _colored_run_groups(record.paragraph)
        if not groups:
            continue

        full_text = record.text or ""
        for start, end, raw in groups:
            if _span_overlaps_claimed(record.ordinal, start, end, whole_claims, span_claims):
                continue
            text = _normalize_space(raw)
            if not text or _looks_like_static_colored_guidance(text):
                continue
            if (
                X_PLACEHOLDER_PATTERN.search(text)
                or UNDERSCORE_PLACEHOLDER_PATTERN.search(text)
                or CURRENCY_PLACEHOLDER_PATTERN.search(text)
                or ZERO_PHONE_PLACEHOLDER_PATTERN.search(text)
                or ZERO_CPF_PLACEHOLDER_PATTERN.search(text)
                or SAMPLE_EMAIL_PLACEHOLDER_PATTERN.search(text)
            ):
                continue

            prefix = _normalize_space(full_text[:start])
            suffix = _normalize_space(full_text[end:])
            whole_red = not prefix and not suffix
            angle = _ANGLE_PLACEHOLDER_PATTERN.match(text)
            choice_parts = _visual_choice_parts(text)

            # Long all-red policy/contract clauses are usually static unless a
            # stronger detector identified them as an instruction or choice.
            if whole_red and len(text) > 220 and angle is None and not choice_parts:
                continue

            label = _visual_label_from_span(record, records, text, start)
            field_id = _unique_field_id(_make_field_id(label), known_ids)
            known_ids.add(field_id)

            if choice_parts:
                options = [
                    {"label": part[:1].upper() + part[1:] if part else part, "value": part}
                    for part in choice_parts
                ]
                result.append(
                    _candidate(
                        field_id=field_id,
                        label=label,
                        field_type="dropdown",
                        confidence=0.82,
                        source="colored_choice_block" if whole_red else "colored_inline_choice",
                        preview=text,
                        location=(
                            {"kind": "paragraph", "paragraph": record.ordinal}
                            if whole_red
                            else {
                                "kind": "text_span",
                                "paragraph": record.ordinal,
                                "start": start,
                                "end": end,
                                "original": raw,
                            }
                        ),
                        options=options,
                        layout="choice",
                        layout_group=f"auto_choice_{field_id}",
                    )
                )
                continue

            section = str(semantic_section(record) or "").strip()
            inference = infer_field_type(label, section=section, preview=text)
            field_type = inference.field_type
            if angle is not None and field_type == "text" and len(text) >= 70:
                field_type = "multiline"
            if whole_red and len(text) >= 110 and field_type == "text":
                field_type = "multiline"

            confidence = 0.84 if angle is not None else (0.74 if not whole_red else 0.68)
            candidate = _candidate(
                field_id=field_id,
                label=label,
                field_type=field_type,
                confidence=confidence,
                source="colored_visual_field",
                preview=text,
                location=(
                    {"kind": "paragraph", "paragraph": record.ordinal}
                    if whole_red
                    else {
                        "kind": "text_span",
                        "paragraph": record.ordinal,
                        "start": start,
                        "end": end,
                        "original": raw,
                    }
                ),
            )
            if section:
                candidate["section"] = section
            candidate["type_inference"] = {
                "confidence": inference.confidence,
                "reasons": list(inference.reasons),
            }
            candidate["visual_intent"] = "angle_placeholder" if angle is not None else "colored_text"
            result.append(candidate)
    return result


def _detect_colored_inline_choice(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect a choice expressed by a colored span inside static prose.

    Government/administrative templates frequently encode author intent using
    color rather than a native Word control.  A typical sentence is::

        Encaminham-se os autos à [área técnica competente OU à equipe ...]

    where only the bracketed text is red.  Whole-paragraph red detectors cannot
    see that shape, so locate contiguous red runs and replace only that span.
    The rule is deliberately narrow: exactly one colored region, two to five
    alternatives separated by the standalone word ``ou``, and meaningful
    static prose around the region.
    """

    paragraph = record.paragraph
    if not paragraph.runs or _contains_authoritative_marker(paragraph):
        return None

    choice_groups: list[tuple[int, int, str, list[str]]] = []
    for start, end, raw in _colored_run_groups(paragraph):
        text = _normalize_space(raw)
        if len(text) < 3 or len(text) > 520:
            continue
        parts = _visual_choice_parts(text)
        if not parts:
            continue
        choice_groups.append((start, end, raw, parts))

    if len(choice_groups) != 1:
        return None

    start, end, raw, parts = choice_groups[0]
    full_text = paragraph.text or ""
    prefix = full_text[:start]
    suffix = full_text[end:]
    # Mixed prose is the safety signal. A fully-red paragraph belongs to the
    # existing long-choice/instruction logic instead.
    if not _normalize_space(prefix) and not _normalize_space(suffix):
        return None

    # If the static prefix already carries a preposition, avoid duplicating it
    # when a later alternative repeats the same preposition (``à A ou à B``).
    prefix_match = re.search(
        r"\b(?P<prep>à|a|ao|aos|às|para|em|no|na|nos|nas|de|do|da|dos|das)\s*$",
        prefix,
        flags=re.IGNORECASE,
    )
    normalized_parts = list(parts)
    if prefix_match:
        prep = prefix_match.group("prep")
        prep_pattern = re.compile(rf"^{re.escape(prep)}\s+", flags=re.IGNORECASE)
        normalized_parts = [prep_pattern.sub("", part, count=1) for part in normalized_parts]

    if len({part.casefold() for part in normalized_parts}) != len(normalized_parts):
        return None

    section = str(semantic_section(record) or "").strip()
    context_label = _context_label_for_record(record, records)
    if re.search(r"\bencaminh", prefix, flags=re.IGNORECASE):
        label = "Destino do encaminhamento"
    else:
        label = _clean_label(context_label or section or "Escolha uma opção")
    field_id = _unique_field_id(_make_field_id(label), known_ids)
    options = [
        {"label": part[:1].upper() + part[1:] if part else part, "value": part}
        for part in normalized_parts
    ]
    candidate = _candidate(
        field_id=field_id,
        label=label,
        field_type="dropdown",
        confidence=0.95,
        source="colored_inline_choice",
        preview=raw,
        location={
            "kind": "text_span",
            "paragraph": record.ordinal,
            "start": start,
            "end": end,
            "original": raw,
        },
        options=options,
        layout="choice",
        layout_group=f"auto_choice_{field_id}",
    )
    if section:
        candidate["section"] = section
    return candidate


def _detect_colored_prompt(
    record: _ParagraphRecord,
    records: list[_ParagraphRecord],
    known_ids: set[str],
) -> dict[str, Any] | None:
    """Detect short colored placeholders such as ``Nome`` / ``Matrícula``.

    Some institutional DOCX files use red words as the fill target itself
    rather than ``XXXX`` or an underscore.  We require a short field-like
    token and strong structural context so ordinary colored instructions stay
    static.
    """

    from app.document.detection.field_inference import infer_field_type
    from app.document.detection.identifiers import _unique_contextual_field_id

    text = _normalize_space(record.text)
    red_text = _normalize_space("".join(raw for _start, _end, raw in _colored_run_groups(record.paragraph)))
    if (
        not text
        or len(text) > 70
        or red_text != text
        or _contains_authoritative_marker(record.paragraph)
        or INSTRUCTION_PATTERN.match(text)
    ):
        return None
    role = str(getattr(getattr(record, "understanding", None), "role", "") or "")
    if role in {"note", "heading", "table_header", "table_reference", "tagged"}:
        return None
    if record.cell is None and _looks_like_section_label(text):
        return None

    semantic, _source, semantic_confidence = semantic_label(record)
    owner = getattr(record, "structure", None)
    section = str(getattr(owner, "section", "") or semantic_section(record) or "")
    normalized = _slug(text)
    if _looks_like_static_colored_guidance(text):
        return None
    if normalized in {"ou", "e_ou"}:
        return None

    # Unknown red prompts are intentionally surfaced for review.  Requiring a
    # fixed vocabulary here caused new SIAGOV fields to disappear until the
    # application itself was patched.  Structural/semantic role filters above
    # still protect headings and reference-table text from automatic guesses.
    table_kind = str(getattr(owner, "table_kind", "") or "")
    if record.cell is None and semantic_confidence < 0.72 and len(text.split()) > 8:
        return None
    label = (
        semantic
        if table_kind in {"fixed_form", "repeatable", "editable_sheet"}
        and semantic and semantic_confidence >= 0.75
        else _clean_label(text)
    )
    semantics = getattr(record, "understanding", None)
    row_label = str(getattr(semantics, "row_label", "") or "")
    column_label = str(getattr(semantics, "column_label", "") or "")
    field_id = _unique_contextual_field_id(
        label,
        known_ids,
        section=section,
        row_label=row_label,
        column_label=column_label,
    )
    inference = infer_field_type(label, section=section, preview=text)
    candidate = _candidate(
        field_id=field_id,
        label=label,
        field_type=inference.field_type,
        confidence=0.91 if semantic_confidence >= 0.80 else 0.83,
        source="colored_prompt",
        preview=text,
        location={
            "kind": "paragraph",
            "paragraph": record.ordinal,
        },
    )
    if section:
        candidate["section"] = section
    candidate["type_inference"] = {
        "confidence": inference.confidence,
        "reasons": list(inference.reasons),
    }
    return candidate
