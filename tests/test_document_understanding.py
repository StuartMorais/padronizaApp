from pathlib import Path

from docx import Document

from app.document.detection.candidates import candidate_field_definitions
from app.document.detection.detector import detect_docx_field_candidates
from app.document.detection.records import _collect_paragraph_records
from app.document.understanding.semantic import (
    annotate_document_records,
    postprocess_candidates,
    semantic_label,
)


def test_relationship_model_prefers_adjacent_label_in_alternating_form_grid(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).merge(table.cell(0, 3)).text = "1. Identificação"
    table.cell(1, 0).text = "Nome:"
    table.cell(1, 1).text = "____________"
    table.cell(1, 2).text = "Matrícula:"
    table.cell(1, 3).text = "________"

    path = tmp_path / "grid.docx"
    document.save(path)

    loaded = Document(path)
    records = _collect_paragraph_records(loaded)
    annotate_document_records(records)
    value_record = next(record for record in records if record.text == "________")

    label, source, confidence = semantic_label(value_record)
    assert label == "Matrícula"
    assert source == "adjacent_left"
    assert confidence >= 0.95


def test_v3_downgrades_anonymous_field_and_explains_confidence() -> None:
    candidate = {
        "field_id": "auto.campo_39",
        "label": "Campo 39",
        "type": "text",
        "confidence": 0.91,
        "source": "inline_placeholder",
        "preview": "________",
        "location": {"kind": "text_span", "paragraph": 0, "start": 0, "end": 8},
        "selected": True,
    }

    processed = postprocess_candidates([candidate], [], source_kind="docx")
    assert len(processed) == 1
    result = processed[0]
    assert result["detector_version"] == 3
    assert result["confidence_band"] == "low"
    assert result["selected"] is False
    assert result["review_priority"] == "required"
    assert result["needs_review"] is True
    assert any(item["code"] == "poor_label" for item in result["evidence"])


def test_detected_fields_keep_v3_evidence_metadata(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("Telefone: (00) 00000-0000")
    path = tmp_path / "phone.docx"
    document.save(path)

    candidates = detect_docx_field_candidates(path)
    phone = next(item for item in candidates if item.get("type") == "phone")
    fields = candidate_field_definitions([phone])

    assert fields[0]["detector_version"] == 3
    assert fields[0]["detection_confidence_band"] in {"high", "medium", "low"}
    assert isinstance(fields[0]["detection_evidence"], list)
    assert fields[0]["detection_evidence"]
    assert fields[0]["detection_review_priority"] in {"ready", "recommended", "required"}


def test_consistency_repair_can_suggest_missing_peer_column_field(tmp_path: Path) -> None:
    from app.document.detection.text_fields import _detect_consistency_repair_fields

    document = Document()
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Observação"
    table.cell(1, 0).text = "Pneus"
    table.cell(1, 1).text = "________"
    table.cell(2, 0).text = "Freios"
    table.cell(2, 1).text = "________"
    table.cell(3, 0).text = "Iluminação"
    table.cell(3, 1).text = ""
    path = tmp_path / "peer-repair.docx"
    document.save(path)

    loaded = Document(path)
    records = _collect_paragraph_records(loaded)
    annotate_document_records(records)
    peer_records = [
        record
        for record in records
        if record.table_index == 0 and record.cell_index == 1 and record.row_index in {1, 2}
    ]
    candidates = [
        {
            "field_id": f"auto.peer_{index}",
            "label": "Observação",
            "type": "text",
            "source": "inline_placeholder",
            "confidence": 0.9,
            "location": {"kind": "paragraph", "paragraph": record.ordinal},
        }
        for index, record in enumerate(peer_records, start=1)
    ]

    repaired = _detect_consistency_repair_fields(records, candidates, set())
    assert len(repaired) == 1
    assert repaired[0]["source"] == "consistency_repair"
    assert "Iluminação" in repaired[0]["label"]
    assert repaired[0]["selected"] is False


def test_region_ownership_suppresses_lower_level_candidate_inside_repeatable_table() -> None:
    repeatable = {
        "field_id": "auto.itens",
        "label": "Itens",
        "type": "repeatable_table",
        "confidence": 0.95,
        "source": "repeatable_table",
        "location": {
            "kind": "repeatable_table",
            "paragraphs": [10, 11, 12, 13],
            "owned_paragraphs": [10, 11, 12, 13],
        },
    }
    stray_header_field = {
        "field_id": "auto.unidade",
        "label": "Unidade",
        "type": "text",
        "confidence": 0.92,
        "source": "sample_value",
        "location": {"kind": "paragraph", "paragraph": 10},
    }

    processed = postprocess_candidates(
        [repeatable, stray_header_field],
        [],
        source_kind="docx",
    )

    assert len(processed) == 1
    assert processed[0]["field_id"] == "auto.itens"
    assert processed[0]["region_owner"] == "repeatable_table"
    assert processed[0]["ownership_suppressed"] == 1
    assert any(item["code"] == "region_ownership" for item in processed[0]["evidence"])
