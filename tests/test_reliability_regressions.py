from __future__ import annotations

import json
import os
import zipfile
from pathlib import Path
from typing import Any

import pytest
from docx import Document

from app.document.detection.selection_policy import apply_review_first_policy
from app.document.detection.extraction import location_signature
from app.document.docx.generator import DocumentGenerationError
from app.repositories.local_data import LocalDataStore
from app.repositories.templates import TemplateRepository
from app.services.generation import GenerationService
from app.services.output_planner import OutputPlanner
from app.services.templates import TemplatePackage, discover_templates_with_issues


def _docx(path: Path, text: str) -> Path:
    document = Document()
    document.add_paragraph(text)
    document.save(path)
    return path


def _minimal_generated_docx(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr("[Content_Types].xml", "<Types/>")
        archive.writestr("word/document.xml", "<document/>")


def _package(tmp_path: Path) -> TemplatePackage:
    source = _docx(tmp_path / "source.docx", "Documento base")
    return TemplatePackage(
        template_id="modelo",
        name="Modelo",
        description="",
        category="",
        version="1.0",
        source_path=source,
        fields=[],
        output_filename="{{sequence}}.docx",
        config={
            "output": {"filename_pattern": "{{sequence}}.docx"},
            "numbering": {"enabled": True, "key": "documents", "padding": 4},
        },
    )


def test_template_replacement_keeps_docx_suffix_and_succeeds(tmp_path: Path) -> None:
    repository = TemplateRepository(tmp_path / "templates")
    original = _docx(tmp_path / "original.docx", "ORIGINAL")
    replacement = _docx(tmp_path / "replacement.docx", "REPLACEMENT")
    template_id = repository.create_template(
        name="Modelo",
        source_docx=original,
        fields=[],
    )

    repository.update_template(
        template_id=template_id,
        name="Modelo",
        fields=[],
        replacement_docx=replacement,
    )

    saved = Document(repository.get_source_path(template_id))
    assert saved.paragraphs[0].text == "REPLACEMENT"


def test_template_update_rolls_back_source_if_config_publish_fails(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.repositories import templates as templates_module

    repository = TemplateRepository(tmp_path / "templates")
    original = _docx(tmp_path / "original.docx", "ORIGINAL")
    replacement = _docx(tmp_path / "replacement.docx", "REPLACEMENT")
    template_id = repository.create_template(
        name="Modelo",
        source_docx=original,
        fields=[],
        description="old-description",
    )
    folder = tmp_path / "templates" / template_id
    original_config = (folder / "template.json").read_bytes()
    live_config = (folder / "template.json").resolve()

    real_replace = os.replace
    failed = False

    def fail_live_config_once(src: Any, dst: Any) -> None:
        nonlocal failed
        if not failed and Path(dst).resolve() == live_config:
            failed = True
            raise OSError("simulated config publish failure")
        real_replace(src, dst)

    monkeypatch.setattr(templates_module.os, "replace", fail_live_config_once)

    with pytest.raises(OSError, match="simulated config publish failure"):
        repository.update_template(
            template_id=template_id,
            name="Modelo",
            fields=[],
            description="new-description",
            replacement_docx=replacement,
        )

    assert Document(repository.get_source_path(template_id)).paragraphs[0].text == "ORIGINAL"
    assert (folder / "template.json").read_bytes() == original_config


def test_template_package_import_rejects_excessive_uncompressed_size(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.repositories import templates as templates_module

    package = tmp_path / "large-template.zip"
    with zipfile.ZipFile(package, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("template.json", "x" * 64)

    monkeypatch.setattr(templates_module, "MAX_TEMPLATE_PACKAGE_UNCOMPRESSED_BYTES", 32)
    repository = TemplateRepository(tmp_path / "templates")
    with pytest.raises(ValueError, match="grande demais"):
        repository.import_template_package(package)


def test_future_template_schema_is_reported_instead_of_silently_disappearing(
    tmp_path: Path,
) -> None:
    folder = tmp_path / "templates" / "future-model"
    folder.mkdir(parents=True)
    (folder / "template.json").write_text(
        json.dumps(
            {
                "schema_version": 999,
                "template": {
                    "id": "future-model",
                    "name": "Modelo futuro",
                    "source_file": "template.docx",
                },
                "fields": [],
            }
        ),
        encoding="utf-8",
    )

    packages, issues = discover_templates_with_issues(tmp_path / "templates")

    assert packages == []
    assert len(issues) == 1
    assert issues[0]["template_id"] == "future-model"
    assert issues[0]["error_type"] == "SchemaVersionError"
    assert "schema 999" in issues[0]["message"]


def test_generation_metadata_failure_restores_existing_output_and_sequence(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import generation as generation_module

    store = LocalDataStore(tmp_path / "data")
    service = GenerationService(store, OutputPlanner(store))
    package = _package(tmp_path)
    output = tmp_path / "0001.docx"
    output.write_bytes(b"ORIGINAL")

    monkeypatch.setattr(
        generation_module,
        "generate_docx",
        lambda _source, destination, _values: _minimal_generated_docx(Path(destination)),
    )
    monkeypatch.setattr(
        store,
        "commit_generated_document",
        lambda *_args, **_kwargs: (_ for _ in ()).throw(OSError("metadata disk failure")),
    )

    with pytest.raises(DocumentGenerationError, match="metadados anteriores foram restaurados"):
        service.generate_docx(package, {}, output)

    assert output.read_bytes() == b"ORIGINAL"
    assert store.peek_sequence("documents") == 1
    assert store.list_recent() == []


def test_local_metadata_transaction_rolls_back_all_stores_on_publish_failure(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.repositories import local_data as local_data_module

    store = LocalDataStore(tmp_path / "data")
    store.add_recent({"id": "before", "filename": "before.docx"})
    store.add_audit("before", "Before")
    assert store.next_sequence("documents", 2026) == 1
    before_recent = store.list_recent()
    before_audit = store.list_audit()
    before_sequence = store.peek_sequence("documents", 2026)
    live_sequence = (store.data_dir / "sequences.json").resolve()

    real_replace = os.replace
    failed = False

    def fail_sequence_once(src: Any, dst: Any) -> None:
        nonlocal failed
        if not failed and Path(dst).resolve() == live_sequence:
            failed = True
            raise OSError("simulated sequence publish failure")
        real_replace(src, dst)

    monkeypatch.setattr(local_data_module.os, "replace", fail_sequence_once)

    with pytest.raises(OSError, match="simulated sequence publish failure"):
        store.commit_generated_document(
            {"id": "after", "filename": "after.docx"},
            audit_action="document_generated",
            audit_description="after.docx",
            audit_details={"document_id": "after"},
            sequence_key="documents",
            sequence_year=2026,
            expected_sequence=2,
        )

    assert store.list_recent() == before_recent
    assert store.list_audit() == before_audit
    assert store.peek_sequence("documents", 2026) == before_sequence


def test_review_first_policy_keeps_ambiguous_prose_unselected() -> None:
    [candidate] = apply_review_first_policy(
        [
            {
                "source": "instruction",
                "confidence": 0.99,
                "confidence_dimensions": {
                    "structure": 1.0,
                    "fillable": 1.0,
                    "label": 1.0,
                    "type": 1.0,
                },
                "review_priority": "ready",
            }
        ]
    )

    assert candidate["selected"] is False
    assert candidate["auto_apply_eligible"] is False
    assert any("confirmação humana" in reason for reason in candidate["auto_apply_reasons"])


def test_review_first_policy_preselects_strong_structural_candidate() -> None:
    [candidate] = apply_review_first_policy(
        [
            {
                "source": "empty_cell",
                "confidence": 0.95,
                "confidence_dimensions": {
                    "structure": 0.90,
                    "fillable": 0.85,
                    "label": 0.80,
                    "type": 0.75,
                },
                "review_priority": "ready",
            }
        ]
    )

    assert candidate["selected"] is True
    assert candidate["auto_apply_eligible"] is True
    assert candidate["pipeline_version"] == 6


def test_location_signature_is_stable_across_mapping_order() -> None:
    assert location_signature({"row": 3, "table": 1, "cell": 2}) == location_signature(
        {"cell": 2, "table": 1, "row": 3}
    )


def test_backup_settings_failure_rolls_back_data_and_previous_settings(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    from app.services import backup as backup_module

    project = tmp_path / "project"
    (project / "templates").mkdir(parents=True)
    (project / "data").mkdir(parents=True)
    (project / "templates" / "current.txt").write_text("old-template", encoding="utf-8")
    (project / "data" / "current.txt").write_text("old-data", encoding="utf-8")

    backup = tmp_path / "backup.zip"
    with zipfile.ZipFile(backup, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "backup_info.json",
            json.dumps({"format": 2, "schema_version": 1}),
        )
        archive.writestr("settings.json", json.dumps({"appearance": "new"}))
        archive.writestr("templates/current.txt", "new-template")
        archive.writestr("data/current.txt", "new-data")

    class FailingSettings:
        def __init__(self) -> None:
            self.values = {"appearance": "old"}
            self.failed_once = False

        def allKeys(self) -> list[str]:
            return list(self.values)

        def value(self, key: str) -> Any:
            return self.values.get(key)

        def clear(self) -> None:
            self.values.clear()

        def setValue(self, key: str, value: Any) -> None:
            self.values[str(key)] = value

        def sync(self) -> None:
            if self.values.get("appearance") == "new" and not self.failed_once:
                self.failed_once = True
                raise OSError("settings backend failure")

        def status(self) -> int:
            return 0

    settings = FailingSettings()
    monkeypatch.setattr(backup_module, "_settings_store", lambda: settings)

    with pytest.raises(backup_module.BackupError, match="pastas restauradas também foram revertidas"):
        backup_module.restore_backup(project, backup)

    assert (project / "templates" / "current.txt").read_text(encoding="utf-8") == "old-template"
    assert (project / "data" / "current.txt").read_text(encoding="utf-8") == "old-data"
    assert settings.values == {"appearance": "old"}
