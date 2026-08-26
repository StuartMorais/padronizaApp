from __future__ import annotations

from pathlib import Path

from docx import Document

from app.services.template_scanning import TemplateScanResult, locate_template_fields


def test_one_localization_pass_returns_authoritative_and_untagged_fields(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("Nome: {{cliente.nome}}")
    document.add_paragraph("Matrícula: XXXXXXXX")
    source = tmp_path / "mixed.docx"
    document.save(str(source))

    result = locate_template_fields(source)

    assert isinstance(result, TemplateScanResult)
    assert {field["id"] for field in result.fields} >= {"cliente.nome"}
    assert any(candidate["label"] == "Matrícula" for candidate in result.candidates)
    assert all(candidate["field_id"] != "cliente.nome" for candidate in result.candidates)
    assert isinstance(result.report, dict)


def test_real_dfd_localization_keeps_binary_choice_and_inline_dispatch_choice() -> None:
    fixture = Path(__file__).parent / "fixtures" / "dfd_licitacao_tradicional_sia13tdr.docx"

    result = locate_template_fields(fixture, repair_repeatable_markers=False)
    sources = {str(candidate.get("source", "")) for candidate in result.candidates}

    assert "long_choice" in sources
    assert "colored_inline_choice" in sources
    assert any(str(field.get("id", "")) == "descricao.demanda" for field in result.fields)
