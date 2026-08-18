from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.automatic_field_detector import (
    apply_docx_field_candidates,
    candidate_field_definitions,
    detect_docx_field_candidates,
)
from app.smart_template import smart_fields_from_docx
from app.docx_engine import generate_docx


def _save(document: Document, path: Path) -> Path:
    document.save(str(path))
    return path


def test_detects_inline_placeholders_and_converts_them_to_tags(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Responsável pela demanda: XXXXXXXX"
    table.cell(0, 1).text = "Matrícula: XXXXXXXX"
    table.cell(1, 0).text = "E-mail: xxxxxxxx@xxxxxxxxxxx"
    table.cell(1, 1).text = "Telefone: (83) XXXX-XXXX"
    source = _save(document, tmp_path / "source.docx")

    candidates = detect_docx_field_candidates(source)
    selected = [item for item in candidates if item["source"] == "inline_placeholder"]

    assert len(selected) == 4
    assert {item["type"] for item in selected} >= {"email", "phone", "text"}

    output = tmp_path / "prepared.docx"
    apply_docx_field_candidates(source, output, selected)
    fields = smart_fields_from_docx(output, candidate_field_definitions(selected))

    by_label = {field["label"]: field for field in fields}
    assert by_label["E-mail"]["type"] == "email"
    assert by_label["Telefone"]["type"] == "phone"
    assert len(fields) == 4


def test_detects_red_instruction_as_multiline_field(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "2. Descrição da Demanda"
    paragraph = table.cell(1, 0).paragraphs[0]
    run = paragraph.add_run(
        "Informar a descrição sucinta do objeto da compra ou serviço, "
        "incluindo todas as informações necessárias."
    )
    run.font.color.rgb = RGBColor(255, 0, 0)
    source = _save(document, tmp_path / "instruction.docx")

    candidates = detect_docx_field_candidates(source)
    instruction = next(item for item in candidates if item["source"] == "instruction")

    assert instruction["type"] == "multiline"
    assert instruction["selected"] is True
    assert "Descrição" in instruction["label"]


def test_detects_ou_block_as_long_single_choice(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "9.1 Justificativa em caso de ausência no PCA"
    cell = table.cell(1, 0)
    cell.text = "Não se aplica"
    for value in (
        "OU",
        "Está dispensada em razão do valor estimado não ser superior ao limite legal.",
        "OU",
        "Justifica-se por demandas supervenientes surgidas após o planejamento.",
        "OU",
        "Será verificado posteriormente pelo setor administrativo.",
    ):
        cell.add_paragraph(value)
    source = _save(document, tmp_path / "choice.docx")

    candidates = detect_docx_field_candidates(source)
    choice = next(item for item in candidates if item["source"] == "long_choice")

    assert choice["type"] == "dropdown"
    assert choice["layout"] == "choice"
    assert len(choice["options"]) == 4
    assert choice["confidence"] >= 0.9

    output = tmp_path / "choice-prepared.docx"
    apply_docx_field_candidates(source, output, [choice])
    fields = smart_fields_from_docx(output, candidate_field_definitions([choice]))

    assert len(fields) == 1
    assert fields[0]["type"] == "dropdown"
    assert fields[0]["layout"] == "choice"
    assert len(fields[0]["options"]) == 4

    generated_template = Document(str(output))
    all_text = "\n".join(
        paragraph.text
        for table in generated_template.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    assert "{{single_choice:" in all_text
    assert "\nOU\n" not in all_text


def test_existing_tags_are_authoritative_and_not_suggested_again(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("Nome: {{responsavel.nome}}")
    document.add_paragraph("Matrícula: XXXXXXXX")
    source = _save(document, tmp_path / "tagged.docx")

    candidates = detect_docx_field_candidates(source)

    assert all(item["field_id"] != "responsavel.nome" for item in candidates)
    assert len(candidates) == 1
    assert candidates[0]["label"] == "Matrícula"


def test_generic_dropdown_prompt_requires_manual_options(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "7. Grau de Prioridade"
    table.cell(1, 0).text = "Escolher um item."
    source = _save(document, tmp_path / "dropdown.docx")

    candidates = detect_docx_field_candidates(source)
    prompt = next(item for item in candidates if item["source"] == "dropdown_prompt")

    assert prompt["type"] == "dropdown"
    assert prompt["selected"] is False
    assert prompt["requires_configuration"] is True
    assert prompt.get("options", []) == []


def test_automatic_tag_insertion_preserves_surrounding_run_formatting(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run("Responsável: ")
    label_run.bold = True
    label_run.font.color.rgb = RGBColor(0, 64, 192)
    placeholder_run = paragraph.add_run("XXXXXXXX")
    placeholder_run.italic = True
    placeholder_run.font.color.rgb = RGBColor(255, 0, 0)
    suffix_run = paragraph.add_run(" / conferido")
    suffix_run.underline = True
    source = _save(document, tmp_path / "formatted.docx")

    candidate = next(
        item
        for item in detect_docx_field_candidates(source)
        if item["source"] == "inline_placeholder"
    )
    output = tmp_path / "formatted-prepared.docx"
    apply_docx_field_candidates(source, output, [candidate])

    result = Document(str(output))
    runs = [run for run in result.paragraphs[0].runs if run.text]
    assert runs[0].text == "Responsável: "
    assert runs[0].bold is True
    assert runs[0].font.color.rgb == RGBColor(0, 64, 192)
    assert "{{" in runs[1].text
    assert runs[1].italic is True
    assert runs[1].font.color.rgb == RGBColor(255, 0, 0)
    assert runs[2].text == " / conferido"
    assert runs[2].underline is True




def test_detected_date_mask_stays_editable(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=3)
    table.cell(0, 0).text = "Unidade visitada: XXXXXXXXXXXX"
    table.cell(0, 1).text = "Data: __/__/____"
    table.cell(0, 2).text = "Horário: __:__"
    source = _save(document, tmp_path / "editable-date.docx")

    candidates = detect_docx_field_candidates(source)
    date_candidate = next(item for item in candidates if item["label"] == "Data")
    assert date_candidate["type"] == "date"

    definitions = candidate_field_definitions([date_candidate])
    assert len(definitions) == 1
    assert definitions[0]["type"] == "date"
    assert definitions[0]["automatic"] is False

    output = tmp_path / "editable-date-prepared.docx"
    apply_docx_field_candidates(source, output, [date_candidate])
    fields = smart_fields_from_docx(output, definitions)
    date_field = next(field for field in fields if field["id"] == date_candidate["field_id"])
    assert date_field["automatic"] is False

def test_detects_full_masks_sample_email_phone_and_inline_dropdown(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=2)
    table.cell(0, 0).text = "CNPJ: __.___.___/____-__"
    table.cell(0, 1).text = "CPF: ___.___.___-__"
    table.cell(1, 0).text = "E-mail principal: contato@empresa.com.br"
    table.cell(1, 1).text = "Telefone: (00) 00000-0000"
    table.cell(2, 0).text = "UF: __"
    table.cell(2, 1).text = "CEP: __.___-___"
    table.cell(3, 0).text = "Tipo de fornecedor: Escolher uma opção"
    table.cell(3, 1).text = "País: Brasil"
    source = _save(document, tmp_path / "masks.docx")

    candidates = detect_docx_field_candidates(source)
    by_label = {item["label"]: item for item in candidates}

    assert by_label["CNPJ"]["preview"] == "__.___.___/____-__"
    assert by_label["CNPJ"]["type"] == "cnpj"
    assert by_label["CPF"]["preview"] == "___.___.___-__"
    assert by_label["CPF"]["type"] == "cpf"
    assert by_label["E-mail principal"]["type"] == "email"
    assert by_label["E-mail principal"]["preview"] == "contato@empresa.com.br"
    assert by_label["Telefone"]["type"] == "phone"
    assert by_label["Telefone"]["preview"] == "(00) 00000-0000"
    assert by_label["UF"]["preview"] == "__"
    assert by_label["CEP"]["type"] == "cep"
    assert by_label["Tipo de fornecedor"]["source"] == "dropdown_prompt"
    assert by_label["Tipo de fornecedor"]["requires_configuration"] is True
    assert by_label["País"]["source"] == "sample_value"
    assert by_label["País"]["preview"] == "Brasil"
    assert by_label["País"]["placeholder"] == "Brasil"
    assert by_label["País"]["selected"] is True

    accepted = [
        item
        for item in candidates
        if item["label"] != "Tipo de fornecedor"
    ]
    output = tmp_path / "masks-prepared.docx"
    apply_docx_field_candidates(source, output, accepted)
    text = "\n".join(
        cell.text
        for table in Document(str(output)).tables
        for row in table.rows
        for cell in row.cells
    )
    assert "CNPJ: {{" in text
    assert "__.___.___/____-__" not in text
    assert "contato@empresa.com.br" not in text
    assert "(00) 00000-0000" not in text
    assert "País: {{" in text
    assert "País: Brasil" not in text

    fields = smart_fields_from_docx(output, candidate_field_definitions(accepted))
    country = next(field for field in fields if field["label"] == "País")
    assert country["placeholder"] == "Brasil"



def test_detects_zero_cpf_mask_in_adjacent_cell_as_editable_cpf(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=4)
    table.cell(0, 0).text = "CPF"
    table.cell(0, 1).text = "000.000.000-00"
    table.cell(0, 2).text = "Data da solicitação"
    table.cell(0, 3).text = "__/__/____"
    source = _save(document, tmp_path / "cpf-zero-mask.docx")

    candidates = detect_docx_field_candidates(source)
    cpf_candidates = [item for item in candidates if item["label"] == "CPF"]

    assert len(cpf_candidates) == 1
    cpf = cpf_candidates[0]
    assert cpf["type"] == "cpf"
    assert cpf["preview"] == "000.000.000-00"
    assert cpf["source"] == "inline_placeholder"

    output = tmp_path / "cpf-zero-mask-prepared.docx"
    apply_docx_field_candidates(source, output, [cpf])
    text = "\n".join(
        cell.text
        for doc_table in Document(str(output)).tables
        for row in doc_table.rows
        for cell in row.cells
    )
    assert "000.000.000-00" not in text
    assert "{{" in text

    fields = smart_fields_from_docx(output, candidate_field_definitions([cpf]))
    cpf_field = next(field for field in fields if field["label"] == "CPF")
    assert cpf_field["type"] == "cpf"

def test_detects_currency_masks_and_preserves_parenthetical_context(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph(
        "Valor estimado: R$ XXX.XXX,XX (valor por extenso)"
    )
    document.add_paragraph(
        "Valor contratual: R$ XXX.XXX.XX"
    )
    document.add_paragraph(
        "Custo previsto: R$ 000.000,00"
    )
    document.add_paragraph(
        "Orçamento: R$ ________,__"
    )
    source = _save(document, tmp_path / "currency-masks.docx")

    candidates = detect_docx_field_candidates(source)
    currency = [item for item in candidates if item.get("type") == "currency"]

    assert len(currency) == 4
    assert {item["label"] for item in currency} == {
        "Valor estimado",
        "Valor contratual",
        "Custo previsto",
        "Orçamento",
    }
    assert {item["preview"] for item in currency} == {
        "R$ XXX.XXX,XX",
        "R$ XXX.XXX.XX",
        "R$ 000.000,00",
        "R$ ________,__",
    }

    output = tmp_path / "currency-masks-prepared.docx"
    apply_docx_field_candidates(source, output, currency)
    text = "\n".join(paragraph.text for paragraph in Document(str(output)).paragraphs)

    assert "R$ XXX.XXX,XX" not in text
    assert "R$ XXX.XXX.XX" not in text
    assert "R$ 000.000,00" not in text
    assert "R$ ________,__" not in text
    assert "(valor por extenso)" in text

    fields = smart_fields_from_docx(output, candidate_field_definitions(currency))
    assert all(field["type"] == "currency" for field in fields)


def test_bare_currency_mask_gets_semantic_value_label(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("R$ XXX.XXX,XX (valor por extenso)")
    source = _save(document, tmp_path / "bare-currency-mask.docx")

    candidates = detect_docx_field_candidates(source)
    currency = [item for item in candidates if item.get("type") == "currency"]

    assert len(currency) == 1
    assert currency[0]["label"] == "Valor"
    assert currency[0]["preview"] == "R$ XXX.XXX,XX"


def test_detects_inline_and_vertical_checkbox_groups_with_correct_semantics(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = (
        "Natureza do atendimento: ☐ Material   ☐ Serviço   ☐ Material e serviço"
    )

    declarations = document.add_table(rows=3, cols=1)
    declarations.cell(0, 0).text = "☐ Declaro que as informações são verdadeiras."
    declarations.cell(1, 0).text = "☐ Autorizo o uso administrativo dos dados."
    paragraph = declarations.cell(2, 0).paragraphs[0]
    paragraph.add_run("Observação do cadastro: ").bold = True
    instruction = paragraph.add_run("Preencher somente se houver informação complementar.")
    instruction.font.color.rgb = RGBColor(255, 0, 0)
    source = _save(document, tmp_path / "checkboxes.docx")

    candidates = detect_docx_field_candidates(source)
    checkbox_groups = [item for item in candidates if item["source"] == "checkbox_choice"]
    assert len(checkbox_groups) == 2

    nature = next(item for item in checkbox_groups if item["label"] == "Natureza do atendimento")
    assert nature["selection"] == "single"
    assert nature["location"]["kind"] == "checkbox_group_inline"
    assert [field["label"] for field in nature["fields"]] == [
        "Material",
        "Serviço",
        "Material e serviço",
    ]
    assert all(field.get("layout") == "choice" for field in nature["fields"])

    declaration = next(item for item in checkbox_groups if item["selection"] == "multiple")
    assert declaration["location"]["kind"] == "checkbox_group"
    assert len(declaration["fields"]) == 2
    assert all(field.get("selection") == "multiple" for field in declaration["fields"])
    assert all(not field.get("choice_required", False) for field in declaration["fields"])

    observation = next(item for item in candidates if item["label"] == "Observação do cadastro")
    assert observation["type"] == "multiline"
    assert observation["source"] == "instruction"

    accepted = [nature, declaration, observation]
    output = tmp_path / "checkboxes-prepared.docx"
    apply_docx_field_candidates(source, output, accepted)
    fields = smart_fields_from_docx(output, candidate_field_definitions(accepted))
    by_id = {field["id"]: field for field in fields}
    assert len([field for field in fields if field.get("selection") == "single"]) == 3
    assert len([field for field in fields if field.get("selection") == "multiple"]) == 2
    assert by_id[observation["field_id"]]["type"] == "multiline"


def test_detects_bare_label_field_without_duplicating_adjacent_empty_cell(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Responsável legal:"
    table.cell(0, 1).text = "CPF: ___.___.___-__"
    table.cell(1, 0).text = "Endereço completo:"
    table.cell(1, 1).text = ""
    source = _save(document, tmp_path / "bare-label.docx")

    candidates = detect_docx_field_candidates(source)
    labels = [item["label"] for item in candidates]
    assert labels.count("Responsável legal") == 1
    assert labels.count("Endereço completo") == 1

    responsible = next(item for item in candidates if item["label"] == "Responsável legal")
    address = next(item for item in candidates if item["label"] == "Endereço completo")
    assert responsible["location"]["kind"] == "append_tag"
    assert address["location"]["kind"] == "empty_cell"
    assert address["selected"] is True
    assert address["confidence"] >= 0.80

    output = tmp_path / "bare-label-prepared.docx"
    apply_docx_field_candidates(source, output, [responsible, address])
    result = Document(str(output))
    assert "Responsável legal: {{" in result.tables[0].cell(0, 0).text
    assert result.tables[0].cell(1, 0).text == "Endereço completo:"
    assert "{{" in result.tables[0].cell(1, 1).text


def test_table_context_uses_row_and_column_labels_for_unlabeled_fill_area(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Item verificado"
    table.cell(0, 1).text = "Situação"
    table.cell(0, 2).text = "Observação curta"
    table.cell(1, 0).text = "Documentação disponível"
    table.cell(1, 1).text = "☐ Conforme   ☐ Não conforme   ☐ Não se aplica"
    table.cell(1, 2).text = "________________"
    source = _save(document, tmp_path / "questionnaire.docx")

    candidates = detect_docx_field_candidates(source)
    observation = next(
        item
        for item in candidates
        if item.get("source") == "inline_placeholder"
        and "Observação curta" in item.get("label", "")
    )

    assert observation["label"] == "Documentação disponível — Observação curta"
    assert observation["field_id"] == "auto.documentacao_disponivel_observacao_curta"
    assert observation["selected"] is True
    assert not observation["field_id"].startswith("auto.campo_")


def test_instruction_uses_label_from_previous_paragraph_in_same_cell(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "Descrição das constatações:"
    instruction_one = table.cell(0, 0).add_paragraph(
        "Descrever os fatos observados e as evidências encontradas."
    )
    instruction_one.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    table.cell(1, 0).text = "Providência recomendada:"
    instruction_two = table.cell(1, 0).add_paragraph(
        "Informar a ação necessária para regularização."
    )
    instruction_two.runs[0].font.color.rgb = RGBColor(255, 0, 0)
    source = _save(document, tmp_path / "same-cell-label.docx")

    candidates = detect_docx_field_candidates(source)
    labels = [item["label"] for item in candidates if item.get("source") == "instruction"]

    assert "Descrição das constatações" in labels
    assert "Providência recomendada" in labels
    assert not any("Descrição das constatações: Descrever" in label for label in labels)


def test_detects_numbered_occurrence_table_as_repeatable_and_suppresses_anonymous_fields(tmp_path: Path) -> None:
    document = Document()
    title = document.add_table(rows=1, cols=1)
    title.cell(0, 0).text = "4. Ocorrências / pendências"

    table = document.add_table(rows=4, cols=4)
    for index, header in enumerate(("Nº", "Descrição", "Prazo", "Responsável")):
        table.cell(0, index).text = header
    table.cell(1, 0).text = "01"
    table.cell(1, 1).text = "________________________________"
    table.cell(1, 2).text = "__/__/____"
    table.cell(1, 3).text = "____________"
    table.cell(2, 0).text = "02"
    table.cell(3, 0).text = "03"
    source = _save(document, tmp_path / "occurrences.docx")

    candidates = detect_docx_field_candidates(source)
    repeatable = next(item for item in candidates if item.get("source") == "repeatable_table")

    assert repeatable["type"] == "repeatable_table"
    assert repeatable["field_id"] == "auto.ocorrencias_pendencias"
    assert repeatable["section"] == "4. Ocorrências / pendências"
    assert repeatable["selected"] is True
    assert [(column["id"], column["type"]) for column in repeatable["columns"]] == [
        ("item", "auto_number"),
        ("descricao", "multiline"),
        ("prazo", "date"),
        ("responsavel", "text"),
    ]
    assert not any(item["field_id"].startswith("auto.campo_") for item in candidates)

    output = tmp_path / "occurrences-prepared.docx"
    apply_docx_field_candidates(source, output, [repeatable])
    prepared = Document(str(output))
    prepared_table = prepared.tables[1]
    assert len(prepared_table.rows) == 2
    assert "{{repeat:auto.ocorrencias_pendencias}}" in prepared_table.cell(1, 0).text
    assert "{{row.number}}" in prepared_table.cell(1, 0).text
    assert "{{auto.ocorrencias_pendencias.descricao}}" in prepared_table.cell(1, 1).text
    assert "{{date:auto.ocorrencias_pendencias.prazo}}" in prepared_table.cell(1, 2).text
    assert "{{auto.ocorrencias_pendencias.responsavel}}" in prepared_table.cell(1, 3).text

    fields = smart_fields_from_docx(output, candidate_field_definitions([repeatable]))
    field = next(item for item in fields if item["id"] == "auto.ocorrencias_pendencias")
    assert field["type"] == "repeatable_table"
    assert field["section"] == "4. Ocorrências / pendências"
    assert [column["id"] for column in field["columns"]] == [
        "item",
        "descricao",
        "prazo",
        "responsavel",
    ]


def test_questionnaire_choices_stay_embedded_with_observation_column(tmp_path: Path) -> None:
    document = Document()
    title = document.add_table(rows=1, cols=1)
    title.cell(0, 0).text = "2. Condições verificadas"

    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Item verificado"
    table.cell(0, 1).text = "Situação"
    table.cell(0, 2).text = "Observação curta"
    table.cell(1, 0).text = "Documentação disponível"
    table.cell(1, 1).text = "☐ Conforme   ☐ Não conforme   ☐ Não se aplica"
    table.cell(1, 2).text = "________________"
    table.cell(2, 0).text = "Condições de segurança"
    table.cell(2, 1).text = "☐ Conforme   ☐ Não conforme   ☐ Não se aplica"
    table.cell(2, 2).text = "________________"
    source = _save(document, tmp_path / "questionnaire-layout.docx")

    candidates = detect_docx_field_candidates(source)
    accepted = [
        item
        for item in candidates
        if item.get("selected", True) and not item.get("requires_configuration", False)
    ]
    output = tmp_path / "questionnaire-layout-prepared.docx"
    apply_docx_field_candidates(source, output, accepted)
    fields = smart_fields_from_docx(output, candidate_field_definitions(accepted))
    by_id = {field["id"]: field for field in fields}

    choice = by_id["auto.documentacao_disponivel.conforme"]
    observation = by_id["auto.documentacao_disponivel_observacao_curta"]

    assert choice["layout"] == "table"
    assert observation["layout"] == "table"
    assert choice["layout_group"] == observation["layout_group"]
    assert choice["layout_row"] == observation["layout_row"] == "row_1"
    assert choice["layout_column"] == "Situação"
    assert observation["layout_column"] == "Observação curta"
    assert choice["layout_row_header_label"] == "Item verificado"
    assert choice["group"].startswith("auto_checkbox_")
    assert choice["selection"] == "single"
    assert choice["compact_choice"] is True
    assert observation["type"] == "text"


def test_inline_yes_no_question_keeps_group_prompt_inside_form_grid(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=1)
    table.cell(0, 0).text = "3. Fundamentação complementar"
    table.cell(1, 0).text = "Justificativa complementar: Preencher quando necessário."
    table.cell(2, 0).text = "Grau de risco: Escolher um item"
    table.cell(3, 0).text = "Há impedimento conhecido? ☐ Sim   ☐ Não"
    source = _save(document, tmp_path / "inline-question.docx")

    candidates = detect_docx_field_candidates(source)
    choice = next(item for item in candidates if item["label"] == "Há impedimento conhecido?")
    assert choice["selection"] == "single"

    output = tmp_path / "inline-question-prepared.docx"
    apply_docx_field_candidates(source, output, [choice])
    fields = smart_fields_from_docx(output, candidate_field_definitions([choice]))

    assert len(fields) == 2
    assert {field["label"] for field in fields} == {"Sim", "Não"}
    assert all(field["layout"] == "form_grid" for field in fields)
    assert all(field["layout_row"] == "row_3" for field in fields)
    assert all(field["choice_group_label"] == "Há impedimento conhecido?" for field in fields)
    assert len({field["group"] for field in fields}) == 1


def test_detects_checkbox_alternatives_split_across_table_cells(tmp_path: Path) -> None:
    document = Document()
    title = document.add_table(rows=1, cols=1)
    title.cell(0, 0).text = "4. Forma de entrega"

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = (
        "☐ Entrega imediata e integral do material ou serviço\n\n"
        "* Considera-se imediata a entrega em até 30 dias."
    )
    table.cell(0, 1).text = (
        "☐ Entrega parcelada do material ou serviço\n\n"
        "Informar cronograma quando aplicável."
    )
    source = _save(document, tmp_path / "delivery-choice.docx")

    candidates = detect_docx_field_candidates(source)
    choice = next(
        item
        for item in candidates
        if item.get("source") == "checkbox_choice"
        and item.get("label") == "Forma de entrega"
    )

    assert choice["selection"] == "single"
    assert choice["location"]["kind"] == "checkbox_group_multi_cell"
    assert len(choice["fields"]) == 2
    labels = [field["label"] for field in choice["fields"]]
    assert labels[0].startswith("Entrega imediata e integral")
    assert "30 dias" in labels[0]
    assert labels[1].startswith("Entrega parcelada")
    assert "cronograma" in labels[1].casefold()

    output = tmp_path / "delivery-choice-prepared.docx"
    apply_docx_field_candidates(source, output, [choice])
    prepared = Document(str(output))
    left = prepared.tables[1].cell(0, 0).text
    right = prepared.tables[1].cell(0, 1).text
    assert "{{checkbox:" in left
    assert "{{checkbox:" in right
    assert "Considera-se imediata a entrega em até 30 dias." in left
    assert "Informar cronograma quando aplicável." in right

    fields = smart_fields_from_docx(output, candidate_field_definitions([choice]))
    assert len(fields) == 2
    assert all(field.get("selection") == "single" for field in fields)
    assert all(field.get("layout") == "choice" for field in fields)
    assert all(field.get("layout_group_label") == "Forma de entrega" for field in fields)


def test_detects_vertical_checkbox_markers_in_separate_cell_with_adjacent_option_text(tmp_path: Path) -> None:
    document = Document()
    intro = document.add_paragraph(
        "Em cumprimento às diretrizes estabelecidas por este órgão, declaro que foi realizada "
        "a verificação de existência de unidades do(s) produto(s) em estoque, apresentando "
        "a seguinte ocorrência:"
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "☐"
    table.cell(0, 1).text = "1. Existência de unidades em estoque;"
    table.cell(0, 1).add_paragraph(
        "Segue em anexo, relatório detalhado com o saldo de unidades em estoque."
    )
    table.cell(0, 1).add_paragraph(
        "Observação/Justificativa em caso de insuficiência ou comprometimento de saldo de estoque"
    )
    table.cell(0, 1).add_paragraph("")
    table.cell(1, 0).text = "☑"
    table.cell(1, 1).text = "2. Unidade(s) do produto não registradas em estoque"
    source = _save(document, tmp_path / "separate-checkbox-cells.docx")

    candidates = detect_docx_field_candidates(source)
    choice = next(
        item
        for item in candidates
        if item.get("source") == "checkbox_choice"
        and item.get("location", {}).get("kind") == "checkbox_group_multi_cell"
    )

    assert choice["label"] == "Ocorrência verificada"
    assert choice["selection"] == "single"
    assert len(choice["fields"]) == 2
    assert choice["fields"][0]["label"].startswith("1. Existência de unidades em estoque")
    assert "Segue em anexo" in choice["fields"][0]["label"]
    assert "Observação/Justificativa" not in choice["fields"][0]["label"]
    assert choice["fields"][1]["label"].startswith(
        "2. Unidade(s) do produto não registradas em estoque"
    )

    followup = next(
        item
        for item in candidates
        if item.get("label", "").startswith("Observação/Justificativa")
    )
    assert followup["type"] == "multiline"
    assert followup["selected"] is True

    output = tmp_path / "separate-checkbox-cells-prepared.docx"
    apply_docx_field_candidates(source, output, [choice, followup])
    prepared = Document(str(output))
    assert "{{checkbox:" in prepared.tables[0].cell(0, 0).text
    assert "{{checkbox:" in prepared.tables[0].cell(1, 0).text
    assert "1. Existência de unidades em estoque" in prepared.tables[0].cell(0, 1).text
    assert "{{auto.observacao_justificativa" in prepared.tables[0].cell(0, 1).text

    fields = smart_fields_from_docx(output, candidate_field_definitions([choice, followup]))
    assert len([field for field in fields if field.get("selection") == "single"]) == 2
    assert any(field["type"] == "multiline" for field in fields)



def _append_unnamed_native_checkbox(paragraph, *, checked: bool) -> None:
    sdt = OxmlElement("w:sdt")
    props = OxmlElement("w:sdtPr")
    checkbox = OxmlElement("w14:checkbox")
    checked_el = OxmlElement("w14:checked")
    checked_el.set(qn("w14:val"), "1" if checked else "0")
    checkbox.append(checked_el)
    props.append(checkbox)
    sdt.append(props)

    content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "☒" if checked else "☐"
    run.append(text)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def test_adjacent_checkbox_group_keeps_second_unnamed_native_checked_control(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph(
        "Foi realizada a verificação, apresentando a seguinte ocorrência:"
    )
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "☐"
    table.cell(0, 1).text = "1. Existência de unidades em estoque"

    second_marker = table.cell(1, 0)
    second_marker.text = ""
    _append_unnamed_native_checkbox(second_marker.paragraphs[0], checked=True)
    table.cell(1, 1).text = "2. Unidade(s) do produto não registradas em estoque"
    source = _save(document, tmp_path / "native-second-checkbox.docx")

    candidates = detect_docx_field_candidates(source)
    choice = next(
        item
        for item in candidates
        if item.get("source") == "checkbox_choice"
        and item.get("location", {}).get("kind") == "checkbox_group_multi_cell"
    )

    assert len(choice["fields"]) == 2
    assert choice["fields"][1]["label"].startswith(
        "2. Unidade(s) do produto não registradas em estoque"
    )
    assert choice["location"]["checkbox_marker_modes"] == ["text_span", "paragraph"]

    output = tmp_path / "native-second-checkbox-prepared.docx"
    apply_docx_field_candidates(source, output, [choice])
    prepared = Document(str(output))
    assert "{{checkbox:" in prepared.tables[0].cell(0, 0).text
    assert "{{checkbox:" in prepared.tables[0].cell(1, 0).text


def test_adjacent_checkbox_group_accepts_bare_check_mark_for_selected_row(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("Apresentando as seguintes situações:")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "☐"
    table.cell(0, 1).text = "Primeira situação"
    table.cell(1, 0).text = "✓"
    table.cell(1, 1).text = "Segunda situação já marcada"
    source = _save(document, tmp_path / "checkmark-second-row.docx")

    candidates = detect_docx_field_candidates(source)
    choice = next(
        item
        for item in candidates
        if item.get("source") == "checkbox_choice"
        and item.get("location", {}).get("kind") == "checkbox_group_multi_cell"
    )
    assert [field["label"] for field in choice["fields"]] == [
        "Primeira situação",
        "Segunda situação já marcada",
    ]

    output = tmp_path / "checkmark-second-row-prepared.docx"
    apply_docx_field_candidates(source, output, [choice])
    prepared = Document(str(output))
    assert "{{checkbox:" in prepared.tables[0].cell(1, 0).text


def test_adjacent_checkbox_group_infers_blank_row_when_word_uses_floating_checked_box(tmp_path: Path) -> None:
    """Mirror a real Word pattern where only one table cell owns its square.

    The first row stores an empty checkbox as DrawingML+VML AlternateContent.
    The second row's narrow marker cell is genuinely blank because Word places
    the visible checked square in an absolutely-positioned text box elsewhere.
    Once the table establishes column 0 as the marker column, the blank second
    row must still be included in the same choice group.
    """
    from lxml import etree

    document = Document()
    document.add_paragraph(
        "Foi realizada a verificação, apresentando a seguinte ocorrência:"
    )
    table = document.add_table(rows=2, cols=2)
    from docx.shared import Inches
    for row in table.rows:
        row.cells[0].width = Inches(0.55)
        row.cells[1].width = Inches(4.5)

    first_marker = table.cell(0, 0)
    first_marker.text = ""
    paragraph = first_marker.paragraphs[0]
    run = paragraph.add_run()
    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    alt = etree.Element(f"{{{mc_ns}}}AlternateContent", nsmap={"mc": mc_ns})
    choice = etree.SubElement(alt, f"{{{mc_ns}}}Choice")
    drawing = OxmlElement("w:drawing")
    choice.append(drawing)
    fallback = etree.SubElement(alt, f"{{{mc_ns}}}Fallback")
    pict = OxmlElement("w:pict")
    fallback.append(pict)
    run._r.append(alt)

    table.cell(0, 1).text = "1. Existência de unidades em estoque"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = "2. Unidade(s) do produto não registradas em estoque"
    source = _save(document, tmp_path / "floating-second-checkbox.docx")

    candidates = detect_docx_field_candidates(source)
    choice_candidate = next(
        item
        for item in candidates
        if item.get("source") == "checkbox_choice"
        and item.get("location", {}).get("kind") == "checkbox_group_multi_cell"
    )

    assert len(choice_candidate["fields"]) == 2
    assert choice_candidate["location"]["checkbox_marker_modes"] == [
        "paragraph",
        "inferred_blank",
    ]
    assert choice_candidate["location"]["inferred_blank_markers"] == 1

    output = tmp_path / "floating-second-checkbox-prepared.docx"
    apply_docx_field_candidates(source, output, [choice_candidate])
    prepared = Document(str(output))
    assert "{{checkbox:" in prepared.tables[0].cell(0, 0).text
    assert "{{checkbox:" in prepared.tables[0].cell(1, 0).text


def test_four_cell_form_grid_avoids_duplicate_label_fields_and_keeps_date_pairs(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=8, cols=4)

    section_one = table.cell(0, 0).merge(table.cell(0, 3))
    section_one.text = "1. Dados do solicitante"
    table.cell(1, 0).text = "Nome do servidor:"
    table.cell(1, 1).text = "XXXXXXXXXXXXXXXXXXXXXXXX"
    table.cell(1, 2).text = "Matrícula:"
    table.cell(1, 3).text = "__________"
    table.cell(2, 0).text = "E-mail:"
    table.cell(2, 1).text = "servidor@orgao.gov.br"
    table.cell(2, 2).text = "Telefone:"
    table.cell(2, 3).text = "(83) 99999-9999"
    table.cell(3, 0).text = "Unidade:"
    table.cell(3, 1).merge(table.cell(3, 3)).text = "Diretoria Administrativa"

    section_two = table.cell(4, 0).merge(table.cell(4, 3))
    section_two.text = "2. Dados da viagem"
    table.cell(5, 0).text = "Destino:"
    table.cell(5, 1).text = "________________________"
    table.cell(5, 2).text = "UF:"
    table.cell(5, 3).text = "__"
    table.cell(6, 0).text = "Data de saída:"
    table.cell(6, 1).text = "__/__/____"
    table.cell(6, 2).text = "Data de retorno:"
    table.cell(6, 3).text = "__/__/____"
    table.cell(7, 0).text = "Meio de transporte:"
    table.cell(7, 1).merge(table.cell(7, 3)).text = "Escolher um item."

    source = _save(document, tmp_path / "four-cell-grid.docx")
    candidates = detect_docx_field_candidates(source)

    labels = [item.get("label") for item in candidates]
    for label in (
        "Nome do servidor",
        "Matrícula",
        "E-mail",
        "Telefone",
        "Unidade",
        "Destino",
        "UF",
        "Data de saída",
        "Data de retorno",
    ):
        assert labels.count(label) == 1

    by_label = {item.get("label"): item for item in candidates}
    assert by_label["E-mail"]["source"] == "sample_value"
    assert by_label["E-mail"]["placeholder"] == "servidor@orgao.gov.br"
    assert by_label["Telefone"]["type"] == "phone"
    assert by_label["Unidade"]["placeholder"] == "Diretoria Administrativa"
    assert by_label["Data de saída"]["type"] == "date"
    assert by_label["Data de retorno"]["type"] == "date"

    accepted = [item for item in candidates if item.get("selected")]
    output = tmp_path / "four-cell-grid-prepared.docx"
    apply_docx_field_candidates(source, output, accepted)
    fields = smart_fields_from_docx(output, candidate_field_definitions(accepted))
    fields_by_label = {field.get("label"): field for field in fields}

    assert fields_by_label["Nome do servidor"]["layout_row"] == fields_by_label["Matrícula"]["layout_row"]
    assert fields_by_label["Nome do servidor"]["layout_column_index"] == 0
    assert fields_by_label["Matrícula"]["layout_column_index"] == 2
    assert fields_by_label["E-mail"]["layout_row"] == fields_by_label["Telefone"]["layout_row"]
    assert fields_by_label["Data de saída"]["layout_row"] == fields_by_label["Data de retorno"]["layout_row"]
    assert fields_by_label["Data de saída"]["layout_column_index"] == 0
    assert fields_by_label["Data de retorno"]["layout_column_index"] == 2


def test_detects_single_declaration_checkbox_inside_form_row(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=4)
    section = table.cell(0, 0).merge(table.cell(0, 3))
    section.text = "5. Ciência e autorização"
    declaration = table.cell(1, 0).merge(table.cell(1, 3))
    declaration.text = "Declaro que as informações acima são verdadeiras.  ☐ Li e concordo"
    source = _save(document, tmp_path / "single-declaration-checkbox.docx")

    candidates = detect_docx_field_candidates(source)
    candidate = next(item for item in candidates if item.get("source") == "checkbox_single")

    assert candidate["type"] == "checkbox"
    assert "Declaro que as informações acima são verdadeiras" in candidate["label"]
    assert "Li e concordo" in candidate["label"]
    assert candidate["selected"] is True

    output = tmp_path / "single-declaration-checkbox-prepared.docx"
    apply_docx_field_candidates(source, output, [candidate])
    prepared = Document(str(output))
    assert "{{checkbox:" in prepared.tables[0].cell(1, 0).text
    assert "Li e concordo" in prepared.tables[0].cell(1, 0).text


def test_inline_placeholders_after_manual_line_breaks_keep_offsets_and_local_labels(tmp_path: Path) -> None:
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("Placa: ABC1D23")
    paragraph.add_run().add_break()
    paragraph.add_run("Data: __/__/____")
    paragraph.add_run().add_break()
    paragraph.add_run("Horário: __:__")
    source = _save(document, tmp_path / "pdf-like-lines.docx")

    candidates = detect_docx_field_candidates(source)
    selected = [item for item in candidates if item.get("source") == "inline_placeholder"]
    by_label = {item.get("label"): item for item in selected}

    assert "Data" in by_label
    assert "Horário" in by_label
    assert by_label["Data"]["type"] == "date"

    output = tmp_path / "pdf-like-lines-prepared.docx"
    apply_docx_field_candidates(source, output, selected)
    prepared_text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)

    assert "{{date:" in prepared_text
    assert "{{" in prepared_text and "horario" in prepared_text.casefold()


def test_colonless_adjacent_examples_become_fields_and_keep_grid_pairs(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("1. Identificação")
    table = document.add_table(rows=2, cols=4)
    table.cell(0, 0).text = "E-mail"
    table.cell(0, 1).text = "servidor@orgao.gov.br"
    table.cell(0, 2).text = "Telefone"
    table.cell(0, 3).text = "(83) 99999-9999"
    table.cell(1, 0).text = "País"
    table.cell(1, 1).text = "Brasil"
    table.cell(1, 2).text = "Endereço completo"
    table.cell(1, 3).text = ""
    source = tmp_path / "colonless-pairs.docx"
    document.save(source)

    candidates = detect_docx_field_candidates(source)
    by_label = {str(item.get("label", "")): item for item in candidates}

    assert by_label["E-mail"]["placeholder"] == "servidor@orgao.gov.br"
    assert by_label["Telefone"]["placeholder"] == "(83) 99999-9999"
    assert by_label["País"]["placeholder"] == "Brasil"
    assert by_label["E-mail"]["selected"] is True
    assert by_label["Telefone"]["selected"] is True
    assert by_label["País"]["selected"] is True

    accepted = [item for item in candidates if item.get("selected")]
    prepared = tmp_path / "colonless-pairs-prepared.docx"
    apply_docx_field_candidates(source, prepared, accepted)
    fields = smart_fields_from_docx(prepared, candidate_field_definitions(accepted))
    mapped = {str(field.get("label", "")): field for field in fields}

    assert mapped["E-mail"]["layout_column_index"] == 0
    assert mapped["E-mail"]["layout_column_span"] == 2
    assert mapped["Telefone"]["layout_column_index"] == 2
    assert mapped["Telefone"]["layout_column_span"] == 2
    assert mapped["País"]["layout_column_index"] == 0
    assert mapped["País"]["layout_column_span"] == 2
    assert mapped["Endereço completo"]["layout_column_index"] == 2
    assert mapped["Endereço completo"]["layout_column_span"] == 2


def test_repeatable_table_owns_header_and_model_rows_without_duplicate_header_fields(tmp_path: Path) -> None:
    """Regression: a repeatable table must not also become header/value fields."""

    document = Document()
    document.add_paragraph("3. Itens / serviços solicitados")
    table = document.add_table(rows=5, cols=5)
    for index, header in enumerate(("Nº", "Descrição", "Unidade", "Quantidade", "Valor estimado")):
        table.cell(0, index).text = header
    for row_index, number in enumerate(("01", "02", "03", "04"), start=1):
        table.cell(row_index, 0).text = number
        table.cell(row_index, 1).text = "________________________"
        table.cell(row_index, 2).text = "UND"
        table.cell(row_index, 3).text = "____"
        table.cell(row_index, 4).text = "R$ ________"

    source = _save(document, tmp_path / "repeatable-owner.docx")
    candidates = detect_docx_field_candidates(source)
    repeatable = next(item for item in candidates if item.get("source") == "repeatable_table")

    assert repeatable["region_owner"] == "repeatable_table"
    assert repeatable["location"]["owned_rows"] == [0, 1, 2, 3, 4]
    assert len(repeatable["location"]["owned_paragraphs"]) == 25
    assert repeatable["location"]["paragraphs"] == repeatable["location"]["owned_paragraphs"]

    # The header row is evidence for the table schema, not another form row.
    forbidden_labels = {"Nº", "Descrição", "Unidade", "Quantidade", "Valor estimado"}
    assert not any(
        item is not repeatable and str(item.get("label", "")).strip() in forbidden_labels
        for item in candidates
    )

    output = tmp_path / "repeatable-owner-prepared.docx"
    apply_docx_field_candidates(source, output, [repeatable])
    fields = smart_fields_from_docx(output, candidate_field_definitions([repeatable]))
    section_fields = [
        field
        for field in fields
        if str(field.get("section", "")).startswith("3.")
    ]
    assert [field["id"] for field in section_fields] == [repeatable["field_id"]]
    assert section_fields[0]["type"] == "repeatable_table"


def test_detects_existing_written_justification_as_prefilled_multiline(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "1. Justificativa da necessidade da Contratação:"
    prose = (
        "Em atendimento ao art. 4º do Decreto nº 44.639/23, a presente demanda "
        "visa elaborar o Plano de Contratação Anual, reunindo as informações "
        "necessárias para o planejamento da contratação."
    )
    table.cell(1, 0).text = prose
    source = _save(document, tmp_path / "prefilled-justification.docx")

    candidates = detect_docx_field_candidates(source)
    written = [item for item in candidates if item.get("source") == "prefilled_text"]

    assert len(written) == 1
    candidate = written[0]
    assert candidate["label"] == "Justificativa da necessidade da Contratação"
    assert candidate["type"] == "multiline"
    assert candidate["default_value"] == prose
    assert candidate["selected"] is True

    output = tmp_path / "prefilled-justification-prepared.docx"
    apply_docx_field_candidates(source, output, [candidate])
    fields = smart_fields_from_docx(output, candidate_field_definitions([candidate]))

    assert len(fields) == 1
    assert fields[0]["type"] == "multiline"
    assert fields[0]["default_value"] == prose
    assert fields[0]["full_width"] is True


def test_detects_full_width_written_table_response_with_existing_value(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=5)
    title = table.cell(0, 0)
    title.merge(table.cell(0, 4))
    title.text = "2. Quantidade a ser contratada, considerada a expectativa de consumo anual"

    for index, value in enumerate(
        ["Item", "Quantidade", "Unidade de medida", "Especificação/Descrição", "Valor"]
    ):
        table.cell(1, index).text = value

    response = table.cell(2, 0)
    response.merge(table.cell(2, 4))
    prose = (
        "Informamos que encontra-se em anexo o demonstrativo do Plano de "
        "Contratação Anual, contendo todas as informações necessárias para "
        "suprimento das informações exigidas neste item."
    )
    response.text = prose
    source = _save(document, tmp_path / "prefilled-merged-response.docx")

    candidates = detect_docx_field_candidates(source)
    written = [item for item in candidates if item.get("source") == "prefilled_text"]

    assert len(written) == 1
    assert written[0]["label"] == "Quantidade a ser contratada, considerada a expectativa de consumo anual"
    assert written[0]["default_value"] == prose
    assert written[0]["type"] == "multiline"


def test_prefilled_written_text_does_not_convert_explicit_fixed_note(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=1)
    table.cell(0, 0).text = "3. Justificativa complementar:"
    table.cell(1, 0).text = (
        "Texto fixo: este conteúdo é apenas uma orientação institucional e não "
        "deve ser convertido em campo editável pelo detector automático."
    )
    source = _save(document, tmp_path / "fixed-note.docx")

    candidates = detect_docx_field_candidates(source)

    assert all(item.get("source") != "prefilled_text" for item in candidates)


def test_spreadsheet_header_creates_editable_rows_and_keeps_merged_note(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=5)
    table.cell(0, 0).merge(table.cell(0, 4))
    table.cell(0, 0).text = "2. Quantidade a ser contratada, considerada a expectativa de consumo anual"
    for index, label in enumerate(
        (
            "Item",
            "Quantidade",
            "Unidade de medida",
            "Especificação/Descrição (Material/Equipamento/Serviço)",
            "Valor",
        )
    ):
        table.cell(1, index).text = label
    table.cell(2, 0).merge(table.cell(2, 4))
    original = (
        "Informamos que encontra-se em anexo, o demonstrativo do Plano de Contratação "
        "Anual PCA 2025, contendo todas as informações necessárias, para suprimento das "
        "informações exigidas neste item."
    )
    table.cell(2, 0).text = original
    source = _save(document, tmp_path / "sheet-editable.docx")

    candidates = detect_docx_field_candidates(source)
    repeatable = [item for item in candidates if item.get("source") == "repeatable_table"]
    prefilled = [item for item in candidates if item.get("source") == "prefilled_text"]

    assert len(repeatable) == 1
    assert len(prefilled) == 1
    sheet = repeatable[0]
    assert sheet["location"]["synthetic_template_row"] is True
    assert [column["label"] for column in sheet["columns"]] == [
        "Item",
        "Quantidade",
        "Unidade de medida",
        "Especificação/Descrição (Material/Equipamento/Serviço)",
        "Valor",
    ]
    assert [column["type"] for column in sheet["columns"]] == [
        "text",
        "integer",
        "text",
        "multiline",
        "currency",
    ]
    assert prefilled[0]["default_value"] == original

    output = tmp_path / "sheet-editable-prepared.docx"
    apply_docx_field_candidates(source, output, [sheet, prefilled[0]])
    prepared = Document(output)
    assert len(prepared.tables[0].rows) == 4
    model_row_text = " ".join(cell.text for cell in prepared.tables[0].rows[2].cells)
    assert "{{repeat:" in model_row_text
    assert ".item}}" in model_row_text
    assert ".quantidade}}" in model_row_text
    assert ".unidade_de_medida}}" in model_row_text
    assert ".valor}}" in model_row_text

    fields = smart_fields_from_docx(
        output,
        candidate_field_definitions([sheet, prefilled[0]]),
    )
    repeat_field = next(field for field in fields if field["type"] == "repeatable_table")
    note_field = next(field for field in fields if field["type"] == "multiline")
    assert [column["type"] for column in repeat_field["columns"]] == [
        "text",
        "integer",
        "text",
        "multiline",
        "currency",
    ]
    assert note_field["default_value"] == original

    generated = tmp_path / "sheet-editable-generated.docx"
    generate_docx(
        output,
        generated,
        {
            repeat_field["id"]: [
                {
                    "item": "Notebook",
                    "quantidade": "2",
                    "unidade_de_medida": "UN",
                    "especificacao_descricao_material_equipamento_servico": "Notebook 16 GB",
                    "valor": "R$ 5.000,00",
                }
            ],
            note_field["id"]: "Texto complementar atualizado.",
        },
    )
    result = Document(generated)
    assert len(result.tables[0].rows) == 4
    generated_row = result.tables[0].rows[2]
    assert [cell.text.strip() for cell in generated_row.cells] == [
        "Notebook",
        "2",
        "UN",
        "Notebook 16 GB",
        "R$ 5.000,00",
    ]
    assert result.tables[0].rows[3].cells[0].text.strip() == "Texto complementar atualizado."



def test_repeatable_child_types_do_not_inherit_parent_quantity_keyword(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Unidade de medida"
    table.cell(0, 2).text = "Quantidade"
    parent = "auto.quantidade_a_ser_contratada"
    table.cell(1, 0).text = f"{{{{repeat:{parent}}}}} {{{{{parent}.item}}}}"
    table.cell(1, 1).text = f"{{{{{parent}.unidade_de_medida}}}}"
    table.cell(1, 2).text = f"{{{{{parent}.quantidade}}}}"
    source = _save(document, tmp_path / "repeat-parent-quantity.docx")

    fields = smart_fields_from_docx(source)
    repeatable = next(field for field in fields if field["type"] == "repeatable_table")
    types = {column["id"]: column["type"] for column in repeatable["columns"]}

    assert types["item"] == "text"
    assert types["unidade_de_medida"] == "text"
    assert types["quantidade"] == "integer"


def test_detects_legacy_single_brace_placeholder_with_adjacent_label(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    header = table.cell(0, 0).merge(table.cell(0, 1))
    header.text = "1.1 Identificação da Demanda:"
    table.cell(1, 0).text = "Descrição da demanda:"
    table.cell(1, 1).text = "{descrição.demanda}"
    source = _save(document, tmp_path / "legacy-braces.docx")

    candidates = detect_docx_field_candidates(source)
    candidate = next(item for item in candidates if item["source"] == "legacy_placeholder")

    assert candidate["label"] == "Descrição da demanda"
    assert candidate["type"] == "multiline"
    assert candidate["field_id"] == "auto.descricao.demanda"
    assert candidate["section"] == "1.1 Identificação da Demanda"
    assert candidate["selected"] is True
    assert candidate["legacy_marker"] == "{descrição.demanda}"

    prepared = tmp_path / "legacy-braces-prepared.docx"
    apply_docx_field_candidates(source, prepared, [candidate])
    fields = smart_fields_from_docx(prepared, candidate_field_definitions([candidate]))

    assert len(fields) == 1
    assert fields[0]["id"] == "auto.descricao.demanda"
    assert fields[0]["label"] == "Descrição da demanda"
    assert fields[0]["type"] == "multiline"


def test_does_not_treat_ordinary_single_brace_prose_as_field(tmp_path: Path) -> None:
    document = Document()
    document.add_paragraph("Use chaves {assim} apenas para ilustrar a sintaxe.")
    source = _save(document, tmp_path / "literal-braces.docx")

    candidates = detect_docx_field_candidates(source)

    assert all(item.get("source") != "legacy_placeholder" for item in candidates)


def test_explicit_adjacent_tag_owns_label_cell_and_suppresses_empty_cell_candidate(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "1.1 Identificação da Demanda"
    table.cell(1, 0).text = "Descrição da demanda:"
    table.cell(1, 1).text = "{{descrição.demanda}}"
    source = _save(document, tmp_path / "explicit-owns-adjacent-cell.docx")

    candidates = detect_docx_field_candidates(source)

    assert all(
        item.get("field_id") != "auto.descricao_da_demanda"
        for item in candidates
    )
    assert all(
        str(item.get("label", "")).casefold() != "descrição da demanda".casefold()
        for item in candidates
    )
