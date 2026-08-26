from __future__ import annotations

import json
from copy import deepcopy
from pathlib import Path

import pytest
from docx import Document

from app.document.detection.application import (
    AutomaticDetectionError,
    apply_docx_field_candidates,
)
from app.document.detection.candidates import candidate_field_definitions
from app.document.detection.detector import (
    clear_detection_cache,
    detect_docx_field_candidates,
    detect_docx_with_report,
)
from app.document.detection.records import _collect_paragraph_records
from app.document.detection.roles import ContentRole, classify_record_role
from app.document.detection.structure import (
    SCANNER_STRUCTURE_VERSION,
    extract_document_structure,
)
from app.document.docx.scanner import clear_docx_scan_cache, scan_docx_fields
from app.document.understanding.smart_template import smart_fields_from_docx


FIXTURES = Path(__file__).resolve().parent / "fixtures"
DFD = FIXTURES / "dfd_licitacao_tradicional_sia13tdr.docx"
EXPECTED = FIXTURES / "dfd_licitacao_tradicional_sia13tdr.expected.json"
MANUAL = FIXTURES / "dfd_licitacao_tradicional_sia13tdr_manual_tags.docx"


def _matches(candidate: dict, spec: dict) -> bool:
    for key in ("source", "type", "label"):
        if key in spec and str(candidate.get(key, "")) != str(spec[key]):
            return False
    prefix = str(spec.get("section_prefix", ""))
    if prefix and not str(candidate.get("section", "")).startswith(prefix):
        return False
    return True


def test_real_dfd_matches_committed_structure_contract() -> None:
    expected = json.loads(EXPECTED.read_text(encoding="utf-8"))
    document = Document(str(DFD))
    records = _collect_paragraph_records(document)
    structure = extract_document_structure(document, records)
    candidates, report = detect_docx_with_report(DFD)

    assert structure.version == expected["scanner_structure_version"] == SCANNER_STRUCTURE_VERSION
    actual_sections = [section.full_title for section in structure.sections]
    for section in expected["required_sections"]:
        assert section in actual_sections

    for spec in expected["required_tables"]:
        matches = []
        for table in structure.tables:
            header_labels = [str(value).upper() for value in table.structure.header_labels]
            title_prefix = str(spec.get("title_prefix", ""))
            required_headers = [str(value).upper() for value in spec.get("header_labels", [])]
            if title_prefix and not table.section.startswith(title_prefix):
                continue
            if required_headers and not all(value in header_labels for value in required_headers):
                continue
            matches.append(table)
        assert matches, spec
        table = matches[0]
        assert table.kind == spec["kind"]
        assert table.structure.total_columns == spec["columns"]

    for spec in expected["required_candidates"]:
        assert any(_matches(candidate, spec) for candidate in candidates), spec

    for spec in expected["forbidden_candidate_sources_by_section_prefix"]:
        assert not any(_matches(candidate, spec) for candidate in candidates), spec

    assert report.scanner_version == SCANNER_STRUCTURE_VERSION
    assert report.candidate_count == len(candidates)
    assert report.ignored_ambiguous_tables == 0


def test_section_four_terminal_prompt_is_owned_typed_and_explained() -> None:
    clear_detection_cache()
    candidates = detect_docx_field_candidates(DFD)
    prompt = next(candidate for candidate in candidates if candidate.get("source") == "terminal_prompt")

    assert prompt["section"].startswith("4. Previsão de data")
    assert prompt["type"] == "date"
    assert prompt["location"]["kind"] == "append_tag"
    assert prompt["scanner_version"] == SCANNER_STRUCTURE_VERSION
    assert set(prompt["confidence_dimensions"]) == {"structure", "fillable", "label", "type"}
    assert float(prompt["confidence_dimensions"]["type"]) >= 0.85
    evidence = " ".join(str(value) for value in prompt.get("evidence", [])).casefold()
    inference = " ".join(
        str(value) for value in (prompt.get("type_inference", {}) or {}).get("reasons", [])
    ).casefold()
    assert "prompt" in evidence or "instru" in evidence
    assert "data" in inference or "previs" in inference


def test_structure_parser_does_not_promote_instruction_list_items_to_sections() -> None:
    document = Document(str(DFD))
    records = _collect_paragraph_records(document)
    structure = extract_document_structure(document, records)
    sections = [section.full_title for section in structure.sections]

    assert not any(section.startswith("2) ") for section in sections)
    assert not any("Não se aplica" in section for section in sections)
    assert not any("Nos termos do art." in section for section in sections)

    section_four_records = [
        record for record in records
        if (owner := structure.owner_for(record.ordinal)) is not None
        and owner.section.startswith("4. Previsão")
    ]
    roles = {record.text.strip(): classify_record_role(record, records, structure) for record in section_four_records}
    assert any(role is ContentRole.FIELD_PROMPT for role in roles.values())
    assert any(role in {ContentRole.INSTRUCTION, ContentRole.NOTE} for role in roles.values())


def test_manually_tagged_repeatable_table_is_authoritative_but_other_sections_still_scan() -> None:
    clear_detection_cache()
    clear_docx_scan_cache()
    candidates, report = detect_docx_with_report(MANUAL)

    assert not any(candidate.get("source") == "repeatable_table" for candidate in candidates)
    assert any(
        candidate.get("source") == "terminal_prompt"
        and str(candidate.get("section", "")).startswith("4. Previsão")
        for candidate in candidates
    )
    assert report.protected_tables >= 1

    fields = scan_docx_fields(MANUAL)
    table = next(field for field in fields if field.get("type") == "repeatable_table")
    assert table["id"] == "itens"
    assert table["section"] == "3. Quantidade a ser contratada"
    assert [column["id"] for column in table["columns"]][3:7] == [
        "quantidade_2023",
        "quantidade_2024",
        "quantidade_2025",
        "quantidade_solicitada",
    ]


def test_reference_history_table_is_owned_and_never_flattened() -> None:
    document = Document(str(DFD))
    records = _collect_paragraph_records(document)
    structure = extract_document_structure(document, records)
    reference = next(table for table in structure.tables if table.kind == "reference")
    assert reference.structure.header_labels[:3] == ["DATA", "VERSÃO", "DESCRIÇÃO"]

    candidates = detect_docx_field_candidates(DFD)
    reference_ordinals = set(reference.record_ordinals)
    assert not any(
        int(candidate.get("location", {}).get("paragraph", -1)) in reference_ordinals
        for candidate in candidates
    )


def test_transactional_apply_rejects_duplicate_repeat_columns_without_publishing(tmp_path: Path) -> None:
    candidates = detect_docx_field_candidates(DFD)
    repeatable = deepcopy(next(item for item in candidates if item.get("source") == "repeatable_table"))
    editable = [column for column in repeatable["columns"] if column.get("type") != "auto_number"]
    assert len(editable) >= 2
    editable[1]["id"] = editable[0]["id"]

    destination = tmp_path / "existing.docx"
    destination.write_bytes(b"ORIGINAL DESTINATION")

    with pytest.raises(AutomaticDetectionError, match="ida-e-volta|coluna|reescaneada"):
        apply_docx_field_candidates(DFD, destination, [repeatable])

    assert destination.read_bytes() == b"ORIGINAL DESTINATION"


def test_every_candidate_has_stable_owner_and_scanner_metadata() -> None:
    candidates = detect_docx_field_candidates(DFD)
    ids = [str(candidate.get("field_id", "")) for candidate in candidates]
    assert len(ids) == len(set(ids))
    for candidate in candidates:
        assert candidate.get("scanner_version") == SCANNER_STRUCTURE_VERSION
        assert str(candidate.get("section", "")).strip()
        dimensions = candidate.get("confidence_dimensions")
        assert isinstance(dimensions, dict)
        assert all(0.0 <= float(value) <= 1.0 for value in dimensions.values())


def test_reviewed_section_and_confidence_metadata_survive_candidate_conversion() -> None:
    candidate = next(
        item for item in detect_docx_field_candidates(DFD)
        if item.get("source") == "terminal_prompt"
    )
    reviewed = deepcopy(candidate)
    reviewed["section"] = "4. Entrega revisada manualmente"
    reviewed["section_source"] = "manual_review"
    reviewed["reviewed_by_user"] = True

    field = candidate_field_definitions([reviewed])[0]
    assert field["section"] == "4. Entrega revisada manualmente"
    assert field["section_source"] == "manual_review"
    assert field["detection_reviewed"] is True
    assert field["detection_confidence_dimensions"] == reviewed["confidence_dimensions"]
    assert field["scanner_version"] == SCANNER_STRUCTURE_VERSION


def test_real_dfd_full_detection_apply_and_strict_rescan_roundtrip(tmp_path: Path) -> None:
    """The real regression fixture must survive the complete automatic-tagging boundary.

    This is intentionally stronger than unit-level detector assertions: it detects all
    current suggestions, writes them to a temporary DOCX, then asks the strict normal
    tag scanner to read the result back. Any writer/scanner disagreement must fail here
    before it reaches a user's template editor work copy.
    """

    clear_detection_cache()
    clear_docx_scan_cache()
    candidates = detect_docx_field_candidates(DFD)
    destination = tmp_path / "dfd-scanner-v4-roundtrip.docx"

    apply_docx_field_candidates(DFD, destination, candidates)
    clear_docx_scan_cache()
    fields = scan_docx_fields(destination)
    by_id = {str(field.get("id", "")): field for field in fields}

    repeatable = next(item for item in candidates if item.get("source") == "repeatable_table")
    table_id = str(repeatable["field_id"])
    assert table_id in by_id
    assert [column["id"] for column in by_id[table_id]["columns"]][3:7] == [
        "quantidade_2023",
        "quantidade_2024",
        "quantidade_2025",
        "quantidade_solicitada",
    ]

    prompt = next(item for item in candidates if item.get("source") == "terminal_prompt")
    prompt_id = str(prompt["field_id"])
    assert by_id[prompt_id]["type"] == "date"

    # The strict tag scanner proves syntax/type compatibility. Semantic section
    # ownership is restored by the normal smart-template/context layer when the
    # tagged document is loaded back into the editor.
    smart_fields = smart_fields_from_docx(destination, [])
    smart_by_id = {str(field.get("id", "")): field for field in smart_fields}
    assert str(smart_by_id[prompt_id].get("section", "")).startswith("4. Previsão")

    candidate_ids = {
        str(item.get("field_id", ""))
        for item in candidates
        if str(item.get("field_id", ""))
    }
    assert candidate_ids.issubset(by_id)
