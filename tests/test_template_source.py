from pathlib import Path

import fitz
from docx import Document

from app.document.source import (
    TemplateSourceError,
    prepare_template_source,
)


def test_docx_source_is_used_directly(tmp_path: Path) -> None:
    source = tmp_path / 'modelo.docx'
    Document().save(source)

    prepared = prepare_template_source(source, tmp_path / 'work')

    assert prepared.original_path == source.resolve()
    assert prepared.docx_path == source.resolve()
    assert prepared.converted_from_pdf is False
    assert prepared.warnings == ()


def test_pdf_source_is_reconstructed_to_docx(tmp_path: Path) -> None:
    source = tmp_path / 'modelo.pdf'
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 90), 'Nome: ____________________')
    page.insert_text((72, 120), 'Data: __/__/____')
    pdf.save(source)
    pdf.close()

    prepared = prepare_template_source(source, tmp_path / 'work')

    assert prepared.original_path == source.resolve()
    assert prepared.converted_from_pdf is True
    assert prepared.docx_path.suffix == '.docx'
    assert prepared.docx_path.exists()
    assert prepared.docx_path.stat().st_size > 0


def test_template_source_rejects_other_extensions(tmp_path: Path) -> None:
    source = tmp_path / 'modelo.txt'
    source.write_text('x', encoding='utf-8')

    try:
        prepare_template_source(source, tmp_path / 'work')
    except TemplateSourceError as exc:
        assert 'DOCX ou PDF' in str(exc)
    else:
        raise AssertionError('expected TemplateSourceError')



def test_pdf_acroForm_fields_become_regular_template_tags(tmp_path: Path) -> None:
    from reportlab.pdfgen import canvas
    from app.document.understanding.smart_template import smart_fields_from_docx

    source = tmp_path / 'formulario.pdf'
    pdf = canvas.Canvas(str(source))
    pdf.drawString(72, 760, 'Nome:')
    pdf.acroForm.textfield(name='pessoa_nome', x=115, y=748, width=180, height=20)
    pdf.drawString(72, 720, 'Tipo:')
    pdf.acroForm.choice(
        name='tipo',
        value='Selecione...',
        options=['Selecione...', 'Interno', 'Externo'],
        x=115,
        y=708,
        width=140,
        height=20,
    )
    pdf.save()

    prepared = prepare_template_source(source, tmp_path / 'work')
    fields = smart_fields_from_docx(prepared.docx_path, [])
    by_id = {field['id']: field for field in fields}

    assert prepared.native_pdf_fields == 2
    assert 'pessoa_nome' in by_id
    assert by_id['tipo']['type'] == 'dropdown'
    assert by_id['tipo']['options'] == ['Interno', 'Externo']


def test_reconstructed_pdf_candidates_can_be_applied_when_lines_share_one_paragraph(tmp_path: Path) -> None:
    from app.document.detection.application import apply_docx_field_candidates
    from app.document.detection.detector import detect_docx_field_candidates

    source = tmp_path / 'inspecao.pdf'
    pdf = fitz.open()
    page = pdf.new_page()
    # Closely stacked lines encourage the PDF->DOCX converter to reconstruct
    # them as a single paragraph with manual line breaks, matching real forms.
    page.insert_text((72, 90), 'Placa: ABC1D23')
    page.insert_text((72, 101), 'Data: __/__/____')
    page.insert_text((72, 112), 'Horário: __:__')
    pdf.save(source)
    pdf.close()

    prepared = prepare_template_source(source, tmp_path / 'work')
    candidates = detect_docx_field_candidates(prepared.docx_path)
    selected = [item for item in candidates if item.get('source') == 'inline_placeholder']

    assert any(item.get('label') == 'Data' for item in selected)
    assert any(item.get('label') == 'Horário' for item in selected)

    output = tmp_path / 'prepared-fields.docx'
    apply_docx_field_candidates(prepared.docx_path, output, selected)
    assert output.exists()


def test_native_pdf_fields_keep_visible_labels_and_editable_dates(tmp_path: Path) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from app.document.understanding.smart_template import smart_fields_from_docx

    source = tmp_path / 'auditorio.pdf'
    pdf = canvas.Canvas(str(source), pagesize=A4)
    form = pdf.acroForm

    # Section 2-like row: internal PDF names deliberately differ from the
    # labels printed for humans on the page.
    pdf.drawString(52, 620, 'Nome do evento:')
    form.textfield(name='evento_nome', x=137, y=608, width=405, height=17)
    pdf.drawString(52, 590, 'Data:')
    form.textfield(name='evento_data', x=87, y=578, width=90, height=17)
    pdf.drawString(202, 590, 'Horário inicial:')
    form.textfield(name='evento_hora_inicio', x=282, y=578, width=65, height=17)
    pdf.drawString(372, 590, 'Horário final:')
    form.textfield(name='evento_hora_fim', x=442, y=578, width=100, height=17)

    # Section 5-like controls.
    form.checkbox(name='aceite_termo', x=52, y=520, size=12)
    pdf.drawString(72, 523, 'Li e concordo com o termo acima.')
    pdf.drawString(52, 490, 'Responsável pela autorização:')
    form.textfield(name='autorizador_nome', x=197, y=478, width=220, height=17)
    pdf.drawString(432, 490, 'Data:')
    form.textfield(name='autorizacao_data', x=462, y=478, width=80, height=17)
    pdf.save()

    prepared = prepare_template_source(source, tmp_path / 'work')
    hints = {item['id']: item for item in prepared.native_pdf_field_hints}

    assert hints['evento_nome']['label'] == 'Nome do evento'
    assert hints['evento_nome']['layout'] == 'full_width'
    assert hints['evento_data']['label'] == 'Data'
    assert hints['evento_data']['type'] == 'date'
    assert hints['evento_data']['automatic'] is False
    assert hints['evento_hora_inicio']['label'] == 'Horário inicial'
    assert hints['evento_hora_fim']['label'] == 'Horário final'
    assert hints['aceite_termo']['label'] == 'Li e concordo com o termo acima.'
    assert hints['autorizador_nome']['label'] == 'Responsável pela autorização'
    assert hints['autorizacao_data']['label'] == 'Data'
    assert hints['autorizacao_data']['automatic'] is False

    # Seeding smart scanning with source hints mirrors the template editor's
    # behavior and guarantees that native dates remain user-editable.
    fields = smart_fields_from_docx(prepared.docx_path, list(prepared.native_pdf_field_hints))
    by_id = {field['id']: field for field in fields}
    assert by_id['evento_data']['automatic'] is False
    assert by_id['autorizacao_data']['automatic'] is False
    assert by_id['aceite_termo']['label'] == 'Li e concordo com o termo acima.'
    assert by_id['autorizador_nome']['label'] == 'Responsável pela autorização'


def test_native_pdf_radio_group_keeps_option_boundaries_and_question_label(tmp_path: Path) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from app.document.understanding.smart_template import smart_fields_from_docx
    from app.document.detection.detector import detect_docx_field_candidates

    source = tmp_path / 'radio-prioridade.pdf'
    pdf = canvas.Canvas(str(source), pagesize=A4)
    form = pdf.acroForm

    y = 610
    pdf.drawString(52, y, 'Prioridade:')
    options = [
        ('normal', 122, 'Normal'),
        ('alta', 204, 'Alta'),
        ('critica', 280, 'Crítica'),
    ]
    for value, x, label in options:
        form.radio(
            name='priority_group_internal',
            value=value,
            selected=(value == 'normal'),
            x=x,
            y=y - 4,
            buttonStyle='circle',
            size=10,
        )
        pdf.drawString(x + 16, y, label)

    pdf.drawString(370, y, 'Requer acesso externo:')
    form.checkbox(name='external_access_flag_x', x=500, y=y - 4, size=11)
    pdf.drawString(518, y, 'Sim')
    pdf.save()

    prepared = prepare_template_source(source, tmp_path / 'work')
    hints = {item['id']: item for item in prepared.native_pdf_field_hints}

    assert hints['priority_group_internal']['label'] == 'Prioridade'
    assert hints['priority_group_internal']['type'] == 'dropdown'
    assert hints['priority_group_internal']['options'] == ['Normal', 'Alta', 'Crítica']
    assert hints['external_access_flag_x']['label'] == 'Requer acesso externo'

    fields = smart_fields_from_docx(prepared.docx_path, list(prepared.native_pdf_field_hints))
    by_id = {field['id']: field for field in fields}
    priority = by_id['priority_group_internal']
    assert priority['type'] == 'dropdown'
    assert priority['options'] == ['Normal', 'Alta', 'Crítica']
    assert priority['layout'] == 'choice'
    assert priority['selection'] == 'single'
    assert priority['layout_group_label'] == 'Prioridade'

    candidates = detect_docx_field_candidates(
        prepared.docx_path,
        existing_field_ids=by_id.keys(),
        existing_fields=fields,
    )
    assert all(str(item.get('label', '')).casefold() != 'prioridade' for item in candidates)
    assert all(str(item.get('semantic_label_suggestion', '')).casefold() != 'requer acesso externo' for item in candidates)


def test_native_pdf_questionnaire_matrix_keeps_row_labels_and_observation_columns(tmp_path: Path) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from app.document.understanding.smart_template import smart_fields_from_docx
    from app.document.understanding.layout import layout_blocks

    source = tmp_path / 'matrix-native.pdf'
    pdf = canvas.Canvas(str(source), pagesize=A4)
    form = pdf.acroForm

    pdf.drawString(42, 720, '3. Avaliação preliminar')
    pdf.drawString(42, 696, 'Item verificado')
    pdf.drawString(230, 696, 'Situação')
    pdf.drawString(455, 696, 'Observação')

    pdf.setFont('Helvetica', 8)
    rows = [
        ('m1', 'Documentação disponível', 670),
        ('m2', 'Compatibilidade técnica', 642),
        ('m3', 'Acesso ao ambiente', 614),
    ]
    option_specs = [
        ('conforme', 230, 'Conforme'),
        ('parcial', 302, 'Parcial'),
        ('nao_conforme', 355, 'Não conforme'),
        ('na', 430, 'N/A'),
    ]
    for suffix, row_label, y in rows:
        pdf.drawString(42, y, row_label)
        for value, x, label in option_specs:
            form.radio(
                name=f'x020_{suffix}',
                value=value,
                selected=(value == 'conforme'),
                x=x,
                y=y - 4,
                buttonStyle='circle',
                size=9,
            )
            pdf.drawString(x + 13, y, label)
        form.textfield(
            name=f'x021_obs_{suffix}',
            x=485,
            y=y - 6,
            width=65,
            height=15,
        )
    pdf.save()

    prepared = prepare_template_source(source, tmp_path / 'work')
    hints = {item['id']: item for item in prepared.native_pdf_field_hints}

    matrix_group = hints['x020_m1']['layout_group']
    assert matrix_group.startswith('native_pdf_matrix_')
    assert hints['x020_m1']['label'] == 'Documentação disponível'
    assert hints['x020_m1']['layout'] == 'table'
    assert hints['x020_m1']['layout_row_label'] == 'Documentação disponível'
    assert hints['x020_m1']['layout_row_header_label'] == 'Item verificado'
    assert hints['x020_m1']['layout_column'] == 'Situação'
    assert hints['x020_m1']['selection'] == 'single'
    assert hints['x021_obs_m1']['label'] == 'Documentação disponível — Observação'
    assert hints['x021_obs_m1']['layout_group'] == matrix_group
    assert hints['x021_obs_m1']['layout_row'] == hints['x020_m1']['layout_row']
    assert hints['x021_obs_m1']['layout_column'] == 'Observação'

    fields = smart_fields_from_docx(prepared.docx_path, list(prepared.native_pdf_field_hints))
    by_id = {field['id']: field for field in fields}
    assert by_id['x020_m1']['layout'] == 'table'
    assert by_id['x020_m1']['label'] == 'Documentação disponível'
    assert by_id['x020_m1']['options'] == ['Conforme', 'Parcial', 'Não conforme', 'N/A']
    assert by_id['x021_obs_m1']['type'] == 'text'
    assert by_id['x021_obs_m1']['label'] == 'Documentação disponível — Observação'

    blocks = layout_blocks(fields)
    matrix = next(block for block in blocks if block.get('group') == matrix_group)
    assert matrix['type'] == 'table'
    assert {field['id'] for field in matrix['fields']} == {
        'x020_m1', 'x021_obs_m1',
        'x020_m2', 'x021_obs_m2',
        'x020_m3', 'x021_obs_m3',
    }


def test_repeated_native_pdf_radio_questions_without_column_headers_are_not_forced_into_matrix(tmp_path: Path) -> None:
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4

    source = tmp_path / 'plain-radio-questions.pdf'
    pdf = canvas.Canvas(str(source), pagesize=A4)
    form = pdf.acroForm
    for index, (question, y) in enumerate([
        ('Possui autorização?', 680),
        ('Possui anexo?', 640),
        ('Necessita revisão?', 600),
    ], start=1):
        pdf.drawString(42, y, question)
        form.radio(name=f'q_{index}', value='sim', x=230, y=y - 4, size=9, buttonStyle='circle')
        pdf.drawString(243, y, 'Sim')
        form.radio(name=f'q_{index}', value='nao', x=302, y=y - 4, size=9, buttonStyle='circle')
        pdf.drawString(315, y, 'Não')
    pdf.save()

    prepared = prepare_template_source(source, tmp_path / 'work')
    hints = {item['id']: item for item in prepared.native_pdf_field_hints}
    assert all(str(hints[f'q_{index}'].get('layout', '')).casefold() != 'table' for index in range(1, 4))


def _append_unnamed_dropdown(paragraph, options: list[str]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sdt = OxmlElement('w:sdt')
    props = OxmlElement('w:sdtPr')
    dropdown = OxmlElement('w:dropDownList')
    for option in options:
        item = OxmlElement('w:listItem')
        item.set(qn('w:displayText'), option)
        item.set(qn('w:value'), option)
        dropdown.append(item)
    props.append(dropdown)
    sdt.append(props)
    content = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = 'Selecione...'
    run.append(text)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def _append_unnamed_date(paragraph) -> None:
    from docx.oxml import OxmlElement

    sdt = OxmlElement('w:sdt')
    props = OxmlElement('w:sdtPr')
    props.append(OxmlElement('w:date'))
    sdt.append(props)
    content = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = 'Selecione uma data'
    run.append(text)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def _append_unnamed_checkbox(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sdt = OxmlElement('w:sdt')
    props = OxmlElement('w:sdtPr')
    checkbox = OxmlElement('w14:checkbox')
    checked = OxmlElement('w14:checked')
    checked.set(qn('w14:val'), '0')
    checkbox.append(checked)
    props.append(checkbox)
    sdt.append(props)
    content = OxmlElement('w:sdtContent')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = '☐'
    run.append(text)
    content.append(run)
    sdt.append(content)
    paragraph._p.append(sdt)


def test_docx_unnamed_word_controls_are_auto_tagged_from_context(tmp_path: Path) -> None:
    from docx.oxml.ns import qn
    from app.document.docx.generator import generate_docx
    from app.document.docx.scanner import scan_docx_fields
    from app.document.understanding.smart_template import smart_fields_from_docx

    source = tmp_path / 'controles-sem-tag.docx'
    document = Document()
    document.add_paragraph('2. Dados da demanda')

    priority = document.add_paragraph('Prioridade: ')
    _append_unnamed_dropdown(priority, ['Normal', 'Alta', 'Crítica'])

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = 'Data de início'
    table.cell(0, 1).text = ''
    _append_unnamed_date(table.cell(0, 1).paragraphs[0])

    access = document.add_paragraph('Requer acesso externo: ')
    _append_unnamed_checkbox(access)
    document.save(source)

    # Even the read-only scanner must no longer fail when Developer metadata
    # is missing; it resolves the same deterministic identities in memory.
    scanned_original = {item['id']: item for item in scan_docx_fields(source)}
    assert set(scanned_original) >= {'prioridade', 'data_de_inicio', 'requer_acesso_externo'}
    assert scanned_original['prioridade']['label'] == 'Prioridade'
    assert scanned_original['prioridade']['type'] == 'dropdown'

    prepared = prepare_template_source(source, tmp_path / 'work')
    assert prepared.prepared_work_copy is True
    assert prepared.docx_path != source.resolve()
    assert prepared.docx_path.exists()
    assert prepared.native_word_field_hints
    assert any('tags automaticamente' in warning for warning in prepared.warnings)

    # Original file is untouched; only the working copy receives Word tags.
    original_doc = Document(str(source))
    assert all(
        sdt.find(qn('w:sdtPr')).find(qn('w:tag')) is None
        for sdt in original_doc.element.iter(qn('w:sdt'))
        if sdt.find(qn('w:sdtPr')) is not None
    )

    prepared_doc = Document(str(prepared.docx_path))
    tags = {
        tag.get(qn('w:val'))
        for sdt in prepared_doc.element.iter(qn('w:sdt'))
        if (props := sdt.find(qn('w:sdtPr'))) is not None
        and (tag := props.find(qn('w:tag'))) is not None
    }
    assert {'prioridade', 'data_de_inicio', 'requer_acesso_externo'} <= tags

    fields = smart_fields_from_docx(
        prepared.docx_path,
        list(prepared.native_word_field_hints),
    )
    by_id = {field['id']: field for field in fields}
    assert by_id['prioridade']['label'] == 'Prioridade'
    assert by_id['prioridade']['options'] == ['Normal', 'Alta', 'Crítica']
    assert by_id['data_de_inicio']['type'] == 'date'
    assert by_id['requer_acesso_externo']['type'] == 'checkbox'
    assert by_id['prioridade']['context_resolver_version'] == 3

    # The generated ID is not metadata-only: the engine can fill the prepared
    # native controls using it.
    output = tmp_path / 'filled.docx'
    generate_docx(
        prepared.docx_path,
        output,
        {
            'prioridade': 'Alta',
            'data_de_inicio': '17/08/2026',
            'requer_acesso_externo': True,
        },
    )
    assert output.exists()


def test_unnamed_word_control_without_label_gets_safe_fallback_instead_of_error(tmp_path: Path) -> None:
    from app.document.docx.scanner import scan_docx_fields

    source = tmp_path / 'controle-isolado.docx'
    document = Document()
    paragraph = document.add_paragraph()
    _append_unnamed_dropdown(paragraph, ['A', 'B'])
    document.save(source)

    scanned = scan_docx_fields(source)
    assert len(scanned) == 1
    assert scanned[0]['id'].startswith('auto_word.dropdown_')
    assert scanned[0]['auto_tagged'] is True

    prepared = prepare_template_source(source, tmp_path / 'work')
    assert prepared.prepared_work_copy is True
    assert prepared.native_word_field_hints[0]['id'].startswith('auto_word.dropdown_')
    assert any('fallback' in warning for warning in prepared.warnings)


def _append_block_unnamed_dropdown(cell, options: list[str], display_text: str = 'Escolher um item.') -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    sdt = OxmlElement('w:sdt')
    props = OxmlElement('w:sdtPr')
    dropdown = OxmlElement('w:dropDownList')
    for option in options:
        item = OxmlElement('w:listItem')
        item.set(qn('w:displayText'), option)
        item.set(qn('w:value'), option)
        dropdown.append(item)
    props.append(dropdown)
    sdt.append(props)

    content = OxmlElement('w:sdtContent')
    paragraph = OxmlElement('w:p')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = display_text
    run.append(text)
    paragraph.append(run)
    content.append(paragraph)
    sdt.append(content)
    cell._tc.append(sdt)


def test_context_resolver_keeps_multiple_block_controls_and_groups_neighboring_choices(tmp_path: Path) -> None:
    """Regression for the real DFD pattern: block dropdowns + PCA checkboxes.

    XML element wrapper identities are not stable enough to use as cache keys;
    every control must keep its own structural context even when Word returns
    fresh wrappers during traversal.
    """
    from app.document.docx.scanner import scan_docx_fields

    source = tmp_path / 'dfd-like-controls.docx'
    document = Document()

    unit_table = document.add_table(rows=1, cols=1)
    _append_block_unnamed_dropdown(
        unit_table.cell(0, 0),
        ['Escolha a Unidade Gestora', '27.0001 - SEDH', '50.0001 - FEAS'],
    )

    modality = document.add_table(rows=2, cols=1)
    modality.cell(0, 0).text = '1.2. Modalidade de licitação Sugerida:'
    _append_block_unnamed_dropdown(
        modality.cell(1, 0),
        ['Escolha a opção', '1. PREGÃO', '2. CONCORRÊNCIA'],
    )

    pca = document.add_table(rows=2, cols=1)
    pca.cell(0, 0).text = '5. Plano de Contratações Anual - PCA 2025:'
    cell = pca.cell(1, 0)
    cell.text = 'Declaramos que todos os itens indicados neste documento de formalização da demanda:'
    alternatives = [
        'Consta(m) no Plano de Contratações Anual – PCA.',
        'Consta(m) parcialmente no Plano de Contratações Anual – PCA.',
        'Não consta(m) no Plano de Contratações Anual – PCA.',
    ]
    for alternative in alternatives:
        paragraph = cell.add_paragraph()
        _append_unnamed_checkbox(paragraph)
        paragraph.add_run('   ' + alternative)

    document.save(source)

    fields = scan_docx_fields(source)
    native = [field for field in fields if field.get('detection_source') == 'native_word']
    assert len(native) == 5
    by_id = {field['id']: field for field in native}

    assert 'auto_word.dropdown_01' not in by_id
    assert by_id['unidade_gestora']['label'] == 'Unidade Gestora'
    assert by_id['unidade_gestora']['options'] == ['27.0001 - SEDH', '50.0001 - FEAS']

    modality_field = by_id['modalidade_de_licitacao_sugerida']
    assert modality_field['label'] == 'Modalidade de licitação Sugerida'
    assert modality_field['section'].startswith('1.2.')
    assert modality_field['options'] == ['1. PREGÃO', '2. CONCORRÊNCIA']

    choice_fields = [field for field in native if field['type'] == 'checkbox']
    assert [field['label'] for field in choice_fields] == alternatives
    assert len({field['id'] for field in choice_fields}) == 3
    assert len({field.get('layout_group') for field in choice_fields}) == 1
    assert all(field.get('layout') == 'choice' for field in choice_fields)
    assert all(field.get('selection') == 'single' for field in choice_fields)
    assert all(field.get('section', '').startswith('5. Plano') for field in choice_fields)
