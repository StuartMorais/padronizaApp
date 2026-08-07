from pathlib import Path

from docx import Document

from app.layout_inference import apply_layout_metadata, infer_docx_layout, layout_blocks
from app.smart_template import smart_fields_from_docx


def _build_structured_template(path: Path) -> None:
    document = Document()

    choice_table = document.add_table(rows=2, cols=2)
    choice_table.cell(0, 0).merge(choice_table.cell(0, 1))
    choice_table.cell(0, 0).text = "13. Prazo de entrega/Execução do serviço:"
    choice_table.cell(1, 0).text = "{{checkbox:prazo.imediata}} Entrega imediata"
    choice_table.cell(1, 1).text = "{{checkbox:prazo.parcelada}} Entrega parcelada"

    people = document.add_table(rows=4, cols=5)
    people.cell(0, 0).merge(people.cell(0, 4))
    people.cell(0, 0).text = "14. Responsáveis pela fiscalização:"
    for column, label in enumerate(("Função", "Nome completo", "Matrícula", "Setor", "Ciência")):
        people.cell(1, column).text = label
    people.cell(2, 0).text = "Gestor do Contrato"
    people.cell(2, 1).text = "{{gestor.nome}}"
    people.cell(2, 2).text = "{{gestor.matricula}}"
    people.cell(2, 3).text = "{{gestor.setor}}"
    people.cell(2, 4).text = "{{gestor.ciencia}}"
    people.cell(3, 0).text = "Fiscal Técnico"
    people.cell(3, 1).text = "{{fiscal.nome}}"
    people.cell(3, 2).text = "{{fiscal.matricula}}"
    people.cell(3, 3).text = "{{fiscal.setor}}"
    people.cell(3, 4).text = "{{fiscal.ciencia}}"
    document.save(path)


def test_infers_choice_and_table_layout(tmp_path: Path) -> None:
    path = tmp_path / "structured.docx"
    _build_structured_template(path)

    metadata = infer_docx_layout(path)
    assert metadata["prazo.imediata"]["layout"] == "choice"
    assert metadata["prazo.parcelada"]["selection"] == "single"
    assert metadata["prazo.imediata"]["section"].startswith("13.")

    assert metadata["gestor.nome"]["layout"] == "table"
    assert metadata["gestor.nome"]["layout_row_label"] == "Gestor do Contrato"
    assert metadata["gestor.matricula"]["layout_column"] == "Matrícula"
    assert metadata["fiscal.setor"]["layout_row_label"] == "Fiscal Técnico"


def test_smart_fields_receive_layout_metadata(tmp_path: Path) -> None:
    path = tmp_path / "structured.docx"
    _build_structured_template(path)
    fields = smart_fields_from_docx(path)
    by_id = {field["id"]: field for field in fields}

    assert by_id["prazo.imediata"]["layout"] == "choice"
    assert by_id["prazo.imediata"]["choice_required"] is True
    assert by_id["gestor.nome"]["layout"] == "table"
    assert by_id["gestor.nome"]["section"].startswith("14.")

    blocks = layout_blocks([by_id["gestor.nome"], by_id["gestor.matricula"]])
    assert len(blocks) == 1
    assert blocks[0]["type"] == "table"


def test_manual_layout_metadata_is_preserved() -> None:
    fields = [
        {
            "id": "prazo.imediata",
            "label": "Prazo definido manualmente",
            "label_source": "manual",
            "section": "Seção manual",
            "layout": "full_width",
        }
    ]
    inferred = {
        "prazo.imediata": {
            "detected_label": "Entrega imediata",
            "section": "13. Prazo",
            "layout": "choice",
            "layout_group": "prazo",
        }
    }

    merged = apply_layout_metadata(fields, inferred)[0]
    assert merged["label"] == "Prazo definido manualmente"
    assert merged["section"] == "Seção manual"
    assert merged["layout"] == "full_width"
    assert merged["layout_group"] == "prazo"


def test_layout_blocks_keep_semantic_groups_in_document_order() -> None:
    fields = [
        {"id": "a", "type": "text", "layout": "grid"},
        {"id": "b", "type": "checkbox", "layout": "choice", "layout_group": "g"},
        {"id": "c", "type": "checkbox", "layout": "choice", "layout_group": "g"},
        {"id": "d", "type": "text", "layout": "full_width"},
    ]

    blocks = layout_blocks(fields)
    assert [block["type"] for block in blocks] == ["grid", "choice", "grid"]
    assert [field["id"] for field in blocks[1]["fields"]] == ["b", "c"]


def test_single_choice_tag_creates_one_checkbox_style_choice_field(tmp_path: Path) -> None:
    path = tmp_path / "single-choice.docx"
    document = Document()
    document.add_paragraph("9. Situação no PCA")
    document.add_paragraph(
        "{{single_choice:pca_2025.situacao|"
        "Consta no PCA => Consta(m) no Plano de Contratações Anual – PCA, conforme comprovação em anexo.|"
        "Não consta no PCA => Não consta(m) no Plano de Contratações Anual – PCA.}}"
    )
    document.save(path)

    fields = smart_fields_from_docx(path)
    assert len(fields) == 1
    field = fields[0]

    assert field["id"] == "pca_2025.situacao"
    assert field["type"] == "dropdown"
    assert field["layout"] == "choice"
    assert field["selection"] == "single"
    assert field["choice_required"] is True
    assert field["layout_group_label"] == "9. Situação no PCA"
    assert len(field["options"]) == 2

    blocks = layout_blocks(fields)
    assert len(blocks) == 1
    assert blocks[0]["type"] == "choice"
    assert [item["id"] for item in blocks[0]["fields"]] == ["pca_2025.situacao"]


def test_template_repository_preserves_single_choice_layout(tmp_path: Path) -> None:
    from app.template_repository import TemplateRepository

    repository = TemplateRepository(tmp_path / "templates")
    normalized = repository._normalize_fields(
        [
            {
                "id": "pca_2025.situacao",
                "label": "Situação no PCA",
                "type": "dropdown",
                "required": True,
                "options": ["Consta no PCA", "Não consta no PCA"],
                "layout": "choice",
                "layout_group": "single_choice_pca_2025.situacao",
                "layout_group_label": "Situação no PCA",
                "selection": "single",
                "choice_required": True,
                "tag_type": "single_choice",
            }
        ],
        strict=True,
    )

    field = normalized[0]
    assert field["layout"] == "choice"
    assert field["layout_group"] == "single_choice_pca_2025.situacao"
    assert field["choice_required"] is True
    assert field["tag_type"] == "single_choice"


def _build_form_grid_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("PREENCHIMENTO PELA ÁREA REQUISITANTE")

    table = document.add_table(rows=5, cols=3)
    table.cell(0, 0).merge(table.cell(0, 2))
    table.cell(0, 0).text = "1. Área Requisitante da Demanda:"

    table.cell(1, 0).merge(table.cell(1, 2))
    table.cell(1, 0).text = (
        "Órgão: Secretaria de Estado do Desenvolvimento Humano - SEDH"
    )

    table.cell(2, 0).merge(table.cell(2, 2))
    table.cell(2, 0).text = "Setor Requisitante (Unidade/Setor/Depto): {{setor.nome}}"

    table.cell(3, 0).merge(table.cell(3, 1))
    table.cell(3, 0).text = "Responsável pela Demanda: {{responsavel.nome}}"
    table.cell(3, 2).text = "Matrícula: {{responsavel.matricula}}"

    table.cell(4, 0).merge(table.cell(4, 1))
    table.cell(4, 0).text = "E-mail: {{requisitante.email}}"
    table.cell(4, 2).text = "Telefone: {{responsavel.numero}}"
    document.save(path)


def test_form_like_word_table_uses_form_grid_and_merged_spans(tmp_path: Path) -> None:
    path = tmp_path / "form-grid.docx"
    _build_form_grid_template(path)

    metadata = infer_docx_layout(path)

    assert metadata["setor.nome"]["layout"] == "form_grid"
    assert metadata["setor.nome"]["layout_column_index"] == 0
    assert metadata["setor.nome"]["layout_column_span"] == 3
    assert metadata["setor.nome"]["layout_grid_columns"] == 3
    assert metadata["setor.nome"]["detected_label"] == (
        "Setor Requisitante (Unidade/Setor/Depto)"
    )

    assert metadata["responsavel.nome"]["layout"] == "form_grid"
    assert metadata["responsavel.nome"]["layout_column_index"] == 0
    assert metadata["responsavel.nome"]["layout_column_span"] == 2
    assert metadata["responsavel.nome"]["detected_label"] == "Responsável pela Demanda"

    assert metadata["responsavel.matricula"]["layout_column_index"] == 2
    assert metadata["responsavel.matricula"]["layout_column_span"] == 1
    assert metadata["responsavel.matricula"]["detected_label"] == "Matrícula"
    assert metadata["requisitante.email"]["detected_label"] == "E-mail"
    assert metadata["responsavel.numero"]["detected_label"] == "Telefone"

    static_rows = metadata["setor.nome"]["layout_static_rows"]
    assert [row["text"] for row in static_rows] == [
        "Órgão: Secretaria de Estado do Desenvolvimento Humano - SEDH"
    ]


def test_form_grid_becomes_one_semantic_block(tmp_path: Path) -> None:
    path = tmp_path / "form-grid.docx"
    _build_form_grid_template(path)
    fields = smart_fields_from_docx(path)
    blocks = layout_blocks(fields)

    assert len(blocks) == 1
    assert blocks[0]["type"] == "form_grid"
    assert [field["id"] for field in blocks[0]["fields"]] == [
        "setor.nome",
        "responsavel.nome",
        "responsavel.matricula",
        "requisitante.email",
        "responsavel.numero",
    ]
    assert blocks[0]["static_rows"][0]["text"].startswith("Órgão:")


def test_old_automatic_table_layout_is_migrated_to_form_grid() -> None:
    fields = [
        {
            "id": "responsavel.nome",
            "label": "Responsável",
            "layout": "table",
            "layout_group": "doc_table_1",
            "layout_row": "row_3",
            "layout_row_label": "Órgão",
            "layout_column": "Órgão",
        }
    ]
    inferred = {
        "responsavel.nome": {
            "layout": "form_grid",
            "layout_group": "doc_table_1_segment_1",
            "layout_row": "row_3",
            "layout_column_index": 0,
            "layout_column_span": 2,
            "layout_grid_columns": 3,
        }
    }

    migrated = apply_layout_metadata(fields, inferred)[0]
    assert migrated["layout"] == "form_grid"
    assert migrated["layout_group"] == "doc_table_1_segment_1"
    # A single field with no same-row context is intentionally expanded to
    # the whole visual row by the defensive layout normalizer.
    assert migrated["layout_column_index"] == 0
    assert migrated["layout_column_span"] == 3
    assert migrated["full_width"] is True
    assert "layout_row_label" not in migrated
    assert "layout_column" not in migrated


def test_template_repository_preserves_form_grid_metadata(tmp_path: Path) -> None:
    from app.template_repository import TemplateRepository

    repository = TemplateRepository(tmp_path / "templates")
    static_rows = [
        {
            "layout_order": 1,
            "layout_column_index": 0,
            "layout_column_span": 3,
            "layout_grid_columns": 3,
            "text": "Órgão: Secretaria de exemplo",
        }
    ]
    normalized = repository._normalize_fields(
        [
            {
                "id": "setor.nome",
                "label": "Setor requisitante",
                "type": "text",
                "required": True,
                "layout": "form_grid",
                "layout_group": "doc_table_1_segment_1",
                "layout_row": "row_2",
                "layout_column_index": 0,
                "layout_column_span": 3,
                "layout_grid_columns": 3,
                "layout_order": 2,
                "layout_static_rows": static_rows,
            }
        ],
        strict=True,
    )

    field = normalized[0]
    assert field["layout"] == "form_grid"
    assert field["layout_column_span"] == 3
    assert field["layout_grid_columns"] == 3
    assert field["layout_static_rows"] == static_rows


def _build_single_orphan_field_template(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "2. Descrição da Demanda"
    table.cell(1, 0).text = ""
    table.cell(1, 1).text = "Descrição da demanda: {{demanda.descricao}}"
    document.save(path)


def test_single_field_in_partial_word_row_expands_to_full_width(tmp_path: Path) -> None:
    path = tmp_path / "single-orphan.docx"
    _build_single_orphan_field_template(path)

    fields = smart_fields_from_docx(path)
    field = next(item for item in fields if item["id"] == "demanda.descricao")

    assert field["type"] == "multiline"
    assert field["layout"] == "form_grid"
    assert field["layout_grid_columns"] == 2
    assert field["layout_column_index"] == 0
    assert field["layout_column_span"] == 2
    assert field["full_width"] is True

    blocks = layout_blocks(fields)
    assert blocks[0]["fields"][0]["layout_column_index"] == 0
    assert blocks[0]["fields"][0]["layout_column_span"] == 2


def test_same_row_static_content_preserves_partial_field_position(tmp_path: Path) -> None:
    path = tmp_path / "static-peer.docx"
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "2. Descrição da Demanda"
    table.cell(1, 0).text = "Orientação para preenchimento"
    table.cell(1, 1).text = "Descrição: {{demanda.descricao}}"
    document.save(path)

    fields = smart_fields_from_docx(path)
    field = next(item for item in fields if item["id"] == "demanda.descricao")

    assert field["layout_column_index"] == 1
    assert field["layout_column_span"] == 1
    assert field["layout_row_static_cells"][0]["text"] == "Orientação para preenchimento"

    blocks = layout_blocks(fields)
    assert blocks[0]["row_static_cells"][0]["layout_column_index"] == 0


def test_exact_partial_position_can_be_locked() -> None:
    from app.layout_inference import normalize_form_layout

    fields = normalize_form_layout(
        [
            {
                "id": "demanda.descricao",
                "type": "multiline",
                "layout": "form_grid",
                "layout_group": "g",
                "layout_row": "r",
                "layout_grid_columns": 2,
                "layout_column_index": 1,
                "layout_column_span": 1,
                "layout_position_locked": True,
            }
        ]
    )

    assert fields[0]["layout_column_index"] == 1
    assert fields[0]["layout_column_span"] == 1
    assert "full_width" not in fields[0]


def test_layout_quality_rejects_overlapping_form_grid_cells() -> None:
    from app.layout_inference import layout_quality_issues

    issues = layout_quality_issues(
        [
            {
                "id": "a",
                "layout": "form_grid",
                "layout_group": "g",
                "layout_row": "r",
                "layout_grid_columns": 3,
                "layout_column_index": 0,
                "layout_column_span": 2,
            },
            {
                "id": "b",
                "layout": "form_grid",
                "layout_group": "g",
                "layout_row": "r",
                "layout_grid_columns": 3,
                "layout_column_index": 1,
                "layout_column_span": 2,
            },
        ]
    )

    assert any("sobrepostas" in issue for issue in issues)


def test_normalize_absorbs_adjacent_label_cell_into_field() -> None:
    from app.layout_inference import normalize_form_layout

    fields = normalize_form_layout(
        [
            {
                "id": "auto.endereco_completo",
                "label": "Endereço completo",
                "type": "text",
                "layout": "form_grid",
                "layout_group": "g",
                "layout_row": "r",
                "layout_grid_columns": 2,
                "layout_column_index": 1,
                "layout_column_span": 1,
                "layout_row_static_cells": [
                    {
                        "layout_row": "r",
                        "layout_column_index": 0,
                        "layout_column_span": 1,
                        "layout_grid_columns": 2,
                        "text": "Endereço completo:",
                    }
                ],
            }
        ]
    )

    field = fields[0]
    assert field["layout_column_index"] == 0
    assert field["layout_column_span"] == 2
    assert field["full_width"] is True
    assert not field.get("layout_row_static_cells")


def test_normalize_removes_stale_dropdown_prompt_overlap() -> None:
    from app.layout_inference import layout_quality_issues, normalize_form_layout

    fields = normalize_form_layout(
        [
            {
                "id": "auto.tipo",
                "label": "Tipo",
                "type": "dropdown",
                "layout": "form_grid",
                "layout_group": "g",
                "layout_row": "r",
                "layout_grid_columns": 4,
                "layout_column_index": 3,
                "layout_column_span": 1,
                "layout_row_static_cells": [
                    {
                        "layout_row": "r",
                        "layout_column_index": 3,
                        "layout_column_span": 1,
                        "layout_grid_columns": 4,
                        "text": "Tipo: Escolher um item",
                    }
                ],
            }
        ]
    )

    assert not fields[0].get("layout_row_static_cells")
    assert layout_quality_issues(fields) == []


def test_data_table_preserves_original_row_header_title(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Item verificado"
    table.cell(0, 1).text = "Situação"
    table.cell(0, 2).text = "Observação curta"
    table.cell(1, 0).text = "Documentação disponível"
    table.cell(1, 1).text = "{{checkbox:item.conforme}} Conforme {{checkbox:item.nao}} Não conforme"
    table.cell(1, 2).text = "{{item.observacao}}"
    table.cell(2, 0).text = "Condições de segurança"
    table.cell(2, 1).text = "{{checkbox:seguranca.conforme}} Conforme {{checkbox:seguranca.nao}} Não conforme"
    table.cell(2, 2).text = "{{seguranca.observacao}}"
    path = tmp_path / "row-header.docx"
    document.save(path)

    metadata = infer_docx_layout(path)
    assert metadata["item.conforme"]["layout_row_header_label"] == "Item verificado"
    assert metadata["item.observacao"]["layout_row_header_label"] == "Item verificado"
