from __future__ import annotations

from pathlib import Path

from docx import Document

from app.document.detection.candidates import candidate_field_definitions
from app.document.detection.detector import clear_detection_cache, detect_docx_field_candidates
from app.document.understanding.semantic import postprocess_candidates
from app.domain.template_quality import field_configuration_issues, issue_counts
from app.domain.validation import sample_values_for_fields


def test_field_configuration_issues_are_row_specific_and_fast() -> None:
    fields = [
        {"id": "cliente.nome", "label": "Nome", "type": "text"},
        {"id": "cliente.nome", "label": "Nome repetido", "type": "text"},
        {"id": "tipo", "label": "Tipo", "type": "dropdown", "options": ["A"]},
        {
            "id": "detalhes",
            "label": "Detalhes",
            "type": "text",
            "visible_when": {"field": "campo.inexistente", "equals": "sim"},
        },
    ]

    issues = field_configuration_issues(fields)
    counts = issue_counts(issues)

    assert counts["error"] >= 4
    assert {issue.row for issue in issues if issue.code == "field.id_duplicate"} == {0, 1}
    assert any(issue.row == 2 and issue.code == "dropdown.options_missing" for issue in issues)
    assert any(issue.row == 3 and issue.code == "condition.field_unknown" for issue in issues)


def test_approved_assisted_field_does_not_keep_live_review_warning() -> None:
    candidate = {
        "field_id": "auto.nome",
        "label": "Nome",
        "type": "text",
        "confidence": 0.74,
        "confidence_band": "medium",
        "source": "inline_placeholder",
        "evidence": [],
        "review_reasons": [],
        "review_priority": "recommended",
        "needs_review": True,
        "reviewed_by_user": True,
    }
    field = candidate_field_definitions([candidate])[0]

    assert field["detection_reviewed"] is True
    assert not field_configuration_issues([field])


def test_unreviewed_medium_confidence_field_is_flagged() -> None:
    field = {
        "id": "auto.nome",
        "label": "Nome",
        "type": "text",
        "detection_source": "automatic",
        "detection_confidence": 0.74,
        "detection_confidence_band": "medium",
    }

    issues = field_configuration_issues([field])
    assert [issue.code for issue in issues] == ["detection.medium_confidence"]


def test_sample_values_respect_exclusive_checkbox_groups() -> None:
    fields = [
        {
            "id": "declaracao.sim",
            "label": "Sim",
            "type": "checkbox",
            "group": "declaracao",
            "selection": "single",
        },
        {
            "id": "declaracao.nao",
            "label": "Não",
            "type": "checkbox",
            "group": "declaracao",
            "selection": "single",
        },
        {"id": "data", "label": "Data", "type": "date"},
        {
            "id": "itens",
            "label": "Itens",
            "type": "repeatable_table",
            "columns": [
                {"id": "descricao", "label": "Descrição", "type": "text"},
                {"id": "quantidade", "label": "Quantidade", "type": "integer"},
            ],
        },
    ]

    values = sample_values_for_fields(fields)

    assert values["declaracao.sim"] is True
    assert values["declaracao.nao"] is False
    assert values["data"] == "19/08/2026"
    assert values["itens"][0]["quantidade"] == "10"


def test_detector_v3_exposes_review_priority() -> None:
    raw = {
        "field_id": "auto.campo_1",
        "label": "Campo 1",
        "type": "text",
        "confidence": 0.91,
        "source": "inline_placeholder",
        "preview": "________",
        "location": {"kind": "text_span", "paragraph": 0, "start": 0, "end": 8},
    }

    processed = postprocess_candidates([raw], [], source_kind="docx")
    assert processed[0]["detector_version"] == 3
    assert processed[0]["review_priority"] == "required"
    assert processed[0]["needs_review"] is True
    assert processed[0]["review_summary"]


def test_assisted_detection_cache_returns_independent_copies_and_invalidates(
    tmp_path: Path,
    monkeypatch,
) -> None:
    import app.document.detection.detector as detector_module
    from docx import Document as WordDocument

    path = tmp_path / "cache.docx"
    document = WordDocument()
    document.add_paragraph("Nome: XXXXXXXX")
    document.save(path)

    clear_detection_cache()
    calls = 0

    def counted_document(*args, **kwargs):
        nonlocal calls
        calls += 1
        return WordDocument(*args, **kwargs)

    monkeypatch.setattr(detector_module, "Document", counted_document)

    first = detect_docx_field_candidates(path)
    second = detect_docx_field_candidates(path)
    assert calls == 1
    assert first == second
    if first:
        first[0]["label"] = "MODIFICADO"
        third = detect_docx_field_candidates(path)
        assert third[0]["label"] != "MODIFICADO"
        assert calls == 1

    changed = WordDocument()
    changed.add_paragraph("Nome: XXXXXXXX")
    changed.add_paragraph("CPF: ___.___.___-__")
    changed.save(path)

    detect_docx_field_candidates(path)
    assert calls == 2


def test_assisted_detection_supports_cooperative_cancel(tmp_path: Path) -> None:
    import pytest

    from app.document.detection.models import AutomaticDetectionCancelled

    path = tmp_path / "cancel.docx"
    document = Document()
    document.add_paragraph("Nome: XXXXXXXX")
    document.save(path)

    clear_detection_cache()
    with pytest.raises(AutomaticDetectionCancelled):
        detect_docx_field_candidates(path, cancel_check=lambda: True)


def test_preflight_rejects_dropdown_with_only_one_configured_choice(tmp_path: Path) -> None:
    from app.document.diagnostics import diagnose_template

    path = tmp_path / "dropdown.docx"
    document = Document()
    document.add_paragraph("Tipo: {{dropdown:tipo|A|B}}")
    document.save(path)

    report = diagnose_template(
        {
            "fields": [
                {"id": "tipo", "label": "Tipo", "type": "dropdown", "options": ["A"]}
            ],
            "output": {"filename_pattern": "{{template.name}}.docx"},
        },
        path,
    )

    assert report["blocking"] is True
    assert any(issue["code"] == "dropdown.no_options" for issue in report["issues"])


def test_live_quality_detects_duplicate_dropdown_options_and_table_columns() -> None:
    fields = [
        {
            "id": "tipo",
            "label": "Tipo",
            "type": "dropdown",
            "options": ["A", "A", "B"],
        },
        {
            "id": "itens",
            "label": "Itens",
            "type": "repeatable_table",
            "columns": [
                {"id": "descricao", "label": "Descrição", "type": "text"},
                {"id": "descricao", "label": "Descrição 2", "type": "text"},
            ],
        },
    ]

    issues = field_configuration_issues(fields)
    codes = {(issue.row, issue.code) for issue in issues}

    assert (0, "dropdown.options_duplicate") in codes
    assert (1, "table.columns_duplicate") in codes


def test_live_quality_detects_visibility_cycles() -> None:
    fields = [
        {
            "id": "a",
            "label": "A",
            "type": "text",
            "visible_when": {"field": "b", "equals": "sim"},
        },
        {
            "id": "b",
            "label": "B",
            "type": "text",
            "visible_when": {"field": "a", "equals": "sim"},
        },
    ]

    issues = field_configuration_issues(fields)
    cycle_rows = {issue.row for issue in issues if issue.code == "condition.cycle"}

    assert cycle_rows == {0, 1}


def test_strong_label_disagreement_becomes_review_evidence() -> None:
    from types import SimpleNamespace

    from app.document.understanding.semantic import RecordSemantics

    record = SimpleNamespace(
        ordinal=0,
        understanding=RecordSemantics(
            label="Razão social",
            label_source="same_text",
            label_confidence=0.96,
        ),
    )
    raw = {
        "field_id": "empresa.endereco",
        "label": "Endereço",
        "type": "text",
        "confidence": 0.90,
        "source": "inline_placeholder",
        "preview": "XXXXXXXX",
        "location": {"kind": "text_span", "paragraph": 0, "start": 0, "end": 8},
    }

    processed = postprocess_candidates([raw], [record], source_kind="docx")
    candidate = processed[0]

    assert any(item["code"] == "label_disagreement" for item in candidate["evidence"])
    assert candidate["review_priority"] == "recommended"
    assert candidate["needs_review"] is True


def test_full_diagnostics_preserve_duplicate_configuration_evidence(tmp_path: Path) -> None:
    from app.document.diagnostics import diagnose_template

    path = tmp_path / "duplicates.docx"
    document = Document()
    document.add_paragraph("Tipo: {{dropdown:tipo|A|B}}")
    document.add_table(rows=2, cols=1).cell(0, 0).text = "{{table:itens}}"
    document.save(path)

    report = diagnose_template(
        {
            "fields": [
                {
                    "id": "tipo",
                    "label": "Tipo",
                    "type": "dropdown",
                    "options": ["A", "A", "B"],
                },
                {
                    "id": "itens",
                    "label": "Itens",
                    "type": "repeatable_table",
                    "columns": [
                        {"id": "descricao", "label": "Descrição", "type": "text"},
                        {"id": "descricao", "label": "Descrição duplicada", "type": "text"},
                    ],
                },
            ],
            "output": {"filename_pattern": "{{template.name}}.docx"},
        },
        path,
    )

    codes = {issue["code"] for issue in report["issues"]}
    assert "dropdown.duplicate_options" in codes
    assert "table.duplicate_columns" in codes


def test_live_quality_rejects_malformed_visibility_rule_representation() -> None:
    issues = field_configuration_issues(
        [
            {
                "id": "detalhes",
                "label": "Detalhes",
                "type": "text",
                "visible_when": "regra incompleta",
            }
        ]
    )

    assert any(issue.code == "condition.invalid" for issue in issues)
