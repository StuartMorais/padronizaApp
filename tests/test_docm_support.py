from __future__ import annotations

import json
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

from app.document.conversion.service import DocumentConverter
from app.document.docx.scanner import scan_docx_fields
from app.document.source import SUPPORTED_TEMPLATE_SUFFIXES, prepare_template_source
from app.core.schema import TEMPLATE_SCHEMA_VERSION
from app.repositories.templates import TemplateRepository
from app.document.word_package import (
    DOCX_MAIN_CONTENT_TYPE,
    normalize_word_input,
)


_CONTENT_TYPES_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
_RELS_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _make_docm(path: Path, text: str = "Nome: {{pessoa.nome}}") -> Path:
    base = path.with_suffix(".base.docx")
    document = Document()
    document.add_paragraph(text)
    document.save(base)

    with zipfile.ZipFile(base, "r") as source, zipfile.ZipFile(
        path, "w", compression=zipfile.ZIP_DEFLATED
    ) as target:
        for info in source.infolist():
            data = source.read(info.filename)
            if info.filename == "[Content_Types].xml":
                root = ET.fromstring(data)
                for child in root:
                    if child.attrib.get("PartName") == "/word/document.xml":
                        child.set(
                            "ContentType",
                            "application/vnd.ms-word.document.macroEnabled.main+xml",
                        )
                ET.SubElement(
                    root,
                    f"{{{_CONTENT_TYPES_NS}}}Override",
                    {
                        "PartName": "/word/vbaProject.bin",
                        "ContentType": "application/vnd.ms-office.vbaProject",
                    },
                )
                ET.register_namespace("", _CONTENT_TYPES_NS)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            elif info.filename == "word/_rels/document.xml.rels":
                root = ET.fromstring(data)
                ET.SubElement(
                    root,
                    f"{{{_RELS_NS}}}Relationship",
                    {
                        "Id": "rIdPadronizaMacroTest",
                        "Type": "http://schemas.microsoft.com/office/2006/relationships/vbaProject",
                        "Target": "vbaProject.bin",
                    },
                )
                ET.register_namespace("", _RELS_NS)
                data = ET.tostring(root, encoding="utf-8", xml_declaration=True)
            target.writestr(info, data)
        target.writestr("word/vbaProject.bin", b"not-a-real-vba-project")

    base.unlink()
    return path


def test_docm_is_a_supported_template_input() -> None:
    assert ".docm" in SUPPORTED_TEMPLATE_SUFFIXES


def test_docm_is_normalized_to_macro_free_docx(tmp_path: Path) -> None:
    source = _make_docm(tmp_path / "macro.docm")
    destination = tmp_path / "safe.docx"

    result = normalize_word_input(source, destination)

    assert result.macros_removed is True
    assert source.exists()
    assert destination.exists()

    with zipfile.ZipFile(destination, "r") as archive:
        names = {name.casefold() for name in archive.namelist()}
        assert not any("vbaproject" in name or "vbadata" in name for name in names)
        content_types = ET.fromstring(archive.read("[Content_Types].xml"))
        document_override = next(
            child
            for child in content_types
            if child.attrib.get("PartName") == "/word/document.xml"
        )
        assert document_override.attrib["ContentType"] == DOCX_MAIN_CONTENT_TYPE
        relationships = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "vbaProject" not in relationships

    reopened = Document(destination)
    assert "pessoa.nome" in "\n".join(paragraph.text for paragraph in reopened.paragraphs)


def test_docm_template_source_becomes_persistent_docx_work_copy(tmp_path: Path) -> None:
    source = _make_docm(tmp_path / "modelo.docm")

    prepared = prepare_template_source(source, tmp_path / "work")

    assert prepared.original_path == source.resolve()
    assert prepared.converted_from_pdf is False
    assert prepared.converted_from_docm is True
    assert prepared.docx_path.suffix.casefold() == ".docx"
    assert prepared.docx_path.exists()
    assert any("macros VBA" in warning for warning in prepared.warnings)
    fields = scan_docx_fields(prepared.docx_path)
    assert any(field["id"] == "pessoa.nome" for field in fields)


class _InspectingBackend:
    name = "fake"
    description = "fake"
    priority = 1

    def __init__(self) -> None:
        self.source: Path | None = None

    def is_available(self) -> bool:
        return True

    def convert(self, source: Path, destination: Path, warnings: list[str]) -> Path:
        self.source = source
        assert source.suffix.casefold() == ".docx"
        with zipfile.ZipFile(source, "r") as archive:
            assert all("vbaproject" not in name.casefold() for name in archive.namelist())
        destination.write_bytes(b"%PDF-1.4\n")
        return destination


def test_docm_to_pdf_never_hands_macros_to_backend(tmp_path: Path) -> None:
    source = _make_docm(tmp_path / "modelo.docm")
    destination = tmp_path / "modelo.pdf"
    backend = _InspectingBackend()
    converter = DocumentConverter([backend])
    warnings: list[str] = []

    converter.docx_to_pdf(source, destination, warnings=warnings)

    assert destination.exists()
    assert backend.source is not None
    assert backend.source != source
    assert any("macros VBA" in warning for warning in warnings)


def test_docm_inside_template_package_is_normalized_on_import(tmp_path: Path) -> None:
    source = _make_docm(tmp_path / "template.docm")
    config = {
        "schema_version": TEMPLATE_SCHEMA_VERSION,
        "template": {
            "id": "macro-template",
            "name": "Macro Template",
            "version": "1.0",
            "source_file": "template.docm",
            "adapter": "docx",
        },
        "fields": [{"id": "pessoa.nome", "type": "text", "label": "Nome"}],
        "sections": [{"title": "Dados", "fields": ["pessoa.nome"]}],
        "output": {"filename_pattern": "{{template.name}}.docx"},
    }
    package = tmp_path / "macro-template.zip"
    with zipfile.ZipFile(package, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.write(source, "template.docm")
        archive.writestr("template.json", json.dumps(config))

    repository = TemplateRepository(tmp_path / "templates")
    template_id = repository.import_template_package(package)
    imported_source = repository.get_source_path(template_id)

    assert imported_source.name == "template.docx"
    assert imported_source.exists()
    with zipfile.ZipFile(imported_source, "r") as archive:
        assert all("vbaproject" not in name.casefold() for name in archive.namelist())
    assert repository.read_config(template_id)["template"]["source_file"] == "template.docx"
