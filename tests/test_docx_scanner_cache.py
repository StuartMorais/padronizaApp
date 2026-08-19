from __future__ import annotations

from docx import Document

from app.document.docx.scanner import clear_docx_scan_cache, scan_docx_fields


def test_scan_cache_returns_isolated_field_objects(tmp_path):
    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{company.name}}")
    doc.save(path)

    clear_docx_scan_cache()
    first = scan_docx_fields(path)
    first[0]["label"] = "mutated"
    second = scan_docx_fields(path)

    assert second[0]["id"] == "company.name"
    assert second[0].get("label") != "mutated"


def test_scan_cache_invalidates_when_docx_changes(tmp_path):
    path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("{{a}}")
    doc.save(path)
    clear_docx_scan_cache()
    assert [field["id"] for field in scan_docx_fields(path)] == ["a"]

    doc = Document()
    doc.add_paragraph("{{b}} {{c}}")
    doc.save(path)
    assert [field["id"] for field in scan_docx_fields(path)] == ["b", "c"]
