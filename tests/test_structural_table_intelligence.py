from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document

from app.document.detection.application import apply_docx_field_candidates
from app.document.detection.candidates import candidate_field_definitions
from app.document.detection.detector import clear_detection_cache, detect_docx_field_candidates
from app.document.detection.table_structure import TableKind, analyze_word_table
from app.document.docx.generator import generate_docx
from app.document.docx.repair import repair_repeatable_table_markers
from app.document.docx.scanner import clear_docx_scan_cache, scan_docx_fields
from app.domain.validation import sample_values_for_fields
from app.document.understanding.smart_template import smart_fields_from_docx


FIXTURE = Path(__file__).parent / "fixtures" / "dfd_licitacao_tradicional_sia13tdr.docx"


def test_real_dfd_quantity_grid_is_repeatable_not_flattened() -> None:
    document = Document(str(FIXTURE))

    quantity = analyze_word_table(document.tables[7])
    assert quantity.kind is TableKind.REPEATABLE
    assert quantity.title == "3. Quantidade a ser contratada"
    assert quantity.header_row == 1
    assert quantity.data_rows == [2]
    assert quantity.continuation_rows == [3]
    assert quantity.header_labels == [
        "Item",
        "Descrição",
        "UND",
        "Quantidade — 2023",
        "Quantidade — 2024",
        "Quantidade — 2025",
        "Quantidade Solicitada",
        "Consta no PCA para 2026?",
        "Justificativa se for o caso",
    ]
    assert quantity.header_options[7] == [
        {"label": "SIM", "value": "SIM"},
        {"label": "NÃO", "value": "NÃO"},
    ]

    # Other real tables in the same document must not be reinterpreted as an
    # editable sheet just because a later row happens to precede a merged note.
    assert analyze_word_table(document.tables[11]).kind is TableKind.FIXED_FORM
    assert analyze_word_table(document.tables[14]).kind is TableKind.REFERENCE


def test_real_dfd_detector_emits_one_owned_table_candidate_and_no_sim_field(tmp_path: Path) -> None:
    clear_detection_cache()
    candidates = detect_docx_field_candidates(FIXTURE)
    repeatables = [item for item in candidates if item.get("source") == "repeatable_table"]

    assert len(repeatables) == 1
    table = repeatables[0]
    assert table["location"]["document_table_index"] == 7
    assert table["region_owner"] == "repeatable_table"
    assert table["location"]["owned_rows"] == [0, 1, 2, 3]
    assert table["location"]["data_rows"] == [2, 3]
    assert not any(str(item.get("label", "")).strip() == "SIM" for item in candidates)

    assert [(column["id"], column["type"]) for column in table["columns"]] == [
        ("item", "auto_number"),
        ("descricao", "multiline"),
        ("und", "text"),
        ("quantidade_2023", "integer"),
        ("quantidade_2024", "integer"),
        ("quantidade_2025", "integer"),
        ("quantidade_solicitada", "integer"),
        ("consta_no_pca_para_2026", "dropdown"),
        ("justificativa_se_for_o_caso", "multiline"),
    ]
    pca = table["columns"][7]
    assert [option["value"] for option in pca["options"]] == ["SIM", "NÃO"]
    assert table["columns"][8]["required"] is False

    prepared = tmp_path / "dfd-prepared.docx"
    apply_docx_field_candidates(FIXTURE, prepared, [table])
    prepared_document = Document(str(prepared))
    prepared_table = prepared_document.tables[7]

    # The visual table remains a table: title + header + one tagged model row.
    assert len(prepared_table.rows) == 3
    assert "{{repeat:auto.quantidade_a_ser_contratada}}" in prepared_table.cell(2, 0).text
    assert "{{auto.quantidade_a_ser_contratada.descricao}}" in prepared_table.cell(2, 1).text
    assert "{{auto.quantidade_a_ser_contratada.quantidade_2023}}" in prepared_table.cell(2, 3).text
    assert "{{dropdown:auto.quantidade_a_ser_contratada.consta_no_pca_para_2026|SIM|NÃO}}" in prepared_table.cell(2, 7).text

    fields = smart_fields_from_docx(prepared, candidate_field_definitions([table]))
    repeat_field = next(field for field in fields if field["type"] == "repeatable_table")
    assert repeat_field["id"] == "auto.quantidade_a_ser_contratada"
    assert len(repeat_field["columns"]) == 9
    assert [column["label"] for column in repeat_field["columns"]][3:6] == [
        "Quantidade — 2023",
        "Quantidade — 2024",
        "Quantidade — 2025",
    ]
    assert [column.get("group_label") for column in repeat_field["columns"]][3:6] == [
        "Quantidade",
        "Quantidade",
        "Quantidade",
    ]


def test_structural_quantity_table_generates_real_rows_end_to_end(tmp_path: Path) -> None:
    document = Document()
    table = document.add_table(rows=4, cols=9)
    table.cell(0, 0).merge(table.cell(0, 8)).text = "3. Quantidade a ser contratada:"
    table.cell(1, 0).text = "Item"
    table.cell(1, 1).text = "Descrição"
    table.cell(1, 2).text = "UND"
    quantity_header = table.cell(1, 3).merge(table.cell(1, 5))
    quantity_header.text = "Quantidade"
    quantity_header.add_paragraph("2023      2024      2025")
    table.cell(1, 6).text = "Quantidade Solicitada"
    table.cell(1, 7).text = "Consta no PCA para 2026?\nSIM / NÃO"
    table.cell(1, 8).text = "Justificativa\nse for o caso"
    table.cell(2, 0).text = "01"
    table.cell(2, 7).text = "SIM"
    table.cell(3, 0).text = "..."
    source = tmp_path / "quantity-structure.docx"
    document.save(source)

    candidates = detect_docx_field_candidates(source)
    repeatable = next(item for item in candidates if item.get("source") == "repeatable_table")
    prepared = tmp_path / "quantity-prepared.docx"
    apply_docx_field_candidates(source, prepared, [repeatable])
    fields = smart_fields_from_docx(prepared, candidate_field_definitions([repeatable]))
    field = next(item for item in fields if item["type"] == "repeatable_table")

    generated = tmp_path / "quantity-generated.docx"
    generate_docx(
        prepared,
        generated,
        {
            field["id"]: [
                {
                    "descricao": "Notebook",
                    "und": "UN",
                    "quantidade_2023": "1",
                    "quantidade_2024": "2",
                    "quantidade_2025": "3",
                    "quantidade_solicitada": "4",
                    "consta_no_pca_para_2026": "NÃO",
                    "justificativa_se_for_o_caso": "Aquisição nova",
                },
                {
                    "descricao": "Monitor",
                    "und": "UN",
                    "quantidade_2023": "0",
                    "quantidade_2024": "1",
                    "quantidade_2025": "1",
                    "quantidade_solicitada": "2",
                    "consta_no_pca_para_2026": "SIM",
                    "justificativa_se_for_o_caso": "",
                },
            ]
        },
    )

    result = Document(str(generated))
    assert len(result.tables[0].rows) == 4
    assert [cell.text.strip() for cell in result.tables[0].rows[2].cells] == [
        "01", "Notebook", "UN", "1", "2", "3", "4", "NÃO", "Aquisição nova"
    ]
    assert [cell.text.strip() for cell in result.tables[0].rows[3].cells] == [
        "02", "Monitor", "UN", "0", "1", "1", "2", "SIM", ""
    ]


def test_real_dfd_generation_preserves_header_background_assets(tmp_path: Path) -> None:
    candidates = detect_docx_field_candidates(FIXTURE)
    repeatable = next(item for item in candidates if item.get("source") == "repeatable_table")
    prepared = tmp_path / "dfd-prepared-with-background.docx"
    apply_docx_field_candidates(FIXTURE, prepared, [repeatable])

    fields = smart_fields_from_docx(
        prepared,
        candidate_field_definitions([repeatable]),
    )
    generated = tmp_path / "dfd-generated-with-background.docx"
    generate_docx(
        prepared,
        generated,
        sample_values_for_fields(fields),
    )

    def package_visual_assets(path: Path) -> tuple[set[str], int]:
        with zipfile.ZipFile(path) as archive:
            media = {
                name.rsplit("/", 1)[-1]
                for name in archive.namelist()
                if name.startswith("word/media/")
            }
            header = archive.read("word/header1.xml").decode("utf-8", errors="ignore")
            return media, header.count("<w:drawing")

    source_media, source_drawings = package_visual_assets(prepared)
    generated_media, generated_drawings = package_visual_assets(generated)

    assert generated_media == source_media
    assert source_drawings >= 1
    assert generated_drawings == source_drawings


def test_malformed_repeatable_work_copy_is_repaired_from_physical_headers(tmp_path: Path) -> None:
    broken = tmp_path / "broken-context.docx"
    document = Document(str(FIXTURE))
    row = document.tables[7].rows[2]
    bad_markers = [
        "{{repeat:itens}} {{row.number}}",
        "{{itens.descricao}}",
        "{{itens.unidade}}",
        "{{itens.quantidade}}",
        "{{itens.quantidade}}",
        "{{itens.quantidade}}",
        "{{item.quatidade}}",
        "{{itens.conta}}",
        "{{itens.justificativa}}",
    ]
    for cell, marker in zip(row.cells, bad_markers, strict=True):
        cell.text = marker
    document.save(str(broken))

    clear_docx_scan_cache()
    try:
        scan_docx_fields(broken)
    except ValueError as exc:
        assert "coluna 'quantidade' aparece mais de uma vez" in str(exc)
    else:
        raise AssertionError("O DOCX malformado deveria ser rejeitado antes do reparo.")

    result = repair_repeatable_table_markers(broken)
    assert result.changed is True
    assert result.marker_count == 4
    assert [(repair.column_index, repair.new_marker) for repair in result.repairs] == [
        (3, "itens.quantidade_2023"),
        (4, "itens.quantidade_2024"),
        (5, "itens.quantidade_2025"),
        (6, "itens.quantidade_solicitada"),
    ]

    clear_docx_scan_cache()
    fields = smart_fields_from_docx(broken, [])
    repeatable = next(field for field in fields if field["type"] == "repeatable_table")
    assert repeatable["id"] == "itens"
    assert repeatable["section"] == "3. Quantidade a ser contratada"
    assert repeatable["section_source"] == "word_table_title"
    assert [column["id"] for column in repeatable["columns"]] == [
        "item",
        "descricao",
        "unidade",
        "quantidade_2023",
        "quantidade_2024",
        "quantidade_2025",
        "quantidade_solicitada",
        "conta",
        "justificativa",
    ]

    repaired_document = Document(str(broken))
    assert [cell.text for cell in repaired_document.tables[7].rows[2].cells] == [
        "{{repeat:itens}} {{row.number}}",
        "{{itens.descricao}}",
        "{{itens.unidade}}",
        "{{itens.quantidade_2023}}",
        "{{itens.quantidade_2024}}",
        "{{itens.quantidade_2025}}",
        "{{itens.quantidade_solicitada}}",
        "{{itens.conta}}",
        "{{itens.justificativa}}",
    ]

    # The repair is a migration, not an ongoing mutation. Once normalized, a
    # second smart scan leaves the work copy untouched.
    second = repair_repeatable_table_markers(broken)
    assert second.changed is False
    assert second.marker_count == 0


def test_manual_repeatable_markers_keep_numbered_word_table_section(tmp_path: Path) -> None:
    source = tmp_path / "manual-repeatable-section.docx"
    document = Document()
    table = document.add_table(rows=3, cols=4)
    table.cell(0, 0).merge(table.cell(0, 3)).text = "3. Quantidade a ser contratada:"
    table.cell(1, 0).text = "Item"
    table.cell(1, 1).text = "Descrição"
    table.cell(1, 2).text = "UND"
    table.cell(1, 3).text = "Quantidade"
    table.cell(2, 0).text = "{{repeat:itens}} {{row.number}}"
    table.cell(2, 1).text = "{{itens.descricao}}"
    table.cell(2, 2).text = "{{itens.unidade}}"
    table.cell(2, 3).text = "{{itens.quantidade}}"
    document.save(str(source))

    clear_docx_scan_cache()
    fields = smart_fields_from_docx(source, [])
    repeatable = next(field for field in fields if field["type"] == "repeatable_table")

    assert repeatable["id"] == "itens"
    assert repeatable["section"] == "3. Quantidade a ser contratada"
    assert repeatable["section_source"] == "word_table_title"

    # Templates created with an older Padroniza version may already have the
    # generic fallback saved. A stronger physical Word section should migrate
    # that fallback automatically instead of keeping the table under
    # ``Dados do documento`` forever.
    migrated = smart_fields_from_docx(
        source,
        [
            {
                "id": "itens",
                "label": "Itens",
                "type": "repeatable_table",
                "required": True,
                "section": "Dados do documento",
            }
        ],
    )
    migrated_repeatable = next(
        field for field in migrated if field["type"] == "repeatable_table"
    )
    assert migrated_repeatable["section"] == "3. Quantidade a ser contratada"
    assert migrated_repeatable["section_source"] == "word_table_title"

    # An explicitly edited section remains authoritative, even if it happens
    # to use the generic wording intentionally.
    manual = smart_fields_from_docx(
        source,
        [
            {
                "id": "itens",
                "label": "Itens",
                "type": "repeatable_table",
                "required": True,
                "section": "Dados do documento",
                "section_source": "manual",
            }
        ],
    )
    manual_repeatable = next(
        field for field in manual if field["type"] == "repeatable_table"
    )
    assert manual_repeatable["section"] == "Dados do documento"
    assert manual_repeatable["section_source"] == "manual"
