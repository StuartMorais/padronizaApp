from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from app.domain.fields import FieldDefinition
from app.domain.field_types import FieldType, normalize_field_type
from app.document.docx.generator import generate_docx
from app.document.docx.scanner import create_default_fields, scan_docx_fields
from app.document.docx.tags import TagKind, parse_tag


def test_field_definition_normalizes_legacy_type_aliases() -> None:
    field = FieldDefinition({
        "id": " contato.email ",
        "label": " E-mail ",
        "type": "input",
        "required": 1,
    })

    assert field.field_id == "contato.email"
    assert field.label == "E-mail"
    assert field.field_type is FieldType.TEXT
    assert field["type"] == "text"
    assert field.required is True
    assert normalize_field_type("money") == "currency"


def test_default_or_text_is_parsed_once_as_a_supported_tag() -> None:
    definition = parse_tag(
        "default_or_text:detalhamento.contratacao|Não há detalhamento adicional."
    )

    assert definition.kind is TagKind.FIELD
    assert definition.field_id == "detalhamento.contratacao"
    assert definition.field_type == "multiline"
    assert definition.default_value == "Não há detalhamento adicional."
    assert definition.metadata["tag_type"] == "default_or_text"


def test_default_or_text_scans_and_generates_default_or_user_text(tmp_path: Path) -> None:
    template = tmp_path / "template.docx"
    default_output = tmp_path / "default.docx"
    custom_output = tmp_path / "custom.docx"
    default_text = "Não há necessidade de detalhamento mais amplo."

    document = Document()
    document.add_paragraph(
        "Detalhamento: "
        f"{{{{default_or_text:detalhamento.contratacao|{default_text}}}}}"
    )
    document.save(template)

    scanned = scan_docx_fields(template)
    assert len(scanned) == 1
    assert isinstance(scanned[0], FieldDefinition)
    assert scanned[0]["id"] == "detalhamento.contratacao"
    assert scanned[0]["type"] == "multiline"
    assert scanned[0]["default_value"] == default_text

    fields = create_default_fields(scanned)
    assert isinstance(fields[0], FieldDefinition)
    assert fields[0]["default_value"] == default_text

    generate_docx(template, default_output, {"detalhamento.contratacao": ""})
    generated_default = Document(default_output)
    assert default_text in generated_default.paragraphs[0].text

    generate_docx(
        template,
        custom_output,
        {"detalhamento.contratacao": "Texto informado pelo usuário."},
    )
    generated_custom = Document(custom_output)
    assert "Texto informado pelo usuário." in generated_custom.paragraphs[0].text
    assert default_text not in generated_custom.paragraphs[0].text


def test_discovered_template_fields_use_canonical_model(tmp_path: Path) -> None:
    from app.services.templates import discover_templates

    template_dir = tmp_path / "templates" / "simple"
    template_dir.mkdir(parents=True)
    document = Document()
    document.add_paragraph("Nome: {{pessoa.nome}}")
    document.save(template_dir / "template.docx")
    (template_dir / "template.json").write_text(
        json.dumps(
            {
                "template": {
                    "id": "simple",
                    "name": "Simple",
                    "source_file": "template.docx",
                    "version": "1.0",
                },
                "fields": [
                    {
                        "id": "pessoa.nome",
                        "label": "Nome",
                        "type": "input",
                    }
                ],
                "output": {"filename_pattern": "{{template.name}}.docx"},
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )

    packages = discover_templates(tmp_path / "templates")

    assert len(packages) == 1
    assert isinstance(packages[0].fields[0], FieldDefinition)
    assert packages[0].fields[0].field_id == "pessoa.nome"
    assert packages[0].fields[0].field_type is FieldType.TEXT
