from __future__ import annotations

import os
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest
from docx import Document
from docx.enum.section import WD_SECTION

from app.document.conversion.pdf import convert_docx_to_pdf
import app.document.letterhead as letterhead_module
from app.document.letterhead import apply_letterhead, default_letterhead_path
from app.repositories.local_data import LocalDataStore
from app.repositories.templates import TemplateRepository
from app.services.generation import GenerationService
from app.services.templates import TemplatePackage


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _source(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Conteúdo de teste")
    document.save(path)
    return path


def _package(source: Path, *, enabled: bool) -> TemplatePackage:
    return TemplatePackage(
        template_id="modelo",
        name="Modelo",
        description="",
        category="",
        version="1.0",
        source_path=source,
        fields=[],
        output_filename="saida.docx",
        config={
            "output": {"filename_pattern": "saida.docx"},
            "numbering": {"enabled": False},
            "letterhead": {"enabled": enabled, "source": "bundled_default"},
        },
    )


def test_bundled_official_letterhead_is_shipped() -> None:
    path = default_letterhead_path()
    assert path.is_file()
    assert path.name == "Timbrado.docx"


def test_letterhead_replaces_headers_and_footers_for_every_section(tmp_path: Path) -> None:
    target = tmp_path / "multi-section.docx"
    document = Document()
    document.add_paragraph("Página um")
    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph("Página dois")
    document.save(target)

    apply_letterhead(target)

    with zipfile.ZipFile(target) as archive:
        names = set(archive.namelist())
        assert "word/header1.xml" in names
        assert "word/footer1.xml" in names
        assert any(name.startswith("word/media/padroniza-letterhead-") for name in names)

        document_root = ET.fromstring(archive.read("word/document.xml"))
        rels_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        relation_targets = {
            child.get("Id"): child.get("Target")
            for child in rels_root.findall(f"{{{REL_NS}}}Relationship")
        }
        sections = list(document_root.iter(f"{{{W_NS}}}sectPr"))
        assert len(sections) == 2
        for section in sections:
            headers = section.findall(f"{{{W_NS}}}headerReference")
            footers = section.findall(f"{{{W_NS}}}footerReference")
            assert {node.get(f"{{{W_NS}}}type") for node in headers} == {"default", "even", "first"}
            assert {node.get(f"{{{W_NS}}}type") for node in footers} == {"default", "even", "first"}
            header_targets = {relation_targets[node.get(f"{{{R_NS}}}id")] for node in headers}
            footer_targets = {relation_targets[node.get(f"{{{R_NS}}}id")] for node in footers}
            assert len(header_targets) == 1
            assert len(footer_targets) == 1
            assert next(iter(header_targets)).startswith("header")
            assert next(iter(header_targets)).endswith(".xml")
            assert next(iter(footer_targets)).startswith("footer")
            assert next(iter(footer_targets)).endswith(".xml")

    loaded = Document(target)
    footer_text = "\n".join(paragraph.text for paragraph in loaded.sections[0].footer.paragraphs)
    assert "Secretaria de Estado do Desenvolvimento Humano" in footer_text
    assert "João Pessoa/PB" in footer_text


def test_letterhead_closes_destination_before_atomic_replace(
    tmp_path: Path, monkeypatch: pytest.MonkeyPatch
) -> None:
    """Regression for Windows WinError 5 when the target ZIP remained open."""

    target = _source(tmp_path / "windows-lock.docx")
    destination = target.resolve()
    real_zipfile = zipfile.ZipFile
    real_replace = os.replace
    active_zip_paths: set[Path] = set()

    class TrackingZipFile(real_zipfile):
        def __init__(self, file, *args, **kwargs):  # noqa: ANN001
            super().__init__(file, *args, **kwargs)
            self._tracked_path = (
                Path(file).resolve()
                if isinstance(file, (str, os.PathLike))
                else None
            )
            if self._tracked_path is not None:
                active_zip_paths.add(self._tracked_path)

        def close(self) -> None:
            try:
                super().close()
            finally:
                tracked = getattr(self, "_tracked_path", None)
                if tracked is not None:
                    active_zip_paths.discard(tracked)

    def checked_replace(source, target_path) -> None:  # noqa: ANN001
        assert destination not in active_zip_paths, (
            "apply_letterhead attempted to replace the destination while its "
            "ZIP handle was still open; this fails with WinError 5 on Windows"
        )
        real_replace(source, target_path)

    monkeypatch.setattr(letterhead_module.zipfile, "ZipFile", TrackingZipFile)
    monkeypatch.setattr(letterhead_module.os, "replace", checked_replace)

    apply_letterhead(target)

    assert target.is_file()
    with real_zipfile(target) as archive:
        assert any(name.startswith("word/header") and name.endswith(".xml") for name in archive.namelist())


def test_template_repository_persists_letterhead_choice(tmp_path: Path) -> None:
    repository = TemplateRepository(tmp_path / "templates")
    source = _source(tmp_path / "source.docx")
    template_id = repository.create_template(
        name="Modelo timbrado",
        source_docx=source,
        fields=[],
        letterhead={"enabled": True},
    )
    assert repository.read_config(template_id)["letterhead"] == {
        "enabled": True,
        "source": "bundled_default",
    }

    repository.update_template(
        template_id=template_id,
        name="Modelo timbrado",
        fields=[],
        letterhead={"enabled": False},
    )
    assert repository.read_config(template_id)["letterhead"]["enabled"] is False


def test_generation_service_applies_letterhead_to_docx(tmp_path: Path) -> None:
    source = _source(tmp_path / "source.docx")
    service = GenerationService(LocalDataStore(tmp_path / "data"))
    output = tmp_path / "saida.docx"
    service.generate_docx(_package(source, enabled=True), {}, output)

    with zipfile.ZipFile(output) as archive:
        assert any(name.startswith("word/header") and name.endswith(".xml") for name in archive.namelist())
        assert any(name.startswith("word/footer") and name.endswith(".xml") for name in archive.namelist())


def test_pdf_generation_converts_the_letterheaded_docx(tmp_path: Path) -> None:
    source = _source(tmp_path / "source.docx")

    class InspectingConverter:
        def __init__(self) -> None:
            self.saw_letterhead = False

        def available_backend(self) -> str:
            return "fake"

        def last_backend(self) -> str:
            return "fake"

        def docx_to_pdf(self, source_path, destination, **_kwargs):
            with zipfile.ZipFile(source_path) as archive:
                self.saw_letterhead = any(name.startswith("word/header") and name.endswith(".xml") for name in archive.namelist())
            Path(destination).write_bytes(b"%PDF-1.4\n%%EOF")
            return Path(destination)

    converter = InspectingConverter()
    service = GenerationService(LocalDataStore(tmp_path / "data"), converter=converter)
    service.generate_pdf(_package(source, enabled=True), {}, tmp_path / "saida.pdf")
    assert converter.saw_letterhead is True


def test_integrated_pdf_converter_renders_letterhead_graphics_on_each_page(tmp_path: Path) -> None:
    fitz = pytest.importorskip("fitz")
    target = tmp_path / "two-pages.docx"
    document = Document()
    document.add_paragraph("Primeira página")
    document.add_page_break()
    document.add_paragraph("Segunda página")
    document.save(target)
    apply_letterhead(target)

    pdf_path = tmp_path / "two-pages.pdf"
    warnings: list[str] = []
    convert_docx_to_pdf(target, pdf_path, warnings=warnings)
    assert not [warning for warning in warnings if "decoração" in warning.casefold()]

    pdf = fitz.open(pdf_path)
    try:
        assert pdf.page_count == 2
        for page in pdf:
            # Official logo + red right-side artwork are real image XObjects.
            assert len(page.get_images(full=True)) >= 2
            text = page.get_text()
            assert "Secretaria de Estado do Desenvolvimento Humano" in text
            assert "João Pessoa/PB" in text
    finally:
        pdf.close()


def test_letterhead_preserves_word_namespace_prefixes_referenced_by_mc_ignorable(tmp_path: Path) -> None:
    """Word/LibreOffice must not receive dangling mc:Ignorable prefixes."""

    from lxml import etree as LET

    target = _source(tmp_path / "namespace-safe.docx")
    apply_letterhead(target)

    with zipfile.ZipFile(target) as archive:
        root = LET.fromstring(archive.read("word/document.xml"))

    mc_ns = "http://schemas.openxmlformats.org/markup-compatibility/2006"
    ignorable = root.get(f"{{{mc_ns}}}Ignorable", "").split()
    for prefix in ignorable:
        assert prefix in root.nsmap, f"mc:Ignorable references undeclared prefix {prefix}"


def test_letterhead_uses_word_standard_part_names(tmp_path: Path) -> None:
    target = _source(tmp_path / "standard-parts.docx")
    apply_letterhead(target)

    with zipfile.ZipFile(target) as archive:
        names = set(archive.namelist())
        assert not any("padronizaLetterhead" in name for name in names)
        rels_root = ET.fromstring(archive.read("word/_rels/document.xml.rels"))
        for relation in rels_root.findall(f"{{{REL_NS}}}Relationship"):
            rel_type = relation.get("Type", "")
            if rel_type.endswith("/header"):
                target_name = relation.get("Target", "")
                assert target_name.startswith("header") and target_name.endswith(".xml")
                assert target_name[6:-4].isdigit()
            elif rel_type.endswith("/footer"):
                target_name = relation.get("Target", "")
                assert target_name.startswith("footer") and target_name.endswith(".xml")
                assert target_name[6:-4].isdigit()


def test_letterhead_removes_legacy_custom_parts_on_reapply(tmp_path: Path) -> None:
    target = _source(tmp_path / "legacy.docx")

    # Simulate the V6.1.8 package naming that Microsoft Word rejected.
    rewritten = tmp_path / "legacy-staged.docx"
    with zipfile.ZipFile(target, "r") as source_zip, zipfile.ZipFile(
        rewritten, "w", compression=zipfile.ZIP_DEFLATED
    ) as output_zip:
        for info in source_zip.infolist():
            output_zip.writestr(info, source_zip.read(info.filename))
        output_zip.writestr("word/padronizaLetterheadHeader.xml", b"<legacy/>")
        output_zip.writestr("word/padronizaLetterheadFooter.xml", b"<legacy/>")
    rewritten.replace(target)

    apply_letterhead(target)

    with zipfile.ZipFile(target) as archive:
        names = set(archive.namelist())
        assert "word/padronizaLetterheadHeader.xml" not in names
        assert "word/padronizaLetterheadFooter.xml" not in names
