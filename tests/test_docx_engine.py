import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from app.document.docx.generator import (
    DocumentGenerationError,
    generate_docx,
)


class DocxEngineTests(unittest.TestCase):
    """Automated tests for the DOCX generation engine."""

    def test_replaces_paragraph_and_table_placeholders(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            document.add_paragraph(
                "Company: {{company.name}}"
            )
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = (
                "Process: {{process.number}}"
            )
            document.save(template_path)

            generate_docx(
                template_path=template_path,
                output_path=output_path,
                values={
                    "company.name": "Example Ltd.",
                    "process.number": "001/2026",
                },
            )

            result = Document(output_path)
            self.assertEqual(
                result.paragraphs[0].text,
                "Company: Example Ltd.",
            )
            self.assertEqual(
                result.tables[0].cell(0, 0).text,
                "Process: 001/2026",
            )

    def test_preserves_surrounding_run_formatting(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            paragraph = document.add_paragraph()

            prefix = paragraph.add_run("Important: ")
            prefix.bold = True

            marker = paragraph.add_run("{{client.name}}")
            marker.italic = True

            suffix = paragraph.add_run(" must sign.")
            suffix.underline = True

            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {"client.name": "Maria Silva"},
            )

            result = Document(output_path)
            runs = result.paragraphs[0].runs

            self.assertEqual(
                result.paragraphs[0].text,
                "Important: Maria Silva must sign.",
            )
            self.assertEqual(runs[0].text, "Important: ")
            self.assertTrue(runs[0].bold)
            self.assertEqual(runs[1].text, "Maria Silva")
            self.assertTrue(runs[1].italic)
            self.assertEqual(runs[2].text, " must sign.")
            self.assertTrue(runs[2].underline)

    def test_replaces_placeholder_split_across_runs(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("Client: ")

            first_marker_run = paragraph.add_run("{{client.")
            first_marker_run.italic = True
            paragraph.add_run("na")
            paragraph.add_run("me}}")

            suffix = paragraph.add_run(" / approved")
            suffix.underline = True
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {"client.name": "João"},
            )

            result = Document(output_path)
            paragraph = result.paragraphs[0]
            nonempty_runs = [
                run for run in paragraph.runs if run.text
            ]

            self.assertEqual(
                paragraph.text,
                "Client: João / approved",
            )
            self.assertEqual(nonempty_runs[1].text, "João")
            self.assertTrue(nonempty_runs[1].italic)
            self.assertEqual(
                nonempty_runs[-1].text,
                " / approved",
            )
            self.assertTrue(nonempty_runs[-1].underline)

    def test_preserves_template_text_colors(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            paragraph = document.add_paragraph()

            static_run = paragraph.add_run("Warning: ")
            static_run.font.color.rgb = RGBColor(
                0xC0,
                0x00,
                0x00,
            )

            marker_run = paragraph.add_run("{{message}}")
            marker_run.font.color.rgb = RGBColor(
                0x00,
                0x40,
                0xC0,
            )
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {"message": "Review required"},
            )

            result = Document(output_path)
            runs = result.paragraphs[0].runs

            self.assertEqual(
                runs[0].font.color.rgb,
                RGBColor(0xC0, 0x00, 0x00),
            )
            self.assertEqual(
                runs[1].font.color.rgb,
                RGBColor(0x00, 0x00, 0x00),
            )
            self.assertEqual(
                runs[1].text,
                "Review required",
            )

    def test_only_generated_text_becomes_black_within_same_run(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            paragraph = document.add_paragraph()
            colored_run = paragraph.add_run(
                "Responsável: {{responsavel.nome}} (confirmado)"
            )
            colored_run.font.color.rgb = RGBColor(
                0x00,
                0x40,
                0xC0,
            )
            colored_run.bold = True
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {"responsavel.nome": "Maria Silva"},
            )

            result = Document(output_path)
            runs = [
                run
                for run in result.paragraphs[0].runs
                if run.text
            ]

            self.assertEqual(
                [run.text for run in runs],
                [
                    "Responsável: ",
                    "Maria Silva",
                    " (confirmado)",
                ],
            )
            self.assertEqual(
                runs[0].font.color.rgb,
                RGBColor(0x00, 0x40, 0xC0),
            )
            self.assertEqual(
                runs[1].font.color.rgb,
                RGBColor(0x00, 0x00, 0x00),
            )
            self.assertEqual(
                runs[2].font.color.rgb,
                RGBColor(0x00, 0x40, 0xC0),
            )
            self.assertTrue(all(run.bold for run in runs))

    def test_supports_multiple_and_multiline_values(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            document.add_paragraph(
                "{{first}} | {{second}} | {{notes}} END"
            )
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {
                    "first": "A",
                    "second": "A much longer value",
                    "notes": "Line 1\nLine 2",
                },
            )

            result = Document(output_path)
            self.assertEqual(
                result.paragraphs[0].text,
                "A | A much longer value | Line 1\nLine 2 END",
            )
            generated_runs = [
                run
                for run in result.paragraphs[0].runs
                if run.text in {"A", "A much longer value", "Line 1\nLine 2"}
            ]
            self.assertTrue(generated_runs)
            self.assertTrue(
                all(
                    run.font.color.rgb == RGBColor(0x00, 0x00, 0x00)
                    for run in generated_runs
                )
            )

    def test_replaces_header_and_footer_placeholders(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            section = document.sections[0]
            section.header.paragraphs[0].text = (
                "Header {{document.code}}"
            )
            section.footer.paragraphs[0].text = (
                "Footer {{document.year}}"
            )
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {
                    "document.code": "ABC-10",
                    "document.year": "2026",
                },
            )

            result = Document(output_path)
            result_section = result.sections[0]
            self.assertEqual(
                result_section.header.paragraphs[0].text,
                "Header ABC-10",
            )
            self.assertEqual(
                result_section.footer.paragraphs[0].text,
                "Footer 2026",
            )

    def test_expands_repeatable_table_rows(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "No."
            table.cell(0, 1).text = "Description"

            number_paragraph = table.cell(1, 0).paragraphs[0]
            number_paragraph.text = (
                "{{repeat:items}}{{row.number}}"
            )

            description_paragraph = (
                table.cell(1, 1).paragraphs[0]
            )
            description_run = description_paragraph.add_run(
                "{{items.description}}"
            )
            description_run.italic = True
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {
                    "items": [
                        {"description": "Notebook"},
                        {"description": "Monitor"},
                    ]
                },
            )

            result = Document(output_path)
            result_table = result.tables[0]

            self.assertEqual(len(result_table.rows), 3)
            self.assertEqual(
                result_table.cell(1, 0).text,
                "01",
            )
            self.assertEqual(
                result_table.cell(2, 0).text,
                "02",
            )
            self.assertEqual(
                result_table.cell(1, 1).text,
                "Notebook",
            )
            self.assertEqual(
                result_table.cell(2, 1).text,
                "Monitor",
            )
            generated_run = result_table.cell(1, 1).paragraphs[0].runs[0]
            self.assertTrue(generated_run.italic)
            self.assertEqual(
                generated_run.font.color.rgb,
                RGBColor(0x00, 0x00, 0x00),
            )

    def test_native_word_control_value_is_black(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            paragraph = document.add_paragraph("Data: ")

            sdt = OxmlElement("w:sdt")
            properties = OxmlElement("w:sdtPr")
            tag = OxmlElement("w:tag")
            tag.set(qn("w:val"), "date:document.date")
            properties.append(tag)
            properties.append(OxmlElement("w:date"))
            sdt.append(properties)

            content = OxmlElement("w:sdtContent")
            run = OxmlElement("w:r")
            run_properties = OxmlElement("w:rPr")
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "0040C0")
            run_properties.append(color)
            run.append(run_properties)
            text_element = OxmlElement("w:t")
            text_element.text = "Selecione uma data"
            run.append(text_element)
            content.append(run)
            sdt.append(content)
            paragraph._p.append(sdt)
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {"document.date": "06/08/2026"},
            )

            result = Document(output_path)
            result_sdt = next(result.element.iter(qn("w:sdt")))
            result_text = next(result_sdt.iter(qn("w:t")))
            result_run = result_text.getparent()
            result_color = result_run.find(qn("w:rPr")).find(
                qn("w:color")
            )

            self.assertEqual(result_text.text, "06/08/2026")
            self.assertEqual(result_color.get(qn("w:val")), "000000")

    def test_unnamed_native_dropdown_uses_same_context_id_as_scanner(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            table = document.add_table(rows=1, cols=1)
            cell = table.cell(0, 0)

            sdt = OxmlElement("w:sdt")
            properties = OxmlElement("w:sdtPr")
            dropdown = OxmlElement("w:dropDownList")
            for option in (
                "Escolha a Unidade Gestora",
                "27.0001 - SEDH",
                "50.0001 - FEAS",
            ):
                item = OxmlElement("w:listItem")
                item.set(qn("w:displayText"), option)
                item.set(qn("w:value"), option)
                dropdown.append(item)
            properties.append(dropdown)
            sdt.append(properties)

            content = OxmlElement("w:sdtContent")
            paragraph = OxmlElement("w:p")
            run = OxmlElement("w:r")
            text = OxmlElement("w:t")
            text.text = "Escolha a Unidade Gestora"
            run.append(text)
            paragraph.append(run)
            content.append(paragraph)
            sdt.append(content)
            cell._tc.append(sdt)
            document.save(template_path)

            generate_docx(
                template_path,
                output_path,
                {"unidade_gestora": "27.0001 - SEDH"},
            )

            result = Document(output_path)
            result_sdt = next(result.element.iter(qn("w:sdt")))
            result_text = "".join(
                node.text or ""
                for node in result_sdt.iter(qn("w:t"))
            )
            self.assertEqual(result_text, "27.0001 - SEDH")

    def test_reports_missing_values(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            template_path = temp_path / "template.docx"
            output_path = temp_path / "output.docx"

            document = Document()
            document.add_paragraph("{{required.field}}")
            document.save(template_path)

            with self.assertRaisesRegex(
                DocumentGenerationError,
                "required.field",
            ):
                generate_docx(
                    template_path,
                    output_path,
                    {},
                )


if __name__ == "__main__":
    unittest.main()


def test_single_choice_tag_inserts_selected_full_text(tmp_path: Path) -> None:
    template_path = tmp_path / "single-choice-template.docx"
    output_path = tmp_path / "single-choice-output.docx"
    selected = "Não consta(m) no Plano de Contratações Anual – PCA."

    document = Document()
    document.add_paragraph(
        "Situação: {{single_choice:pca_2025.situacao|"
        "Consta no PCA => Consta(m) no Plano de Contratações Anual – PCA, conforme comprovação em anexo.|"
        "Não consta no PCA => Não consta(m) no Plano de Contratações Anual – PCA.}}"
    )
    document.save(template_path)

    generate_docx(
        template_path,
        output_path,
        {"pca_2025.situacao": selected},
    )

    result = Document(output_path)
    assert result.paragraphs[0].text == f"Situação: {selected}"
    selected_run = next(
        run for run in result.paragraphs[0].runs if run.text == selected
    )
    assert selected_run.font.color.rgb == RGBColor(0x00, 0x00, 0x00)
