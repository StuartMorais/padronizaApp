import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.docx_engine import generate_docx
from app.placeholder_scanner import scan_docx_fields
from app.smart_template import scan_docx_health, smart_fields_from_docx


class UnicodeFieldIdTests(unittest.TestCase):
    def test_accented_tag_is_valid_and_keeps_section_context(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "modelo.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).merge(table.cell(0, 1))
            table.cell(0, 0).text = "1.1 Identificação da Demanda:"
            table.cell(1, 0).text = "Descrição da demanda:"
            table.cell(1, 1).text = "{{descrição.demanda}}"
            document.save(path)

            scanned = scan_docx_fields(path)
            self.assertIn("descrição.demanda", {field["id"] for field in scanned})

            health = scan_docx_health(path)
            self.assertEqual(health["malformed_placeholders"], [])
            self.assertIn("descrição.demanda", health["field_ids"])

            fields = smart_fields_from_docx(path)
            field = next(item for item in fields if item["id"] == "descrição.demanda")
            self.assertEqual(field["label"], "Descrição da demanda")
            self.assertEqual(field["section"], "1.1 Identificação da Demanda")
            self.assertEqual(field["type"], "multiline")

    def test_generation_replaces_accented_tag(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template = temp / "template.docx"
            output = temp / "output.docx"
            document = Document()
            document.add_paragraph("Descrição: {{descrição.demanda}}")
            document.save(template)

            generate_docx(
                template,
                output,
                {"descrição.demanda": "Aquisição de equipamentos"},
            )

            result = Document(output)
            self.assertEqual(
                result.paragraphs[0].text,
                "Descrição: Aquisição de equipamentos",
            )

    def test_repeatable_table_accepts_accented_identifier(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            temp = Path(temp_dir)
            template = temp / "template.docx"
            output = temp / "output.docx"
            document = Document()
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Nº"
            table.cell(0, 1).text = "Descrição"
            table.cell(1, 0).text = "{{repeat:itens.solicitação}}{{row.number}}"
            table.cell(1, 1).text = "{{itens.solicitação.descrição}}"
            document.save(template)

            fields = scan_docx_fields(template)
            repeat = next(field for field in fields if field["id"] == "itens.solicitação")
            self.assertEqual(repeat["type"], "repeatable_table")

            generate_docx(
                template,
                output,
                {
                    "itens.solicitação": [
                        {"descrição": "Notebook"},
                        {"descrição": "Monitor"},
                    ]
                },
            )
            result = Document(output)
            self.assertEqual(len(result.tables[0].rows), 3)
            self.assertEqual(result.tables[0].cell(1, 1).text, "Notebook")
            self.assertEqual(result.tables[0].cell(2, 1).text, "Monitor")


if __name__ == "__main__":
    unittest.main()
