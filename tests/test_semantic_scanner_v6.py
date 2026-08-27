from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document

from app.core.schema import TEMPLATE_SCHEMA_VERSION
from app.document.diagnostics import diagnose_template
from app.document.detection.application import apply_docx_field_candidates
from app.document.detection.detector import clear_detection_cache, detect_docx_field_candidates
from app.document.docx.generator import generate_docx
from app.document.docx.scanner import clear_docx_scan_cache, scan_docx_fields
from app.domain.field_metadata import source_anchor_errors
from app.domain.field_metadata import preserved_editor_field_metadata
from app.domain.template_quality import field_configuration_issues
from app.repositories.semantic_learning import SemanticLearningStore
from app.repositories.templates import TemplateRepository
from app.services.template_scanning import record_semantic_reviews
from tools.check_semantic_benchmark import run_benchmark


FIXTURE = (
    Path(__file__).resolve().parent
    / "fixtures"
    / "semantic_v6"
    / "justificativa_vantajosidade_adesao_ata.docx"
)
EXPECTED_IDS = {
    "procurement.items",
    "procurement.ata_number",
    "procurement.managing_agency",
    "organization.acronym",
    "justification.celerity",
    "justification.economy",
    "justification.efficiency",
    "justification.legal_security",
    "justification.continuity",
    "justification.conclusion",
    "person.name",
    "person.role",
    "person.registration",
}


def _semantic_candidates(path: Path = FIXTURE, *, memory: dict | None = None):
    clear_detection_cache()
    return detect_docx_field_candidates(path, semantic_memory=memory, semantic_enabled=True)


def _review(candidate: dict, accepted: bool) -> dict:
    value = deepcopy(candidate)
    value["accepted_by_user"] = accepted
    return value


def _replace_paragraph_text(document: Document, old: str, new: str) -> None:
    replaced = False
    for paragraph in document.paragraphs:
        if old not in paragraph.text:
            continue
        paragraph.text = paragraph.text.replace(old, new)
        replaced = True
    assert replaced, old


def _changed_family_document(destination: Path) -> Path:
    document = Document(str(FIXTURE))
    replacements = {
        "0006/2026": "0099/2027",
        "Secretaria de Estado da Administração – SEAD/PB": "Secretaria de Estado da Saúde – SES/PB",
        "• Cadernos;\n• Canetas hidrográficas;\n• Pilhas.": "• Computadores;\n• Monitores;\n• Nobreaks.",
        "A adesão possibilita a contratação imediata do fornecedor registrado, reduzindo prazos e assegurando o pronto atendimento das demandas relacionadas ao fornecimento de materiais de expediente essenciais.": "A adesão atualizada permite iniciar o fornecimento com menor tempo administrativo e atendimento imediato da nova necessidade institucional.",
        "LUCILEIDE DA SILVA NASCIMENTO": "MARIA CLARA DE OLIVEIRA SOUZA",
        "Assistente Administrativa": "Coordenadora de Compras Governamentais",
        "917.108-8": "123.456-7",
    }
    for old, new in replacements.items():
        _replace_paragraph_text(document, old, new)
    document.save(str(destination))
    return destination


def test_semantic_benchmark_contract_is_green() -> None:
    result = run_benchmark()
    assert result == {
        "required_total": 13,
        "required_found": 13,
        "unexpected": 0,
        "fresh_semantic_preselected": 0,
    }


def test_real_narrative_document_discovers_all_review_first_semantic_regions() -> None:
    candidates = _semantic_candidates()
    assert {candidate["field_id"] for candidate in candidates} == EXPECTED_IDS
    assert all(candidate["selected"] is False for candidate in candidates)
    assert all(candidate["auto_apply_eligible"] is False for candidate in candidates)

    item_list = next(item for item in candidates if item["field_id"] == "procurement.items")
    assert item_list["type"] == "repeatable_list"
    assert item_list["default_value"] == ["Cadernos", "Canetas hidrográficas", "Pilhas"]

    for candidate in candidates:
        assert candidate["pipeline_version"] == 6
        assert candidate.get("source_anchor")
        assert not source_anchor_errors(
            candidate["source_anchor"], expected_scope=str(candidate.get("dynamic_scope", ""))
        )


def test_all_semantic_regions_apply_and_strictly_rescan(tmp_path: Path) -> None:
    candidates = _semantic_candidates()
    destination = tmp_path / "semantic-tagged.docx"
    apply_docx_field_candidates(FIXTURE, destination, candidates)

    clear_docx_scan_cache()
    scanned = scan_docx_fields(destination)
    by_id = {field["id"]: field for field in scanned}
    assert EXPECTED_IDS <= set(by_id)
    assert by_id["procurement.items"]["type"] == "repeatable_list"
    assert by_id["procurement.items"]["list_style"] == "bullet"
    assert by_id["procurement.items"]["list_punctuation"] == "semicolon"


def test_accepted_family_mappings_relocate_changed_values_and_preselect(tmp_path: Path) -> None:
    original = _semantic_candidates()
    assert record_semantic_reviews(tmp_path, [_review(item, True) for item in original]) == 13
    memory = SemanticLearningStore(tmp_path).snapshot()
    changed = _changed_family_document(tmp_path / "changed-family.docx")

    candidates = _semantic_candidates(changed, memory=memory)
    by_id = {item["field_id"]: item for item in candidates}
    assert set(by_id) == EXPECTED_IDS
    assert all(item["source"] == "learned_mapping" for item in candidates)
    assert all(item["selected"] is True for item in candidates)
    assert by_id["procurement.ata_number"]["default_value"] == "0099/2027"
    assert by_id["procurement.managing_agency"]["default_value"] == "Secretaria de Estado da Saúde"
    assert by_id["organization.acronym"]["default_value"] == "SES/PB"
    assert by_id["procurement.items"]["default_value"] == ["Computadores", "Monitores", "Nobreaks"]
    assert by_id["person.name"]["default_value"] == "MARIA CLARA DE OLIVEIRA SOUZA"
    assert by_id["person.role"]["default_value"] == "Coordenadora de Compras Governamentais"
    assert by_id["person.registration"]["default_value"] == "123.456-7"
    assert "atendimento imediato" in by_id["justification.celerity"]["default_value"]


def test_rejected_family_region_is_not_suggested_again(tmp_path: Path) -> None:
    ata = next(item for item in _semantic_candidates() if item["field_id"] == "procurement.ata_number")
    assert record_semantic_reviews(tmp_path, [_review(ata, False)]) == 1
    memory = SemanticLearningStore(tmp_path).snapshot()

    candidates = _semantic_candidates(memory=memory)
    assert "procurement.ata_number" not in {item["field_id"] for item in candidates}
    # Rejection is local to the reviewed region; unrelated semantic fields remain.
    assert "procurement.managing_agency" in {item["field_id"] for item in candidates}


def test_invalid_repeatable_list_and_semantic_anchor_are_blocked_before_save(tmp_path: Path) -> None:
    field = {
        "id": "procurement.items",
        "label": "Materiais",
        "type": "repeatable_list",
        "default_value": "not-a-list",
        "minimum_items": 4,
        "maximum_items": 2,
        "list_style": "mystery",
        "list_punctuation": "random",
        "dynamic_scope": "magic",
        "source_anchor": {"version": 99, "scope": "inline", "spans": []},
    }
    issues = field_configuration_issues([field])
    codes = {issue.code for issue in issues}
    assert "list.invalid" in codes
    assert "semantic.scope_invalid" in codes
    assert "semantic.anchor_invalid" in codes

    repository = TemplateRepository(tmp_path / "templates")
    try:
        repository._normalize_fields([field], strict=True)  # noqa: SLF001 - schema boundary regression
    except ValueError as exc:
        assert "lista repetível" in str(exc).casefold() or "âncora" in str(exc).casefold()
    else:  # pragma: no cover - defensive assertion
        raise AssertionError("Invalid semantic/list metadata was accepted in strict template normalization")


def test_diagnostics_blocks_invalid_v6_list_and_semantic_metadata(tmp_path: Path) -> None:
    source = tmp_path / "semantic-invalid.docx"
    document = Document()
    document.add_paragraph("Materiais: {{repeat_list:procurement.items|bullet|semicolon}}")
    document.save(str(source))

    report = diagnose_template(
        {
            "fields": [
                {
                    "id": "procurement.items",
                    "label": "Materiais",
                    "type": "repeatable_list",
                    "default_value": ["Papel"],
                    "minimum_items": 3,
                    "maximum_items": 1,
                    "list_style": "invalid",
                    "list_punctuation": "invalid",
                    "dynamic_scope": "list",
                    "source_anchor": {
                        "version": 99,
                        "scope": "list",
                        "spans": [],
                    },
                }
            ]
        },
        source,
    )

    codes = {issue["code"] for issue in report["issues"]}
    assert report["blocking"] is True
    assert "list.invalid" in codes
    assert "semantic.metadata_invalid" in codes
    assert report["invalid_repeatable_lists"]["procurement.items"]
    assert report["invalid_semantic_metadata"]["procurement.items"]


def test_template_schema_v1_loads_as_v2_and_v6_metadata_survives_normalization(tmp_path: Path) -> None:
    templates_dir = tmp_path / "templates"
    folder = templates_dir / "legacy"
    folder.mkdir(parents=True)
    document = Document()
    document.add_paragraph("Ata nº {{procurement.ata_number}}")
    document.save(str(folder / "template.docx"))

    repository = TemplateRepository(templates_dir)
    legacy = repository._normalize_config(  # noqa: SLF001 - migration boundary regression
        {
            "schema_version": 1,
            "template": {
                "id": "legacy",
                "name": "Legacy",
                "source_file": "template.docx",
            },
            "fields": [
                {
                    "id": "procurement.ata_number",
                    "label": "Número da Ata",
                    "type": "text",
                }
            ],
        },
        canonical_id="legacy",
        folder=folder,
    )
    assert legacy["schema_version"] == TEMPLATE_SCHEMA_VERSION == 2

    semantic = next(
        item for item in _semantic_candidates() if item["field_id"] == "procurement.ata_number"
    )
    normalized = repository._normalize_fields(  # noqa: SLF001 - schema boundary regression
        [
            {
                "id": semantic["field_id"],
                "label": semantic["label"],
                "type": semantic["type"],
                "dynamic_scope": semantic["dynamic_scope"],
                "semantic_concept_id": semantic["semantic_concept_id"],
                "semantic_model_version": semantic["semantic_model_version"],
                "source_anchor": semantic["source_anchor"],
                "source_context": semantic["source_context"],
                "family_fingerprint": semantic["family_fingerprint"],
            }
        ],
        strict=True,
    )[0]
    assert normalized["dynamic_scope"] == "inline"
    assert normalized["semantic_concept_id"] == "procurement.ata_number"
    assert normalized["source_anchor"] == semantic["source_anchor"]
    assert normalized["family_fingerprint"] == semantic["family_fingerprint"]


def test_repeatable_list_generation_formats_items_and_punctuation(tmp_path: Path) -> None:
    template = tmp_path / "list-template.docx"
    output = tmp_path / "list-output.docx"
    document = Document()
    document.add_paragraph("Materiais:\n{{repeat_list:procurement.items|bullet|semicolon}}")
    document.save(str(template))

    generate_docx(
        template,
        output,
        {"procurement.items": ["Papel A4", "Toner", "Pastas"]},
    )

    result = Document(str(output))
    assert result.paragraphs[0].text == "Materiais:\n• Papel A4;\n• Toner;\n• Pastas."


def test_editor_preserves_v6_semantic_and_repeatable_list_metadata() -> None:
    candidate = next(
        item for item in _semantic_candidates() if item["field_id"] == "procurement.items"
    )
    field = {
        "id": candidate["field_id"],
        "type": candidate["type"],
        "dynamic_scope": candidate["dynamic_scope"],
        "semantic_concept_id": candidate["semantic_concept_id"],
        "semantic_model_version": candidate["semantic_model_version"],
        "source_anchor": candidate["source_anchor"],
        "source_context": candidate["source_context"],
        "family_fingerprint": candidate["family_fingerprint"],
        "default_value": candidate["default_value"],
        "list_style": candidate["list_style"],
        "list_punctuation": candidate["list_punctuation"],
        "minimum_items": candidate["minimum_items"],
    }
    kept = preserved_editor_field_metadata(field)
    for key in (
        "dynamic_scope",
        "semantic_concept_id",
        "semantic_model_version",
        "source_anchor",
        "source_context",
        "family_fingerprint",
        "default_value",
        "list_style",
        "list_punctuation",
        "minimum_items",
    ):
        assert kept[key] == field[key]
