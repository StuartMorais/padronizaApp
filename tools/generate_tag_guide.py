from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "GUIA_DE_TAGS_PADRONIZA.docx"
LOGO = ROOT / "assets" / "padroniza-256x256.png"

BLUE = "1D5F91"
LIGHT_BLUE = "EAF3F9"
DARK = "1E2B36"
LIGHT_GRAY = "F2F4F6"
MID_GRAY = "D7DDE3"
GREEN = "E9F6EC"
YELLOW = "FFF7D6"
RED = "FDEBEC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_code(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(30, 70, 100)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "EEF4F8")
    run._r.get_or_add_rPr().append(shading)


def add_callout(document: Document, title: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 130, 150, 130, 150)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title + "\n")
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(BLUE)
    paragraph.add_run(text)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_tag_table(document: Document, rows: list[tuple[str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header = table.rows[0]
    set_repeat_table_header(header)
    for index, label in enumerate(("Uso", "Tag no DOCX", "Resultado no formulário")):
        cell = header.cells[index]
        set_cell_shading(cell, BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for use, tag, result in rows:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
        cells[0].text = use
        cells[1].paragraphs[0].clear()
        add_code(cells[1].paragraphs[0], tag)
        cells[2].text = result
    document.add_paragraph()


def add_step(document: Document, number: int, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(1.1)
    table.columns[1].width = Cm(15.5)
    number_cell, text_cell = table.rows[0].cells
    set_cell_shading(number_cell, BLUE)
    set_cell_shading(text_cell, LIGHT_GRAY)
    set_cell_margins(number_cell, 90, 90, 90, 90)
    set_cell_margins(text_cell, 100, 130, 100, 130)
    number_p = number_cell.paragraphs[0]
    number_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    number_run = number_p.add_run(str(number))
    number_run.bold = True
    number_run.font.color.rgb = RGBColor(255, 255, 255)
    text_p = text_cell.paragraphs[0]
    title_run = text_p.add_run(title + " — ")
    title_run.bold = True
    text_p.add_run(text)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def add_example_static_table(document: Document) -> None:
    table = document.add_table(rows=4, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Função", "Nome completo", "Matrícula", "Setor", "Ciência")
    for index, label in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    rows = (
        ("Gestor do contrato", "{{gestor.nome}}", "{{gestor.matricula}}", "{{gestor.setor}}", "{{checkbox:gestor.ciencia}}"),
        ("Gestor substituto", "{{gestor_substituto.nome}}", "{{gestor_substituto.matricula}}", "{{gestor_substituto.setor}}", "{{checkbox:gestor_substituto.ciencia}}"),
        ("Fiscal técnico", "{{fiscal.nome}}", "{{fiscal.matricula}}", "{{fiscal.setor}}", "{{checkbox:fiscal.ciencia}}"),
    )
    for row_index, values in enumerate(rows, start=1):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            if column_index == 0:
                cell.text = value
                set_cell_shading(cell, LIGHT_GRAY)
            else:
                cell.paragraphs[0].clear()
                add_code(cell.paragraphs[0], value)
            set_cell_margins(cell)
    document.add_paragraph()


def add_example_repeatable_table(document: Document) -> None:
    table = document.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Item", "Descrição", "Unidade", "Quantidade", "Aceito")
    for index, label in enumerate(headers):
        cell = table.cell(0, index)
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    values = (
        "{{repeat:itens}} {{row.number}}",
        "{{itens.descricao}}",
        "{{itens.unidade}}",
        "{{itens.quantidade}}",
        "{{checkbox:itens.aceito}}",
    )
    for index, value in enumerate(values):
        cell = table.cell(1, index)
        cell.paragraphs[0].clear()
        add_code(cell.paragraphs[0], value)
        set_cell_margins(cell)
    document.add_paragraph()


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(DARK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, BLUE),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 13, DARK),
        ("Heading 3", 11.5, BLUE),
    ):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)


def build_guide() -> Path:
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("PADRONIZA • Guia de tags e organização de formulários")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(110, 125, 138)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run("Use este guia junto com Criar modelo > Campos e seções > Prévia do formulário")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(110, 125, 138)

    cover = document.add_table(rows=1, cols=2)
    cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    cover.autofit = False
    cover.columns[0].width = Cm(3.2)
    cover.columns[1].width = Cm(13.4)
    left, right = cover.rows[0].cells
    set_cell_shading(left, BLUE)
    set_cell_shading(right, LIGHT_BLUE)
    set_cell_margins(left, 250, 200, 250, 200)
    set_cell_margins(right, 250, 300, 250, 300)
    if LOGO.exists():
        p = left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Cm(2.3))
    title_p = right.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run("Guia de Tags do Padroniza")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string(BLUE)
    subtitle = right.add_paragraph("Como preparar modelos DOCX, organizar campos e gerar formulários claros")
    subtitle.style = document.styles["Subtitle"]
    subtitle.runs[0].font.name = "Arial"
    subtitle.runs[0].font.color.rgb = RGBColor(65, 84, 99)

    document.add_paragraph()
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = intro.add_run(
        "Versão de referência para o editor Campos e seções, grupos de escolha, "
        "tabelas de responsáveis e tabelas repetíveis."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(90, 105, 118)

    document.add_page_break()

    document.add_heading("1. Fluxo rápido", level=1)
    add_step(document, 1, "Prepare o DOCX", "Mantenha o texto e a formatação final no Word. Substitua somente os pontos variáveis por tags.")
    add_step(document, 2, "Crie o modelo", "No Padroniza, abra Gerenciar modelos > Criar modelo e selecione o DOCX.")
    add_step(document, 3, "Localize os campos", "Use Ferramentas DOCX > Localizar campos. O aplicativo identifica tags, rótulos, seções e estruturas de tabela.")
    add_step(document, 4, "Revise Campos e seções", "Ajuste tipo, obrigatoriedade, seção e layout. Use a prévia para conferir a organização.")
    add_step(document, 5, "Salve e teste", "Gere um documento com dados de exemplo e confirme a aparência do DOCX final.")

    add_callout(
        document,
        "Regra principal",
        "A tag controla o conteúdo que será substituído. A tela Campos e seções controla como esse conteúdo aparece no formulário.",
        GREEN,
    )

    document.add_heading("2. Regras para IDs de campo", level=1)
    bullets = (
        "Comece com uma letra e use letras, números, ponto, hífen ou sublinhado.",
        "Prefira IDs estáveis e sem espaços, por exemplo: gestor.nome e processo.numero.",
        "Use o mesmo ID em todas as ocorrências que devem receber o mesmo valor.",
        "Não reutilize o mesmo ID para informações diferentes.",
        "Organize por assunto: empresa.cnpj, empresa.nome, contrato.numero, fiscal.setor.",
    )
    for text in bullets:
        document.add_paragraph(text, style="List Bullet")

    add_callout(
        document,
        "Bom padrão",
        "empresa.razao_social, processo.numero, gestor.nome, gestor.matricula, fiscal.ciencia",
        LIGHT_BLUE,
    )
    add_callout(
        document,
        "Evite",
        "campo1, valor2, nome novo, informação-final! — IDs genéricos ou com caracteres especiais dificultam manutenção e perfis.",
        RED,
    )

    document.add_heading("3. Referência de tags", level=1)
    add_tag_table(
        document,
        [
            ("Texto simples", "{{empresa.nome}}", "Campo de texto. O tipo pode ser refinado no editor."),
            ("Data explícita", "{{date:documento.data}}", "Seletor de data; por padrão usa a data atual automaticamente."),
            ("Caixa de seleção", "{{checkbox:declaracao.aceita}}", "Opção marcada/desmarcada; imprime ☑ ou ☐."),
            ("Lista suspensa", "{{dropdown:contratacao.modalidade|Pregão|Concorrência|Dispensa}}", "Lista pesquisável com as opções informadas."),
            ("Lista com texto longo", "{{dropdown:parecer.tipo|Favorável => Texto completo favorável|Desfavorável => Texto completo desfavorável}}", "Mostra um título curto e insere o texto completo."),
            ("Escolha única em caixas", "{{single_choice:pca.situacao|Consta no PCA|Não consta no PCA}}", "Exibe alternativas completas como caixas grandes; somente uma pode ser selecionada."),
            ("Tabela repetível", "{{repeat:itens}}", "Transforma uma linha modelo do Word em várias linhas preenchíveis."),
            ("Número da linha", "{{row.number}}", "Numeração automática dentro da mesma linha de uma tabela repetível."),
        ],
    )

    document.add_heading("Tipos escolhidos no Editor de Modelo", level=2)
    document.add_paragraph(
        "Os prefixos especiais são date:, checkbox:, dropdown:, single_choice: e repeat:. Os demais tipos são definidos na coluna Tipo da tela Campos: texto com várias linhas, moeda, inteiro, decimal, porcentagem, CNPJ, CPF, CEP, telefone e e-mail."
    )
    add_callout(
        document,
        "Exemplo",
        "A tag {{empresa.cnpj}} pode ser escrita como texto normal. Depois da análise, escolha o tipo CNPJ para receber máscara e validação.",
        YELLOW,
    )

    document.add_heading("4. Como criar seções claras", level=1)
    document.add_paragraph(
        "O Padroniza tenta identificar o título da seção imediatamente anterior aos campos. Para melhorar o resultado automático:"
    )
    for text in (
        "Use estilos de título do Word, como Título 1 ou Título 2.",
        "Títulos numerados curtos também são reconhecidos, por exemplo: 14. Responsáveis pela fiscalização:",
        "Dentro de tabelas, use uma linha mesclada em toda a largura para o título da seção.",
        "Evite colocar duas seções diferentes dentro da mesma linha do Word.",
    ):
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("Exemplo de seção no DOCX", level=3)
    example = document.add_table(rows=2, cols=1)
    example.style = "Table Grid"
    set_cell_shading(example.cell(0, 0), BLUE)
    p = example.cell(0, 0).paragraphs[0]
    run = p.add_run("14. Responsáveis pela fiscalização:")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    example.cell(1, 0).paragraphs[0].clear()
    add_code(example.cell(1, 0).paragraphs[0], "{{gestor.nome}}   {{gestor.matricula}}   {{gestor.setor}}")
    document.add_paragraph()

    document.add_heading("5. Opções exclusivas sem desorganização", level=1)
    document.add_paragraph(
        "Quando o usuário deve escolher somente uma alternativa, prefira single_choice:. "
        "O formulário mostra cada opção como uma caixa grande, com o texto visível e toda "
        "a linha clicável. A aparência é de caixa de seleção, mas somente uma alternativa "
        "pode permanecer marcada."
    )
    p = document.add_paragraph()
    add_code(
        p,
        "{{single_choice:pca_2025.situacao|Consta no PCA => Consta(m) no Plano de Contratações Anual – PCA, conforme comprovação em anexo.|Não consta no PCA => Não consta(m) no Plano de Contratações Anual – PCA.}}",
    )
    add_callout(
        document,
        "Como aparece no formulário",
        "O título curto aparece em destaque e o texto completo logo abaixo. O usuário pode clicar no quadrado, no título ou em qualquer parte da alternativa.",
        LIGHT_BLUE,
    )
    document.add_paragraph("Na tela Campos e seções:", style="Heading 3")
    for text in (
        "O campo é identificado como Lista suspensa com Layout = Grupo de escolha.",
        "As opções são editadas na coluna Configuração; use Título curto => Texto completo para textos longos.",
        "Marque Obrigatório quando uma alternativa tiver de ser escolhida antes da geração.",
        "Na Prévia do formulário, confirme se todo o texto está legível e se as opções aparecem empilhadas.",
    ):
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("Alternativa com caixas separadas", level=2)
    document.add_paragraph(
        "Modelos antigos podem manter duas ou mais tags checkbox: no DOCX. Defina o mesmo "
        "Grupo do layout e escolha Layout = Grupo de escolha para todas elas. O Padroniza "
        "também as exibirá como caixas grandes exclusivas e continuará imprimindo ☑ ou ☐ "
        "em cada posição do documento."
    )
    choice = document.add_table(rows=2, cols=2)
    choice.style = "Table Grid"
    choice.cell(0, 0).merge(choice.cell(0, 1))
    set_cell_shading(choice.cell(0, 0), BLUE)
    p = choice.cell(0, 0).paragraphs[0]
    run = p.add_run("13. Prazo de entrega")
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    for index, value in enumerate((
        "{{checkbox:prazo.imediata}} Entrega imediata",
        "{{checkbox:prazo.parcelada}} Entrega parcelada",
    )):
        choice.cell(1, index).paragraphs[0].clear()
        add_code(choice.cell(1, index).paragraphs[0], value)
        set_cell_margins(choice.cell(1, index), 120, 140, 120, 140)
    document.add_paragraph()

    document.add_heading("6. Tabelas de campos fixos", level=1)
    document.add_paragraph(
        "Para responsáveis, fiscais ou assinaturas com colunas repetidas, use tags normais em uma tabela do Word. Não use repeat: quando a quantidade de linhas é fixa."
    )
    add_example_static_table(document)
    document.add_paragraph(
        "O analisador usa a primeira coluna como rótulo da linha e o cabeçalho como rótulo das colunas. Na prévia, os campos aparecem na mesma estrutura: uma linha por função e uma coluna por informação."
    )
    add_callout(
        document,
        "Ajuste manual",
        "Em Layout > Detalhes, campos da mesma tabela devem compartilhar o Grupo do layout. Campos do mesmo responsável compartilham a Chave da linha; Nome, Matrícula, Setor e Ciência usam colunas diferentes.",
        LIGHT_BLUE,
    )

    document.add_heading("7. Tabelas repetíveis", level=1)
    repeat_intro = document.add_paragraph(
        "Use uma tabela repetível quando o usuário precisa adicionar uma quantidade variável de itens. A linha que contém repeat: é a linha modelo que será duplicada."
    )
    repeat_intro.paragraph_format.keep_with_next = True
    add_example_repeatable_table(document)
    for text in (
        "O marcador repeat: e todas as colunas devem estar na mesma linha modelo.",
        "Cada coluna deve começar com o ID da tabela, por exemplo itens.descricao.",
        "Depois do prefixo, use um ID de coluna simples; evite itens.dados.descricao.",
        "Não use células mescladas verticalmente na linha repetível.",
        "row.number só é válido dentro da linha que contém repeat:.",
    ):
        document.add_paragraph(text, style="List Bullet")

    document.add_heading("8. Tela Criar modelo > Campos e seções", level=1)
    document.add_heading("Aba Campos", level=2)
    add_tag_table(
        document,
        [
            ("ID do campo", "gestor.nome", "Deve ser igual ao ID usado entre {{ }}."),
            ("Rótulo", "Nome completo", "Texto mostrado ao usuário no formulário."),
            ("Tipo", "Texto / CNPJ / Data / ...", "Controle, máscara e validação."),
            ("Obrigatório", "Marcado", "Impede a geração quando o campo visível está vazio."),
            ("Seção", "Responsáveis", "Define o grupo recolhível do formulário."),
            ("Layout", "Automático / Grade / Largura total / Escolha / Tabela", "Define a posição e o agrupamento visual."),
        ],
    )
    document.add_paragraph(
        "O Modo simples oculta regras avançadas para reduzir a rolagem horizontal. Desmarque-o para editar Chave do perfil, Grupo, Escolha única e Visível quando."
    )

    document.add_heading("Aba Seções e layout", level=2)
    document.add_paragraph(
        "A árvore resume a ordem do formulário e mostra subgrupos de escolha e tabela. Use Nova seção, Renomear seção e Atribuir seleção à seção para reorganizar campos sem alterar as tags do DOCX."
    )

    document.add_heading("Aba Prévia do formulário", level=2)
    document.add_paragraph(
        "Revise a organização antes de salvar: seções recolhíveis, opções exclusivas em caixas clicáveis, campos em grade, campos de largura total e tabelas. A prévia não altera o DOCX e não salva dados."
    )

    document.add_heading("9. Regras condicionais", level=1)
    conditional_intro = document.add_paragraph(
        "Use Visível quando para mostrar um campo somente quando outro campo tiver um valor específico."
    )
    conditional_intro.paragraph_format.keep_with_next = True
    p = document.add_paragraph()
    p.paragraph_format.keep_with_next = True
    add_code(p, "contratacao.modalidade=Dispensa")
    document.add_paragraph(
        "Nesse exemplo, o campo configurado fica visível apenas quando contratacao.modalidade tiver o valor Dispensa. Para caixas, também são aceitos true/false, sim/não e 1/0."
    )

    document.add_heading("10. Tags no nome do arquivo e nas pastas", level=1)
    add_tag_table(
        document,
        [
            ("Nome do modelo", "{{template.name}}", "Nome cadastrado no Gerenciador de Modelos."),
            ("Ano atual", "{{year}}", "Ano com quatro dígitos."),
            ("Sequência", "{{sequence}}", "Contador do modelo quando a numeração está ativada."),
            ("Qualquer campo", "{{processo.numero}}", "Usa o valor preenchido no formulário."),
        ],
    )
    p = document.add_paragraph("Exemplo de arquivo: ")
    add_code(p, "{{processo.numero}} - {{empresa.nome}} - {{sequence}}.docx")
    p = document.add_paragraph("Exemplo de pastas: ")
    add_code(p, "{{year}}/{{processo.numero}}/{{empresa.nome}}")

    document.add_heading("11. Controles nativos do Word", level=1)
    document.add_paragraph(
        "O Padroniza também reconhece controles modernos do Word: caixa de seleção, seletor de data e lista suspensa. Em Desenvolvedor > Propriedades, informe uma Marca exclusiva igual ao ID do campo."
    )
    add_callout(
        document,
        "Importante",
        "Um controle nativo sem Marca não pode ser associado com segurança a um campo. O diagnóstico do modelo informa esse problema.",
        YELLOW,
    )

    document.add_heading("12. Diagnóstico rápido", level=1)
    troubleshooting = [
        ("O campo não foi localizado", "Verifique se a tag usa exatamente duas chaves de abertura e fechamento e se o ID não contém espaços."),
        ("A lista não mostra opções", "Use dropdown:campo.id|Opção A|Opção B e confirme que existe pelo menos uma opção."),
        ("Os campos ficaram fora da seção", "Aplique um estilo de título no Word ou preencha a coluna Seção no editor."),
        ("Duas caixas aparecem separadas", "Defina Layout = Grupo de escolha e use o mesmo grupo. As alternativas serão empilhadas como caixas grandes e exclusivas."),
        ("Responsáveis aparecem como cartões soltos", "Use Layout = Tabela e configure o mesmo grupo, linha e colunas, ou reorganize as tags em uma tabela do Word."),
        ("A tabela repetível não funciona", "Confirme repeat:, prefixo das colunas, ausência de mesclagem vertical e row.number na mesma linha."),
        ("O documento perdeu formatação", "Evite editar a formatação dentro das próprias chaves. A versão atual preserva os estilos ao redor das tags."),
    ]
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for index, label in enumerate(("Sintoma", "Correção")):
        cell = table.cell(0, index)
        set_cell_shading(cell, BLUE)
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
    for symptom, fix in troubleshooting:
        cells = table.add_row().cells
        cells[0].text = symptom
        cells[1].text = fix
        for cell in cells:
            set_cell_margins(cell)
    document.add_paragraph()

    document.add_heading("13. Checklist antes de salvar o modelo", level=1)
    checks = (
        "Todos os IDs são únicos e correspondem às tags do DOCX.",
        "Os rótulos são curtos e compreensíveis para o cliente.",
        "As seções seguem a ordem do documento.",
        "Escolhas exclusivas aparecem em um único grupo.",
        "Campos fixos de responsáveis aparecem em tabela.",
        "Campos longos usam largura total ou texto com várias linhas.",
        "A prévia do formulário está organizada em telas pequenas e grandes.",
        "O diagnóstico do modelo não apresenta erros bloqueadores.",
        "Um documento de teste foi gerado e revisado no Word.",
    )
    for text in checks:
        p = document.add_paragraph(style="List Bullet")
        p.add_run("☐ ").bold = True
        p.add_run(text)

    add_callout(
        document,
        "Resultado esperado",
        "O formulário deve seguir a mesma lógica do documento: títulos viram seções, alternativas viram um grupo de escolha e tabelas do Word viram tabelas de preenchimento.",
        GREEN,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_guide())
